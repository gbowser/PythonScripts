#!/usr/bin/env python3
"""Build a visual Word guide to the toy-object optimisation objective."""

from __future__ import annotations

from pathlib import Path
from datetime import date
import math

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import Circle, Rectangle, FancyBboxPatch
import numpy as np
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT_DIR = ROOT / "Foreground Masking" / "documentation" / "objective_metrics_artifact"
ASSET_DIR = OUT_DIR / "assets"
DOCX_PATH = OUT_DIR / "MTObjects Toy Object Optimisation - Objective Function and Metrics.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
PALE_BLUE = "EAF2F8"
PALE_GREEN = "EAF5EE"
PALE_RED = "FBECEE"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "687386"
BLACK = "111111"


def set_font(run, name="Calibri", size=11, bold=None, italic=None, color=BLACK):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths_inches):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    total_dxa = int(round(sum(widths_inches) * 1440))
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_dxa)); tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd"); tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120"); tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_inches:
        col = OxmlElement("w:gridCol"); col.set(qn("w:w"), str(int(round(width * 1440)))); grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths_inches):
            cell.width = Inches(width)
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW"); cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(int(round(width * 1440)))); tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_alt_text(shape, title, description):
    inline = shape._inline
    doc_pr = inline.docPr
    doc_pr.set("title", title)
    doc_pr.set("descr", description)


def add_page_field(paragraph):
    """Append a live Word PAGE field to a paragraph."""
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText"); instruction.set(qn("xml:space"), "preserve"); instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar"); separate.set(qn("w:fldCharType"), "separate")
    display = OxmlElement("w:t"); display.text = "1"
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, display, end])
    set_font(run, size=8.5, color=MID_GRAY)


def add_para(doc, text="", *, size=11, bold=False, italic=False, color=BLACK,
             align=None, before=0, after=6, keep=False):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    p.paragraph_format.keep_with_next = keep
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, italic=italic, color=color)
    return p


def add_mixed_para(doc, parts, *, before=0, after=6, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    if align is not None:
        p.alignment = align
    for text, kwargs in parts:
        r = p.add_run(text)
        set_font(r, **kwargs)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.add_run(text)
    return p


def add_callout(doc, heading, body, fill=PALE_BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [6.5])
    cell = table.cell(0, 0); set_cell_shading(cell, fill)
    p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(3)
    r = p.add_run(heading); set_font(r, size=11, bold=True, color=DARK_BLUE)
    p2 = cell.add_paragraph(); p2.paragraph_format.space_after = Pt(0); p2.paragraph_format.line_spacing = 1.10
    r2 = p2.add_run(body); set_font(r2, size=10.5)
    add_para(doc, "", after=2)


def add_equation(doc, lines):
    table = doc.add_table(rows=len(lines), cols=1)
    set_table_geometry(table, [6.5])
    for index, line in enumerate(lines):
        cell = table.cell(index, 0); set_cell_shading(cell, "F7F9FB")
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
        r = p.add_run(line); set_font(r, name="Cambria Math", size=11.5, color=DARK_BLUE)
    add_para(doc, "", after=1)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    hdr = table.rows[0]; set_repeat_header(hdr)
    for cell, text in zip(hdr.cells, headers):
        set_cell_shading(cell, LIGHT_GRAY)
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(str(text)); set_font(r, size=9.5, bold=True, color=DARK_BLUE)
    for values in rows:
        cells = table.add_row().cells
        for index, (cell, value) in enumerate(zip(cells, values)):
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if index == 0 else WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(value)); set_font(r, size=9.5)
    set_table_geometry(table, widths)
    add_para(doc, "", after=2)
    return table


def metric_figures():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)

    fig, ax = plt.subplots(figsize=(8.8, 3.2))
    ax.set_xlim(-.5, 9.5); ax.set_ylim(-.8, 2.0); ax.axis("off")
    detected = {0, 1, 3, 5, 7, 8}
    for i in range(10):
        color = "#2CA25F" if i in detected else "#B8C2CC"
        ax.add_patch(Circle((i, .75), .28, facecolor=color, edgecolor="#1B4D3E" if i in detected else "#687386", lw=2))
        ax.text(i, .1, str(i + 1), ha="center", va="center", fontsize=9, color="#555")
    ax.text(2.5, 1.55, "6 recovered toys", ha="center", fontsize=14, weight="bold", color="#237A46")
    ax.text(7.5, 1.55, "4 missed toys", ha="center", fontsize=14, weight="bold", color="#687386")
    ax.text(4.5, -.5, "Toy detection rate = 6 / 10 = 60%", ha="center", fontsize=15, weight="bold", color="#1F4D78")
    fig.tight_layout(); fig.savefig(ASSET_DIR / "toy_detection_rate.png", dpi=180, transparent=False); plt.close(fig)

    truth = np.zeros((11, 15), bool); yy, xx = np.indices(truth.shape)
    truth[((xx - 7) / 5) ** 2 + ((yy - 5) / 3.5) ** 2 <= 1] = True
    recovered = truth & (xx <= 8)
    canvas = np.zeros((*truth.shape, 3), float) + 1
    canvas[truth] = [0.75, 0.86, 0.97]
    canvas[recovered] = [0.20, 0.68, 0.42]
    fig, ax = plt.subplots(figsize=(8.8, 3.5))
    ax.imshow(canvas, interpolation="nearest")
    ax.contour(truth, levels=[.5], colors=["#2E74B5"], linewidths=2)
    ax.set_xticks([]); ax.set_yticks([])
    total = int(truth.sum()); overlap = int(recovered.sum())
    ax.set_title(f"Recovered truth pixels / injected truth pixels = {overlap} / {total} = {overlap/total:.0%}", fontsize=14, weight="bold", color="#1F4D78")
    ax.text(1, 9.6, "green = recovered", color="#237A46", fontsize=11, weight="bold")
    ax.text(10, 9.6, "blue = missed truth", color="#2E74B5", fontsize=11, weight="bold")
    fig.tight_layout(); fig.savefig(ASSET_DIR / "mean_toy_recall.png", dpi=180); plt.close(fig)

    fig, axes = plt.subplots(1, 3, figsize=(9.4, 3.2))
    fractions = [0.06, 0.10, 0.14]
    rng = np.random.default_rng(42)
    for ax, frac, label in zip(axes, fractions, ["Galaxy A", "Galaxy B", "Galaxy C"]):
        base = np.linspace(.85, .25, 400).reshape(20, 20)
        mask = np.zeros((20, 20), bool)
        count = int(round(frac * mask.size))
        mask.flat[rng.choice(mask.size, count, replace=False)] = True
        ax.imshow(base, cmap="gray", vmin=0, vmax=1)
        overlay = np.zeros((*mask.shape, 4)); overlay[mask] = [0.9, 0.15, 0.18, .75]
        ax.imshow(overlay)
        ax.set_xticks([]); ax.set_yticks([]); ax.set_title(f"{label}\n{frac:.0%} masked", fontsize=12)
    fig.suptitle("Maximum displayed-frame masking = max(6%, 10%, 14%) = 14%", fontsize=14, weight="bold", color="#1F4D78")
    fig.tight_layout(rect=[0, 0, 1, .84]); fig.savefig(ASSET_DIR / "maximum_masking.png", dpi=180); plt.close(fig)

    # Pixel-classification illustration used to explain F and FP.
    truth = np.zeros((12, 18), bool); yy, xx = np.indices(truth.shape)
    truth[((xx - 7) / 5.2) ** 2 + ((yy - 6) / 3.5) ** 2 <= 1] = True
    candidate = np.zeros_like(truth)
    candidate[((xx - 8.5) / 5.0) ** 2 + ((yy - 5.5) / 3.2) ** 2 <= 1] = True
    tp = truth & candidate; fn = truth & ~candidate; fp = candidate & ~truth; tn = ~truth & ~candidate
    canvas = np.ones((*truth.shape, 3), float)
    canvas[fn] = [0.65, 0.80, 0.95]
    canvas[fp] = [0.94, 0.35, 0.35]
    canvas[tp] = [0.20, 0.68, 0.42]
    fig, ax = plt.subplots(figsize=(8.8, 3.7))
    ax.imshow(canvas, interpolation="nearest")
    ax.contour(truth, levels=[.5], colors=["#2E74B5"], linewidths=2)
    ax.set_xticks([]); ax.set_yticks([])
    n_tp, n_fn, n_fp, n_tn = int(tp.sum()), int(fn.sum()), int(fp.sum()), int(tn.sum())
    precision = n_tp / (n_tp + n_fp); recall = n_tp / (n_tp + n_fn)
    f_score = 2 * precision * recall / (precision + recall)
    ax.set_title(f"TP={n_tp}, FN={n_fn}, FP={n_fp}  |  precision={precision:.0%}, recall={recall:.0%}, F={f_score:.0%}", fontsize=13, weight="bold", color="#1F4D78")
    fig.text(.10, .06, "green = TP: truth masked", color="#237A46", fontsize=9.5, weight="bold", ha="left")
    fig.text(.53, .06, "blue = FN: truth missed", color="#2E74B5", fontsize=9.5, weight="bold", ha="left")
    fig.text(.10, .015, "red = FP: non-truth masked", color="#C0392B", fontsize=9.5, weight="bold", ha="left")
    fig.text(.53, .015, f"white = TN: non-truth unmasked ({n_tn} pixels)", color="#555555", fontsize=9.5, weight="bold", ha="left")
    fig.tight_layout(rect=[0, .14, 1, 1]); fig.savefig(ASSET_DIR / "f_score_false_positive.png", dpi=180); plt.close(fig)

    fig, ax = plt.subplots(figsize=(10, 4.2)); ax.set_xlim(0, 10); ax.set_ylim(0, 4); ax.axis("off")
    boxes = [
        (0.2, 2.55, 2.7, 1.0, "Recovery floors", "Detection ≥ 25%\nMean recall ≥ 20%", "#EAF2F8", "#2E74B5"),
        (3.65, 2.55, 2.7, 1.0, "Masking cap", "Maximum mask ≤ 15%", "#FBECEE", "#C0392B"),
        (7.1, 2.55, 2.7, 1.0, "Quality score", "Recovery − data loss", "#EAF5EE", "#237A46"),
    ]
    for x, y, w, h, title, body, fill, edge in boxes:
        ax.add_patch(FancyBboxPatch((x, y), w, h, boxstyle="round,pad=.08", facecolor=fill, edgecolor=edge, linewidth=2))
        ax.text(x+w/2, y+.68, title, ha="center", fontsize=12, weight="bold", color=edge)
        ax.text(x+w/2, y+.28, body, ha="center", fontsize=10, color="#222")
    ax.annotate("", xy=(3.48,3.05), xytext=(3.0,3.05), arrowprops=dict(arrowstyle="->", lw=2, color="#687386"))
    ax.annotate("", xy=(6.95,3.05), xytext=(6.45,3.05), arrowprops=dict(arrowstyle="->", lw=2, color="#687386"))
    ax.add_patch(FancyBboxPatch((2.25, .45), 5.5, 1.05, boxstyle="round,pad=.08", facecolor="#F7F9FB", edgecolor="#1F4D78", linewidth=2))
    ax.text(5, 1.08, "Piecewise objective supplied to Optuna", ha="center", fontsize=13, weight="bold", color="#1F4D78")
    ax.text(5, .72, "large penalties for infeasibility; otherwise minimise −score", ha="center", fontsize=10.5)
    for x in (1.55, 5.0, 8.45):
        ax.annotate("", xy=(5,1.52), xytext=(x,2.48), arrowprops=dict(arrowstyle="->", lw=1.6, color="#687386"))
    fig.tight_layout(); fig.savefig(ASSET_DIR / "objective_flow.png", dpi=180); plt.close(fig)


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.78); section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(1.0); section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.3); section.footer_distance = Inches(0.35)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"; normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri"); normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0); normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.10
    for level, size, color, before, after in ((1,16,BLUE,16,8),(2,13,BLUE,12,6),(3,12,DARK_BLUE,8,4)):
        style = styles[f"Heading {level}"]
        style.font.name = "Calibri"; style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri"); style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size); style.font.bold = True; style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before); style.paragraph_format.space_after = Pt(after); style.paragraph_format.keep_with_next = True
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = header.add_run("Foreground Masking Research  |  Technical Reference")
    set_font(r, size=8.5, color=MID_GRAY)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("MTObjects toy-object optimisation objective  |  Page ")
    set_font(r, size=8.5, color=MID_GRAY)
    add_page_field(footer)


def build_document():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    metric_figures()
    doc = Document(); configure_document(doc)

    add_para(doc, "TECHNICAL REFERENCE", size=10, bold=True, color=BLUE, after=3)
    add_para(doc, "MTObjects Toy-Object Optimisation", size=25, bold=True, color=BLACK, after=2)
    add_para(doc, "Objective Function and Core Metrics", size=15, color=MID_GRAY, after=14)
    add_mixed_para(doc, [
        ("Purpose: ", {"size":10.5,"bold":True,"color":DARK_BLUE}),
        ("Explain how recovery and masking are measured, combined and supplied to Optuna.", {"size":10.5,"color":BLACK}),
    ], after=2)
    add_mixed_para(doc, [
        ("Methodology basis: ", {"size":10.5,"bold":True,"color":DARK_BLUE}),
        ("paired-toy-metrics-displayed-frame-v2; MTObjects detection on the original centred image.", {"size":10.5,"color":BLACK}),
    ], after=2)
    add_mixed_para(doc, [
        ("Prepared: ", {"size":10.5,"bold":True,"color":DARK_BLUE}),
        (date.today().strftime("%d %B %Y"), {"size":10.5,"color":BLACK}),
    ], after=14)
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(16)
    p_pr = p._p.get_or_add_pPr(); p_bdr = OxmlElement("w:pBdr"); bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "18"); bottom.set(qn("w:color"), BLUE); p_bdr.append(bottom); p_pr.append(p_bdr)

    add_callout(doc, "In one sentence", "Optuna searches for MTObjects parameters that recover injected toys, cover their truth pixels, and avoid masking too much of any displayed galaxy frame.", PALE_BLUE)
    add_heading(doc, "1. Overview", 1)
    add_para(doc, "Synthetic foreground objects (toys) are inserted at known positions in clean galaxy images. MTObjects produces a mask, and the mask is compared with the known injected truth. The optimisation is deliberately constrained: a configuration is not considered scientifically useful merely because it masks every toy if it also removes substantial galaxy structure.")
    add_para(doc, "The calculation uses incremental masking: the mask obtained from the toy-injected image is compared with the baseline mask from the corresponding image without toys. Only newly masked pixels contribute to toy recovery. This prevents pre-existing detections from being credited as successful toy recovery.")
    add_table(doc,
        ["Symbol", "Meaning", "Aggregation"],
        [
            ("D", "Toy detection rate", "All toys across all training galaxy/seed cases"),
            ("R̄toy", "Mean per-toy recall", "Mean of individual toy recalls"),
            ("Mmax", "Maximum displayed-frame masked fraction", "Worst galaxy/seed case"),
            ("M̄", "Mean displayed-frame masked fraction", "Mean across cases"),
            ("F̄", "Mean pixel-level F-score", "Mean across cases"),
            ("FP̄", "Mean false-positive masked fraction", "Mean across cases"),
        ], [0.75, 3.05, 2.70])

    add_heading(doc, "2. Metric 1 — Toy detection rate", 1)
    add_equation(doc, ["D = number of toys with incremental recall ≥ 0.50  /  total number of toys"])
    add_para(doc, "A toy counts as detected only when at least half of its known truth-mask pixels are covered by the incremental MTObjects mask. Detection is therefore a per-object pass/fail measure; a toy with 49% recall is missed, while one with 50% recall is detected.")
    inline_shape = doc.add_picture(str(ASSET_DIR / "toy_detection_rate.png"), width=Inches(6.25))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_alt_text(inline_shape, "Toy detection illustration", "Ten toy symbols: six green recovered toys and four grey missed toys, giving a 60 percent detection rate.")
    add_para(doc, "Figure 1. Toy detection rate counts recovered objects, not recovered pixels.", size=9.5, italic=True, color=MID_GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    add_callout(doc, "Interpretation", "Detection rate answers: “How many injected objects were recovered well enough to count?” It does not distinguish a barely detected toy (50% recall) from a nearly complete one (95% recall).", PALE_GREEN)

    add_heading(doc, "3. Metric 2 — Mean toy recall", 1)
    add_equation(doc, [
        "Recalli = incremental mask pixels inside toy i truth  /  truth pixels belonging to toy i",
        "R̄toy = (Recall1 + Recall2 + … + RecallN) / N",
    ])
    add_para(doc, "Recall is continuous from 0 to 1. It measures how completely each injected toy is covered. The mean is calculated across toys, so every toy contributes equally regardless of its pixel area. This complements detection rate by rewarding more complete masks even after the 50% detection boundary has been crossed.")
    inline_shape = doc.add_picture(str(ASSET_DIR / "mean_toy_recall.png"), width=Inches(6.2)); doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_alt_text(inline_shape, "Toy recall illustration", "An elliptical toy truth mask outlined in blue. Green pixels are recovered and pale blue pixels are missed; recall is recovered truth pixels divided by all truth pixels.")
    add_para(doc, "Figure 2. Per-toy recall measures completeness of coverage within the known injected truth.", size=9.5, italic=True, color=MID_GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    add_callout(doc, "Why both D and R̄toy?", "Detection rate prevents many weak partial overlaps from appearing successful. Mean recall then distinguishes parameter sets that detect the same number of toys but recover different proportions of their shapes.", PALE_GREEN)

    add_heading(doc, "4. Metric 3 — Maximum displayed-frame masking", 1)
    add_equation(doc, [
        "Mj = incremental masked pixels inside displayed analysis frame j  /  finite pixels in that frame",
        "Mmax = maximum(M1, M2, …, MJ)",
    ])
    add_para(doc, "Masking is measured only inside the displayed, deprojected, galaxy-centred analysis square—the same area used for visual cleanliness review and toy placement. Pixels elsewhere in a larger FITS mosaic do not dilute the percentage.")
    add_para(doc, "The objective uses the maximum across galaxy/seed cases, not the mean, as its hard 15% constraint. A method cannot compensate for severe damage to one galaxy by masking very little in the others.")
    inline_shape = doc.add_picture(str(ASSET_DIR / "maximum_masking.png"), width=Inches(6.3)); doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_alt_text(inline_shape, "Maximum masking illustration", "Three displayed galaxy frames with red masked pixels at 6, 10 and 14 percent. The maximum metric is the worst value, 14 percent.")
    add_para(doc, "Figure 3. The worst displayed-frame masking controls feasibility.", size=9.5, italic=True, color=MID_GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    add_callout(doc, "Current constraint", "Training configurations must satisfy Mmax ≤ 0.15. This is a safeguard against an apparently good average result that destroys structure in a single galaxy.", PALE_RED)

    add_heading(doc, "5. Supporting recovery metric — Pixel-level F-score", 1)
    add_equation(doc, [
        "Precision P = TP / (TP + FP)",
        "Pixel recall R = TP / (TP + FN)",
        "F = 2PR / (P + R) = 2TP / (2TP + FP + FN)",
        "F̄ = mean of the per-case F-scores",
    ])
    add_para(doc, "For one galaxy/seed case, TP is an incremental masked pixel that lies inside the combined toy-truth mask. FP is an incremental masked pixel outside that truth. FN is a toy-truth pixel that the incremental mask missed—the blue region in Figure 4. TN is a non-toy pixel that correctly remains unmasked—the white area outside both the blue truth boundary and the red/green mask. Precision asks how much of the new mask is relevant; pixel recall asks how much of the combined toy truth was covered.")
    add_callout(doc, "Why true negatives are absent from F", "TN does not appear in precision, recall or F. The displayed frame contains very many ordinary non-toy pixels, so a model could obtain an apparently excellent accuracy simply by leaving most of the frame unmasked while detecting no toys. F deliberately concentrates on the positive class: toy pixels that should be masked.", PALE_BLUE)
    add_para(doc, "The F-score is their harmonic mean. It is high only when both precision and recall are high: aggressive masking cannot obtain a good F merely by covering all toy pixels, and a tiny but perfectly placed mask cannot obtain a good F while missing most of the toys. F is calculated separately for every galaxy/seed case and then averaged, so F̄ gives each case equal weight.")
    inline_shape = doc.add_picture(str(ASSET_DIR / "f_score_false_positive.png"), width=Inches(5.65)); doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_alt_text(inline_shape, "Pixel classification for F-score and false positives", "Green pixels are true positives inside toy truth, blue pixels are false negatives where toy truth was missed, red pixels are false positives outside toy truth, and white pixels are true negatives outside the toy and mask.")
    add_para(doc, "Figure 4. The same pixel classification supplies precision, pixel recall, F and FP.", size=9.5, italic=True, color=MID_GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    add_callout(doc, "Difference from mean toy recall", "Pixel recall pools all toy-truth pixels in a case, so larger toys contribute more pixels. Mean toy recall first calculates recall for each toy and then averages the toys equally. F additionally includes precision, whereas mean toy recall does not.", PALE_GREEN)

    add_heading(doc, "6. Supporting data-loss metric — False-positive masked fraction", 1)
    add_equation(doc, [
        "FP fraction = incremental masked pixels outside toy truth / non-toy pixels in the displayed frame",
        "FP̄ = mean of the per-case FP fractions",
    ])
    add_para(doc, "The numerator is the red area in Figure 4: newly masked pixels that cannot be attributed to an injected toy. The denominator is every non-toy pixel in the displayed analysis frame, equal to displayed-frame pixels minus toy-truth pixels. The metric therefore estimates the fraction of available galaxy/background area removed unnecessarily.")
    add_para(doc, "FP fraction is not the same as 1 − precision. One minus precision divides false-positive pixels by all incremental masked pixels; the implemented FP fraction divides them by all non-toy pixels in the displayed frame. For example, if 30 of 100 new mask pixels are outside the toys, precision is 70% and 1 − precision is 30%. If the frame contains 10,000 non-toy pixels, however, the implemented FP fraction is 30/10,000 = 0.3%.")
    add_callout(doc, "Why include both M̄ and FP̄?", "M̄ measures the total incremental mask as a fraction of the displayed frame, including useful toy coverage. FP̄ isolates masking outside the known toys. Their penalties overlap deliberately but answer different questions: total image loss versus demonstrably irrelevant image loss.", PALE_RED)

    add_heading(doc, "7. The complete objective supplied to Optuna", 1)
    inline_shape = doc.add_picture(str(ASSET_DIR / "objective_flow.png"), width=Inches(6.45)); doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_alt_text(inline_shape, "Piecewise objective flow", "Recovery floors, the maximum masking cap and the quality score feed a piecewise objective minimised by Optuna.")
    add_para(doc, "Figure 5. Feasibility is established before the ordinary quality score can win.", size=9.5, italic=True, color=MID_GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)

    add_heading(doc, "7.1 Recovery and data-loss scores", 2)
    add_equation(doc, [
        "Recovery score = 0.45 F̄ + 0.35 R̄toy + 0.20 D",
        "Data loss = 0.50 M̄ + 0.10 min(FP̄, 1)",
        "Score S = recovery score − data loss",
    ])
    add_para(doc, "The recovery score gives its largest weight (45%) to F̄ because F rewards accurate and complete pixel coverage simultaneously. Mean toy recall contributes 35%, ensuring every toy matters equally, and detection rate contributes 20%, rewarding the number that cross the 50% recovery threshold. In data loss, M̄ has coefficient 0.50 and FP̄ has coefficient 0.10; FP̄ is capped at 1 defensively. These terms refine the ranking once the recovery floors and masking cap have been considered.")

    add_heading(doc, "7.2 Piecewise minimisation objective", 2)
    add_equation(doc, [
        "Let δD = max(0, 0.25 − D),  δR = max(0, 0.20 − R̄toy),  δM = max(0, Mmax − 0.15)",
        "If no incremental pixels, D < 0.25, or R̄toy < 0.20:  Objective = 50 + 20δD + 20δR  (+50 if no incremental pixels)",
        "Else if Mmax > 0.15:  Objective = 10 + 100δM + data loss − recovery score",
        "Else:  Objective = −S",
    ])
    add_para(doc, "Optuna minimises this objective. The large constants create three regimes. Recovery-infeasible configurations are worst. Configurations that recover enough toys but exceed the masking cap are still penalised. Only fully feasible configurations receive the negative quality score, so a good feasible solution is preferred to any cap-violating one.")

    add_table(doc,
        ["Regime", "Condition", "Effect on search"],
        [
            ("Recovery infeasible", "D < 25%, R̄toy < 20%, or no incremental mask", "Large objective beginning at 50"),
            ("Mask infeasible", "Recovery floors met but Mmax > 15%", "Objective beginning at 10 plus 100× cap excess"),
            ("Feasible", "Recovery floors and cap all met", "Objective = −score; best score becomes most negative"),
        ], [1.35, 2.65, 2.50])

    add_heading(doc, "8. Worked example", 1)
    add_para(doc, "Assume a parameter set produces D = 0.32, R̄toy = 0.29, F̄ = 0.24, M̄ = 0.07, FP̄ = 0.05 and Mmax = 0.14. It meets both recovery floors and the 15% cap.")
    add_equation(doc, [
        "Recovery score = 0.45(0.24) + 0.35(0.29) + 0.20(0.32) = 0.2735",
        "Data loss = 0.50(0.07) + 0.10(0.05) = 0.0400",
        "S = 0.2735 − 0.0400 = 0.2335; therefore Objective = −0.2335",
    ])
    add_para(doc, "If the same result had Mmax = 0.18, it would enter the mask-infeasible regime. Its cap excess would be 0.03 and the objective would become 10 + 100(0.03) + 0.0400 − 0.2735 = 12.7665. This is deliberately much worse than any feasible negative objective.")

    add_heading(doc, "9. Training objective versus held-out acceptance", 1)
    add_callout(doc, "Important distinction", "The optimisation’s recovery floors are not the final scientific success criteria. They prevent unproductive trials; the untouched validation seeds test a stricter target.", PALE_RED)
    add_table(doc,
        ["Measure", "Training feasibility floor", "Held-out acceptance target"],
        [
            ("Toy detection rate", "≥ 25%", "≥ 50%"),
            ("Mean toy recall", "≥ 20%", "≥ 30%"),
            ("Maximum displayed-frame masking", "≤ 15%", "≤ 15%"),
            ("Galaxy/seed cases with ≥1 recovered toy", "Not a training gate", "≥ 80% of cases"),
        ], [2.55, 1.85, 2.10])
    add_para(doc, "The distinction explains how Optuna can find a valid training configuration while the final experiment still fails the 50% held-out detection goal. A configuration that generalises at approximately 25% detection can satisfy the optimisation floor but not the scientific acceptance test.")

    add_heading(doc, "10. Practical interpretation", 1)
    bullets = [
        "High detection with excessive Mmax indicates aggressive masking rather than a usable solution.",
        "Low detection and low recall indicate that MTObjects is not recovering enough injected structure.",
        "Adequate detection but low mean recall indicates marginal, incomplete toy coverage.",
        "A low mean masked fraction does not override a cap violation in a single displayed frame.",
        "High recall with low precision produces a modest F-score and indicates an over-broad mask.",
        "A high FP̄ means the method is removing a substantial fraction of non-toy image area.",
        "A negative objective value indicates a fully feasible training result; more negative is better.",
    ]
    for text in bullets:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(.5); p.paragraph_format.first_line_indent = Inches(-.25)
        p.paragraph_format.space_after = Pt(8); p.paragraph_format.line_spacing = 1.167
        r = p.add_run(text); set_font(r, size=11)

    add_heading(doc, "11. Methodological source", 1)
    add_para(doc, "Definitions and coefficients in this guide are taken from the implemented code used for the revised 22-clean-galaxy, displayed-frame optimisation:")
    source_rows = [
        ("Metric implementation", "Foreground Masking/Optimisation/paired_toy_common.py — evaluate_mask"),
        ("Objective implementation", "Foreground Masking/Optimisation/optimise_toy_objects_MTObjects.py — aggregate_score"),
        ("Held-out acceptance", "Foreground Masking/Optimisation/evaluate_mtobjects_multiseed_winner.py"),
        ("Metric version", "paired-toy-metrics-displayed-frame-v2"),
        ("Recorded software revision", "ebf8733bf8988733346e4797b0e7498bf03caf8f"),
    ]
    add_table(doc, ["Item", "Implementation source"], source_rows, [1.75, 4.75])
    # add_table leaves a spacer paragraph; remove it at document end so Word
    # cannot push a mandatory trailing paragraph onto an otherwise blank page.
    trailing = doc.paragraphs[-1]
    if not trailing.text:
        trailing._element.getparent().remove(trailing._element)
    core = doc.core_properties
    core.title = "MTObjects Toy-Object Optimisation — Objective Function and Core Metrics"
    core.subject = "Technical explanation of toy detection, recall, displayed-frame masking and the Optuna objective"
    core.author = "MSc Research — Foreground Masking"
    core.keywords = "MTObjects, Optuna, toy objects, masking, recall, detection rate"
    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build_document()
