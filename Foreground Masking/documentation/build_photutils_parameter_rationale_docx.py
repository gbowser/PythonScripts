from __future__ import annotations

import importlib.util
import shutil
import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


SCRIPT_DIR = Path(__file__).resolve().parents[1]
PROJECT_ROOT = SCRIPT_DIR.parent
if str(PROJECT_ROOT) not in sys.path:
    sys.path.append(str(PROJECT_ROOT))

from machine_paths import remove_foreground_folder  # noqa: E402

DOC_DIR = SCRIPT_DIR / "documentation"
OUTPUT_DOCX = DOC_DIR / "Photutils Parameter Rationale.docx"
DROPBOX_DOC_DIR = remove_foreground_folder("Laptop") / "documentation"

template_path = (
    SCRIPT_DIR.parent
    / "Shoulder Recognition Erwin"
    / "documentation"
    / "build_shoulder_quantification_docs.py"
)
spec = importlib.util.spec_from_file_location("shoulder_docs_template", template_path)
template = importlib.util.module_from_spec(spec)
assert spec.loader is not None
spec.loader.exec_module(template)

style_document = template.style_document
set_paragraph_border_bottom = template.set_paragraph_border_bottom
add_text = template.add_text
add_bullets = template.add_bullets
add_table = template.add_table
MUTED = template.MUTED


def add_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = "Photutils Parameter Rationale"
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = MUTED
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.text = "Generated documentation"
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = MUTED


def add_intro(doc: Document) -> None:
    doc.add_paragraph("Photutils Parameter Rationale", style="Title")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run(f"Updated {date.today().isoformat()}").italic = True
    set_paragraph_border_bottom(subtitle)
    add_text(
        doc,
        "This note explains why the current Photutils masking parameters exist and how they should be interpreted in the global and spike-gated workflows. "
        "It replaces the older PDF-only rationale document with a Word document so the documentation set is DOCX-only.",
    )


def add_methods(doc: Document) -> None:
    doc.add_heading("Masking methods", level=1)
    add_table(
        doc,
        ["Method", "Masking rule", "Consequence"],
        [
            [
                "Global",
                "Photutils residual detections that pass compact-source filters become mask pixels.",
                "More complete foreground removal, but higher risk of masking real galaxy structure.",
            ],
            [
                "Spike-gated",
                "Photutils detections are candidates only; a candidate must intersect detected bar-major profile spike samples.",
                "More conservative for bar-profile science and safer for no-spike control galaxies.",
            ],
        ],
        [1500, 4300, 3560],
    )


def add_parameters(doc: Document) -> None:
    doc.add_heading("Parameter rationale", level=1)
    add_table(
        doc,
        ["Parameter", "Role", "Optimisation logic"],
        [
            ["nsigma", "Residual detection threshold.", "Lower values find fainter candidates; global mode usually needs stricter values than spike-gated mode."],
            ["dilations", "Expands accepted candidate masks.", "Increase when candidate cores are found but wings remain; avoid excessive profile bridging."],
            ["max_area", "Rejects segments larger than the compact-source limit.", "Keep low enough to avoid galaxy structure; raise only for broad contaminant footprints."],
            ["max_elongation", "Rejects elongated segments.", "Protects against bars, arms, and residual ridges being treated as compact foreground objects."],
            ["smooth_sigma_pixels", "Builds the smooth galaxy model for residual detection.", "Too small follows galaxy structure; too large leaves broad galaxy residuals."],
            ["npixels", "Minimum connected-pixel detection size.", "Suppresses single-pixel noise and tiny artifacts; not converted to arcsec in the GUI."],
            ["central_exclusion", "Excludes candidates near the deprojected galaxy centre.", "Protects nuclei, rings, and inner bar structure from being masked."],
        ],
        [1900, 3300, 4160],
    )


def add_units(doc: Document) -> None:
    doc.add_heading("Pixels versus arcsec", level=1)
    add_text(
        doc,
        "The interactive tester can display pixel-based controls in Pixels or Arcsec. "
        "This is a display/input convenience only: before masking, the values are converted back to pixels using the selected galaxy's pixel scale.",
    )
    add_bullets(
        doc,
        [
            "Linear controls use value_pixels x pixel_scale when displayed in arcsec.",
            "max_area uses value_pixels x pixel_scale^2 when displayed in square arcsec.",
            "npixels remains a connected-pixel count because that is the Photutils detection definition.",
            "Spike-neighbour and spike-centre controls are already arcsec quantities and are not affected by the unit switch.",
        ],
    )


def add_output_names(doc: Document) -> None:
    doc.add_heading("Current output naming", level=1)
    add_table(
        doc,
        ["Workflow", "Filename stem"],
        [
            ["Global production reports", "{galaxy_name}_fg_removed_global"],
            ["Spike-gated production reports", "{galaxy_name}_fg_removed_sp-gated"],
            ["Interactive tester PNGs", "{galaxy}_{pc}_{method}_nsigma{value}_dil{value}_area{value}"],
        ],
        [2800, 6560],
    )


def add_practical_guidance(doc: Document) -> None:
    doc.add_heading("Practical optimisation guidance", level=1)
    add_bullets(
        doc,
        [
            "Optimise against both spike-positive galaxies and no-spike controls; no-spike controls are the main guard against overmasking.",
            "Prefer the most conservative parameter set that removes the visible profile spike without creating long artificial bridge sections.",
            "Treat global mode as a comparison and stress test unless it has been visually checked for the target science sample.",
            "Use the interactive tester to inspect individual failures before expanding the global optimisation grid.",
        ],
    )


def build() -> Path:
    DOC_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_document(doc)
    add_footer(doc)
    add_intro(doc)
    add_methods(doc)
    add_parameters(doc)
    add_units(doc)
    add_output_names(doc)
    add_practical_guidance(doc)
    for style_name in ("Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3"):
        doc.styles[style_name].paragraph_format.keep_with_next = True
    doc.core_properties.title = "Photutils Parameter Rationale"
    doc.core_properties.subject = "Rationale for global and spike-gated Photutils masking parameters"
    doc.core_properties.author = ""
    doc.save(OUTPUT_DOCX)
    return OUTPUT_DOCX


if __name__ == "__main__":
    docx = build()
    print(docx)
    DROPBOX_DOC_DIR.mkdir(parents=True, exist_ok=True)
    shutil.copy2(docx, DROPBOX_DOC_DIR / docx.name)
    print(DROPBOX_DOC_DIR)
