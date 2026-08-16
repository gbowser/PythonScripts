from pathlib import Path
from copy import deepcopy
import shutil

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parent / "folder_wide_review"
SUMMARY = Path(__file__).resolve().parent / "SEP and MTObjects Toy Objects Optimisation Summary and Conclusions.docx"
DATE = "16 August 2026"
BLUE = RGBColor(46, 116, 181)
GREY = RGBColor(90, 90, 90)


def set_font(run, size=None, bold=None, color=None, italic=None):
    name = "Calibri"
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    if italic is not None:
        run.italic = italic


def replace_paragraph_text(paragraph, old, new):
    if old not in paragraph.text:
        return False
    text = paragraph.text.replace(old, new)
    for run in paragraph.runs:
        run.text = ""
    run = paragraph.add_run(text)
    set_font(run, 11)
    return True


def replace_everywhere(doc, old, new):
    count = 0
    for paragraph in doc.paragraphs:
        count += int(replace_paragraph_text(paragraph, old, new))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    count += int(replace_paragraph_text(paragraph, old, new))
    return count


def add_review_section(doc, classification, changes, implications):
    heading = doc.add_paragraph(style="Heading 1")
    heading.paragraph_format.page_break_before = False
    heading.paragraph_format.keep_with_next = True
    run = heading.add_run(f"Systematic documentation review - {DATE}")
    set_font(run, 16, True, BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run("Review classification: ")
    set_font(r, 10.5, True)
    r = p.add_run(classification)
    set_font(r, 10.5)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run("Updates made: ")
    set_font(r, 10.5, True)
    r = p.add_run(changes)
    set_font(r, 10.5)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Current interpretation: ")
    set_font(r, 10.5, True)
    r = p.add_run(implications)
    set_font(r, 10.5)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(
        "Authoritative combined outcome: SEP and MTObjects Toy Objects Optimisation Summary and Conclusions.docx. "
        "Full-precision results: Foreground Masking Optimisation Results.xlsx."
    )
    set_font(r, 9.5, False, GREY, True)


REVIEWS = {
    "Azure Optimisation Setup and Performance Report 20260730.docx": (
        "Reviewed - historical infrastructure and performance report",
        "No historical measurements were altered. A dated scope note was added so the Azure benchmark is not mistaken for the August 2026 local four-fold Toy Objects run.",
        "Use this document for Azure setup, Linux compatibility, transfer and historical performance only. Use the combined outcome document for the final SEP and MTObjects Toy Objects conclusions.",
    ),
    "Bar Spike-Gated Foreground Candidate Report Documentation.docx": (
        "Reviewed - method reference remains valid",
        "No algorithmic changes were required. The document describes the bar Spike Gate candidate-report workflow, not the Toy Objects optimisation or its eight-panel batches.",
        "The Spike Gate definitions remain applicable. PhotUtils-folder material was excluded from this review as requested; this root-level report was reviewed because it is part of the main documentation set.",
    ),
    "Foreground Mask SourceXtractor Documentation.docx": (
        "Reviewed - separate SourceXtractor++ method reference",
        "No Toy Objects result claims or obsolete MTObjects conclusions were found. A scope note was added.",
        "This document remains a SourceXtractor++ implementation reference and should not be used as evidence for the SEP/MTObjects Toy Objects parameter selection.",
    ),
    "Foreground Masking Canonical Tools Reference.docx": (
        "Updated - canonical launch paths and batch-output conventions",
        "Example commands were corrected to include the Interactive tools and Batch tools folders. The automatic `_clean` suffix for the validated 40 calibration galaxies and the aligned eight-panel Toy Objects batch layout were documented.",
        "Future SEP and MTObjects Toy Objects batches read CleanGalaxies.txt, require 40 unique names, and add `_clean` before `.png` for those calibration galaxies.",
    ),
    "Foreground Masking Four Optimisation Objective Functions.docx": (
        "Reviewed - current objective-function authority",
        "The recovery feasibility gate, empty-mask penalty, 15% cap and final MTObjects cross-fold release gate were verified against the implemented follow-up design. No formula change was required.",
        "This is the authoritative mathematical description of the four objective families. Numeric objective values are comparable only within the same method and weight configuration.",
    ),
    "interactive_mtobjects_no_spike_gate_algorithm_and_parameters.docx": (
        "Reviewed - interactive baseline reference remains valid",
        "No optimisation-result claims were found. A dated scope note was added without changing the algorithm or parameter descriptions.",
        "This document describes interactive MTObjects without Spike Gate. The selected Toy Objects production parameters are recorded in the combined outcome document, not here.",
    ),
    "interactive_mtobjects_no_spike_gate_code_process_and_flow.docx": (
        "Reviewed - code-flow reference remains valid",
        "No stale zero-mask conclusion, six-panel batch claim or obsolete optimisation result was found. A scope note was added.",
        "The code-flow description remains useful for maintenance; final Toy Objects parameter selection and batch status are maintained separately.",
    ),
    "interactive_toy_objects_MTObjects.docx": (
        "Updated - interactive/batch output distinction clarified",
        "The interactive scientific workflow remains valid. A note now distinguishes its live diagnostic panels from the canonical aligned eight-panel batch PNG and records the selected recovery follow-up result.",
        "The recovery follow-up restores non-zero detections, but the selected MTObjects candidate has low pixel precision and substantial collateral masking. Visual review remains mandatory.",
    ),
    "MTObjects Toy Objects Four Fold Optimisation Documentation.docx": (
        "Updated - recovery objective, release gates and final winner",
        "The obsolete soft score was replaced by the implemented recovery reward/loss, infeasibility penalty and final release gates. Fold-3 selected parameters and common-40 metrics were recorded.",
        "MTObjects is no longer a zero-mask solution. It is a feasible secondary method, but mean precision of 0.31% and false-positive fraction of 9.05% prevent recommendation as the primary production mask.",
    ),
    "SEP and MTObjects Optuna Optimisation Process - Updated.docx": (
        "Updated - current Toy Objects optimisation process",
        "The shared pre-follow-up Toy Objects score was replaced by separate SEP and MTObjects objective definitions. Four-fold candidate evaluation, recovery feasibility and release-gate behaviour were added.",
        "SEP remains the primary result. MTObjects follow-up optimisation is valid but should be used as a review-controlled comparison because recovery is accompanied by substantial collateral masking.",
    ),
    "SEP Spike Gate Optuna Optimisation.docx": (
        "Reviewed - SEP Spike Gate reference remains valid",
        "No Toy Objects objective or batch-layout claim required correction. A dated scope note was added.",
        "This document remains specific to SEP Spike Gate optimisation and does not supersede the Toy Objects four-fold results.",
    ),
    "Spike Gate Samples Explained.docx": (
        "Reviewed - explanatory reference remains valid",
        "Definitions of spike samples, side-drop tests, neighbour baselines and central exclusion were reviewed; no Toy Objects-specific correction was required.",
        "Use this document to interpret Spike Gate samples only. It is independent of the final Toy Objects method comparison.",
    ),
    "Toy Object Parameters Guide.docx": (
        "Reviewed - Toy Object construction reference remains valid",
        "Object morphology, placement, truth-mask and recovery definitions were reviewed. A dated link to the final optimisation outcome was added.",
        "The guide defines injected objects and metrics; it does not prescribe the final SEP or MTObjects production parameter set.",
    ),
    "Toy Objects Four Fold Cross Validation Documentation.docx": (
        "Updated - current objectives, selection gates and eight-panel output",
        "The MTObjects recovery feasibility penalty and final release gate were added, and the obsolete six-panel statement was replaced by the required aligned eight-panel report specification. Final selected folds were recorded.",
        "The design is complete and statistically defensible: four rotations of 30 training and 10 held out, followed by common-40 candidate comparison. SEP fold 4 and MTObjects fold 3 are the selected solutions.",
    ),
}


def update_canonical(doc):
    for method in ("interactive_spike_gate_SEP.py", "interactive_toy_objects_SEP.py", "interactive_spike_gate_MTObjects.py", "interactive_toy_objects_MTObjects.py"):
        replace_everywhere(doc, f'Foreground Masking\\{method}', f'Foreground Masking\\Interactive tools\\{method}')
    for method in ("batch_spike_gate_SEP.py", "batch_toy_objects_SEP.py", "batch_spike_gate_MTObjects.py", "batch_toy_objects_MTObjects.py"):
        replace_everywhere(doc, f'Foreground Masking\\{method}', f'Foreground Masking\\Batch tools\\{method}')
    add_review_section(doc, *REVIEWS["Foreground Masking Canonical Tools Reference.docx"])
    p = doc.add_paragraph(style="List Bullet")
    p.add_run("Aligned Toy Objects batch layout: original / original plus toys; mask / recovered image with green correct and red incorrect outlines; original / processed isophotes; original / processed bar-major profile.")
    p = doc.add_paragraph(style="List Bullet")
    p.add_run("Calibration filenames: reports for the 40 validated CleanGalaxies.txt entries end in `_clean.png`; the remaining reports keep their ordinary suffix.")


def update_mt_four_fold(doc):
    replace_everywhere(
        doc,
        "score = recovery - 2.0 mean masked fraction - 0.5 false-positive fraction",
        "Recovery reward: R_MT = 0.45 mean F-score + 0.35 mean toy recall + 0.20 toy detection rate. Data-loss term: L_MT = 0.50 mean masked fraction + 0.10 min(false-positive fraction, 1). Feasible score: S_MT = R_MT - L_MT; Optuna minimises objective = -S_MT.",
    )
    replace_everywhere(
        doc,
        "A toy is counted as detected when at least 50% of its truth pixels are recovered. A hard worst-image masked fraction limit of 15% prevents a high recovery result from winning through excessive masking. Trials over the limit receive a large cap-excess objective penalty.",
        "A toy is counted as detected when at least 50% of its truth pixels are recovered. A trial is recovery-infeasible when the incremental mask is empty, toy detection rate is below 25%, or mean toy recall is below 20%; it receives a large positive objective, with an empty mask receiving at least 100. Feasible trials retain the 15% worst-image masking cap. Final release requires common-40 toy detection of at least 50%, common-40 mean toy recall of at least 30%, and non-zero recovery in every held-out fold.",
    )
    replace_everywhere(
        doc,
        "The smallest all-40 objective is selected for the final batch.",
        "The best feasible common-40 candidate is selected for the final batch only after the cross-fold recovery release gates are satisfied.",
    )
    add_review_section(doc, *REVIEWS["MTObjects Toy Objects Four Fold Optimisation Documentation.docx"])
    doc.add_paragraph("Selected fold: 3. Common-40 score: 0.2325; mean toy recall: 50.7%; toy detection rate: 53.3%; mean masked fraction: 9.07%; maximum masked fraction: 14.63%; false-positive fraction: 9.05%; mean pixel precision: 0.31%.")
    doc.add_paragraph("Selected parameters: detect_on=original; move_factor=0.805164; min_distance=0.215158; gaussian_fwhm=0.429846; bg_variance=0.000821033; minarea=5; dilation_radius=3; max_area=2987; max_elongation=12.3811; exclude_center_pixels=8; alpha=1e-6.")


def update_optuna_process(doc):
    replace_everywhere(
        doc,
        "For MTObjects toy-object and SEP toy-object runs, Optuna minimises negative score. A larger score is better, so objective = -score:",
        "SEP and MTObjects now use separate Toy Objects objectives. Both score only incremental masking beyond the unaltered baseline and retain a 15% worst-image masking cap. Optuna minimises the negative feasible score; MTObjects additionally applies explicit recovery infeasibility and final cross-fold release gates:",
    )
    replace_everywhere(
        doc,
        "score = 0.45 * mean_f_score\n      + 0.35 * mean_toy_recall\n      + 0.20 * toy_detection_rate\n      - 0.15 * min(false_positive_fraction, 1.0)\nobjective = -score",
        "SEP: recovery = 0.45 mean recall + 0.20 mean F-score + 0.25 mean toy recall + 0.20 toy detection rate; loss = 0.35 mean masked fraction + 0.05 false-positive fraction. MTObjects: recovery = 0.45 mean F-score + 0.35 mean toy recall + 0.20 toy detection rate; loss = 0.50 mean masked fraction + 0.10 false-positive fraction. MTObjects trials are infeasible for an empty incremental mask, toy detection below 25%, or mean toy recall below 20%.",
    )
    add_review_section(doc, *REVIEWS["SEP and MTObjects Optuna Optimisation Process - Updated.docx"])
    doc.add_paragraph("Current four-fold design: 40 validated galaxies; four rotations of 30 training and 10 held out; 40 trials per fold (8 startup + 32 TPE); candidates compared on the same common-40 injection set. Selected winners: SEP fold 4 and MTObjects fold 3.")


def update_cross_validation(doc):
    replace_everywhere(
        doc,
        "score = recovery - 2.0 mean masked fraction - 0.5 false-positive fraction",
        "MTObjects recovery = 0.45 mean F-score + 0.35 mean toy recall + 0.20 toy detection rate; loss = 0.50 mean masked fraction + 0.10 false-positive fraction; feasible score = recovery - loss.",
    )
    replace_everywhere(
        doc,
        "For both algorithms, any trial whose worst individual image masks more than 15% receives a large cap-excess penalty. Optuna minimises objective = -score for feasible trials.",
        "For both algorithms, any trial whose worst individual image masks more than 15% receives a large cap-excess penalty. Optuna minimises objective = -score for feasible trials. MTObjects additionally rejects an empty incremental mask, toy detection below 25%, or mean toy recall below 20%. A candidate is released only when its common-40 detection is at least 50%, common-40 mean toy recall is at least 30%, and every held-out fold has non-zero recovery.",
    )
    replace_everywhere(
        doc,
        "Once SEP and MTObjects batches both finish, the established comparison utility can assemble the requested six-panel, one-PNG-per-galaxy comparison set.",
        "Each method produces one aligned eight-panel PNG per galaxy: top row original and original plus toys; second row mask and recovered image with green correct and red incorrect outlines; third row original and processed isophotes; fourth row original and processed bar-major profiles. The profile x-axis limits and physical panel widths match the image panels above. Calibration-galaxy filenames receive the `_clean` suffix automatically.",
    )
    add_review_section(doc, *REVIEWS["Toy Objects Four Fold Cross Validation Documentation.docx"])
    doc.add_paragraph("Final candidate selection: SEP fold 4 (common-40 score 0.5416) and MTObjects fold 3 (common-40 score 0.2325, the only candidate passing every release gate).")


def update_interactive_toy(doc):
    add_review_section(doc, *REVIEWS["interactive_toy_objects_MTObjects.docx"])
    doc.add_paragraph("The interactive display and the batch report serve different purposes. The canonical batch PNG has eight panels in four rows and uses aligned profile/image widths. Batch filenames for the 40 calibration galaxies end in `_clean.png`.")
    doc.add_paragraph("Selected recovery-follow-up parameters are documented in the combined outcome report. They achieve non-zero recovery but remain review-controlled because common-40 pixel precision is 0.31% and false-positive fraction is 9.05%.")


def update_results_document(path):
    shutil.copy2(SUMMARY, path)
    doc = Document(path)
    for paragraph in doc.paragraphs:
        if paragraph.text == "SEP and MTObjects Toy Objects":
            replace_paragraph_text(paragraph, paragraph.text, "SEP and MTObjects Toy Objects Results - Updated")
            break
    doc.core_properties.title = "SEP and MTObjects Toy Objects Results - Updated"
    doc.save(path)


def main():
    for path in sorted(ROOT.glob("*.docx")):
        name = path.name
        if name == "SEP and MTObjects Toy Objects Optimisation Results.docx":
            update_results_document(path)
            continue
        doc = Document(path)
        if name == "Foreground Masking Canonical Tools Reference.docx":
            update_canonical(doc)
        elif name == "MTObjects Toy Objects Four Fold Optimisation Documentation.docx":
            update_mt_four_fold(doc)
        elif name == "SEP and MTObjects Optuna Optimisation Process - Updated.docx":
            update_optuna_process(doc)
        elif name == "Toy Objects Four Fold Cross Validation Documentation.docx":
            update_cross_validation(doc)
        elif name == "interactive_toy_objects_MTObjects.docx":
            update_interactive_toy(doc)
        else:
            add_review_section(doc, *REVIEWS[name])
        doc.save(path)
        print(f"updated: {name}")


if __name__ == "__main__":
    main()
