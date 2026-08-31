from __future__ import annotations

from pathlib import Path
import math

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
from matplotlib.patches import Circle, Rectangle

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


HERE = Path(__file__).resolve().parent
OUT_DIR = HERE / "haigh_aligned_methodology_artifact"
OUT_DIR.mkdir(parents=True, exist_ok=True)
DOCX_PATH = OUT_DIR / "Foreground Masking Methodology - Clean Galaxies, Haigh-aligned Contaminants and Optimisation.docx"

BLUE = "2E74B5"
DARK = "1F2937"
MID = "52606D"
LIGHT = "F2F4F7"
PALE_BLUE = "EAF2F8"
GREEN = "228B22"
RED = "C62828"
AMBER = "C47F00"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def add_field(paragraph, field_code: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, end])


def style_document(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Inches(0.78)
    sec.bottom_margin = Inches(0.72)
    sec.left_margin = Inches(0.82)
    sec.right_margin = Inches(0.82)
    sec.header_distance = Inches(0.28)
    sec.footer_distance = Inches(0.28)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color, before, after in (
        ("Title", 26, DARK, 0, 12),
        ("Heading 1", 16, BLUE, 16, 7),
        ("Heading 2", 13, BLUE, 12, 5),
        ("Heading 3", 11, MID, 9, 3),
    ):
        s = doc.styles[name]
        s.font.name = "Calibri"
        s.font.size = Pt(size)
        s.font.color.rgb = RGBColor.from_string(color)
        s.font.bold = True
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(after)
        s.paragraph_format.keep_with_next = True

    if "Equation" not in doc.styles:
        eq = doc.styles.add_style("Equation", WD_STYLE_TYPE.PARAGRAPH)
        eq.font.name = "Cambria Math"
        eq.font.size = Pt(10)
        eq.font.color.rgb = RGBColor.from_string(DARK)
        eq.paragraph_format.left_indent = Inches(0.25)
        eq.paragraph_format.right_indent = Inches(0.25)
        eq.paragraph_format.space_before = Pt(4)
        eq.paragraph_format.space_after = Pt(5)
        eq.paragraph_format.keep_together = True

    if "CaptionText" not in doc.styles:
        cap = doc.styles.add_style("CaptionText", WD_STYLE_TYPE.PARAGRAPH)
        cap.font.name = "Calibri"
        cap.font.size = Pt(8.5)
        cap.font.italic = True
        cap.font.color.rgb = RGBColor.from_string(MID)
        cap.paragraph_format.space_before = Pt(2)
        cap.paragraph_format.space_after = Pt(8)
        cap.paragraph_format.keep_with_next = False


def add_header_footer(doc: Document) -> None:
    sec = doc.sections[0]
    header = sec.header.paragraphs[0]
    header.text = "MSc Research | Foreground-object masking methodology"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(MID)

    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Page ")
    run.font.size = Pt(8)
    add_field(footer, "PAGE")
    run = footer.add_run(" of ")
    run.font.size = Pt(8)
    add_field(footer, "NUMPAGES")


def p(doc: Document, text: str = "", *, bold_prefix: str | None = None, style=None):
    para = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        para.add_run(bold_prefix).bold = True
        para.add_run(text[len(bold_prefix):])
    else:
        para.add_run(text)
    return para


def bullet(doc: Document, text: str, level: int = 0):
    para = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    para.add_run(text)
    para.paragraph_format.space_after = Pt(3)
    return para


def number(doc: Document, text: str):
    para = doc.add_paragraph(style="List Number")
    para.add_run(text)
    para.paragraph_format.space_after = Pt(3)
    return para


def callout(doc: Document, title: str, body: str, fill: str = PALE_BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.8)
    prevent_row_split(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, 110, 150, 110, 150)
    para = cell.paragraphs[0]
    para.paragraph_format.space_after = Pt(2)
    run = para.add_run(title)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(BLUE)
    para2 = cell.add_paragraph(body)
    para2.paragraph_format.space_after = Pt(0)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None, font_size=8.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    prevent_row_split(hdr)
    for i, label in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, LIGHT)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        run = cell.paragraphs[0].add_run(label)
        run.bold = True
        run.font.size = Pt(font_size)
    for row_data in rows:
        row = table.add_row()
        prevent_row_split(row)
        for i, text in enumerate(row_data):
            cell = row.cells[i]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            para = cell.paragraphs[0]
            para.paragraph_format.space_after = Pt(0)
            run = para.add_run(str(text))
            run.font.size = Pt(font_size)
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Inches(width)
    return table


def add_compact_grid(doc: Document, rows: list[list[str]], widths: list[float] | None = None, font_size=8.0):
    table = doc.add_table(rows=0, cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    for row_data in rows:
        row = table.add_row()
        prevent_row_split(row)
        for i, text in enumerate(row_data):
            cell = row.cells[i]
            set_cell_margins(cell, top=25, start=80, bottom=25, end=80)
            para = cell.paragraphs[0]
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.line_spacing = 1.0
            run = para.add_run(str(text))
            run.font.size = Pt(font_size)
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Inches(width)
    return table


def add_caption(doc: Document, text: str):
    para = doc.add_paragraph(style="CaptionText")
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.add_run(text)


def selection_figure(path: Path) -> None:
    fig, ax = plt.subplots(figsize=(10.8, 3.2))
    ax.axis("off")
    boxes = [
        (0.02, "182 galaxy\npopulation", "#E8EEF4"),
        (0.21, "Automated\ncleanliness ranking", "#DCEBF7"),
        (0.43, "78-candidate union\n(original 40 + ranked 50)", "#DCEBF7"),
        (0.68, "Three-panel\nvisual re-review", "#FFF2CC"),
    ]
    for x, label, color in boxes:
        w = 0.16 if x != 0.43 else 0.20
        ax.add_patch(Rectangle((x, 0.34), w, 0.34, transform=ax.transAxes,
                               facecolor=color, edgecolor="#2E74B5", lw=1.8))
        ax.text(x + w / 2, 0.51, label, transform=ax.transAxes,
                ha="center", va="center", fontsize=10, weight="bold")
    for x1, x2 in ((0.18, 0.21), (0.37, 0.43), (0.63, 0.68)):
        ax.annotate("", xy=(x2, 0.51), xytext=(x1, 0.51), xycoords=ax.transAxes,
                    arrowprops=dict(arrowstyle="-|>", color="#52606D", lw=1.8))
    outcomes = [(0.84, 0.68, "22 CLEAN", "#D9EAD3", GREEN),
                (0.84, 0.40, "5 AMBIGUOUS", "#FFF2CC", AMBER),
                (0.84, 0.12, "51 POLLUTED", "#F4CCCC", RED)]
    for x, y, label, face, edge in outcomes:
        ax.add_patch(Rectangle((x, y), 0.14, 0.18, transform=ax.transAxes,
                               facecolor=face, edgecolor=f"#{edge}", lw=1.8))
        ax.text(x + 0.07, y + 0.09, label, transform=ax.transAxes,
                ha="center", va="center", fontsize=9, weight="bold")
        ax.annotate("", xy=(x, y + 0.09), xytext=(0.84, 0.51), xycoords=ax.transAxes,
                    arrowprops=dict(arrowstyle="-|>", color="#7A869A", lw=1.0))
    fig.tight_layout()
    fig.savefig(path, dpi=180, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def toy_figure(path: Path) -> None:
    x = np.linspace(-7, 7, 241)
    X, Y = np.meshgrid(x, x)
    core_sigma = 1.66 / 2.355
    star = np.exp(-(X**2 + Y**2)/(2*core_sigma**2)) + 0.04*np.exp(-(X**2 + Y**2)/(2*(3.2*core_sigma)**2))
    q, re, n, pa = 0.55, 2.0, 3.0, np.deg2rad(30)
    xr = X*np.cos(pa) + Y*np.sin(pa)
    yr = -X*np.sin(pa) + Y*np.cos(pa)
    r = np.sqrt(xr**2 + (yr/q)**2)
    bn = 2*n - 1/3
    gal = np.exp(-bn*((np.maximum(r, 1e-4)/re)**(1/n)-1))
    gal = np.clip(gal, 0, np.percentile(gal, 99.5))
    gal /= gal.max()
    fig, axes = plt.subplots(1, 2, figsize=(9.4, 3.7))
    for ax, img, title in zip(axes, (star, gal), ("IRAC-PSF foreground star", "PSF-convolved Sérsic background galaxy")):
        ax.imshow(img, origin="lower", extent=[-7,7,-7,7], cmap="magma")
        ax.contour(img, levels=[0.05, 0.2, 0.5], colors=["#00FF33"], linewidths=[1.0,1.2,1.4], origin="lower", extent=[-7,7,-7,7])
        ax.set_title(title, fontsize=11, weight="bold")
        ax.set_xlabel("arcsec")
        ax.set_ylabel("arcsec")
    axes[0].text(0.03, 0.05, "FWHM 1.66 arcsec\n4% broad wing", transform=axes[0].transAxes,
                 color="white", fontsize=9, weight="bold", va="bottom")
    axes[1].text(0.03, 0.05, "Example: Re=2.0 arcsec, n=3, q=0.55\nAllowed: Re 0.5–3.5, n 2–4, q 0.3–1", transform=axes[1].transAxes,
                 color="white", fontsize=8.5, weight="bold", va="bottom")
    fig.tight_layout()
    fig.savefig(path, dpi=180, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def metrics_figure(path: Path) -> None:
    fig, ax = plt.subplots(figsize=(8.8, 4.2))
    ax.set_aspect("equal")
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 5.4)
    ax.axis("off")
    truth = Circle((4.2, 2.7), 1.65, facecolor="#4A78C2", alpha=0.70, edgecolor="#244A84", lw=2)
    mask = Circle((5.8, 2.7), 1.65, facecolor="#E95B54", alpha=0.65, edgecolor="#A61D18", lw=2)
    ax.add_patch(truth); ax.add_patch(mask)
    ax.text(3.35, 2.7, "FN\nmissed truth", ha="center", va="center", color="white", weight="bold", fontsize=11)
    ax.text(5.0, 2.7, "TP\nrecovered truth", ha="center", va="center", color="white", weight="bold", fontsize=10)
    ax.text(6.65, 2.7, "FP\nextra mask", ha="center", va="center", color="white", weight="bold", fontsize=11)
    ax.text(0.4, 0.35, "TN = displayed-frame pixels that are neither toy truth nor incrementally masked", fontsize=9.5, color="#52606D")
    ax.text(0.4, 4.9, "Blue circle: toy truth (T)     Red circle: incremental mask (MΔ)", fontsize=10.5, weight="bold")
    fig.tight_layout()
    fig.savefig(path, dpi=180, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def folds_figure(path: Path) -> None:
    fig, ax = plt.subplots(figsize=(10.2, 4.0))
    ax.axis("off")
    for i in range(22):
        row = i // 11
        col = i % 11
        x = 0.03 + col * 0.071
        y = 0.70 - row * 0.22
        color = "#F4CCCC" if i == 7 else "#DCEBF7"
        edge = f"#{RED}" if i == 7 else f"#{BLUE}"
        ax.add_patch(Rectangle((x, y), 0.053, 0.13, transform=ax.transAxes,
                               facecolor=color, edgecolor=edge, lw=1.2))
        ax.text(x+0.0265, y+0.065, str(i+1), transform=ax.transAxes,
                ha="center", va="center", fontsize=8, weight="bold")
        if i == 7:
            ax.plot([x+0.005,x+0.048],[y+0.015,y+0.115], transform=ax.transAxes, color=f"#{RED}", lw=2)
            ax.plot([x+0.048,x+0.005],[y+0.015,y+0.115], transform=ax.transAxes, color=f"#{RED}", lw=2)
    ax.text(0.03, 0.90, "One fold (example: galaxy 8 held out)", transform=ax.transAxes, fontsize=12, weight="bold")
    ax.text(0.03, 0.32, "Training", transform=ax.transAxes, fontsize=10.5, weight="bold", color=f"#{BLUE}")
    ax.text(0.15, 0.32, "21 galaxies × Training 1, 2, 3 = 63 cases per trial", transform=ax.transAxes, fontsize=10.5)
    ax.text(0.03, 0.19, "Validation", transform=ax.transAxes, fontsize=10.5, weight="bold", color=f"#{RED}")
    ax.text(0.15, 0.19, "held-out galaxy × Validation 1, 2 = 2 unseen cases", transform=ax.transAxes, fontsize=10.5)
    ax.text(0.03, 0.06, "Repeat until every galaxy has served once as the held-out fold.", transform=ax.transAxes, fontsize=9.5, color=f"#{MID}")
    fig.tight_layout()
    fig.savefig(path, dpi=180, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def add_picture(doc: Document, path: Path, width: float, caption: str) -> None:
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.keep_with_next = True
    para.add_run().add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)


def build() -> Path:
    fig_selection = OUT_DIR / "selection_flow.png"
    fig_toys = OUT_DIR / "toy_models.png"
    fig_metrics = OUT_DIR / "metric_regions.png"
    fig_folds = OUT_DIR / "folds.png"
    selection_figure(fig_selection)
    toy_figure(fig_toys)
    metrics_figure(fig_metrics)
    folds_figure(fig_folds)

    doc = Document()
    style_document(doc)
    add_header_footer(doc)

    # Editorial cover
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(72)
    para.paragraph_format.space_after = Pt(8)
    run = para.add_run("METHODOLOGY REPORT")
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(BLUE)
    title = doc.add_paragraph(style="Title")
    title.add_run("Foreground-object masking with clean galaxies and Haigh-aligned contaminants")
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(22)
    r = subtitle.add_run("Selection of the 22-galaxy truth set · synthetic-source design · SEP and MTObjects · objective metrics · nested validation")
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor.from_string(MID)

    callout(doc, "Purpose", "To define the current experimental method precisely enough for scientific review and reproducible rerunning. This document describes the active Haigh-aligned source experiment; it does not report final optimisation results.")
    p(doc, "Prepared for MSc research on removal of foreground objects from galaxy images")
    p(doc, "Method version: haigh-aligned-s4g-injections-v1 | 31 August 2026")

    doc.add_page_break()

    doc.add_heading("Executive method summary", level=1)
    p(doc, "Twenty-two visually clean galaxy images form empirical ‘truth’ backgrounds. Reproducible artificial contaminants are inserted into each galaxy’s displayed area of interest (AOI): IRAC-like point sources and compact PSF-convolved Sérsic galaxies. SEP and MTObjects are independently optimised to recover the added sources while limiting collateral masking of the underlying galaxy.")
    p(doc, "Optimisation uses leave-one-galaxy-out cross-validation. In each of 22 folds, three source-placement seeds for 21 galaxies are used for Optuna training; two different seeds for the held-out galaxy are used only for validation. The final 182-galaxy application is deferred until the optimisation design and winning parameter set have been accepted.")
    callout(doc, "Interpretation", "A clean galaxy is not assumed to be empty sky. It supplies the real astronomical background against which incremental masking caused by known injected contaminants can be measured.")

    doc.add_heading("Scope and terminology", level=2)
    add_table(doc, ["Term", "Meaning in this study"], [
        ["Clean galaxy", "A centred galaxy AOI judged to contain no convincing bright foreground contaminant likely to compromise the target-galaxy analysis."],
        ["Toy / synthetic contaminant", "A reproducible injected source with a known truth footprint; either a foreground star model or a compact background-galaxy model."],
        ["Displayed AOI", "The centred frame shown and classified in the interactive reviewer. All percentages and truth accounting in the current experiment refer to this region."],
        ["Incremental mask", "Pixels masked after injection but not already masked in the clean baseline: MΔ = M1 \\ M0."],
        ["Fold", "One leave-one-galaxy-out training/validation cycle."],
    ], widths=[1.6, 5.1])

    doc.add_heading("1. Identification of the clean-galaxy set", level=1)
    doc.add_heading("1.1 Automated candidate ranking", level=2)
    p(doc, "The starting population comprised 182 centred galaxy images. A two-dimensional cleanliness score was used to prioritise visual review rather than to make the final decision automatically:")
    bullet(doc, "A Gaussian-smoothed model (σ = 5 pixels) was subtracted from the centred image to emphasise compact positive residuals.")
    bullet(doc, "The galaxy centre was excluded from compact-source scoring.")
    bullet(doc, "2MASS candidates were quality filtered (including usable-source, contamination, photometric-quality and signal-to-noise checks). Gaia astrometric evidence was treated as useful but weaker supporting evidence.")
    bullet(doc, "Compact peaks coincident with strong galaxy structure were downweighted, but not automatically ignored.")
    bullet(doc, "The hybrid score combined catalogue evidence with a small image-only component.")
    p(doc, "This ranking exposed likely low-contamination images, but early results also demonstrated that morphology and display scale can mislead an automated score. It therefore remained a screening device, not the ground-truth label.")

    doc.add_heading("1.2 Candidate union and visual re-review", level=2)
    p(doc, "The final review population was the union of the original visually selected 40 galaxies and the 50 least-polluted galaxies in the revised ranking. After overlap, this produced 78 unique candidates. Every candidate was reviewed in the same displayed AOI using three aligned panels:")
    bullet(doc, "galaxy-centred original in a negative display;")
    bullet(doc, "centred Gaussian residual in a diverging colour scale; and")
    bullet(doc, "centred original with catalogue candidates overplotted (red for 2MASS; blue/yellow variants were used for Gaia evidence during tool development).")
    add_picture(doc, fig_selection, 6.85, "Figure 1. Candidate reduction and final visual classifications.")

    doc.add_heading("1.3 Decision rubric", level=2)
    add_table(doc, ["Class", "Operational decision"], [
        ["Clean", "No convincing bright, unrelated compact source is likely to affect scientific use of the target galaxy in the displayed AOI."],
        ["Ambiguous", "A feature may be foreground, a background object, or intrinsic galaxy structure; its impact cannot be assigned confidently."],
        ["Polluted", "One strong or several moderate unrelated compact sources overlap, or lie close enough to, the scientifically relevant galaxy area to matter."],
    ], widths=[1.15, 5.55])
    p(doc, "Intrinsic structures—nucleus, bar, spiral arms, rings and star-forming knots—were not counted as pollution merely because they were bright. Catalogue circles were supporting evidence rather than automatic labels. The completed review yielded 22 clean, 5 ambiguous and 51 polluted candidates.")

    doc.add_heading("1.4 Final 22 clean galaxies", level=2)
    clean = ["IC1954", "NGC0289", "NGC0578", "NGC0986", "NGC1097", "NGC1367", "NGC2903", "NGC3227", "NGC3359", "NGC3486", "NGC3627", "NGC3681", "NGC3684", "NGC4102", "NGC4405", "NGC4450", "NGC4579", "NGC4639", "NGC4765", "NGC4981", "NGC7531", "PGC013821"]
    rows = [[clean[i], clean[i+1] if i+1 < len(clean) else "", clean[i+2] if i+2 < len(clean) else ""] for i in range(0, len(clean), 3)]
    add_compact_grid(doc, rows, widths=[2.2, 2.2, 2.2], font_size=8.0)

    doc.add_heading("2. Synthetic contaminant design", level=1)
    doc.add_heading("2.1 Rationale and relation to Haigh et al.", level=2)
    p(doc, "Haigh et al. (2021) evaluated source-extraction methods against images with known truth, including real EFIGI galaxies placed at relatively quiet positions in Fornax Deep Survey fields and simulated populations of stars and galaxies. The present study uses the 22 clean observed galaxies as the empirical background (‘truth’) and injects artificial contaminants whose locations and footprints are exactly known.")
    p(doc, "The alignment is conceptual rather than a literal reproduction. In particular, this experiment is at IRAC 3.6 μm and uses an IRAC-like PSF and local-noise brightness scale, whereas Haigh et al. worked with different imaging and source populations.")

    doc.add_heading("2.2 Source types and physical parameters", level=2)
    add_table(doc, ["Source", "Model", "Current parameterisation", "Purpose"], [
        ["Foreground star", "IRAC-PSF approximation", "Gaussian core, FWHM 1.66 arcsec; 4% broader wing with 3.2× the core FWHM.", "Tests compact point-source recovery at the survey resolution."],
        ["Background galaxy", "PSF-convolved Sérsic profile", "Effective radius Re 0.5–3.5 arcsec; Sérsic index n 2–4; axis ratio q 0.3–1.0; random position angle 0–180°.", "Tests compact extended contaminants in the ranges reported by Haigh et al."],
    ], widths=[1.1, 1.45, 2.7, 1.5], font_size=8.2)
    add_picture(doc, fig_toys, 6.75, "Figure 2. Illustrative source morphologies. Green contours indicate model-footprint levels, not segmentation results.")
    callout(doc, "Deliberate exclusion", "The current injection set does not include Haigh et al.’s larger ‘cluster-galaxy’ population (Re 2.5–40 arcsec, n 0.5–2). Those objects could dominate a small displayed AOI and test a different scientific question from foreground-star/compact-background removal.", fill="FFF2CC")

    doc.add_heading("2.3 Number, mixture and brightness", level=2)
    bullet(doc, "Each displayed frame receives 1–5 sources in total—not 1–5 of each type. The count is round(Nfinite / 5000), clipped to [1, 5].")
    bullet(doc, "The target mixture is 75% stars and 25% background galaxies. For frames with at least three sources, at least one background galaxy is enforced.")
    bullet(doc, "Peak amplitude is sampled uniformly from 6–30 times the robust local image noise σ. This controls detectability relative to each frame.")
    bullet(doc, "Brightness is a peak signal-to-noise prescription, not an integrated AB/Vega magnitude distribution and not a direct copy of the Haigh/FDS magnitude function.")
    p(doc, "Across the immutable v1 manifest there are 110 galaxy/seed cases and 375 injected sources: 263 stars and 112 compact background galaxies. No placement fallback was required when the manifest was generated.")

    doc.add_heading("2.4 Placement and truth footprints", level=2)
    number(doc, "Insert sources in observed-sky image pixels, within the finite part of the displayed AOI.")
    number(doc, "Keep injected sources separated from one another and require their truth footprints to remain inside the AOI.")
    number(doc, "Construct truth pixels using the local robust noise: galaxy-model pixels at or above 1σ; star truth is the intersection of that threshold with the region containing 95% of model flux.")
    number(doc, "Save the injected delta image, combined truth mask and labelled per-source truth masks in immutable NPZ files, together with seeds and SHA-256 hashes.")
    callout(doc, "Important v1 limitation", "The active manifest does not exclude the observed target-galaxy structure or pre-existing compact objects when choosing toy positions. Some contaminants may therefore overlap spiral arms, bars or bright knots. A quiet-placement version should be a planned sensitivity experiment, not silently substituted into this run.", fill="FCE8E6")

    doc.add_heading("3. Masking methods", level=1)
    doc.add_heading("3.1 SEP / Source Extractor", level=2)
    p(doc, "SEP (Barbary, 2016) exposes the core Source Extractor algorithm (Bertin & Arnouts, 1996) as a library. In this study it estimates the image background and noise, thresholds the detection image, filters and groups connected pixels, deblends overlapping detections, rejects implausible segments by size/shape, and dilates accepted segments into a practical mask.")
    p(doc, "Optimised controls include detection threshold, minimum area, deblending thresholds and contrast, background-mesh size, filter size, dilation radius, maximum area and maximum elongation.")

    doc.add_heading("3.2 MTObjects", level=2)
    p(doc, "MTObjects is a morphology-based detector built around max-tree representations of connected intensity components and statistical attribute filtering (Teeninga et al., 2016). The tree encodes nested structures from local maxima toward their surrounding background; significant nodes are selected, filtered and dilated to form the mask. Haigh et al. (2021) compared MTObjects with Source Extractor and other extraction methods using objective segmentation-quality criteria.")
    p(doc, "Optimised controls include movement/significance factor, minimum distance, Gaussian smoothing FWHM, background variance, minimum area, dilation radius, maximum area and maximum elongation.")

    doc.add_heading("3.3 Matched implementation principle", level=2)
    p(doc, "The interactive reviewer and the later 182-galaxy batch must call the same masking functions, coordinate transforms, AOI definition and selected parameter files. Display stretch may change visibility to the reviewer, but must never alter the numerical image supplied to SEP or MTObjects.")

    doc.add_heading("4. Evaluation masks and elementary metrics", level=1)
    doc.add_heading("4.1 Baseline subtraction", level=2)
    p(doc, "Each parameter set is evaluated twice on a clean galaxy: first without toys to obtain baseline mask M0, then with exactly the same image plus toys to obtain M1. The incremental response attributable to the toys is:")
    p(doc, "MΔ = M1 \\ M0", style="Equation")
    p(doc, "Only MΔ is scored against injected truth T inside the displayed AOI A. This prevents a method from receiving credit for pixels it would have masked even before a toy was inserted.")
    add_picture(doc, fig_metrics, 6.4, "Figure 3. Pixel accounting for the incremental mask. Missed truth is a false negative (FN), not a true negative.")

    doc.add_heading("4.2 Pixel and object metrics", level=2)
    add_table(doc, ["Metric", "Equation", "Meaning"], [
        ["True positive (TP)", "|MΔ ∩ T|", "Toy-truth pixels correctly added to the mask."],
        ["False positive (FP)", "|MΔ \\ T|", "Incremental mask pixels outside toy truth; collateral data loss."],
        ["False negative (FN)", "|T \\ MΔ|", "Toy-truth pixels missed by the method."],
        ["Recall R", "TP / (TP + FN)", "Fraction of all toy-truth pixels recovered."],
        ["Precision P", "TP / (TP + FP)", "Fraction of incremental mask pixels that are genuinely toy truth."],
        ["F score", "2PR / (P + R)", "Harmonic balance of precision and recall."],
        ["False-positive fraction", "FP / (|A| − |T|)", "Fraction of available non-truth AOI unnecessarily masked."],
        ["Displayed-frame masking", "|MΔ| / |A|", "Incremental proportion of the reviewed frame removed."],
    ], widths=[1.25, 1.65, 3.75], font_size=8.2)
    p(doc, "For toy j, per-toy recall is rj = |MΔ ∩ Tj| / |Tj|. Mean toy recall is the arithmetic mean of rj, giving every source equal weight regardless of apparent area. A toy is classed as detected when rj ≥ 0.50; toy detection rate is the number detected divided by the number injected.")

    doc.add_heading("5. Objective functions", level=1)
    p(doc, "Optuna minimises an objective. For feasible trials, the code first forms an interpretable score (higher is better), then returns its negative (lower is better). Large positive values identify infeasible trials. SEP and MTObjects use related but not identical formulas, so their raw objective values should not be compared directly.")

    doc.add_heading("5.1 SEP objective", level=2)
    p(doc, "RecoverySEP = 0.45 mean(R) + 0.20 mean(F) + 0.25 mean(toy recall) + 0.20 toy detection rate", style="Equation")
    p(doc, "LossSEP = 0.35 mean(masked fraction) + 0.05 min(mean FP fraction, 1)", style="Equation")
    p(doc, "ScoreSEP = RecoverySEP − LossSEP", style="Equation")
    p(doc, "If the largest individual displayed-frame masked fraction is at most 15%, ObjectiveSEP = −ScoreSEP. Otherwise the trial receives the constraint penalty 10 + 100(max masked − 0.15) + LossSEP − RecoverySEP.")

    doc.add_heading("5.2 MTObjects objective", level=2)
    p(doc, "RecoveryMTO = 0.45 mean(F) + 0.35 mean(toy recall) + 0.20 toy detection rate", style="Equation")
    p(doc, "LossMTO = 0.50 mean(masked fraction) + 0.10 min(mean FP fraction, 1) + 1.00 mean galaxy excess above 15%", style="Equation")
    p(doc, "ScoreMTO = RecoveryMTO − LossMTO", style="Equation")
    p(doc, "MTObjects also applies feasibility safeguards:")
    bullet(doc, "toy detection rate must be at least 25%;")
    bullet(doc, "mean toy recall must be at least 20%;")
    bullet(doc, "no more than 20% of physical galaxies may have a worst-seed displayed-frame mask above 15%; and")
    bullet(doc, "no individual galaxy/seed case may exceed 30% masking.")
    p(doc, "Repeated seed cases for a galaxy are treated as correlated for the 15% rule: take that galaxy’s worst seed, then count galaxies equally. Thus E15 = (number of physical galaxies whose worst seed exceeds 15%) / (number of physical galaxies evaluated). Feasibility requires E15 ≤ 0.20.")
    p(doc, "If recovery is infeasible, the objective is 50 + 20(detection deficit) + 20(recall deficit), with a further 50 if the incremental mask is empty. If masking is infeasible, the objective is 10 + 100(E15 − 0.20, if positive) + 100(max mask − 0.30, if positive) + LossMTO − RecoveryMTO. Otherwise ObjectiveMTO = −ScoreMTO.")

    doc.add_heading("5.3 Why these metrics are combined", level=2)
    add_table(doc, ["Failure mode", "Metric that exposes it"], [
        ["Masks almost everything", "Precision, false-positive fraction and masked-fraction penalties fall sharply."],
        ["Detects only the brightest core of each source", "Pixel recall and mean toy recall remain low."],
        ["Recovers a few large sources but misses small ones", "Toy detection rate and equal-weight mean toy recall reveal the imbalance."],
        ["Usually acceptable but catastrophic on one or more galaxies", "Displayed-frame feasibility constraints expose the tail risk."],
        ["Masks the same galaxy structure with or without toys", "Baseline subtraction removes M0 from the scored response."],
    ], widths=[2.5, 4.2])

    doc.add_heading("6. Optimisation and cross-validation", level=1)
    doc.add_heading("6.1 Leave-one-galaxy-out folds", level=2)
    add_picture(doc, fig_folds, 6.85, "Figure 4. One of 22 leave-one-galaxy-out folds. The held-out galaxy changes in every fold.")
    p(doc, "The split is by physical galaxy, not by image case. This is essential: allowing different seeds of the same galaxy into both training and validation would leak the galaxy morphology and overstate generalisation.")

    doc.add_heading("6.2 Training and validation seeds", level=2)
    add_table(doc, ["Stage", "Cases used", "Role"], [
        ["Optuna training in each fold", "21 galaxies × Training 1–3 = 63 cases", "Every parameter trial is scored on all three reproducible training placements for every retained galaxy."],
        ["Held-out validation", "1 galaxy × Validation 1–2 = 2 cases", "The selected fold parameters are tested on different placements for a galaxy never used during that fold’s fitting."],
        ["Cross-fold candidate comparison", "22 galaxies × Validation 1–2 = 44 cases", "Candidate winners can be compared on a common validation set before a final configuration is chosen."],
    ], widths=[1.55, 1.85, 3.3])
    callout(doc, "Answer to a common ambiguity", "One Optuna evaluation uses the three training seeds only. The two validation seeds are not included in the trial objective; they remain held out until validation.")

    doc.add_heading("6.3 Optuna search and convergence", level=2)
    p(doc, "For SEP and MTObjects separately, Optuna explores the defined hyperparameter space. The active run permits up to 80 trials per fold, beginning with 8 initial points followed by adaptive trials. Eight process workers evaluate different image cases within a trial; they are not eight independent Optuna searches.")
    p(doc, "Objective convergence control becomes eligible after 40 completed trials. A fold may stop when 20 further completed trials produce no meaningful improvement, where meaningful is judged using the larger of an absolute tolerance of 1×10⁻⁵ and a relative tolerance of 0.001. Trial and fold wall times are recorded for later local-versus-cloud benchmarking.")
    p(doc, "Parameter stability is assessed separately from objective convergence. The audit examines near-optimal trials and variation across folds to detect cases where similar objective values arise from materially different parameter combinations. At present this audit supports model selection and interpretation; it is not an automatic early-stopping gate.")

    doc.add_heading("6.4 Selection and final deployment gate", level=2)
    number(doc, "Complete all 22 SEP folds and all 22 MTObjects folds, including held-out validation.")
    number(doc, "Check feasibility, objective convergence, parameter stability, per-galaxy failure modes and sensitivity to source placements.")
    number(doc, "Select one defensible parameter configuration per method using validation performance—not training objective alone.")
    number(doc, "Load those exact parameters into the interactive tool for visual spot checks and parity testing with the batch implementation.")
    number(doc, "Only after acceptance, run SEP and MTObjects over the complete 182-galaxy population and produce SEP, MTObjects and combined diagnostic PNG sets.")

    doc.add_heading("7. Reproducibility and quality controls", level=1)
    bullet(doc, "The paired injection manifest fixes source types, positions, parameters and training/validation roles for both masking methods.")
    bullet(doc, "Saved delta and truth arrays carry cryptographic hashes, so the evaluated injection set can be verified byte-for-byte.")
    bullet(doc, "SEP and MTObjects receive the same images, source placements, truth masks and folds, enabling paired comparisons.")
    bullet(doc, "Per-trial CSV/JSON outputs retain parameters, objective components, case metrics, runtime, errors and convergence metadata.")
    bullet(doc, "Displayed-frame percentages use the same AOI used for cleanliness classification and toy placement.")
    bullet(doc, "Interactive/batch parity should be regression-tested on selected galaxy/seed/parameter combinations before deployment.")

    doc.add_page_break()
    doc.add_heading("8. Limitations and planned sensitivity experiments", level=1)
    add_table(doc, ["Current limitation", "Recommended follow-up"], [
        ["v1 toys may overlap the target galaxy or pre-existing objects.", "Create a quiet-centred v2 manifest with configurable exclusion from observed galaxy structure; retain v1 as an overlap/stress experiment."],
        ["Displayed AOI provides limited placement space.", "Run a separate whole-FITS empty-field experiment. Do not mix it into the centred scientific test because it answers a different question."],
        ["Star PSF is an analytic approximation.", "Test an empirical IRAC PRF/PSF if one is available for the relevant products."],
        ["Brightness is local-σ based rather than magnitude based.", "Report results by brightness tier and test a survey-informed flux distribution."],
        ["Only compact Sérsic background galaxies are included.", "If scientifically relevant, add a separately labelled extended-galaxy stress set rather than enlarging the core contaminant set silently."],
        ["Clean labels are visual and catalogue-assisted.", "Retain reviewer decisions, notes and panels; consider blinded repeat review or inter-rater checking for a subset."],
        ["22 galaxies limit population diversity.", "Report per-galaxy distributions and uncertainty; avoid claiming universal performance from means alone."],
        ["SEP and MTO objectives differ.", "Compare methods on the same held-out physical metrics and images, not raw objective values."],
    ], widths=[3.0, 3.7], font_size=8.1)

    doc.add_heading("9. Additional reporting recommended", level=1)
    p(doc, "The following items were not explicit in the requested outline but are important for a defensible MSc methods chapter and results section:")
    bullet(doc, "A versioned parameter-space table, including units, bounds and whether each parameter is continuous, integer or categorical.")
    bullet(doc, "Per-galaxy and per-source-type results, not only global averages; include brightness and size strata.")
    bullet(doc, "Uncertainty intervals across held-out galaxies and seeds, plus sensitivity to the 50% toy-detection threshold.")
    bullet(doc, "A failure-mode gallery showing false positives on galaxy structure, missed faint stars, missed extended sources and catastrophic overmasking.")
    bullet(doc, "A computational benchmark: trial time, fold time, worker count, CPU/RAM utilisation and reproducible software/environment versions.")
    bullet(doc, "A pre-deployment decision table stating the criteria by which one SEP and one MTObjects configuration will be accepted for the 182-galaxy run.")

    doc.add_heading("10. References", level=1)
    refs = [
        "Akiba, T., Sano, S., Yanase, T., Ohta, T. & Koyama, M. (2019). Optuna: A next-generation hyperparameter optimization framework. Proceedings of KDD 2019, 2623–2631. https://doi.org/10.1145/3292500.3330701",
        "Barbary, K. (2016). SEP: Source Extractor as a library. Journal of Open Source Software, 1(6), 58. https://doi.org/10.21105/joss.00058",
        "Bertin, E. & Arnouts, S. (1996). SExtractor: Software for source extraction. Astronomy and Astrophysics Supplement Series, 117, 393–404. https://doi.org/10.1051/aas:1996164",
        "Haigh, C., Chamba, N., Venhola, A., Peletier, R., Doorenbos, L., Watkins, M. & Wilkinson, M. H. F. (2021). Optimising and comparing source-extraction tools using objective segmentation quality criteria. Astronomy & Astrophysics, 645, A107. https://doi.org/10.1051/0004-6361/201936561",
        "Teeninga, P., Moschini, U., Trager, S. C. & Wilkinson, M. H. F. (2016). Statistical attribute filtering to detect faint extended astronomical sources. Mathematical Morphology—Theory and Applications, 1, 100–115. https://doi.org/10.1515/mathm-2016-0006",
    ]
    for ref in refs:
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.25)
        para.paragraph_format.first_line_indent = Inches(-0.25)
        para.paragraph_format.space_after = Pt(4)
        run = para.add_run(ref)
        run.font.size = Pt(9)

    doc.add_page_break()
    doc.add_heading("Appendix A. Current experiment at a glance", level=1)
    add_table(doc, ["Element", "Current setting"], [
        ["Truth backgrounds", "22 visually clean centred galaxy AOIs"],
        ["Manifest", "haigh-aligned-s4g-injections-v1"],
        ["Saved cases", "110 = 22 galaxies × 5 seeds"],
        ["Seeds", "Training 1–3; Validation 1–2"],
        ["Sources", "375 total: 263 stars; 112 compact background galaxies"],
        ["Folds", "22 leave-one-galaxy-out folds per method"],
        ["Training load per trial", "63 cases = 21 galaxies × 3 training seeds"],
        ["Held-out load per fold", "2 cases = 1 galaxy × 2 validation seeds"],
        ["Optuna budget", "80 trials maximum; 8 initial points; convergence eligible after 40 trials"],
        ["Workers", "8 case-level processes within each trial"],
        ["Deployment", "182-galaxy application deferred until winner selection and parity checks"],
    ], widths=[2.25, 4.45])

    doc.core_properties.title = "Foreground-object masking methodology"
    doc.core_properties.subject = "Clean-galaxy selection, Haigh-aligned contaminants, SEP/MTObjects objectives and cross-validation"
    doc.core_properties.author = "MSc Research Project"
    doc.core_properties.keywords = "SEP, Source Extractor, MTObjects, Optuna, toy objects, cross-validation, foreground masking"
    doc.save(DOCX_PATH)
    return DOCX_PATH


if __name__ == "__main__":
    print(build())
