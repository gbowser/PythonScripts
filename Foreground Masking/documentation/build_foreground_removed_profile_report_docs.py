from __future__ import annotations

import importlib.util
import shutil
import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


SCRIPT_DIR = Path(__file__).resolve().parents[1]
PROJECT_ROOT = SCRIPT_DIR.parent
if str(PROJECT_ROOT) not in sys.path:
    sys.path.append(str(PROJECT_ROOT))

from machine_paths import remove_foreground_folder  # noqa: E402

DOC_DIR = SCRIPT_DIR / "documentation"
SCRIPT_NAME = "bar_spike_gated_foreground_report.py"
OUTPUT_DOCX = DOC_DIR / "Bar Spike-Gated Foreground Candidate Report Documentation.docx"

DROPBOX_DOC_DIR = remove_foreground_folder("Laptop") / "documentation"

# Reuse the established compact reference style so this document sits beside
# the existing foreground-mask, S4G, and shoulder-quantification manuals.
template_path = (
    SCRIPT_DIR.parent
    / "Shoulder Recognition Erwin"
    / "documentation"
    / "build_shoulder_quantification_docs.py"
)
spec = importlib.util.spec_from_file_location("shoulder_docs_template", template_path)
template = importlib.util.module_from_spec(spec)
assert spec.loader is not None
spec.loader.exec_module(template)

style_document = template.style_document
set_paragraph_border_bottom = template.set_paragraph_border_bottom
add_text = template.add_text
add_bullets = template.add_bullets
add_numbered = template.add_numbered
add_table = template.add_table
MUTED = template.MUTED


def add_metadata(doc: Document) -> None:
    add_table(
        doc,
        ["Field", "Value"],
        [
            ["Script", SCRIPT_NAME],
            ["Location", str(SCRIPT_DIR / SCRIPT_NAME)],
            ["Document date", date.today().isoformat()],
            [
                "Default output directory",
                r"D:\Dropbox\Public Documents\UCLAN\MSc Research\Remove foreground objects\calibrated spike_rule",
            ],
            [
                "Documentation scope",
                "Purpose, data inputs, Photutils candidate-source role, spike-gated masking model, interpolation method, command-line use, calibration set, outputs, validation, and limitations.",
            ],
        ],
        [2500, 6860],
    )


def add_intro(doc: Document) -> None:
    doc.add_paragraph("Bar Spike-Gated Foreground Candidate Report Documentation", style="Title")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run(SCRIPT_NAME).italic = True
    set_paragraph_border_bottom(subtitle)
    add_metadata(doc)

    doc.add_heading("High-Level Overview", level=1)
    add_text(
        doc,
        "This script creates one-page portrait PDF reports that compare S4G bar and minor-axis intensity profiles before and after spike-gated compact-source masking. "
        "It was developed to remove narrow foreground-object-like spikes in bar-major-axis profiles without applying an aggressive global mask that can suppress valid galaxy light.",
    )
    add_text(
        doc,
        "The current default mode is spike-gated. It first identifies narrow positive spikes in the original bar-major-axis intensity profile, then only allows Photutils residual-image segments to become mask pixels if they intersect those spike samples. "
        "When no profile spikes are detected, the report is still generated but no galaxy pixels are masked.",
    )
    add_text(
        doc,
        "This means the method is not a complete foreground-object removal pipeline for making clean galaxy images. "
        "Photutils proposes compact residual-source candidates, but the one-dimensional bar profile decides whether any candidate is relevant to the shoulder-recognition science product.",
    )
    add_bullets(
        doc,
        [
            "The report keeps the original, masked, and interpolated intensity curves visible on the same portrait page.",
            "Masked objects are circled in red on the isophote image panel.",
            "Interpolated bridge segments are drawn as fine dotted line sections in the same colour as the affected profile.",
            "The masking model and key parameters are written into a table at the bottom of each PDF.",
        ],
    )


def add_data_inputs(doc: Document) -> None:
    doc.add_heading("Inputs and Dependencies", level=1)
    add_table(
        doc,
        ["Input or dependency", "Purpose"],
        [
            [
                "S4G geometry manifest",
                "CSV manifest at Erwin_s4g_image_downloader/geometry_output/s4g_image_geometry_manifest.csv. Supplies galaxy name, FITS image path, centre, disk PA, inclination, bar PA, bar semi-major axis, and pixel scale.",
            ],
            [
                "S4G 3.6 micron FITS image",
                "Science image loaded through astropy.io.fits. The script expects a two-dimensional image after squeezing singleton axes.",
            ],
            [
                "foreground_mask_photutils.py",
                "Provides the Photutils residual segmentation helpers used to locate compact candidate sources. In spike-gated mode these candidates are not automatically masked.",
            ],
            [
                "plot_s4g_isophote_axes.py",
                "Provides profile extraction, deprojected radius calculation, image cutout helpers, log image scaling, and PA-line drawing.",
            ],
            [
                "angle_utils.py",
                "Provides bar/minor-axis geometry helpers and deprojection factors.",
            ],
            [
                "Python packages",
                "numpy, scipy, matplotlib, astropy, photutils, and the local project modules listed above.",
            ],
        ],
        [2600, 6760],
    )


def add_algorithm(doc: Document) -> None:
    doc.add_heading("Processing Workflow", level=1)
    add_numbered(
        doc,
        [
            "Select galaxies from the manifest using --names, --all, or the default single-galaxy ESO120-012 mode.",
            "Load each galaxy FITS image and required geometry fields.",
            "Build a broad Gaussian-smoothed galaxy model, subtract it from the science image, and use the residual image for Photutils segmentation.",
            "Extract the original bar-major-axis profile and detect narrow positive profile spikes using the calibrated spike rule.",
            "Run Photutils compact-source detection on the residual image at the selected detection threshold.",
            "Filter residual segments by area, elongation, central exclusion radius, and standard Photutils segmentation criteria.",
            "Sample the residual-segment mask along the bar-major profile and keep only source labels that intersect detected spike samples.",
            "Dilate the retained source mask to include compact-source wings.",
            "Compute original, masked, and log-linear bridge-filled major/minor-axis profiles.",
            "Create a portrait PDF with the isophote image, original profile graph, masked profile graph, interpolated profile graph, and parameter table.",
        ],
    )

    doc.add_heading("Why spike-gated masking?", level=2)
    add_text(
        doc,
        "Earlier global masking runs could remove substantial valid galaxy light in systems with no visible profile spikes. "
        "The spike-gated model reverses the logic: the original intensity profile must first show a narrow spike, and only residual-image objects that spatially intersect that spike are eligible for removal. "
        "This makes no-spike galaxies such as NGC1187, NGC1640, NGC3726, and ESO420-009 pass through unchanged under the current calibration.",
    )


def add_photutils_role(doc: Document) -> None:
    doc.add_heading("How Photutils Is Used Now", level=1)
    add_text(
        doc,
        "Photutils is currently used as a candidate-source locator, not as the final authority on which pixels should be removed. "
        "The final mask is controlled by the bar-major intensity profile: a Photutils segment is retained only if it intersects samples that the spike detector has marked as suspicious.",
    )
    add_table(
        doc,
        ["Stage", "Decision made"],
        [
            [
                "Photutils residual segmentation",
                "Find compact positive residual-source candidates after subtracting a broad smoothed galaxy model.",
            ],
            [
                "Segment filtering",
                "Reject candidates that are too large, too elongated, too central, or otherwise inconsistent with compact contamination.",
            ],
            [
                "Profile spike detection",
                "Detect narrow positive peaks in the original deprojected bar-major profile using the calibrated local-neighbour rule.",
            ],
            [
                "Spike gate",
                "Keep only Photutils candidate labels whose sampled mask footprint intersects the detected bar-major spike samples.",
            ],
            [
                "No-spike behaviour",
                "If the bar-major profile has no detected spikes, no Photutils candidates are applied to the final mask in spike-gated mode.",
            ],
        ],
        [2800, 6560],
    )
    add_text(
        doc,
        "The practical interpretation is therefore: this code removes compact residual-source candidates only when they appear to be causing a profile spike. "
        "A foreground object that is visible in the image but does not disturb the bar-major profile may be left unmasked, because the immediate purpose is to protect the shoulder-recognition profile rather than to produce a visually cleaned science image.",
    )


def add_spike_model(doc: Document) -> None:
    doc.add_heading("Spike Identification Model", level=1)
    add_text(
        doc,
        "The calibrated rule looks for narrow positive peaks in the deprojected bar-major profile. "
        "It uses local profile shape rather than the total image mask area, because the scientific problem is the contamination of the one-dimensional shoulder-recognition profile.",
    )
    add_table(
        doc,
        ["Parameter", "Default", "Meaning"],
        [
            [
                "--spike-excess-fraction",
                "0.25",
                "Peak must be at least 25 percent above the median of neighbouring positive profile samples.",
            ],
            [
                "--spike-neighbour-inner-arcsec",
                "4.0",
                "Inner radius of the local comparison annulus around the candidate peak.",
            ],
            [
                "--spike-neighbour-outer-arcsec",
                "15.0",
                "Outer radius of the local comparison annulus. This keeps the comparison local to the profile shape.",
            ],
            [
                "--spike-side-offset-samples",
                "3",
                "Sample offset used for the immediate side-drop test around the peak.",
            ],
            [
                "--spike-side-drop-fraction",
                "0.4",
                "Peak must be at least 40 percent above the median of the samples +/-3 positions away.",
            ],
            [
                "--spike-center-exclusion-arcsec",
                "8.0",
                "Candidate peaks inside this central radius are ignored to avoid confusing nuclear structure with foreground objects.",
            ],
            [
                "--spike-window-samples",
                "2",
                "Detected spike samples are expanded by this many profile samples when checking whether a residual segment intersects the spike.",
            ],
            [
                "Edge handling",
                "automatic",
                "If a peak is near the profile edge and does not have two-sided neighbours, the code accepts one-sided neighbours when at least four local comparison samples are available.",
            ],
        ],
        [2600, 1200, 5560],
    )


def add_masking_model(doc: Document) -> None:
    doc.add_heading("Photutils Candidate-Source Model", level=1)
    add_table(
        doc,
        ["Parameter", "Default", "Meaning"],
        [
            [
                "--masking-mode",
                "spike-gated",
                "Default mode. Photutils segments are candidates only; final masking requires intersection with bar-major profile spikes. The alternative global mode masks all retained residual segments.",
            ],
            [
                "--smooth-sigma-pixels",
                "15.0",
                "Gaussian smoothing width used for the broad galaxy model before residual-image detection.",
            ],
            [
                "--detection-nsigma",
                "3.5",
                "Fallback residual detection threshold when --auto-tune is not used.",
            ],
            [
                "--auto-tune",
                "off",
                "When enabled, chooses the most conservative threshold from --auto-tune-nsigmas that covers the detected profile spikes.",
            ],
            [
                "--auto-tune-nsigmas",
                "5.0, 4.5, 4.0, 3.5",
                "Candidate residual thresholds, tested from conservative to more aggressive.",
            ],
            [
                "--npixels",
                "8",
                "Minimum connected pixels for Photutils source detection.",
            ],
            [
                "--dilation-radius-pixels",
                "3",
                "Circular dilation radius applied to retained source labels.",
            ],
            [
                "--max-area",
                "500",
                "Rejects broad residual segments that are more likely galaxy structure or background residuals than compact foreground objects.",
            ],
            [
                "--max-elongation",
                "6.0",
                "Rejects highly elongated detections.",
            ],
            [
                "--exclude-center-radius-pixels",
                "8.0",
                "Rejects candidates whose dilated footprint touches the deprojected central exclusion zone.",
            ],
        ],
        [2500, 1700, 5160],
    )


def add_interpolation(doc: Document) -> None:
    doc.add_heading("Profile Interpolation", level=1)
    add_text(
        doc,
        "The shoulder-recognition algorithm expects continuous profile curves. Masked samples therefore need a conservative bridge in the diagnostic and downstream profile representation. "
        "The script fills masked runs using straight log-linear interpolation between the nearest valid profile samples on either side.",
    )
    add_bullets(
        doc,
        [
            "Interpolation is only applied where the profile mask marks samples as affected.",
            "Neighbouring masked runs can be merged when separated by no more than --bridge-merge-gap-samples.",
            "The default bridge merge gap is 12 profile samples.",
            "Bridge sections are plotted as fine dotted line segments, preserving visibility of which data were interpolated.",
            "The bridge is deliberately straight in log-intensity space to avoid spline overshoot and artificial bumps near the bar centre.",
        ],
    )


def add_report_layout(doc: Document) -> None:
    doc.add_heading("PDF Report Layout", level=1)
    add_table(
        doc,
        ["Panel", "Contents"],
        [
            [
                "Observed image panel",
                "S4G 3.6 micron isophote image in the observed sky plane with bar-major and bar-minor axes overlaid. Retained foreground candidates are circled in red.",
            ],
            [
                "Deprojected image panel",
                "Second isophote panel sampled in the face-on disk plane and rotated so the bar major axis is horizontal. The same retained red-circle candidates are transformed into this coordinate system.",
            ],
            [
                "Original profile graph",
                "Original bar-major and bar-minor intensity profiles on log y-axis, before masking.",
            ],
            [
                "Masked profile graph",
                "Profiles sampled after retained source pixels are set to NaN.",
            ],
            [
                "Interpolated profile graph",
                "Masked profiles with straight log-linear bridges over affected samples. Dotted bridge segments show exactly what was replaced.",
            ],
            [
                "Parameter table",
                "Masking model, residual image model, detection threshold, spike rule, smoothing, segment filtering, interpolation, and auto-tune grid.",
            ],
        ],
        [2400, 6960],
    )


def add_code_flow(doc: Document) -> None:
    doc.add_heading("Code Flow", level=1)
    add_table(
        doc,
        ["Function", "Role"],
        [
            ["read_row / read_rows", "Read one or all manifest rows from the S4G geometry CSV."],
            ["profile_radius_pixels", "Choose a safe profile radius from the image geometry, bar size, pixel scale, and image boundaries."],
            ["build_mask_from_residual", "Run Photutils residual detection and filtering without profile spike gating. Used by global mode."],
            ["detect_profile_spikes", "Detect calibrated narrow positive spikes in the deprojected bar-major intensity profile."],
            ["build_spike_gated_mask_from_residual", "Run residual segmentation, then keep only segments whose profile footprint intersects detected spike samples."],
            ["plot_profile", "Draw one original or masked major/minor profile panel."],
            ["plot_profile_with_bridges", "Draw valid data as solid lines and interpolated replacement segments as dotted lines."],
            ["profile_mask_at_pa", "Sample the two-dimensional mask along a profile axis to determine which profile samples are affected."],
            ["fill_masked_profile_with_log_linear_bridges", "Fill affected profile samples with straight bridges in log-intensity space."],
            ["choose_spike_gated_detection_nsigma", "Auto-tune the residual detection threshold so all detected spike samples are covered while staying conservative."],
            ["make_report", "Generate one portrait PDF for one galaxy."],
            ["parse_args / main", "Parse command-line options, select galaxies, and write one output PDF per selected galaxy."],
        ],
        [3200, 6160],
    )


def add_cli_examples(doc: Document) -> None:
    doc.add_heading("Command-Line Use", level=1)
    add_table(
        doc,
        ["Task", "Command"],
        [
            [
                "Run the current calibration set",
                'python "Foreground Masking/bar_spike_gated_foreground_report.py" --names ESO120-012 ESO357-012 ESO358-020 ESO359-031 ESO440-044 NGC1187 NGC1640 NGC3726 ESO420-009 --auto-tune',
            ],
            [
                "Run one galaxy",
                'python "Foreground Masking/bar_spike_gated_foreground_report.py" --names ESO120-012 --auto-tune',
            ],
            [
                "Run every manifest galaxy",
                'python "Foreground Masking/bar_spike_gated_foreground_report.py" --all --detection-nsigma 3.5 --output-dir "D:\\Dropbox\\Public Documents\\UCLAN\\MSc Research\\Remove foreground objects\\spike_gated_3p5_sigma"',
            ],
            [
                "Run global Photutils at 3.5 sigma",
                'python "Foreground Masking/bar_spike_gated_foreground_report.py" --all --masking-mode global --detection-nsigma 3.5 --output-dir "D:\\Dropbox\\Public Documents\\UCLAN\\MSc Research\\Remove foreground objects\\global_photutils_3p5_sigma"',
            ],
            [
                "Override output directory",
                'Append --output-dir "D:\\Dropbox\\Public Documents\\UCLAN\\MSc Research\\Remove foreground objects\\some_test_folder".',
            ],
            [
                "Tighten spike detection",
                "Increase --spike-excess-fraction or --spike-side-drop-fraction.",
            ],
            [
                "Loosen spike detection",
                "Decrease --spike-excess-fraction or --spike-side-drop-fraction, then validate no-spike controls carefully.",
            ],
        ],
        [2500, 6860],
    )


def add_calibration(doc: Document) -> None:
    doc.add_heading("Current Calibration Set", level=1)
    add_text(
        doc,
        "The current spike rule was tuned against a small visual calibration set. "
        "The intent is not to prove the rule for all 182 galaxies, but to establish a conservative behaviour before expanding the run.",
    )
    add_table(
        doc,
        ["Galaxy", "Expected profile spikes", "Detected by current defaults"],
        [
            ["ESO120-012", "3", "3"],
            ["ESO357-012", "5", "5"],
            ["ESO358-020", "2", "2"],
            ["ESO359-031", "1", "1"],
            ["ESO440-044", "2", "2"],
            ["NGC1187", "0", "0"],
            ["NGC1640", "0", "0"],
            ["NGC3726", "0", "0"],
            ["ESO420-009", "0", "0"],
        ],
        [2700, 2900, 3760],
    )


def add_outputs(doc: Document) -> None:
    doc.add_heading("Outputs", level=1)
    add_table(
        doc,
        ["Output", "Meaning"],
        [
            [
                "{galaxy_name}_fg_removed_sp-gated.pdf and .png",
                "Portrait report and PNG preview for one galaxy. The method suffix identifies the spike-gated masking run.",
            ],
            [
                "Default output folder",
                r"D:\Dropbox\Public Documents\UCLAN\MSc Research\Remove foreground objects\spike_gated_auto_tuned",
            ],
            [
                "One-galaxy default output",
                "Foreground Masking/ESO120-012_portrait_mask_report/ESO120-012_fg_removed_sp-gated.pdf when no --names or --all selection is supplied.",
            ],
        ],
        [2800, 6560],
    )


def add_validation(doc: Document) -> None:
    doc.add_heading("Validation and Limitations", level=1)
    add_bullets(
        doc,
        [
            "Always validate a small mixed set before running all galaxies: include known spike cases and known no-spike controls.",
            "Check that no-spike galaxies report zero applied source segments and zero ignored pixels.",
            "Check that bridge segments are short and dotted only over affected samples, not over valid profile data.",
            "Do not interpret a clean no-spike report as proof that the image contains no foreground objects; it means no candidate was allowed to affect the profile mask under the current spike gate.",
            "Auto-tune currently chooses among residual segmentation thresholds; it does not tune the spike rule itself for each galaxy.",
            "The method is deliberately profile-driven and may miss contaminants that are visually obvious in the image but do not disturb the bar-major-axis intensity profile.",
            "The output is best described as bar-profile spike-gated compact-source masking, not general foreground-object removal.",
            "Very broad, saturated, or blended foreground objects may need manual review or a separate masking strategy.",
            "The calibrated defaults should not be treated as final until a wider sample of the 182 galaxies has been visually checked.",
        ],
    )


def add_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = "Bar Spike-Gated Foreground Candidate Report Documentation"
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = MUTED
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.text = "Generated documentation"
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = MUTED


def build() -> Path:
    DOC_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_document(doc)
    add_footer(doc)
    add_intro(doc)
    add_data_inputs(doc)
    add_algorithm(doc)
    add_photutils_role(doc)
    add_spike_model(doc)
    add_masking_model(doc)
    add_interpolation(doc)
    add_report_layout(doc)
    add_code_flow(doc)
    add_cli_examples(doc)
    add_calibration(doc)
    doc.add_page_break()
    add_outputs(doc)
    add_validation(doc)

    for style_name in ("Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3"):
        doc.styles[style_name].paragraph_format.keep_with_next = True
    doc.core_properties.title = "Bar Spike-Gated Foreground Candidate Report Documentation"
    doc.core_properties.subject = "Technical documentation for spike-gated foreground removal profile reports"
    doc.core_properties.author = ""
    doc.save(OUTPUT_DOCX)
    return OUTPUT_DOCX


if __name__ == "__main__":
    docx = build()
    print(docx)
    DROPBOX_DOC_DIR.mkdir(parents=True, exist_ok=True)
    shutil.copy2(docx, DROPBOX_DOC_DIR / docx.name)
    print(DROPBOX_DOC_DIR)
