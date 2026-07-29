#!/usr/bin/env python3
"""Build the DOCX reference for canonical optimisation, interactive, and batch tools."""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


OUT_PATHS = [
    Path(__file__).resolve().parent / "Foreground Masking Canonical Tools Reference.docx",
    Path(__file__).resolve().parent / "Foreground Masking Interactive Tools Matrix.docx",
]


COMBINATIONS = [
    ("SEP", "spike_gate"),
    ("SEP", "toy_objects"),
    ("MTObjects", "spike_gate"),
    ("MTObjects", "toy_objects"),
]


OPTIMISATION_TRIALS = [
    ("SEP", "Spike Gate", "optimise_spike_gate_SEP.py", 16, 64),
    ("MTObjects", "Spike Gate", "optimise_spike_gate_MTObjects.py", 12, 48),
    ("SEP", "Toy Objects", "optimise_toy_objects_SEP.py", 8, 32),
    ("MTObjects", "Toy Objects", "optimise_toy_objects_MTObjects.py", 8, 32),
]


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)
    for style_name in ["Heading 1", "Heading 2"]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style.font.bold = True


def add_code(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def main() -> None:
    doc = Document()
    style_document(doc)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Foreground Masking Canonical Tools Reference")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(18)
    subtitle = doc.add_paragraph(f"Updated {date.today().isoformat()}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "The maintained tools form a three-by-four set: optimise, interactive, and batch entry points for SEP and "
        "MTObjects with Spike Gate and Toy Objects. Each launcher auto-detects Laptop or Desktop from the configured "
        "hostname, with the code drive as a fallback. --pc remains available as an explicit override."
    )

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["Engine", "Method", "Role", "Canonical filename"]
    for cell, header in zip(table.rows[0].cells, headers):
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for engine, method in COMBINATIONS:
        for role in ("optimise", "interactive", "batch"):
            cells = table.add_row().cells
            values = [engine, method.replace("_", " ").title(), role.title(), f"{role}_{method}_{engine}.py"]
            for cell, value in zip(cells, values):
                cell.text = value

    doc.add_heading("Launch Commands", level=1)
    add_code(doc, 'python "Foreground Masking\\optimise_spike_gate_SEP.py"')
    add_code(doc, 'python "Foreground Masking\\interactive_spike_gate_SEP.py"')
    add_code(doc, 'python "Foreground Masking\\batch_spike_gate_SEP.py"')

    doc.add_heading("Parameter hand-off", level=1)
    doc.add_paragraph(
        "Every optimiser writes its best parameter dictionary to a method-specific JSON file in its timestamped "
        "output directory. The matching canonical interactive and batch launchers automatically locate and load the "
        "newest JSON for the detected device. Pass --best-json (or --params-json where supported) to override it."
    )

    doc.add_heading("Parameter optimisation process", level=1)
    doc.add_paragraph(
        "A trial is one candidate parameter combination evaluated by Optuna; it is not one galaxy. By default, "
        "each optimiser run reproducibly selects 20 galaxies from the usable manifest rows using its --seed value. "
        "Every trial in that run is evaluated against the same selected set of 20 galaxies. This allows candidate "
        "parameter sets to be compared on identical data within a run."
    )

    trial_table = doc.add_table(rows=1, cols=6)
    trial_table.style = "Table Grid"
    trial_headers = ["Engine", "Mask method", "Optimiser", "Initial trials", "Further trials", "Total trials"]
    for cell, header in zip(trial_table.rows[0].cells, trial_headers):
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for engine, method, filename, initial, further in OPTIMISATION_TRIALS:
        cells = trial_table.add_row().cells
        values = [engine, method, filename, str(initial), str(further), str(initial + further)]
        for cell, value in zip(cells, values):
            cell.text = value

    doc.add_paragraph(
        "The stability study repeats all four optimiser runs three times with three different reproducible seeds. "
        "Consequently, each optimiser is tested on three differently selected 20-galaxy samples. The three runs "
        "are compared to determine whether the selected best parameters and objective scores are stable across "
        "different galaxy samples. Changing the seed changes the selected sample; it does not change the number "
        "of Optuna trials."
    )
    doc.add_paragraph(
        "Per run, the SEP Spike Gate optimiser performs 80 trials (16 initial plus 64 further), the MTObjects "
        "Spike Gate optimiser performs 60 trials (12 plus 48), and each Toy Objects optimiser performs 40 trials "
        "(8 plus 32). With 20 galaxies, these correspond to 1,600, 1,200, 800, and 800 per-galaxy trial "
        "evaluations respectively, subject to any failed or skipped galaxy evaluations. Command-line overrides "
        "for --initial-points, --max-iter, --max-images, or --names alter these defaults and should be recorded "
        "with the run results."
    )

    doc.add_heading("Device detection", level=1)
    doc.add_paragraph(
        "The configured hostname selects the matching research tree; the drive containing the running code is used "
        "as a fallback. Set FOREGROUND_MASKING_PC to Laptop or Desktop, or pass --pc, when an explicit override is needed."
    )

    doc.add_heading("Support-code folders", level=1)
    doc.add_paragraph(
        "Only the twelve canonical launchers are kept in the Foreground Masking root. Supporting programs are grouped "
        "under Batch tools, PhotUtils, Interactive tools, Shared, Utilities, and Automation. The canonical launchers "
        "add these folders to the Python search path automatically."
    )

    doc.add_heading("Documentation Rule", level=1)
    doc.add_paragraph(
        "Maintained documentation deliverables in this folder are DOCX files only. Render folders, PDFs, and "
        "standalone PNG documentation intermediates are disposable QA artifacts and should not be retained."
    )

    for out_path in OUT_PATHS:
        doc.save(out_path)
        print(out_path)


if __name__ == "__main__":
    main()
