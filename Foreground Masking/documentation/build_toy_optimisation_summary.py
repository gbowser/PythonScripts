from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(__file__).resolve().parent / "SEP and MTObjects Toy Objects Optimisation Summary and Conclusions.docx"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
PALE_BLUE = "EAF2F8"
LIGHT_GREY = "F2F4F7"
MID_GREY = "666666"
WHITE = "FFFFFF"
BLACK = "000000"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths_dxa):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run_font(run, size=None, bold=None, color=BLACK, italic=None):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def style_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.4)
    section.footer_distance = Inches(0.45)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_style in ("List Bullet", "List Number"):
        style = doc.styles[list_style]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167


def add_header_footer(doc):
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("TOY OBJECTS FOREGROUND MASKING | FINAL RESULTS")
    set_run_font(run, 8.5, True, MID_GREY)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("SEP and MTObjects optimisation summary | 16 August 2026")
    set_run_font(run, 8.5, False, MID_GREY)


def add_title_block(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("OPTIMISATION RESULTS")
    set_run_font(r, 10, True, BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run("SEP and MTObjects Toy Objects")
    set_run_font(r, 24, True, BLACK)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("Four-fold cross-validation, final parameter selection, comparative interpretation and conclusions")
    set_run_font(r, 13, False, MID_GREY)

    for label, value in (
        ("Calibration sample", "40 low-foreground galaxies; four folds of 30 training and 10 held out"),
        ("Optimisation", "40 trials per fold: 8 startup trials followed by 32 TPE-guided trials"),
        ("Application target", "182-galaxy sample with six injected toy objects per galaxy"),
        ("Report date", "16 August 2026"),
    ):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"{label}: ")
        set_run_font(r, 10.5, True)
        r = p.add_run(value)
        set_run_font(r, 10.5)


def add_callout(doc, title, text, fill=PALE_BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, 11, True, DARK_BLUE)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_run_font(r, 10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, value in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, LIGHT_GREY)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(value)
        set_run_font(r, 9, True)
    set_repeat_table_header(table.rows[0])
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            p = cells[index].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(value))
            set_run_font(r, 9)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def build():
    doc = Document()
    style_document(doc)
    add_header_footer(doc)
    add_title_block(doc)

    doc.add_heading("Executive summary", level=1)
    add_callout(
        doc,
        "Overall conclusion",
        "SEP is the preferred primary method for the present Toy Objects workflow. Its cross-validated score was materially higher and its toy recovery was stronger. The revised MTObjects optimisation is nevertheless a genuine recovery from the earlier zero-mask failure: it now detects and removes toy objects, but does so with much lower pixel precision and substantial collateral masking. MTObjects should therefore remain a secondary, review-controlled method until its false-positive behaviour is reduced.",
    )
    add_bullets(doc, [
        "SEP winner: fold 4; common-40 score 0.5416. Mean held-out score across four folds was 0.5026 (standard deviation 0.0880).",
        "MTObjects recovery winner: fold 3; common-40 score 0.2325, mean toy recall 50.7%, toy detection rate 53.3%, and mean masked fraction 9.07%.",
        "Only the MTObjects fold-3 candidate met every cross-fold feasibility gate. This supports its selection, but also shows that MTObjects parameter performance is less robust across folds than SEP.",
        "The very low MTObjects pixel precision (0.31%) is not a contradiction of toy detection: toy pixels occupy a very small part of each image, while the selected mask covers many non-toy pixels. The method recovers toys but with high collateral masking.",
    ])

    doc.add_heading("1. Experimental design", level=1)
    doc.add_paragraph(
        "The same set of 40 galaxies with few foreground objects was used for both methods. In each of four rotations, 30 galaxies were used for parameter optimisation and the remaining 10 were held out for validation. Each fold used 40 Optuna trials: eight startup trials for broad exploration and 32 trials guided by the Tree-structured Parzen Estimator. Final fold candidates were also evaluated on a common 40-galaxy injection set to support a like-for-like selection."
    )
    doc.add_paragraph(
        "The objective rewards recovery of injected toy objects while penalising false-positive masking, loss of usable galaxy data and excessive total mask coverage. A hard maximum masked-fraction limit of 15% was applied in the MTObjects recovery follow-up. The revised MTObjects objective also explicitly makes non-detection infeasible, preventing the optimiser from preferring an empty mask merely because it avoids false-positive costs."
    )

    doc.add_heading("2. SEP optimisation results", level=1)
    add_table(doc, ["Fold", "Held-out score", "Common-40 score", "Maximum masked"], [
        ("1", "0.5177", "0.4991", "13.6%"),
        ("2", "0.5809", "0.5289", "14.0%"),
        ("3", "0.3553", "0.4862", "12.3%"),
        ("4 (selected)", "0.5566", "0.5416", "12.3%"),
    ], [1200, 2400, 2600, 3160])
    doc.add_paragraph(
        "The selected fold-4 candidate did not have the single highest held-out score, but it had the best common-40 score. This is the appropriate selection criterion because every candidate was compared against the same 40-galaxy evaluation set. Fold 3 was notably weaker on its held-out group, but the remaining three held-out scores were close, suggesting generally useful transfer with one more difficult partition."
    )
    doc.add_heading("Selected SEP parameters", level=2)
    add_table(doc, ["Parameter", "Selected value", "Parameter", "Selected value"], [
        ("Detection image", "Residual", "Detection threshold", "1.8577"),
        ("Minimum area", "5 pixels", "Deblend thresholds", "32"),
        ("Deblend contrast", "0.00179745", "Background size", "48"),
        ("Filter size", "1", "Dilation radius", "4 pixels"),
        ("Maximum area", "4,309 pixels", "Maximum elongation", "26.6013"),
        ("Central exclusion", "8 pixels", "Winning fold", "4"),
    ], [1900, 2780, 1900, 2780])
    add_callout(
        doc,
        "SEP interpretation",
        "SEP provides the better balance of recovery and restraint. Its common-40 score is more than twice the MTObjects score under their respective recovery objectives. The result supports SEP as the default automated foreground-mask generator, subject to visual review of outliers and galaxies with complex intrinsic structure.",
    )

    doc.add_heading("3. MTObjects recovery optimisation results", level=1)
    add_table(doc, ["Fold", "Held-out score", "Common-40 score", "Feasible across folds"], [
        ("1", "0.1837", "0.1958", "No"),
        ("2", "0.2617", "0.2070", "No"),
        ("3 (selected)", "0.2442", "0.2325", "Yes"),
        ("4", "0.2135", "0.2294", "No"),
    ], [1200, 2300, 2600, 3260])
    doc.add_paragraph(
        "Fold 2 achieved the highest individual held-out score, while fold 4 approached the selected candidate on the common-40 score. Fold 3 was selected because it was the only candidate to satisfy the minimum toy-detection and toy-recall gates in every fold while remaining below the 15% maximum masked-fraction limit. This selection favours robustness over a single favourable split."
    )
    doc.add_heading("Selected MTObjects parameters", level=2)
    add_table(doc, ["Parameter", "Selected value", "Parameter", "Selected value"], [
        ("Detection image", "Original", "Move factor", "0.80516"),
        ("Minimum distance", "0.21516", "Gaussian FWHM", "0.42985"),
        ("Background variance", "0.00082103", "Minimum area", "5 pixels"),
        ("Dilation radius", "3 pixels", "Maximum area", "2,987 pixels"),
        ("Maximum elongation", "12.3811", "Central exclusion", "8 pixels"),
        ("Alpha", "1 x 10^-6", "Winning fold", "3"),
    ], [1900, 2780, 1900, 2780])
    doc.add_heading("Common-40 performance of the selected candidate", level=2)
    add_table(doc, ["Metric", "Result", "Interpretation"], [
        ("Optimisation score", "0.2325", "Best feasible common-set candidate"),
        ("Mean toy recall", "50.7%", "About half of injected toy pixels recovered"),
        ("Toy detection rate", "53.3%", "Just over half of toys detected"),
        ("Mean pixel recall", "45.5%", "Moderate recovery of toy truth pixels"),
        ("Mean pixel precision", "0.31%", "Most masked pixels are outside toy truth"),
        ("Mean masked fraction", "9.07%", "Material image area removed"),
        ("Maximum masked fraction", "14.63%", "Close to the 15% constraint"),
        ("False-positive fraction", "9.05%", "Collateral mask dominates total mask"),
    ], [2600, 1700, 5060])

    doc.add_heading("4. Comparative assessment", level=1)
    add_table(doc, ["Criterion", "SEP", "MTObjects", "Assessment"], [
        ("Common-40 score", "0.5416", "0.2325", "SEP clearly stronger"),
        ("Cross-fold robustness", "Three strong held-out folds; one weaker", "Only one candidate passed all gates", "SEP more stable"),
        ("Recovery behaviour", "Strong recovery/restraint balance", "Moderate recovery", "SEP preferred"),
        ("Collateral masking", "Controlled below evaluated limits", "False-positive fraction about 9%", "MTObjects needs reduction"),
        ("Recommended role", "Primary automated method", "Secondary diagnostic/review method", "Do not combine blindly"),
    ], [1900, 2100, 2200, 3160])
    doc.add_paragraph(
        "Scores should be compared as decision evidence rather than as universal physical quantities: each score is a composite of recovery and penalty terms. The direction and magnitude nevertheless support the same operational conclusion because SEP combines a higher objective score with a more favourable recovery/collateral balance."
    )

    doc.add_heading("5. Conclusions", level=1)
    add_bullets(doc, [
        "The four-fold design is materially more defensible than four independent ten-galaxy optimisations because every galaxy contributes to validation while each parameter set is trained on 75% of the calibration sample.",
        "SEP has produced a credible, cross-validated parameter set for application to the full sample and should be treated as the baseline method.",
        "The MTObjects recovery follow-up solved the objective-function pathology: an empty mask is no longer considered a satisfactory result when toys are present.",
        "MTObjects now demonstrates genuine toy recovery, but its low precision and high false-positive area show that recovery is purchased at substantial cost to unaffected pixels.",
        "Neither optimisation should be interpreted as proof of perfect foreground removal on real contaminants; injected toys approximate foreground sources but cannot span all morphologies, surface-brightness contrasts and galaxy structures.",
    ])

    doc.add_page_break()
    doc.add_heading("6. Recommended next steps", level=1)
    add_bullets(doc, [
        "Adopt SEP as the primary production mask and review the 40 filenames marked '_clean' first, because these are the calibration galaxies and therefore the most important check for overfitting or visually implausible masks.",
        "Retain MTObjects outputs for comparison, targeted difficult cases and method-agreement analysis, rather than replacing SEP across all galaxies.",
        "For a further MTObjects run, increase the penalty on non-toy masked pixels or impose a stricter mean-mask constraint, while preserving the explicit non-detection penalty and minimum recovery gates.",
        "Report both toy-level detection and pixel-level precision/recall. Toy detection alone can look satisfactory even when the mask includes a large amount of unrelated image area.",
        "Before final scientific use, visually classify failures and compare the original versus processed isophotes and bar-major profiles to ensure foreground masking has not altered the galaxy structure being measured.",
    ])

    doc.add_heading("7. Current application status", level=1)
    add_callout(
        doc,
        "Status at document preparation",
        "The optimisation results above are complete. The earlier SEP 182-galaxy application completed successfully. The newly aligned eight-panel MTObjects batch created 181 of 182 reports; NGC3185 failed because no valid toy-injection candidates were available inside its investigated galaxy area. The visible launcher therefore stopped before starting the corresponding aligned SEP batch. Consequently, conclusions about the optimisation are final, but the newest aligned 182-galaxy comparative PNG set is not yet complete.",
        fill="FFF4E5",
    )
    doc.add_paragraph(
        "This application failure is a toy-placement edge case, not evidence that the selected MTObjects parameters failed on NGC3185. It should be handled separately by allowing a controlled injection fallback or explicitly producing a no-toy diagnostic for that galaxy, followed by resumption of both batch stages."
    )

    doc.add_heading("8. Result provenance", level=1)
    doc.add_paragraph(
        "MTObjects final selection: mtobjects toy recovery followup\\20260816_063455\\mtobjects_toy_cross_validation_best.json. SEP final parameters were reconstructed from the preserved optimisation workbook and fold results after the original run directory was removed. Detailed numerical records are retained in documentation\\Foreground Masking Optimisation Results.xlsx."
    )
    doc.add_paragraph(
        "Metric values are rounded for readability. Full-precision values remain in the JSON, CSV and workbook artefacts."
    )

    doc.core_properties.title = "SEP and MTObjects Toy Objects Optimisation Summary and Conclusions"
    doc.core_properties.subject = "Four-fold cross-validation results and scientific interpretation"
    doc.core_properties.author = "UCLAN MSc Research"
    doc.core_properties.keywords = "SEP, MTObjects, Toy Objects, optimisation, cross-validation, foreground masking"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
