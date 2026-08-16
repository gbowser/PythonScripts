#!/usr/bin/env python3
"""Build dedicated MTObjects/Toy Objects cross-validation documentation."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from build_toy_objects_cross_validation_docs import (
    DARK_BLUE,
    add_header_footer,
    add_table,
    configure_styles,
)


ROOT = Path(__file__).resolve().parents[2]
OUTPUT = ROOT / "Foreground Masking" / "documentation" / "MTObjects Toy Objects Four Fold Optimisation Documentation.docx"


def build() -> None:
    doc = Document()
    configure_styles(doc)
    add_header_footer(doc)
    header = doc.sections[0].header.paragraphs[0]
    header.text = "Foreground Masking | MTObjects Toy Objects optimisation"

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)
    run = title.add_run("MTObjects / Toy Objects Optimisation")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Four-fold cross-validation on 40 low-foreground galaxies").italic = True

    doc.add_heading("1. Aim", level=1)
    doc.add_paragraph(
        "The MTObjects optimisation calibrates foreground-object masking against known injected Toy Objects while "
        "penalising unnecessary removal of galaxy pixels. It repeats the SEP experimental design on the same 40 "
        "validated galaxies, allowing the two algorithms to be compared under a common sampling framework."
    )

    doc.add_heading("2. Data and Toy Objects", level=1)
    doc.add_paragraph(
        "The input list contains exactly 40 unique galaxies drawn from the 182-galaxy manifest. Each has complete "
        "bar/disc geometry and an accessible 3.6-micron FITS image. Six stars, compact clusters or elliptical galaxies "
        "are injected per image within the investigated region. Peaks span 5-25 robust-sigma and truth masks include "
        "model pixels above 8% of peak followed by one-pixel dilation."
    )
    add_table(
        doc,
        ["Toy property", "Rule"],
        [
            ["Type mixture", "50% star, 20% compact cluster, 30% galaxy"],
            ["Star/cluster FWHM", "2-10 pixels"],
            ["Galaxy FWHM", "5-22 pixels"],
            ["Galaxy axis ratio", "0.35-0.95"],
            ["Placement", "Non-overlapping and wholly inside the investigated region"],
        ],
        [2800, 6560],
    )

    doc.add_heading("3. Four-fold design", level=1)
    doc.add_paragraph(
        "A fixed shuffle divides the same 40 names used by SEP into four disjoint folds of 10. In each rotation, "
        "MTObjects is optimised on 30 galaxies and evaluated on the excluded 10. Every galaxy therefore appears in a "
        "held-out set once. Each fold uses 40 Optuna trials: eight startup samples followed by 32 adaptive TPE trials."
    )
    add_table(
        doc,
        ["Phase", "Images", "Influences optimisation?"],
        [
            ["Training", "30 per fold", "Yes"],
            ["Held-out validation", "10 per fold", "No"],
            ["Common candidate evaluation", "All 40, independent injections", "Selects among four trained candidates"],
            ["Final batch", "All 182 manifest galaxies", "No"],
        ],
        [2600, 2600, 4160],
    )

    doc.add_heading("4. MTObjects processing model", level=1)
    doc.add_paragraph(
        "MTObjects is run on the original science image for Toy Objects calibration. A baseline mask is first computed "
        "from the unmodified image. The same parameters are then applied to the injected image, and scoring uses only "
        "incremental pixels that were absent from the baseline mask. This separates Toy Object response from objects "
        "already detected in the galaxy."
    )
    doc.add_paragraph("incremental mask = injected-image mask AND NOT baseline mask", style="Equation")

    doc.add_heading("5. Objective", level=1)
    doc.add_paragraph(
        "recovery = 0.45 mean F-score + 0.35 mean toy recall + 0.20 toy detection rate",
        style="Equation",
    )
    doc.add_paragraph(
        "score = recovery - 2.0 mean masked fraction - 0.5 false-positive fraction",
        style="Equation",
    )
    doc.add_paragraph(
        "A toy is counted as detected when at least 50% of its truth pixels are recovered. A hard worst-image masked "
        "fraction limit of 15% prevents a high recovery result from winning through excessive masking. Trials over the "
        "limit receive a large cap-excess objective penalty."
    )

    doc.add_heading("6. Optimised parameters and constraints", level=1)
    doc.add_paragraph(
        "The search domains are the established constraints already implemented by the project MTObjects optimiser. "
        "Parameters not listed remain fixed at the processing-model defaults."
    )
    add_table(
        doc,
        ["Parameter", "Constraint", "Interpretation"],
        [
            ["move_factor", "0.0-1.0", "Tree-node movement control"],
            ["min_distance", "0.0-1.0", "Minimum hierarchy distance"],
            ["gaussian_fwhm", "0.0-5.0 pixels", "Gaussian smoothing applied before segmentation"],
            ["bg_variance", "0.0001-10000; step 0.0001", "Background variance supplied to MTObjects"],
            ["minarea", "1-80 pixels", "Minimum retained segment area"],
            ["dilation_radius", "0-8 pixels", "Expansion of the accepted foreground mask"],
            ["max_area", "20-3000 pixels", "Upper area gate protecting galaxy structure"],
            ["max_elongation", "1.5-20", "Shape gate rejecting extreme segments"],
        ],
        [2200, 2900, 4260],
    )

    doc.add_heading("7. Candidate evaluation and selection", level=1)
    doc.add_paragraph(
        "Each fold contributes its best training parameter set. That candidate is scored on the fold's held-out 10 "
        "galaxies and on a common independent Toy Object realization across all 40 galaxies. The common set removes "
        "differences caused by comparing candidates on different held-out galaxies. The smallest all-40 objective is "
        "selected for the final batch."
    )

    doc.add_heading("8. Outputs", level=1)
    add_table(
        doc,
        ["Artifact", "Review purpose"],
        [
            ["Fold trial summaries", "Inspect every tested parameter combination and objective component"],
            ["Held-out details", "Review per-galaxy generalisation and failures"],
            ["Candidate comparison", "Compare training, held-out and common all-40 metrics"],
            ["Best JSON", "Preserve winning parameters and full provenance"],
            ["182-galaxy summary", "Confirm PNG status, runtime, mask fraction and errors"],
        ],
        [3100, 6260],
    )

    doc.add_heading("9. Restart and monitoring", level=1)
    doc.add_paragraph(
        "The workflow stores every Optuna study in SQLite and writes fold summaries incrementally. Supervisors detect "
        "unexpected launcher exits, reuse folds with all 40 trials, and restart incomplete folds or batches. Batch resume "
        "mode skips galaxies already present in the summary. Progress, remaining trials, rough ETA and expected completion "
        "are retained in terminal logs."
    )

    doc.add_heading("10. Interpretation", level=1)
    doc.add_paragraph(
        "The preferred solution should combine stable parameters across folds with strong held-out recovery, controlled "
        "false-positive masking and no repeated dependence on extreme parameter bounds. The common 40-galaxy evaluation "
        "uses new injected objects but not a wholly external galaxy sample, so it measures injection-level generalisation "
        "within the selected low-foreground population."
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()

