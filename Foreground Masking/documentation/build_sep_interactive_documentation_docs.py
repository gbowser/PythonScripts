#!/usr/bin/env python3
"""Build documentation for SEP + Spike Gate Interactive."""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path(__file__).resolve().parent

DOCS = {
    "spike_method": OUT_DIR / "SEP_Spike_Gate_Interactive_Methodology_and_Parameters.docx",
    "spike_program": OUT_DIR / "SEP_Spike_Gate_Interactive_Program_Documentation.docx",
}

ACCENT = "2E74B5"
DARK = "1F4D78"
INK = "0B2545"
HEADER_FILL = "E8EEF5"
CALLOUT_FILL = "F4F6F9"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[float]) -> None:
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(int(width * 1440) for width in widths)))
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    if grid is not None:
        for child in list(grid):
            grid.remove(child)
        for width in widths:
            col = OxmlElement("w:gridCol")
            col.set(qn("w:w"), str(int(width * 1440)))
            grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def style_document(doc: Document, title_text: str, subtitle: str) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in (
        ("Heading 1", 16, ACCENT, 16, 8),
        ("Heading 2", 13, ACCENT, 12, 6),
        ("Heading 3", 12, DARK, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.text = title_text
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = RGBColor(90, 90, 90)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.text = f"Generated {date.today().isoformat()}"
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = RGBColor(90, 90, 90)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run(title_text)
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(INK)

    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(14)
    sub_run = sub.add_run(subtitle)
    sub_run.font.size = Pt(11)
    sub_run.font.color.rgb = RGBColor(85, 85, 85)


def add_callout(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT_FILL)
    para = cell.paragraphs[0]
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run(text)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(INK)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for cell, text in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, HEADER_FILL)
        para = cell.paragraphs[0]
        para.paragraph_format.space_after = Pt(0)
        run = para.add_run(text)
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string(INK)
    for row_values in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row_values):
            set_cell_margins(cell)
            para = cell.paragraphs[0]
            para.paragraph_format.space_after = Pt(0)
            para.add_run(text)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        para = doc.add_paragraph(style="List Bullet")
        para.paragraph_format.space_after = Pt(4)
        para.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        para = doc.add_paragraph(style="List Number")
        para.paragraph_format.space_after = Pt(4)
        para.add_run(item)


def add_code(doc: Document, command: str) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.25)
    run = para.add_run(command)
    run.font.name = "Consolas"
    run.font.size = Pt(10)


SEP_PARAMETERS = [
    ["Detect on", "residual", "Detection image selector. Residual mode subtracts a broad smooth galaxy model before SEP detection.", "Original mode follows raw galaxy light more strongly; residual mode favours compact excesses."],
    ["Detection threshold", "3.0 sigma", "Threshold passed to sep.extract for the standard SEP mask.", "Higher values retain only stronger peaks; lower values detect more candidate pixels."],
    ["Minimum area", "5 px", "Minimum connected pixel count for a detection.", "Higher values reject smaller detections; lower values accept smaller compact sources."],
    ["Deblend thresholds", "32", "Number of internal levels used to split blended detections.", "Higher values can split complex blends more finely."],
    ["Deblend contrast", "0.005", "Minimum contrast for retaining a deblended child object.", "Higher values merge more subcomponents; lower values split more aggressively."],
    ["Background box", "64 px", "SEP background mesh size.", "Higher values produce a smoother, less local background estimate."],
    ["Filter size", "3 px", "Gaussian-like convolution kernel size used before detection.", "Higher values smooth more strongly and can broaden compact peaks."],
    ["Mask dilation", "2 px", "Circular dilation radius applied after filtering accepted segments.", "Higher values remove more surrounding pixels around each object footprint."],
    ["Max segment area", "230 px", "Rejects accepted segments larger than this area.", "Higher values permit broader detections; lower values protect extended galaxy structure more strongly."],
    ["Max elongation", "6.0", "Rejects objects whose major/minor axis ratio is too high.", "Higher values allow more elongated detections to survive."],
    ["Central exclusion", "8 px", "Rejects detections whose centroids lie inside the central galaxy region.", "Higher values protect more of the nucleus/bar centre."],
]


SPIKE_PARAMETERS = [
    ["Low-threshold SEP nsigma", "0.5", "Detection threshold for the second, spike-gate candidate SEP pass.", "Lower values give the gate more candidate footprints; higher values require stronger candidate segments."],
    ["Spike excess fraction", "0.25", "A profile sample must exceed its local neighbour level by this fraction.", "Higher values make spike detection stricter."],
    ["Neighbour inner", "4.0 arcsec", "Inner radius of the comparison region around a candidate profile peak.", "Higher values ignore closer neighbours."],
    ["Neighbour outer", "15.0 arcsec", "Outer radius of the comparison region around a candidate profile peak.", "Higher values use a broader profile context."],
    ["Side offset", "3 samples", "Profile offset used for the immediate side-drop test.", "Higher values test a wider, less local drop around the spike."],
    ["Side drop fraction", "0.4", "A candidate must stand above samples at +/- side offset by this fraction.", "Higher values require a sharper, more isolated spike."],
    ["Spike window", "2 samples", "Expands detected spike samples along the profile before intersecting with SEP footprints.", "Higher values allow nearby segment footprints to count as spike-associated."],
    ["Central exclusion", "Shared with SEP", "The spike gate uses SEP Central exclusion converted from pixels to arcsec.", "Only one central exclusion setting controls both standard SEP filtering and spike-profile detection."],
]


def add_common_sep_method(doc: Document) -> None:
    doc.add_heading("Method Overview", level=1)
    add_bullets(
        doc,
        [
            "Load a 2D S4G FITS image and manifest geometry for the selected galaxy.",
            "Prepare either the original image or a residual image made by subtracting a Gaussian-smoothed galaxy model.",
            "Estimate and subtract the background with sep.Background.",
            "Detect sources with sep.extract using the selected threshold, area, filter, and deblend settings.",
            "Measure candidate area, centroid, elongation, peak residual significance, and distance from the galaxy centre.",
            "Filter candidates by maximum area, maximum elongation, and central exclusion.",
            "Dilate retained object footprints to include nearby affected pixels.",
            "Project the diagnostics into deprojected bar-aligned coordinates for image, isophote, and profile checks.",
        ],
    )


def build_spike_method_doc() -> Path:
    doc = Document()
    style_document(
        doc,
        "SEP + Spike Gate Interactive Methodology and Parameter Guide",
        "Methodology and parameter descriptions for interactive_sep_spike_gate_parameter_tester.py",
    )
    add_callout(
        doc,
        "Purpose: compare ordinary SEP profile masking against a spike-gated variant that only removes low-threshold SEP segments associated with narrow bar-profile spikes.",
    )
    doc.add_heading("Combined Method", level=1)
    add_common_sep_method(doc)
    doc.add_heading("Spike-Gate Extension", level=1)
    add_bullets(
        doc,
        [
            "First, the standard SEP products are generated with the normal SEP controls. This is the baseline comparison profile.",
            "Second, the program runs a low-threshold SEP pass. This pass deliberately produces more candidate footprints.",
            "The original bar-major profile is searched for narrow positive spike samples outside the shared central exclusion.",
            "Detected spike samples are expanded by the Spike window control.",
            "Each low-threshold SEP segment is deprojected into the bar-aligned display frame and converted to a profile footprint.",
            "A segment is kept only when its profile footprint intersects a detected spike sample.",
            "The resulting SEP + Spike Gate profile is plotted beneath the standard SEP processed profile using the same intensity scale.",
        ],
    )
    doc.add_heading("SEP Parameter Descriptions", level=1)
    add_table(doc, ["Parameter", "Default", "Meaning", "Effect of increasing"], SEP_PARAMETERS, [1.45, 0.9, 2.25, 1.9])
    doc.add_heading("Spike Gate Parameter Descriptions", level=1)
    add_table(doc, ["Parameter", "Default", "Meaning", "Effect of increasing"], SPIKE_PARAMETERS, [1.45, 0.9, 2.25, 1.9])
    doc.add_heading("Comparison Logic", level=1)
    doc.add_paragraph(
        "The standard SEP and SEP + Spike Gate profiles are intentionally displayed with the same semilog y-axis limits. This makes it easier to judge whether the spike gate removes isolated profile artifacts while leaving the broader galaxy profile intact."
    )
    doc.add_heading("Recommended Use", level=1)
    add_numbered(
        doc,
        [
            "Use normal SEP settings to establish the baseline profile behaviour.",
            "Set Low-threshold SEP nsigma low enough that candidate segments exist around possible profile spikes.",
            "Tune Spike excess and Side drop to avoid labelling broad galaxy structure as a spike.",
            "Use the shared SEP Central exclusion to protect the central galaxy and to suppress central spike-gate decisions.",
            "Compare the standard SEP and SEP + Spike Gate profiles directly; the spike gate adds value only if it removes narrow artifacts with less collateral profile masking.",
        ],
    )
    doc.add_heading("Limitations", level=1)
    add_bullets(
        doc,
        [
            "The gate is profile-driven and can miss contaminants that do not intersect the sampled bar-major profile.",
            "A real narrow galaxy feature can look spike-like if the criteria are too loose.",
            "The low-threshold SEP pass still depends on SEP footprint quality; the gate chooses among SEP segments rather than inventing its own object geometry.",
        ],
    )
    doc.save(DOCS["spike_method"])
    return DOCS["spike_method"]


def build_spike_program_doc() -> Path:
    doc = Document()
    style_document(
        doc,
        "SEP + Spike Gate Interactive Program Documentation",
        "Program documentation for interactive_sep_spike_gate_parameter_tester.py",
    )
    add_callout(doc, "Program: Foreground Masking/Interactive tools/interactive_sep_spike_gate_parameter_tester.py")
    doc.add_heading("How to Run", level=1)
    add_code(doc, 'python "Foreground Masking/Interactive tools/interactive_sep_spike_gate_parameter_tester.py"')
    add_bullets(doc, ["--manifest PATH overrides the default manifest.", "--pc Desktop or --pc Laptop selects configured machine-specific image and output paths."])
    doc.add_heading("Inputs", level=1)
    add_table(
        doc,
        ["Input", "Source", "Use"],
        [
            ["Manifest and FITS image", "Shared foreground manifest sources", "Select and load a galaxy image plus geometry."],
            ["SEP controls", "SEP labelled box", "Build the standard SEP baseline and define filtering/dilation shared by the low-threshold pass."],
            ["Spike Gate controls", "Spike Gate labelled box", "Define low-threshold candidate detection and profile-spike acceptance criteria."],
            ["Shared central exclusion", "SEP Central exclusion [px]", "Used both for SEP candidate filtering and spike-profile exclusion after conversion to arcsec."],
        ],
        [1.55, 2.2, 2.75],
    )
    doc.add_heading("Control Panel", level=1)
    add_table(
        doc,
        ["Group", "Controls"],
        [
            ["Machine/Galaxy", "Machine path selector and galaxy selector."],
            ["Detection display", "Detect on and Parameter units selectors."],
            ["Spike Gate", "Low-threshold SEP nsigma, spike excess fraction, neighbour inner/outer arcsec, side offset samples, side drop fraction, and spike window samples."],
            ["SEP", "Standard SEP threshold, min area, deblend, background, filter, dilation, max area, max elongation, and central exclusion."],
            ["Actions", "Calculate, Reset, and Open PNG Folder."],
        ],
        [1.5, 5.0],
    )
    doc.add_heading("Processing Pipeline", level=1)
    add_numbered(
        doc,
        [
            "Load image and geometry for the selected galaxy.",
            "Build standard SEP products using the SEP box controls.",
            "Build low-threshold SEP candidate products using Low-threshold SEP nsigma and the same SEP footprint/filtering controls.",
            "Detect spike samples on the original deprojected bar-major profile.",
            "Expand spike samples by Spike window.",
            "Retain only low-threshold SEP segment labels whose deprojected profile footprint intersects expanded spike samples.",
            "Draw image diagnostics, isophotes, original profile, SEP processed profile, and SEP + Spike Gate profile.",
            "Save a timestamped PNG diagnostic to the spike-gate output folder.",
        ],
    )
    doc.add_heading("Outputs", level=1)
    add_table(
        doc,
        ["Output", "Location", "Description"],
        [
            ["PNG diagnostic", r"{Remove foreground objects}\interactive_sep_spike_gate_parameter_tester", "Saved after each calculation. Filename currently follows the SEP naming pattern with galaxy, normal SEP threshold, area, deblend contrast, dilation, and timestamp."],
            ["Status line", "Control panel", "Reports standard SEP kept segments/masked fraction and spike-gate kept low-threshold segments/spike-sample count/masked fraction."],
            ["No FITS output", "Not written", "The program compares methods visually and does not save cleaned FITS products."],
        ],
        [1.45, 2.2, 2.85],
    )
    doc.add_heading("Figure Layout", level=1)
    add_table(
        doc,
        ["Panel", "Purpose"],
        [
            ["Centered original", "Original bar-aligned image for orientation."],
            ["Residual detection image", "Residual image used by residual-mode SEP detection."],
            ["Original isophotes", "Contour view of the original image."],
            ["Original bar-major profile", "Profile reference, placed in the right column with the processed profiles."],
            ["SEP processed isophotes", "Contour view of the standard SEP median-filled preview."],
            ["SEP processed bar-major profile", "Baseline SEP profile with masked samples bridged."],
            ["SEP & Spike Gate bar profile", "Spike-gated profile below the baseline profile, using the same intensity y-axis scale."],
        ],
        [2.25, 4.25],
    )
    doc.add_heading("Implementation Notes", level=1)
    add_bullets(
        doc,
        [
            "The spike gate does not create its own area geometry; SEP provides candidate segment footprints.",
            "The shared central exclusion is stored as pixels in the SEP controls and converted to arcsec for spike-profile detection.",
            "Profile panels share semilog y-axis limits for direct visual comparison.",
            "The standard SEP code path remains present in this script so the comparison is generated in one Calculate action.",
        ],
    )
    doc.save(DOCS["spike_program"])
    return DOCS["spike_program"]


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    paths = [
        build_spike_method_doc(),
        build_spike_program_doc(),
    ]
    for path in paths:
        print(path)


if __name__ == "__main__":
    main()
