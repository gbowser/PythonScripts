#!/usr/bin/env python3
"""Build the native DOCX report for clean-galaxy selection and toy optimisation."""

from __future__ import annotations

import argparse
import csv
import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open(newline="", encoding="utf-8") as handle:
        return list(csv.DictReader(handle))


def shade(cell, fill: str) -> None:
    props = cell._tc.get_or_add_tcPr()
    element = OxmlElement("w:shd")
    element.set(qn("w:fill"), fill)
    props.append(element)


def add_table(doc: Document, headers: list[str], rows: list[list[object]], widths=None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        shade(cell, "D9EAF7")
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = str(value)
    if widths:
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Inches(width)
    doc.add_paragraph()


def add_callout(doc: Document, text: str, color: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    shade(table.cell(0, 0), color)
    table.cell(0, 0).text = text
    doc.add_paragraph()


def add_params(doc: Document, params: dict) -> None:
    add_table(doc, ["Parameter", "Value"], [[key, value] for key, value in params.items()], [2.5, 4.0])


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--root", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    args = parser.parse_args()
    ranking_path = args.root.parent / "final_cleanest20_severity_review" / "final_cleanest20_ranking.csv"
    sep_run = args.root / "SEP" / "20260828_064912"
    mto_run = args.root / "MTObjects_log_variance" / "20260828_071110"
    ranking = [row for row in read_csv(ranking_path) if row["selected_top20"] == "yes"]
    sep = json.loads((sep_run / "sep_toy_object_optimisation_best.json").read_text(encoding="utf-8"))
    mto = json.loads((mto_run / "mtobjects_parameter_optimisation_best.json").read_text(encoding="utf-8"))

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65); section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7); section.right_margin = Inches(0.7)
    styles = doc.styles
    styles["Normal"].font.name = "Calibri"; styles["Normal"].font.size = Pt(10.5)
    for name, size, color in (("Title", 24, "17365D"), ("Heading 1", 18, "17365D"), ("Heading 2", 13, "365F91")):
        styles[name].font.name = "Calibri"; styles[name].font.size = Pt(size); styles[name].font.color.rgb = RGBColor.from_string(color)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Selection of Low-Contamination Galaxies and Toy-Object Mask Optimisation")
    subtitle = doc.add_paragraph("Final report — 28 August 2026")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("Executive summary", level=1)
    doc.add_paragraph("All 182 S4G galaxy fields were reviewed. Under the final blind visual criterion, 10 were Clean, 4 Ambiguous and 168 Polluted. Because fewer than 20 were strictly Clean, a final severity exercise selected the 20 least polluted fields. These 20 are calibration fields; they are not 20 equally uncontaminated galaxies.")
    doc.add_paragraph(f"SEP completed 40 paired-toy trials and produced a feasible candidate with toy detection rate {sep['toy_detection_rate']:.1%} and mean toy recall {sep['mean_toy_recall']:.1%}. After correcting background variance to logarithmic sampling, MTObjects also produced a feasible candidate with toy detection rate {mto['toy_detection_rate']:.1%} and mean toy recall {mto['mean_toy_recall']:.1%}.")

    doc.add_heading("1. Objective and definitions", level=1)
    doc.add_paragraph("The objective was to find fields with the fewest bright unrelated foreground objects for foreground-mask calibration. Galaxy nuclei, bars, rings, arms and coherent star-forming structure were not contaminants. Global visual cleanliness and bar-profile-specific impact were treated separately.")
    doc.add_heading("2. Selection process", level=1)
    doc.add_heading("2.1 Automated evidence", level=2)
    doc.add_paragraph("A Gaussian-smoothed 3.6 μm image was subtracted from the centred original, the central galaxy region was excluded, and positive compact residuals were measured. Photutils/DAO-style detections were supplemented with Gaia DR3 astrometric evidence and 2MASS point-source evidence. Detections on strong galaxy structure were downweighted, not discarded.")
    doc.add_paragraph("Image-only scoring remained sensitive to spiral arms, rings and star-forming knots. NGC1097 demonstrated that automated evidence could not replace visual assessment.")
    doc.add_heading("2.2 Review stages", level=2)
    add_table(doc, ["Stage", "N", "Clean", "Ambiguous", "Polluted", "Basis"], [
        ["Catalogue batch 1",49,11,7,31,"Gaia-zero/hybrid"],
        ["Catalogue batch 2",30,0,1,29,"Next positive-score fields"],
        ["Catalogue batch 3",30,0,0,30,"Clean-reference similarity"],
        ["Blind consistency audit",21,10,3,8,"Original-only shuffled audit"],
        ["Remaining blind review",71,0,1,70,"All previously unreviewed fields"],
    ])
    doc.add_paragraph("Catalogue panels contained the original, Gaussian residual and catalogue overlay. Red marked scored 2MASS sources and yellow marked weaker Gaia evidence. Overlays were supporting evidence rather than automatic labels.")
    doc.add_heading("2.3 Ambiguous-profile experiment", level=2)
    doc.add_paragraph("Eight ambiguous fields were tested with manually positioned circular masks and before/after bar-major-axis profiles. Three were labelled Clean and five Polluted. All measured profile changes were zero because the masks did not intersect the narrow profile aperture; the labels described global field contamination, not demonstrated profile contamination.")
    doc.add_heading("2.4 Blind consistency and full-population review", level=2)
    doc.add_paragraph("A 21-field identity-hidden, original-only audit gave 10 Clean, 3 Ambiguous and 8 Polluted, with 14/19 (74%) agreement with previous decisions. The remaining 71 fields were reviewed identically: 0 Clean, 1 Ambiguous and 70 Polluted. Final full-population totals were 10 Clean, 4 Ambiguous and 168 Polluted.")
    doc.add_heading("2.5 Severity shortlist", level=2)
    doc.add_paragraph("A blind shortlist combined 10 Clean, 4 Ambiguous and 16 low-score Polluted fields. Severity was 0 none, 1 minor, 2 moderate and 3 severe. Counts were 12, 2, 4 and 12. Severity was primary; prior blind group and then clean-reference similarity broke ties.")
    add_callout(doc, "Only 18 fields scored 0–2. Positions 19–20 cross the severity-3 boundary: NGC4102 was retained from the Ambiguous group and NGC0918 was the nearest remaining Polluted field to the Clean feature pattern.", "EAF2F8")

    doc.add_heading("3. Final cleanest 20", level=1)
    add_table(doc, ["Rank", "Galaxy", "Severity", "Prior group", "Selection basis"], [[r["final_rank"],r["name"],r["severity"],r["prior_blind_group"],r["selection_basis"]] for r in ranking])
    doc.add_paragraph("Ranks within a severity tier are tie-break order, not a continuous contamination measurement. The first 10 are blind-confirmed Clean; ranks 11–20 broaden the calibration population to the requested 20.")

    doc.add_heading("4. Paired Toy Objects design", level=1)
    add_table(doc, ["Design item", "Value"], [
        ["Calibration galaxies",20],["Injection sets","Cross-validation and winner-selection"],["Toys per galaxy per set",6],
        ["Toys per set",120],["Total materialised toys",240],["Type mixture","50% stars; 20% clusters; 30% galaxies"],
        ["Peak amplitude","6–30 robust image sigma"],["Truth dilation","1 pixel"],["Trials per algorithm","40 (8 initial + 32 adaptive)"],
        ["Workers",4],["Detection image","Original science image"],
    ])
    doc.add_paragraph("One immutable paired manifest supplied identical toys to SEP and MTObjects. It records per-galaxy seeds and SHA-256 checksums for science images, injection payloads, deltas and truth masks. The independent winner-selection set was generated but not used in these initial runs.")

    doc.add_heading("5. SEP result", level=1)
    add_table(doc, ["Metric", "SEP optimum"], [
        ["Status",sep["status"]],["Objective (minimised)",f'{sep["objective"]:.6f}'],["Recovery score",f'{sep["score"]:.6f}'],
        ["Toy detection rate",f'{sep["toy_detection_rate"]:.2%}'],["Mean toy recall",f'{sep["mean_toy_recall"]:.2%}'],
        ["Mean pixel recall",f'{sep["mean_recall"]:.2%}'],["Mean pixel precision",f'{sep["mean_precision"]:.3%}'],
        ["Mean F-score",f'{sep["mean_f_score"]:.4f}'],["Mean masked fraction",f'{sep["mean_masked_fraction"]:.2%}'],
        ["Worst masked fraction",f'{sep["max_masked_fraction"]:.2%}'],["False-positive fraction",f'{sep["false_positive_fraction"]:.2%}'],
    ])
    doc.add_heading("Best SEP parameters", level=2); add_params(doc, sep["params"])
    doc.add_paragraph("The SEP optimum respected the 15% worst-image ceiling (12.93%) and recovered about 41% of toys. Pixel precision was low because much incremental mask area lay outside toy truth. This is a feasible candidate requiring independent winner-selection and visual validation.")

    doc.add_heading("6. MTObjects result", level=1)
    add_table(doc, ["Metric", "MTObjects optimum"], [
        ["Status",mto["status"]],["Objective (minimised)",f'{mto["objective"]:.6f}'],["Recovery score",f'{mto["recovery_score"]:.6f}'],["Recovery-infeasible flag",mto["recovery_infeasible"]],
        ["Toy detection rate",f'{mto["toy_detection_rate"]:.2%}'],["Mean toy recall",f'{mto["mean_toy_recall"]:.2%}'],
        ["Mean pixel recall",f'{mto["mean_recall"]:.2%}'],["Mean pixel precision",f'{mto["mean_precision"]:.3%}'],
        ["Mean F-score",f'{mto["mean_f_score"]:.4f}'],["Mean masked fraction",f'{mto["mean_masked_fraction"]:.2%}'],
        ["Worst masked fraction",f'{mto["max_masked_fraction"]:.2%}'],["False-positive fraction",f'{mto["false_positive_fraction"]:.2%}'],
    ])
    doc.add_heading("Best MTObjects parameters", level=2); add_params(doc, mto["params"])
    add_callout(doc, "Resolved run failure: uniform background-variance sampling from 0.0001 to 10,000 produced zero recovery because almost all trials suppressed detections. Re-running with logarithmic sampling found bg_variance = 0.000583 and a feasible optimum. This candidate still requires galaxy-fold validation before production use.", "EAF2F8")

    doc.add_heading("7. Recommended next step", level=1)
    doc.add_paragraph("Run the planned 20-fold leave-one-galaxy-out validation. Each fold optimises on 19 galaxies and validates on one held-out galaxy using the independent winner-selection injections. Compare all 20 fold candidates on the common 20-galaxy, 120-toy winner-selection set. Inspect the low pixel precision of both preflight candidates before production use.")
    doc.add_heading("8. Reproducibility paths", level=1)
    add_table(doc, ["Artifact", "Path"], [
        ["Final ranking",ranking_path],["Clean-list input",Path(__file__).parents[1]/"Optimisation"/"clean_galaxies_final20.txt"],
        ["Paired manifest",args.root/"paired_injections"/"paired_toy_injection_manifest.json"],["SEP run",sep_run],["MTObjects run",mto_run],
    ])
    doc.sections[-1].footer.paragraphs[0].text = "Clean-galaxy selection and paired Toy Objects optimisation — 28 August 2026"
    args.output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(args.output)
    print(f"DOCX={args.output}")
    return 0


if __name__ == "__main__": raise SystemExit(main())
