#!/usr/bin/env python3
"""Build a short DOCX note explaining SEP vs Spike-Gate detections."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


DOC_DIR = Path(__file__).resolve().parent
OUT_PATH = DOC_DIR / "SEP Spike Gate Detection Note.docx"


def set_run_font(run, size_pt: float | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = "Calibri"
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    run = paragraph.add_run(text)
    set_run_font(run, 11)


def add_heading(doc: Document, text: str, level: int) -> None:
    paragraph = doc.add_heading(text, level=level)
    paragraph.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    paragraph.paragraph_format.space_after = Pt(7 if level == 1 else 5)
    for run in paragraph.runs:
        set_run_font(run, 16 if level == 1 else 13, color="2E74B5" if level == 1 else "1F4D78")


def add_body(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.25
    run = paragraph.add_run(text)
    set_run_font(run, 11)


def build_doc() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(10)
    title_run = title.add_run("SEP and Spike-Gate Detection Note")
    set_run_font(title_run, 20, True, "0B2545")

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    subtitle_run = subtitle.add_run("Why the +48 bar-profile spike can be missed by normal SEP while Spike Gate catches it")
    set_run_font(subtitle_run, 11, False, "555555")

    add_heading(doc, "Summary", 1)
    add_body(
        doc,
        "The short version is that normal SEP is not looking at the bar-profile peak directly. "
        "With the screenshot settings, normal SEP detects objects on the 2D residual image, then applies the configured segment filters.",
    )
    add_body(
        doc,
        "The Spike Gate path is different: it first finds candidate spike positions in the 1D bar-major profile, then asks a low-threshold SEP pass which 2D segments intersect those spike radii.",
    )

    add_heading(doc, "Why the +48 Feature Can Be Missed", 1)
    add_bullet(doc, "The normal SEP pass uses detect_on=residual, so SEP sees the residual image rather than the original 1D bar-major profile.")
    add_bullet(doc, "The normal SEP Detection Threshold in the screenshot is high: thresh=7.6.")
    add_bullet(doc, "A feature can look large in the 1D profile but still fail to become a strong enough 2D residual segment after background/model subtraction.")
    add_bullet(doc, "SEP may detect the object initially but then reject it through the configured filters: maximum area, maximum elongation, or central exclusion.")

    add_heading(doc, "Why the -25 Feature Is Different", 1)
    add_body(
        doc,
        "The -25-ish feature appears to be strong enough in the 2D residual detection image to pass the normal SEP detection and filtering route. "
        "That is why it appears in the normal SEP processed profile.",
    )

    add_heading(doc, "What the Bottom Panel Shows", 1)
    add_body(
        doc,
        "The SEP & Spike Gate panel does catch the +48 region because that path uses the low-threshold SEP setting, thresh=0.5 in the screenshot, "
        "then keeps only low-threshold SEP segments that overlap radii flagged by the 1D spike detector.",
    )
    add_body(
        doc,
        "That behaviour is the intended value of the Spike Gate: it can catch profile-damaging spikes that the stricter normal SEP pass misses.",
    )

    add_heading(doc, "Useful Future Diagnostic", 1)
    add_body(
        doc,
        "A useful next step would be a small diagnostic readout for each spike candidate with: radius, profile value, neighbour level, side level, "
        "SEP label found, and kept/rejected reason. That would make it clear whether a feature was missed by normal SEP thresholding or rejected by a later filter.",
    )

    doc.save(OUT_PATH)


if __name__ == "__main__":
    build_doc()
    print(OUT_PATH)
