#!/usr/bin/env python3
"""Build documentation for the interactive MTObjects parameter tester."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path(__file__).resolve().parent
CODE_PATH = OUT_DIR.parent / "interactive_mtobjects_parameter_tester.py"
ALGORITHM_DOC = OUT_DIR / "interactive_mtobjects_algorithm_and_parameters.docx"
FLOW_DOC = OUT_DIR / "interactive_mtobjects_code_process_and_flow.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[float]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    total_dxa = 9360
    tbl_w.set(qn("w:w"), str(total_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    width_dxa = [int(round(width * 1440)) for width in widths]
    for width in width_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_w = cell._tc.get_or_add_tcPr().tcW
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(width_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")


def style_doc(doc: Document, title: str, subtitle: str) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    title_para = doc.add_paragraph()
    title_para.paragraph_format.space_after = Pt(4)
    run = title_para.add_run(title)
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("0B2545")

    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(12)
    sub_run = sub.add_run(subtitle)
    sub_run.font.name = "Calibri"
    sub_run.font.size = Pt(11)
    sub_run.font.color.rgb = RGBColor.from_string("555555")


def add_note(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    para = cell.paragraphs[0]
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run(text)
    run.font.bold = True


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        para = doc.add_paragraph(style="List Bullet")
        para.paragraph_format.left_indent = Inches(0.375)
        para.paragraph_format.first_line_indent = Inches(-0.188)
        para.paragraph_format.space_after = Pt(4)
        para.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        para = doc.add_paragraph(style="List Number")
        para.paragraph_format.left_indent = Inches(0.375)
        para.paragraph_format.first_line_indent = Inches(-0.188)
        para.paragraph_format.space_after = Pt(4)
        para.add_run(item)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for index, header in enumerate(headers):
        cell = table.cell(0, index)
        set_cell_shading(cell, "E8EEF5")
        run = cell.paragraphs[0].add_run(header)
        run.font.bold = True
    for row_values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_values):
            cells[index].text = value
    set_table_geometry(table, widths)


def build_algorithm_doc() -> None:
    doc = Document()
    style_doc(
        doc,
        "Interactive MTObjects Tester: Algorithm and Parameters",
        "Reference guide for the MTObjects foreground-mask backend and tuning controls.",
    )
    add_note(
        doc,
        "Scope: this document describes the new MTObjects tester, not the SEP implementation. "
        "It keeps the same manifest, geometry, profile, and PNG review workflow so future comparisons are straightforward.",
    )

    doc.add_heading("Algorithm Summary", level=1)
    doc.add_paragraph(
        "MTObjects detects astronomical sources by building a max tree of the preprocessed image and applying statistical "
        "attribute filtering to identify significant connected structures. The project README describes it as a tool for "
        "detecting sources, creating segmentation maps, and producing parameter tables; the implementation is based on the "
        "C work by Teeninga et al. and wrapped by Caroline Haigh's Python package."
    )
    add_numbered(
        doc,
        [
            "Prepare either the residual image or original image as the detection surface.",
            "Estimate or accept supplied background mean, variance, soft bias, and gain values.",
            "Subtract the background, optionally smooth with a Gaussian kernel, truncate negative values, and replace NaNs for MTObjects.",
            "Build a max tree over the preprocessed image.",
            "Run MTObjects statistical filtering using alpha, move_factor, and min_distance.",
            "Relabel the object map, restore non-finite pixels to background, and convert MTObjects labels to a boolean foreground mask.",
            "Apply comparison-layer filters: minimum area, maximum area, elongation, central exclusion, and mask dilation.",
        ],
    )

    doc.add_heading("MTObjects Parameters", level=1)
    add_table(
        doc,
        ["Parameter", "Default", "Effect"],
        [
            ["alpha", "1e-6", "Significance level used by the MTObjects statistical test. Smaller values are stricter."],
            ["move_factor", "0.5", "Controls how far object markers move upward in the tree. Higher values reduce spread of large objects."],
            ["min_distance", "0.0", "Minimum brightness separation required between objects."],
            ["gaussian_fwhm", "2.0 px", "FWHM used by MTObjects preprocessing before tree construction; 0 disables smoothing."],
            ["soft_bias", "0.0", "Constant bias term used when estimating gain."],
            ["gain", "-1.0", "Electrons per ADU. A negative value asks MTObjects to estimate it."],
            ["bg_mean", "NaN", "Background mean. NaN asks MTObjects to estimate it from flat tiles."],
            ["bg_variance", "-1.0", "Background variance. A negative value asks MTObjects to estimate it."],
        ],
        [1.55, 1.0, 3.95],
    )

    doc.add_heading("Post-Detection Parameters", level=1)
    add_table(
        doc,
        ["Parameter", "Default", "Effect"],
        [
            ["minarea", "5 px", "Rejects small detected segments after MTObjects relabeling."],
            ["max_area", "230 px", "Rejects broad segments that are more likely galaxy/bar structure than compact foreground objects."],
            ["max_elongation", "6.0", "Rejects very elongated detections using weighted second moments."],
            ["exclude_center_pixels", "8 px", "Rejects detections close to the galaxy centre, with arcsec display conversion available."],
            ["dilation_radius", "2 px", "Expands accepted segments before replacing/masking pixels."],
        ],
        [1.75, 0.9, 3.85],
    )

    doc.add_heading("Spike-Gate Parameters", level=1)
    doc.add_paragraph(
        "The spike-gate run is retained as an optional comparison view. It runs MTObjects with a usually more inclusive "
        "move_factor, then keeps only segments intersecting bar-major-axis spike samples."
    )
    add_table(
        doc,
        ["Parameter", "Default", "Effect"],
        [
            ["spike_gate_move_factor", "0.3", "Move factor for the gate-run segmentation."],
            ["spike_excess_fraction", "0.25", "Required fractional excess above neighbouring bar-profile levels."],
            ["spike_neighbour_inner_arcsec", "4.0", "Inner radius of the local comparison annulus along the profile."],
            ["spike_neighbour_outer_arcsec", "15.0", "Outer radius of the local comparison annulus along the profile."],
            ["spike_side_offset_samples", "3", "Sample offset used to verify a local side drop."],
            ["spike_side_drop_fraction", "0.4", "Required excess over side samples."],
            ["spike_window_samples", "2", "Number of neighbouring profile samples added around detected spikes."],
        ],
        [2.35, 0.7, 3.45],
    )

    doc.add_heading("Operational Notes", level=1)
    add_bullets(
        doc,
        [
            "MTObjects currently expects compiled shared libraries in the MTObjects checkout; use --mtobjects-root or MTOBJECTS_ROOT.",
            "The tester treats MTObjects as the detector and keeps comparison filters outside the library adapter.",
            "Mask replacement currently uses the median of finite unmasked pixels, matching the SEP tester's simple visualization-oriented cleaning.",
            "The MTObjects README warns the project is a work in progress, so parameter sweeps should record library version or commit when used for comparisons.",
        ],
    )
    doc.add_heading("Sources", level=1)
    add_bullets(
        doc,
        [
            "CarolineHaigh/mtobjects README: source detection, segmentation maps, dependencies, and command-line parameters.",
            "CarolineHaigh/mtobjects source files: main.py, preprocessing.py, tree_filtering.py, maxtree.py, and postprocessing.py.",
            f"Local implementation: {CODE_PATH.name}.",
        ],
    )
    doc.save(ALGORITHM_DOC)


def build_flow_doc() -> None:
    doc = Document()
    style_doc(
        doc,
        "Interactive MTObjects Tester: Code Process and Flow",
        "Implementation walkthrough for maintaining and extending the new tester.",
    )
    add_note(
        doc,
        "Design intent: MTObjects is isolated behind a product-producing adapter so SEP, MTObjects, and future detectors can be compared through common masks, rows, profiles, and PNG outputs.",
    )

    doc.add_heading("Runtime Entry Point", level=1)
    add_numbered(
        doc,
        [
            "parse_args reads --manifest, --pc, and --mtobjects-root.",
            "main creates MTObjectsTester with the selected manifest, machine, and MTObjects checkout path.",
            "MTObjectsTester loads the local manifest through foreground_display_helpers.",
            "The Tkinter UI exposes machine, galaxy, detection surface, units, MTObjects parameters, spike-gate controls, and post-filters.",
        ],
    )

    doc.add_page_break()
    doc.add_heading("Calculation Flow", level=1)
    add_table(
        doc,
        ["Stage", "Function", "Responsibility"],
        [
            ["Load FITS", "load_fits / _load_galaxy", "Read image data and geometry, with a per-galaxy cache."],
            ["Detection image", "prepare_detection_image", "Create original or smoothed-residual detection image and non-finite mask."],
            ["MTObjects setup", "mtobjects_context", "Temporarily switch into the MTObjects checkout so its compiled libraries load correctly."],
            ["Parameter object", "mtobjects_parameter_namespace", "Build the SimpleNamespace expected by MTObjects internals."],
            ["Tree detection", "mtobjects_products", "Preprocess image, build max tree, filter tree, relabel segments, and produce raw rows."],
            ["Post-filtering", "measure_segments / filter_segmentation", "Measure labels and apply min/max area, elongation, centre exclusion, and dilation."],
            ["Spike gate", "spike_gated_mtobjects_products", "Run the gate MTObjects pass and keep labels intersecting detected bar-profile spikes."],
            ["Visual output", "draw_products", "Render original, residual, isophotes, profiles, masks, and spike samples; save a PNG."],
        ],
        [1.25, 2.55, 2.7],
    )

    doc.add_page_break()
    doc.add_heading("Product Contract", level=1)
    doc.add_paragraph(
        "Both the normal MTObjects run and spike-gated run return a dictionary shaped like the SEP product boundary. "
        "This keeps downstream drawing and future benchmark code simple."
    )
    add_table(
        doc,
        ["Key", "Type", "Meaning"],
        [
            ["raw_segmentation", "2D int array", "Relabeled MTObjects segmentation before comparison-layer filtering."],
            ["filtered_segmentation", "2D int array", "Segmentation after geometry and size filters."],
            ["mask", "2D bool array", "Dilated foreground mask used for visual/profile cleaning."],
            ["cleaned", "2D float array", "Original image with mask pixels replaced by the finite unmasked median."],
            ["residual", "2D float array", "Original minus smoothed model, used for display and measurements."],
            ["rows", "list[dict]", "Per-label measurements and keep/reject state."],
            ["background_level / background_rms", "float", "MTObjects estimated or supplied background values for audit display."],
        ],
        [1.75, 1.35, 3.4],
    )

    doc.add_heading("Comparison Readiness", level=1)
    add_bullets(
        doc,
        [
            "Detector-specific work is constrained to mtobjects_products and spike_gated_mtobjects_products.",
            "Rows contain detector-neutral fields: label, area, centroid, elongation, peak residual significance, centre distance, and kept.",
            "The output folder and PNG naming include MTObjects-specific alpha and move_factor values.",
            "The display uses the same bar-aligned profiles and isophote helpers as the existing foreground-mask testers.",
        ],
    )

    doc.add_heading("Error Handling and Setup", level=1)
    add_bullets(
        doc,
        [
            "If MTObjects cannot import or load its shared libraries, the adapter raises a setup-focused error message.",
            "Non-finite pixels are replaced only for MTObjects preprocessing, then restored to background in the segmentation.",
            "The max-tree C memory is released in a finally block after filtering.",
            "The script syntax-checks without MTObjects installed because imports occur inside the runtime adapter.",
        ],
    )

    doc.add_heading("Maintenance Notes", level=1)
    add_bullets(
        doc,
        [
            "For a formal SEP versus MTObjects benchmark, call detector product functions directly and compare mask fraction, object counts, profile residuals, runtime, and manual review flags.",
            "Keep new detector adapters returning the same product contract rather than branching the plotting code.",
            "If MTObjects is packaged differently later, update mtobjects_context only; the rest of the tester should not need to change.",
            "If replacement semantics change from median fill to interpolation/inpainting, implement that after product creation so detectors remain comparable.",
        ],
    )
    doc.save(FLOW_DOC)


def main() -> None:
    build_algorithm_doc()
    build_flow_doc()
    print(ALGORITHM_DOC)
    print(FLOW_DOC)


if __name__ == "__main__":
    main()
