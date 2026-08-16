#!/usr/bin/env python3
"""Build a Word summary after SEP and MTObjects Toy Objects runs complete."""

from __future__ import annotations

import argparse
import csv
import json
from pathlib import Path
from statistics import mean, pstdev

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from build_toy_objects_cross_validation_docs import DARK_BLUE, add_header_footer, add_table, configure_styles


def csv_rows(path: Path) -> list[dict[str, str]]:
    with path.open(newline="", encoding="utf-8-sig") as handle:
        return list(csv.DictReader(handle))


def number(row: dict[str, str], key: str) -> float:
    try:
        return float(row.get(key, "nan"))
    except (TypeError, ValueError):
        return float("nan")


def algorithm_data(label: str, run: Path, batch_summary: Path, best_name: str) -> dict:
    candidates = csv_rows(run / "cross_validation_candidates.csv")
    details = csv_rows(run / "held_out_details.csv")
    best = json.loads((run / best_name).read_text(encoding="utf-8"))
    batch = csv_rows(batch_summary)
    return {"label": label, "run": run, "candidates": candidates, "details": details, "best": best, "batch": batch}


def fmt(value: object, digits: int = 4) -> str:
    try:
        return f"{float(value):.{digits}f}"
    except (TypeError, ValueError):
        return "n/a"


def fmt_parameter(value: object) -> str:
    if isinstance(value, float):
        return f"{value:.6g}"
    return str(value)


def add_algorithm_section(doc: Document, data: dict) -> None:
    label = data["label"]
    candidates = data["candidates"]
    best = data["best"]
    batch = data["batch"]
    doc.add_heading(label, level=1)
    add_table(
        doc,
        ["Fold", "Held-out score", "All-40 score", "All-40 max masked"],
        [[
            row.get("fold", ""),
            fmt(row.get("held_out_score")),
            fmt(row.get("all40_score")),
            fmt(row.get("all40_max_masked_fraction"), 3),
        ] for row in candidates],
        [1000, 2600, 2600, 3160],
    )
    doc.add_paragraph(
        f"Winning fold: {best.get('winning_fold', 'n/a')}. "
        f"Selection: {best.get('selection_method', 'common all-40 evaluation')}."
    )
    params = best.get("params", {})
    optimised = {
        "SEP / Toy Objects": {
            "detect_thresh", "minarea", "deblend_nthresh", "deblend_cont", "back_size",
            "filter_size", "dilation_radius", "max_area", "max_elongation",
        },
        "MTObjects / Toy Objects": {
            "move_factor", "min_distance", "gaussian_fwhm", "bg_variance", "minarea",
            "dilation_radius", "max_area", "max_elongation",
        },
    }[label]
    selected = [(str(key), fmt_parameter(value)) for key, value in params.items() if key in optimised]
    paired_rows = []
    for index in range(0, len(selected), 2):
        left = selected[index]
        right = selected[index + 1] if index + 1 < len(selected) else ("", "")
        paired_rows.append([left[0], left[1], right[0], right[1]])
    add_table(
        doc,
        ["Selected parameter", "Value", "Selected parameter", "Value"],
        paired_rows,
        [2300, 2380, 2300, 2380],
    )
    ok = sum(str(row.get("status", "")).casefold() == "ok" for row in batch)
    failed = len(batch) - ok
    masked = [number(row, "masked_fraction") for row in batch if row.get("masked_fraction") not in (None, "")]
    doc.add_paragraph(
        f"Final batch: {len(batch)} records; {ok} successful; {failed} failed. "
        f"Mean masked fraction: {fmt(mean(masked) if masked else None, 3)}."
    )
    if label.startswith("MTObjects") and (not masked or mean(masked) < 0.0005):
        doc.add_heading("MTObjects diagnostic finding", level=2)
        doc.add_paragraph(
            "The selected MTObjects solution is degenerate: recovery scores are effectively zero and the final batch "
            "masks almost no pixels. The selected bg_variance lies close to its upper search bound, while minarea is high "
            "and dilation is zero. This is computational completion, but it is not evidence of useful foreground recovery. "
            "The MTObjects PNGs should be treated as diagnostic output and the objective/search domain should be revised "
            "before these parameters are adopted scientifically."
        )


def build(args: argparse.Namespace) -> None:
    sep = algorithm_data("SEP / Toy Objects", args.sep_run, args.sep_batch, "sep_toy_cross_validation_best.json")
    mto = algorithm_data("MTObjects / Toy Objects", args.mtobjects_run, args.mtobjects_batch, "mtobjects_toy_cross_validation_best.json")
    doc = Document()
    configure_styles(doc)
    add_header_footer(doc)
    doc.sections[0].header.paragraphs[0].text = "Foreground Masking | Completed Toy Objects results"
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("SEP and MTObjects Toy Objects Results")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Four-fold cross-validation and 182-galaxy application summary").italic = True

    doc.add_heading("Executive summary", level=1)
    doc.add_paragraph(
        "Both algorithms were trained and evaluated using identical four-fold membership: 30 training galaxies and "
        "10 held out in each rotation. Candidate winners were selected on a common independent Toy Object realization "
        "across all 40 low-foreground galaxies before application to the 182-galaxy manifest."
    )
    add_algorithm_section(doc, sep)
    add_algorithm_section(doc, mto)

    doc.add_heading("Cross-algorithm comparison", level=1)
    comparison = []
    for data in (sep, mto):
        winning = data["best"].get("cross_validation_metrics", {})
        scores = [number(row, "held_out_score") for row in data["candidates"]]
        comparison.append([
            data["label"],
            str(data["best"].get("winning_fold", "")),
            fmt(winning.get("all40_score")),
            fmt(mean(scores)),
            fmt(pstdev(scores) if len(scores) > 1 else 0.0),
        ])
    add_table(
        doc,
        ["Algorithm", "Winner", "Winner all-40 score", "Mean held-out score", "Held-out SD"],
        comparison,
        [2200, 1100, 2100, 2100, 1860],
    )
    doc.add_paragraph(
        "Scores are algorithm-specific because SEP and MTObjects objective weights differ. Comparison should therefore "
        "consider fold stability, false-positive masking, worst-image mask fraction and the final PNGs, not score alone."
    )
    doc.add_heading("Overall conclusion", level=1)
    doc.add_paragraph(
        "SEP produced a non-zero, stable candidate and completed all 182 PNGs without failures. MTObjects also completed "
        "all requested computations and PNGs, but its optimisation selected a near-zero-mask solution in every fold. "
        "Consequently, SEP is the only scientifically actionable result from this run; MTObjects requires a constrained "
        "follow-up optimisation or a recovery objective that penalises non-detection explicitly."
    )
    doc.add_heading("Provenance", level=1)
    doc.add_paragraph(f"SEP run: {args.sep_run}")
    doc.add_paragraph(f"SEP batch: {args.sep_batch}")
    doc.add_paragraph(f"MTObjects run: {args.mtobjects_run}")
    doc.add_paragraph(f"MTObjects batch: {args.mtobjects_batch}")
    args.output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(args.output)
    print(args.output)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--sep-run", type=Path, required=True)
    parser.add_argument("--sep-batch", type=Path, required=True)
    parser.add_argument("--mtobjects-run", type=Path, required=True)
    parser.add_argument("--mtobjects-batch", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    return parser.parse_args()


if __name__ == "__main__":
    build(parse_args())
