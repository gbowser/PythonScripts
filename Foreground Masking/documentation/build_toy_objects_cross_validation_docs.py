#!/usr/bin/env python3
"""Build methodology documentation for Toy Objects four-fold optimisation."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUTPUT = ROOT / "Foreground Masking" / "documentation" / "Toy Objects Four Fold Cross Validation Documentation.docx"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GREY = "F2F4F7"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            mar = tc_pr.find(qn("w:tcMar"))
            if mar is None:
                mar = OxmlElement("w:tcMar")
                tc_pr.append(mar)
            for side, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
                node = mar.find(qn(f"w:{side}"))
                if node is None:
                    node = OxmlElement(f"w:{side}")
                    mar.append(node)
                node.set(qn("w:w"), str(value))
                node.set(qn("w:type"), "dxa")


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, value in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = value
        set_cell_shading(cell, LIGHT_BLUE)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            cells[idx].text = str(value)
            if row_index % 2:
                set_cell_shading(cells[idx], LIGHT_GREY)
    set_table_geometry(table, widths)
    doc.add_paragraph()
    return table


def configure_styles(doc: Document) -> None:
    doc.settings.odd_and_even_pages_header_footer = False
    section = doc.sections[0]
    section.different_first_page_header_footer = False
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    if "Equation" not in [style.name for style in doc.styles]:
        eq = doc.styles.add_style("Equation", WD_STYLE_TYPE.PARAGRAPH)
        eq.font.name = "Cambria Math"
        eq.font.size = Pt(10.5)
        eq.paragraph_format.left_indent = Inches(0.35)
        eq.paragraph_format.space_before = Pt(4)
        eq.paragraph_format.space_after = Pt(8)


def add_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = "Foreground Masking | Toy Objects cross-validation"
    header.style = doc.styles["Header"]
    header.runs[0].font.color.rgb = RGBColor.from_string("667085")
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("MSc Research - reproducible optimisation methodology")
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = RGBColor.from_string("667085")


def build() -> None:
    doc = Document()
    configure_styles(doc)
    add_header_footer(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)
    run = title.add_run("Toy Objects Foreground-Masking Optimisation")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Four-fold cross-validation for SEP and MTObjects").italic = True
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Dataset: 40 low-foreground S4G galaxies | Application set: 182 galaxies")
    doc.add_paragraph()

    doc.add_heading("1. Purpose and scope", level=1)
    doc.add_paragraph(
        "This document defines the reproducible optimisation design used to calibrate foreground-object masking "
        "with injected Toy Objects. SEP is optimised first; the complete procedure is then repeated for MTObjects. "
        "The design separates parameter fitting from held-out evaluation and retains a common independent injection "
        "set for comparing the four candidate solutions."
    )
    doc.add_paragraph("The workflow produces:")
    for text in (
        "four optimisation runs per algorithm, each trained on 30 galaxies;",
        "held-out evaluation on the remaining 10 galaxies, rotating once through every galaxy;",
        "cross-evaluation of all four candidate parameter sets on a common 40-galaxy Toy Object realization;",
        "a selected best-parameter JSON compatible with the canonical 182-galaxy batch tools; and",
        "CSV/workbook evidence supporting parameter and performance review.",
    ):
        doc.add_paragraph(text, style="List Bullet")

    doc.add_heading("2. Input sample validation", level=1)
    doc.add_paragraph(
        "CleanGalaxies.txt is treated as the authoritative list. Before optimisation, the driver requires exactly "
        "40 nonblank, unique names. Every name must occur in the 182-row geometry manifest, have all required bar and "
        "disc geometry fields, and resolve to an accessible 3.6-micron FITS image. Validation failure stops the run."
    )
    add_table(
        doc,
        ["Validation", "Acceptance criterion"],
        [
            ["Count", "Exactly 40 unique galaxy identifiers"],
            ["Manifest", "Every identifier matches one manifest row"],
            ["Geometry", "Centre, position angles, inclination, bar size and pixel scale are finite"],
            ["Image", "The local or manifest FITS path exists"],
        ],
        [2600, 6760],
    )

    doc.add_heading("3. Toy Object construction", level=1)
    doc.add_paragraph(
        "Six synthetic objects are injected into each science image inside the investigated, bar-aligned region. "
        "Objects are sampled as stars, compact clusters and elliptical galaxies. Their peak brightness is scaled to "
        "the robust image noise, making the task comparable across galaxies with different backgrounds."
    )
    add_table(
        doc,
        ["Property", "Sampling rule"],
        [
            ["Object type", "Star 50%; cluster 20%; galaxy 30%"],
            ["Peak", "Uniform from 5 to 25 robust-sigma"],
            ["Star/cluster FWHM", "Uniform from 2 to 10 pixels"],
            ["Galaxy FWHM", "Uniform from 5 to 22 pixels"],
            ["Galaxy axis ratio", "Uniform from 0.35 to 0.95"],
            ["Truth footprint", "Model pixels at least 8% of peak, then one-pixel circular dilation"],
        ],
        [2600, 6760],
    )

    doc.add_heading("4. Cross-validation design", level=1)
    doc.add_paragraph(
        "A fixed random seed shuffles the 40 names and distributes them into four disjoint folds of 10. For fold k, "
        "the other 30 galaxies form the training set and fold k is held out. Thus every galaxy is excluded from "
        "optimisation exactly once and used for held-out validation exactly once."
    )
    add_table(
        doc,
        ["Stage", "Images", "Role"],
        [
            ["Fold optimisation", "30", "Choose parameters using the Toy Objects objective"],
            ["Held-out validation", "10", "Measure generalisation without influencing Optuna"],
            ["Common evaluation", "40", "Compare all four candidates on identical independent injections"],
            ["Final application", "182", "Generate the canonical per-galaxy diagnostic outputs"],
        ],
        [2100, 1200, 6060],
    )
    doc.add_paragraph(
        "Each fold runs 40 Optuna trials: eight broad startup samples followed by 32 Tree-structured Parzen Estimator "
        "trials. Ten image workers evaluate galaxies concurrently within a trial; trials themselves remain sequential "
        "so the adaptive sampler sees every completed result."
    )

    doc.add_heading("5. Objective function", level=1)
    doc.add_paragraph(
        "The injected-image mask is compared with the baseline mask from the same unmodified galaxy. Only incremental "
        "mask pixels are scored. This isolates the response to the known Toy Objects from pre-existing detections."
    )
    doc.add_paragraph("incremental mask = injected-image mask AND NOT baseline mask", style="Equation")
    doc.add_paragraph(
        "Per-image measurements include truth-pixel recall, precision, F-score, mean per-toy recall, the proportion "
        "of toys with at least 50% recall, incremental masked fraction, false-positive fraction and segment count."
    )
    doc.add_heading("5.1 SEP score", level=2)
    doc.add_paragraph(
        "recovery = 0.45 mean recall + 0.20 mean F-score + 0.25 mean toy recall + 0.20 toy detection rate",
        style="Equation",
    )
    doc.add_paragraph(
        "score = recovery - 0.35 mean masked fraction - 0.05 false-positive fraction",
        style="Equation",
    )
    doc.add_heading("5.2 MTObjects score", level=2)
    doc.add_paragraph(
        "recovery = 0.45 mean F-score + 0.35 mean toy recall + 0.20 toy detection rate",
        style="Equation",
    )
    doc.add_paragraph(
        "score = recovery - 2.0 mean masked fraction - 0.5 false-positive fraction",
        style="Equation",
    )
    doc.add_paragraph(
        "For both algorithms, any trial whose worst individual image masks more than 15% receives a large cap-excess "
        "penalty. Optuna minimises objective = -score for feasible trials."
    )

    doc.add_heading("6. SEP parameter constraints", level=1)
    doc.add_paragraph(
        "SEP ranges follow the Source Extractor guide supplied with the project. The guide describes relative detection "
        "thresholds around 1.2 sigma, examples spanning approximately 0.6-2 sigma, minimum areas commonly between 5 "
        "and 35 pixels, 32 deblend levels as a standard starting point, deblend contrast of order 0.01, and background "
        "meshes larger than the objects of interest. Project-specific post-detection filters retain their established ranges."
    )
    add_table(
        doc,
        ["Parameter", "Optimised range/choices", "Purpose"],
        [
            ["detect_thresh", "0.6-2.0 sigma", "Relative detection threshold"],
            ["minarea", "5-35 pixels", "Minimum connected area"],
            ["deblend_nthresh", "16, 32 or 64", "Deblending levels"],
            ["deblend_cont", "0.001-0.03, log scale", "Minimum branch contrast"],
            ["back_size", "32, 48, 64, 96, 128, 192 or 256", "Background mesh size"],
            ["filter_size", "1, 3, 5, 7 or 9", "Background filter size"],
            ["dilation_radius", "1-6 pixels", "Final mask expansion"],
            ["max_area", "20-8000 pixels", "Reject overlarge segments"],
            ["max_elongation", "1.5-30", "Reject extreme segment shapes"],
        ],
        [2200, 3000, 4160],
    )

    doc.add_heading("7. MTObjects parameter constraints", level=1)
    doc.add_paragraph(
        "MTObjects uses the parameter domains established by the existing project optimiser and interactive model. "
        "Algorithm defaults not listed below remain fixed, including alpha, soft bias, gain and background mean."
    )
    add_table(
        doc,
        ["Parameter", "Optimised range", "Purpose"],
        [
            ["move_factor", "0.0-1.0", "Tree-node movement control"],
            ["min_distance", "0.0-1.0", "Minimum hierarchy distance"],
            ["gaussian_fwhm", "0.0-5.0 pixels", "Optional Gaussian smoothing"],
            ["bg_variance", "0.0001-10000, step 0.0001", "Background variance supplied to MTObjects"],
            ["minarea", "1-80 pixels", "Minimum accepted segment area"],
            ["dilation_radius", "0-8 pixels", "Final mask expansion"],
            ["max_area", "20-3000 pixels", "Reject overlarge segments"],
            ["max_elongation", "1.5-20", "Reject extreme segment shapes"],
        ],
        [2200, 3000, 4160],
    )

    doc.add_heading("8. Candidate selection", level=1)
    doc.add_paragraph(
        "The best training solution from each fold is first reported against its own held-out 10 galaxies. Because the "
        "four held-out sets contain different galaxies, those four scores are diagnostic rather than directly interchangeable. "
        "For a fair final comparison, every candidate is therefore scored on the same 40 galaxies using a separate, fixed "
        "Toy Object injection seed. The candidate with the smallest all-40 objective is selected."
    )
    doc.add_paragraph(
        "This common evaluation still uses the same underlying science galaxies, so it is an independent injection test "
        "rather than a completely external astronomical sample. That limitation should be retained when interpreting results."
    )

    doc.add_heading("9. Recorded outputs and review", level=1)
    add_table(
        doc,
        ["Output", "Contents"],
        [
            ["cross_validation_config.json", "Seeds, folds, constraints and run settings"],
            ["training_names.txt / held_out_names.txt", "Exact membership for each fold"],
            ["fold optimisation summary CSV", "All 40 trials and parameter combinations"],
            ["cross_validation_candidates.csv", "Training, held-out and common all-40 metrics for each candidate"],
            ["held_out_details.csv", "Per-galaxy held-out performance"],
            ["cross-validation best JSON", "Selected parameters and provenance for the 182-galaxy batch"],
            ["results workbook", "Review-oriented sheets for assignments, candidates, details and trial histories"],
        ],
        [3200, 6160],
    )
    doc.add_paragraph(
        "Review should focus on both performance and stability: the spread of parameters across folds, held-out score "
        "variation, false-positive masking, worst-image masked fraction, and whether Optuna repeatedly selects a bound. "
        "A high recovery score is not acceptable if it is achieved through excessive removal of galaxy data."
    )

    doc.add_heading("10. Final 182-galaxy processing", level=1)
    doc.add_paragraph(
        "After candidate selection, the canonical batch runner applies the winning parameters to all 182 manifest galaxies. "
        "Each algorithm writes per-galaxy diagnostic PNGs and a summary CSV with status, segment counts, masked fraction, "
        "runtime and errors. Once SEP and MTObjects batches both finish, the established comparison utility can assemble "
        "the requested six-panel, one-PNG-per-galaxy comparison set."
    )

    doc.add_heading("11. Reproducibility and restart behaviour", level=1)
    doc.add_paragraph(
        "All random seeds and fold assignments are saved. Optuna studies use SQLite storage, and the cross-validation "
        "drivers detect completed fold summaries before rerunning. A stopped workflow can therefore reuse completed folds. "
        "Terminal logs record progress, rough ETA and expected completion. Output folders are timestamped to prevent one "
        "study from overwriting another."
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
