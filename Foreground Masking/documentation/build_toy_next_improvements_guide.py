from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = Path(r"C:\Users\gordo\Documents\Github\PythonScripts\Foreground Masking\documentation\toy_comparison_doc_qa\Toy Objects SEP and MTObjects Next Improvements - Detailed Guide.docx")
NAVY='0B2545'; BLUE='2E74B5'; DARK='1F4D78'; GRAY='555555'; LIGHT='F2F4F7'; CALLOUT='E8EEF5'

def rgb(s): return RGBColor.from_string(s)
def shade(cell, fill):
    p=cell._tc.get_or_add_tcPr(); x=OxmlElement('w:shd'); x.set(qn('w:fill'),fill); p.append(x)
def keep(p): p._p.get_or_add_pPr().append(OxmlElement('w:keepNext'))
def setup(d):
    s=d.sections[0]; s.page_width=Inches(8.5); s.page_height=Inches(11)
    s.top_margin=s.bottom_margin=s.left_margin=s.right_margin=Inches(1)
    st=d.styles; n=st['Normal']; n.font.name='Calibri'; n.font.size=Pt(11); n.font.color.rgb=rgb(GRAY); n.paragraph_format.space_after=Pt(6); n.paragraph_format.line_spacing=1.1
    for name,size,col,b,a in [('Heading 1',16,BLUE,16,8),('Heading 2',13,BLUE,12,6),('Heading 3',12,DARK,8,4)]:
        q=st[name]; q.font.name='Calibri'; q.font.size=Pt(size); q.font.bold=True; q.font.color.rgb=rgb(col); q.paragraph_format.space_before=Pt(b); q.paragraph_format.space_after=Pt(a)
    for name in ['List Bullet','List Number']:
        q=st[name]; q.font.name='Calibri'; q.font.size=Pt(11); q.font.color.rgb=rgb(GRAY); q.paragraph_format.space_after=Pt(6); q.paragraph_format.line_spacing=1.1
def heading(d,t,l=1): p=d.add_heading(t,level=l); keep(p); return p
def bullet(d,items,number=False):
    num_id=None
    if number:
        root=d.part.numbering_part.element
        nums=root.findall(qn('w:num')); num_id=max([int(n.get(qn('w:numId'))) for n in nums]+[0])+1
        style_num=d.styles['List Number']._element.pPr.numPr.numId.val
        source=next(n for n in nums if int(n.get(qn('w:numId')))==style_num)
        abstract=source.find(qn('w:abstractNumId')).get(qn('w:val'))
        num=OxmlElement('w:num'); num.set(qn('w:numId'),str(num_id)); a=OxmlElement('w:abstractNumId'); a.set(qn('w:val'),abstract); num.append(a)
        over=OxmlElement('w:lvlOverride'); over.set(qn('w:ilvl'),'0'); start=OxmlElement('w:startOverride'); start.set(qn('w:val'),'1'); over.append(start); num.append(over); root.append(num)
    for x in items:
        p=d.add_paragraph(x,style='List Number' if number else 'List Bullet')
        if number:
            numPr=p._p.get_or_add_pPr().get_or_add_numPr(); numPr.get_or_add_ilvl().val=0; numPr.get_or_add_numId().val=num_id
def callout(d,label,text):
    p=d.add_paragraph(); pPr=p._p.get_or_add_pPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),CALLOUT); pPr.append(shd)
    p.paragraph_format.left_indent=Pt(6); p.paragraph_format.right_indent=Pt(6); p.paragraph_format.space_before=Pt(4); p.paragraph_format.space_after=Pt(10)
    r=p.add_run(label+': '); r.bold=True; r.font.color.rgb=rgb(NAVY); p.add_run(text)
def table(d,rows,widths):
    t=d.add_table(rows=len(rows),cols=len(rows[0])); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
    for i,row in enumerate(rows):
        trPr=t.rows[i]._tr.get_or_add_trPr(); cant=OxmlElement('w:cantSplit'); trPr.append(cant)
        for j,v in enumerate(row):
            c=t.cell(i,j); c.width=Inches(widths[j]); c.text=str(v); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in c.paragraphs:
                for r in p.runs: r.font.name='Calibri'; r.font.size=Pt(9.5); r.font.color.rgb=rgb(GRAY)
            if i==0:
                shade(c,LIGHT)
                keep(c.paragraphs[0])
                for r in c.paragraphs[0].runs: r.bold=True; r.font.color.rgb=rgb(NAVY)
    hdr=OxmlElement('w:tblHeader'); hdr.set(qn('w:val'),'1'); t.rows[0]._tr.get_or_add_trPr().append(hdr)
    d.add_paragraph(); return t
def labelpara(d,label,text):
    p=d.add_paragraph(); r=p.add_run(label+': '); r.bold=True; r.font.color.rgb=rgb(DARK); p.add_run(text)

d=Document(); setup(d)
p=d.add_paragraph(); r=p.add_run('FOREGROUND MASKING RESEARCH'); r.bold=True; r.font.size=Pt(10); r.font.color.rgb=rgb(BLUE)
p=d.add_paragraph(); p.paragraph_format.space_after=Pt(4); r=p.add_run('Toy Objects: Detailed Next Improvements'); r.bold=True; r.font.size=Pt(25); r.font.color.rgb=rgb(NAVY)
p=d.add_paragraph('Implementation guide for improving SEP and MTObjects optimisation, validation and production selection'); p.runs[0].font.size=Pt(13)
d.add_paragraph('Prepared: 24 August 2026\nStatus: Proposed research and engineering programme')
callout(d,'Purpose','This guide expands the five improvements proposed in the Toy Objects SEP-versus-MTObjects results report. It explains what each change would achieve, how to implement it, how to judge success and what evidence it would produce. It does not change the current results; it defines the next experimental programme.')

heading(d,'1. Current evidence and remaining limitations')
d.add_paragraph('The matched 182-galaxy evaluation found that SEP recovered more toy pixels and had higher F1 at essentially the same mean masked area, while MTObjects had a lower extreme masked-area tail. The comparison is useful but does not fully resolve algorithm selection because the original optimisation runs used different toy seeds, the scalar objectives only approximated scientific preservation, and synthetic truth labels only the injected objects.')
table(d,[('Observed issue','Why it matters','Improvement addressing it'),('Different optimisation seeds','Fold-level SEP and MTObjects scores are not strictly paired.','1. Identical injections'),('Galaxy structure can be masked','Area alone cannot distinguish a compact foreground mask from a coherent arm or ring.','2. Morphology protection'),('One scalar score hides trade-offs','Weight choices can select a solution that is not preferred scientifically.','3. Pareto optimisation'),('Synthetic truth is incomplete','Existing real foreground objects and galaxy features are unlabelled.','4. Real labelled subset'),('Neither method dominates every galaxy','SEP recovery and MTObjects conservatism are useful in different cases.','5. Guarded hybrid')],[1.65,2.55,2.3])
heading(d,'Recommended order')
bullet(d,['Implement identical injections first; this is the minimum requirement for a defensible re-optimisation.','Add morphology metrics before changing the optimiser, so the new objectives measure the science risk that matters.','Run Pareto optimisation after the paired evaluation and morphology metrics are stable.','Develop the real labelled subset in parallel, then use it as an external validation set rather than another tuning set.','Evaluate the guarded hybrid last, using thresholds learned from the preceding evidence.'],number=True)

heading(d,'2. Improvement 1 - Identical toy placements and seeds')
callout(d,'Aim','Make every SEP-versus-MTObjects optimisation comparison paired at the galaxy, fold, toy and pixel levels.')
heading(d,'Why this is necessary',2)
d.add_paragraph('The same 40 galaxies and fold assignments were used previously, but SEP and MTObjects used different injection seeds. A difference in score can therefore reflect both the masking algorithm and the particular toy positions, brightnesses or morphologies. Identical injection manifests remove this nuisance variation and increase statistical power because each method is tested against exactly the same challenge.')
heading(d,'Proposed implementation',2)
bullet(d,['Generate a single immutable injection manifest before optimisation. Store galaxy identifier, image checksum, global seed, per-galaxy seed, toy type, centre, amplitude, FWHM, axis ratio, position angle and truth-mask checksum.','Materialise or deterministically reconstruct the toy image and truth mask from the manifest. Both algorithms must read the same files or verify the same checksums.','Retain the existing four fixed folds of 10 galaxies. Within each held-out fold, compare algorithms galaxy by galaxy using identical truth.','Use identical evaluation code and denominators for both methods. Keep baseline and injected masks so incremental and final-mask metrics are both available.','Record software version, parameter set, worker count and runtime environment in every result row.'])
heading(d,'Statistical analysis',2)
labelpara(d,'Primary paired outcomes','Toy-pixel recall, per-toy recall, toy detection, F1 and total masked fraction for SEP minus MTObjects on each galaxy.')
labelpara(d,'Uncertainty','Report paired bootstrap 95% confidence intervals across galaxies. Also report fold-level estimates to reveal sensitivity to the particular calibration group.')
labelpara(d,'Significance','Use a paired non-parametric test such as Wilcoxon signed-rank when distributional assumptions are doubtful. Emphasise effect sizes and confidence intervals rather than a p-value alone.')
labelpara(d,'Stratification','Summarise performance by toy class, brightness, FWHM, radial position and local galaxy surface brightness. This identifies where one method fails rather than only whether its overall mean is lower.')
heading(d,'Acceptance criteria and outputs',2)
table(d,[('Deliverable','Acceptance check'),('Injection manifest','All 40 calibration galaxies and the 182-galaxy evaluation have reproducible toy/truth checksums.'),('Paired fold workbook','One row per galaxy, fold and method; no unmatched toys or missing metrics.'),('Comparison report','Paired effect sizes, confidence intervals, fold variability and stratified performance.'),('Reproducibility test','A clean rerun regenerates identical truth masks and summary metrics within numerical tolerance.')],[2.0,4.5])

heading(d,'3. Improvement 2 - Protect coherent galaxy morphology')
callout(d,'Aim','Penalise masks that remove spatially coherent galaxy structure, even when their total pixel area appears acceptable.')
heading(d,'Why masked area is insufficient',2)
d.add_paragraph('Two masks can cover 8% of an image but have very different scientific consequences. Compact isolated masks may remove foreground stars with limited impact, whereas a long connected component following a spiral arm, bar, ring or dust lane may alter isophotes and the bar-major profile. A structure-protection term should therefore assess where the mask lies and how it is organised.')
heading(d,'Candidate morphology-risk measures',2)
table(d,[('Measure','Definition or proxy','Risk indicated'),('Protected-zone overlap','Fraction of mask inside the bar, central exclusion, ring or user-defined science zone.','Direct loss in scientifically important regions.'),('Component scale','Largest component area and maximum component diameter relative to image/galaxy size.','One extended mask dominating the result.'),('Azimuthal span','Angular coverage of a component in deprojected polar coordinates.','Mask following an arm, ring or broad sector.'),('Annular coherence','Longest continuous angular run within radial annuli.','Ring-like or arm-like coherent removal.'),('Radial continuity','Number of adjacent radial bins occupied at similar angle.','Spoke, dust lane or arm tracing.'),('Low-frequency overlap','Mask overlap with a smoothed galaxy model or high-confidence galaxy segmentation.','Removal of diffuse galaxy light rather than compact sources.'),('Profile relevance','Fraction of bar-major samples masked and length of the longest bridged interval.','Potential distortion of the scientific profile.')],[1.25,2.65,2.6])
heading(d,'How to construct the penalty',2)
d.add_paragraph('Create a morphology-risk score scaled from 0 to 1. Each component-level risk is first normalised using the calibration distribution, then aggregated using either the maximum component risk or a high percentile. A conservative starting formulation is:')
callout(d,'Illustrative formula','risk = 0.25 protected-zone overlap + 0.20 annular coherence + 0.20 low-frequency overlap + 0.20 profile relevance + 0.15 largest-component scale. The weights are initial engineering values and must be tested, not treated as established scientific constants.')
d.add_paragraph('Use the risk as both a soft optimisation objective and a hard gate. For example, reject trials that mask more than a specified fraction of the protected bar region or create a component spanning an implausibly large azimuthal angle. Soft terms then rank the remaining feasible trials.')
heading(d,'Validation and failure controls',2)
bullet(d,['Build a small review set containing obvious arms, rings, bars and smooth early-type galaxies. A human reviewer labels whether the mask threatens coherent morphology.','Check that the risk score orders severe examples above acceptable compact masks. Report sensitivity and false-alarm rate for the review labels.','Do not calculate protection features from residual images for SEP or MTObjects detection. Detection remains on science images; auxiliary models are used only to evaluate risk.','Preserve the eight-panel PNG and add the morphology-risk value, largest-component fraction and protected-zone overlap to its title or summary metadata.'])

heading(d,'4. Improvement 3 - Multi-objective Pareto optimisation')
callout(d,'Aim','Expose the trade-off between recovery, data loss and morphology preservation instead of hiding it inside one weighted score.')
heading(d,'Why a single scalar can mislead',2)
d.add_paragraph('A scalar objective assumes that the relative value of an additional unit of recovery, mask area and morphology risk is known in advance. Small weight changes can select a different parameter set, and a zero-detection or over-masking solution can be attractive if the penalties are unbalanced. Multi-objective optimisation retains a set of non-dominated solutions so the scientific decision is made after the trade-off is visible.')
heading(d,'Proposed objectives and constraints',2)
table(d,[('Element','Direction','Suggested definition'),('Toy recovery','Maximise','Mean toy-pixel recall plus per-toy recall/detection, reported separately as diagnostics.'),('Mask burden','Minimise','Total or incremental masked fraction, with non-toy mask reported separately.'),('Morphology risk','Minimise','Composite risk from Improvement 2.'),('Overlap quality','Maximise','Toy-associated precision or F1, interpreted with synthetic-truth limitations.'),('Feasibility constraints','Pass/fail','Non-zero credible detection; minimum recovery; maximum mask; protected-zone and component limits.')],[1.5,1.0,4.0])
heading(d,'Optimisation process',2)
bullet(d,['Run a multi-objective sampler for each of the four folds, using the same documentation-constrained parameter ranges and the shared injection manifest.','Retain all non-dominated trials, not only one winner. Plot recovery versus mask and recovery versus morphology risk, coloured by held-out performance.','Apply pre-declared constraints before selecting a candidate. Do not rescue an infeasible high-recovery solution merely because it lies on the Pareto front.','Select a small set of operating points: recovery-priority, balanced and preservation-priority. Evaluate all three on held-out folds and on the external labelled subset.','Choose the production point using an explicit decision rule, such as the highest held-out recovery among candidates meeting mask and morphology limits.'])
heading(d,'Decision outputs',2)
table(d,[('Output','Purpose'),('Pareto plots by fold','Shows whether the apparent optimum is stable or fold-specific.'),('Candidate parameter cards','Documents parameters and expected recovery/mask/risk for each operating point.'),('Constraint audit','Explains why candidates were accepted or rejected.'),('Sensitivity analysis','Repeats selection under plausible thresholds to show whether the recommendation changes.')],[2.0,4.5])

heading(d,'5. Improvement 4 - Manually labelled real-foreground validation subset')
callout(d,'Aim','Measure performance on real foreground objects and real galaxy structure, addressing the central limitation of synthetic-only truth.')
heading(d,'Sample design',2)
d.add_paragraph('Select a stratified subset large enough to cover easy and difficult galaxies without making annotation unmanageable. A practical pilot is 30-50 galaxies sampled across angular size, inclination, morphology, stellar density, foreground burden and image quality. Keep these galaxies external to parameter tuning wherever possible.')
heading(d,'Annotation schema',2)
bullet(d,['Object-level catalogue: centre, approximate boundary, type (star, diffraction feature, compact galaxy, artefact or uncertain) and confidence.','Pixel-level or region-level foreground truth for objects where a defensible boundary can be drawn.','Protected galaxy-structure regions: bar, ring, spiral arm, dust lane, nucleus and other scientifically relevant features.','Ambiguous regions marked explicitly rather than forced into foreground or galaxy classes. Exclude or sensitivity-test them in scoring.','At least two independent annotators for a representative subset, followed by adjudication. Report agreement before consensus.'])
heading(d,'Metrics',2)
table(d,[('Metric family','Examples'),('Foreground recovery','Object detection rate, object-level recall by class, pixel recall and boundary overlap.'),('False masking','Masked galaxy pixels, false components per image and false-mask area within protected structures.'),('Scientific preservation','Change in isophote ellipticity/position angle, bar-major profile deviation and bridged-profile length.'),('Calibration','Performance versus annotator confidence and object brightness/size.'),('Reliability','Inter-annotator agreement and sensitivity to ambiguous-region handling.')],[1.75,4.75])
heading(d,'Governance and leakage prevention',2)
bullet(d,['Freeze the labelled validation set and do not repeatedly tune against it. If it informs a redesign, establish a second untouched confirmation set.','Version the annotations and record the image version/checksum.','Keep synthetic and real-truth metrics separate in reports; they answer related but different questions.','Publish clear inclusion/exclusion rules so uncertain real objects are not silently reclassified to improve scores.'])

heading(d,'6. Improvement 5 - Guarded SEP/MTObjects hybrid')
callout(d,'Aim','Use SEP where its stronger recovery is safe, but fall back to MTObjects or manual review when SEP shows high mask burden or morphology risk.')
heading(d,'Conceptual workflow',2)
bullet(d,['Run SEP using the selected recovery-oriented parameters on the science image.','Calculate quality-control features: masked fraction, morphology-risk score, largest-component fraction, protected-zone overlap, profile-bridge length and detection plausibility.','Accept SEP when every guardrail passes.','If SEP fails a soft guardrail, run MTObjects and compare both masks using the same quality-control features. Select the feasible result under a pre-declared rule.','If neither result is feasible, do not automatically choose the numerically lower score; route the galaxy to manual review and preserve both diagnostics.'],number=True)
heading(d,'Initial guardrails to test',2)
table(d,[('Guardrail','Initial candidate rule','Rationale'),('Total SEP mask','Review above 15%; mandatory review above 20%.','Directly addresses the observed SEP high-mask tail.'),('Largest component','Flag above a calibrated fraction of investigated area.','Finds one coherent over-mask hidden inside an acceptable total.'),('Protected-zone overlap','Reject above a science-defined threshold.','Protects the bar, nucleus and other target regions.'),('Morphology risk','Fallback when composite risk exceeds its validated threshold.','Combines spatial and profile evidence.'),('Recovery plausibility','Flag zero/near-zero response when credible targets exist.','Prevents the earlier non-detection loophole.'),('Method disagreement','Review when mask area or structure-risk difference is extreme.','Large disagreement is evidence of uncertainty.')],[1.35,2.55,2.6])
d.add_paragraph('These thresholds are starting hypotheses. They should be calibrated from paired cross-validation and the labelled real subset, then frozen before the final 182-galaxy assessment.')
heading(d,'Ways to combine outputs',2)
bullet(d,['Selection hybrid (recommended first): choose either the complete SEP mask or complete MTObjects mask. This is easiest to audit and preserves component logic.','Component hybrid: accept individual components from either method if they pass a classifier or rule set. This may improve performance but creates a more complex validation burden.','Consensus mask: use intersection for high precision or union for high recall. Both can be useful diagnostics, but neither should be the default without evidence because intersection can miss objects and union can over-mask.'])
heading(d,'Hybrid success criteria',2)
labelpara(d,'Primary','Improves toy and real-object recovery relative to MTObjects while reducing SEP high-mask and morphology-risk outliers.')
labelpara(d,'Safety','Manual-review rate is operationally manageable and severe morphology failures are rarely auto-accepted.')
labelpara(d,'Transparency','Every selection records which method was chosen, which guardrail fired and the alternative method metrics.')

heading(d,'7. Integrated experimental roadmap')
table(d,[('Phase','Work','Gate to proceed'),('1. Pairing foundation','Shared injection manifest, checksums, common evaluator and fold rerun.','No unmatched toys; reproducible reruns; paired workbook complete.'),('2. Morphology metric pilot','Implement component, annular, protected-zone and profile-risk features.','Risk score discriminates reviewer-labelled safe/severe masks.'),('3. Pareto optimisation','Four-fold multi-objective runs for SEP and MTObjects.','Stable feasible operating points across folds.'),('4. Real validation','Annotate pilot subset and assess both algorithms without tuning leakage.','Real-object and structure-preservation evidence supports or revises selection.'),('5. Hybrid evaluation','Freeze guardrails; compare SEP, MTObjects and selection hybrid.','Hybrid improves agreed primary metrics without unacceptable review burden.'),('6. Production rerun','Apply frozen pipeline to 182 galaxies and regenerate diagnostics/workbooks.','182 successes, no silent failures, all flags documented.')],[1.05,3.25,2.2])
heading(d,'Minimum reporting package')
bullet(d,['A per-galaxy XLSX/CSV containing method, parameters, injection/truth identifiers, recovery, mask burden, morphology risk, profile impact, guardrail result and runtime.','Fold-level and pooled paired statistics with confidence intervals and stratified toy results.','Pareto plots and candidate parameter cards.','A labelled-validation report separating foreground recovery from galaxy-structure preservation.','182 paired diagnostic PNGs and a review queue containing every guardrail failure.','A final decision memo identifying the frozen operating point and explaining any change from the current SEP recommendation.'])
heading(d,'8. Recommended immediate next action')
callout(d,'Recommendation','Begin with Improvement 1 and implement the shared injection manifest plus a common paired evaluator. In parallel, prototype three morphology-risk measures: protected-zone overlap, largest-component fraction and bar-profile bridge length. These changes provide the evidence foundation needed before another expensive optimisation run.')
heading(d,'9. Source document and related evidence')
d.add_paragraph('Source: D:\\Dropbox\\Public Documents\\UCLAN\\MSc Research\\Remove foreground objects\\documentation\\Toy Objects SEP versus MTObjects Results and Recommendations.docx\nMethodology: D:\\Dropbox\\Public Documents\\UCLAN\\MSc Research\\Remove foreground objects\\documentation\\Toy Objects SEP and MTObjects Methodology.docx\nMatched statistics: D:\\Dropbox\\Public Documents\\UCLAN\\MSc Research\\Remove foreground objects\\documentation\\Toy Objects SEP vs MTObjects Statistical Comparison.xlsx')

OUT.parent.mkdir(parents=True,exist_ok=True); d.save(OUT); print(OUT)
