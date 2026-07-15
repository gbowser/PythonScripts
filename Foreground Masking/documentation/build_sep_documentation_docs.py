#!/usr/bin/env python3
"""Build SEP methodology and program documentation DOCX files."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path(__file__).resolve().parent
METHOD_DOC = OUT_DIR / "SEP_Methodology_and_Parameters.docx"
PROGRAM_DOC = OUT_DIR / "SEP_Parameter_Tester_Program_Documentation.docx"

ACCENT = "2E74B5"
DARK = "1F4D78"
HEADER_FILL = "E8EEF5"
CALLOUT_FILL = "F4F6F9"
BORDER = "B8C6D9"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[float]) -> None:
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(int(w * 1440) for w in widths)))
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def style_document(doc: Document, title_text: str, subtitle: str) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, ACCENT, 18, 10),
        ("Heading 2", 13, ACCENT, 14, 7),
        ("Heading 3", 12, DARK, 10, 5),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.text = title_text
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = RGBColor(90, 90, 90)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("Page ")
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    footer._p.append(fld_begin)
    footer._p.append(instr)
    footer._p.append(fld_end)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run(title_text)
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("0B2545")

    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(14)
    sub_run = sub.add_run(subtitle)
    sub_run.font.size = Pt(11)
    sub_run.font.color.rgb = RGBColor(85, 85, 85)


def add_callout(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT_FILL)
    para = cell.paragraphs[0]
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run(text)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("0B2545")


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for cell, text in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, HEADER_FILL)
        para = cell.paragraphs[0]
        para.paragraph_format.space_after = Pt(0)
        run = para.add_run(text)
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string("0B2545")
    for row_values in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row_values):
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            para = cell.paragraphs[0]
            para.paragraph_format.space_after = Pt(0)
            para.add_run(text)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        para = doc.add_paragraph(style="List Bullet")
        para.paragraph_format.space_after = Pt(4)
        para.add_run(item)


def build_method_doc() -> None:
    doc = Document()
    style_document(
        doc,
        "Methodology of SEP and Parameter Guide",
        "Foreground object masking workflow used by interactive_sep_parameter_tester.py",
    )

    add_callout(
        doc,
        "Purpose: detect compact foreground sources in S4G galaxy images while protecting the galaxy centre, bar morphology, and bar-major-axis profile.",
    )

    doc.add_heading("1. What SEP Is", level=1)
    doc.add_paragraph(
        "SEP is a Python library that exposes the core Source Extractor algorithms as functions that operate directly on NumPy arrays. "
        "In this project it is used as a practical replacement for the classic external SExtractor executable on Windows."
    )
    doc.add_paragraph(
        "The original Source Extractor method detects sources above a locally estimated sky background, separates blended detections, "
        "and returns a segmentation map in which detected pixels are assigned to object labels."
    )

    doc.add_heading("2. Method Used in This Routine", level=1)
    add_bullets(
        doc,
        [
            "Load the selected FITS image from the local manifest and machine-specific S4G image folder.",
            "Build a smooth galaxy model with a broad Gaussian filter and subtract it to form a residual detection image.",
            "Estimate the sky/background on the selected detection image with sep.Background.",
            "Subtract the background map and detect connected objects with sep.extract.",
            "Deblend overlapping detections with the SEP/SExtractor deblending parameters.",
            "Reject detections that are too large, too elongated, or inside the central exclusion radius.",
            "Dilate the retained segmentation mask to include object wings and nearby affected pixels.",
            "Create a preview image by replacing masked pixels with a median unmasked value.",
            "Deproject and rotate the display so that the measured bar major axis lies along the x-axis.",
            "Draw original and processed diagnostics, including isophotes and bar-major-axis profiles.",
            "For the processed profile, hide masked samples and draw a dashed blue log-linear interpolation across those samples.",
        ],
    )

    doc.add_heading("3. Parameter Descriptions", level=1)
    add_table(
        doc,
        ["Parameter", "Default", "Role in the method", "Increasing the value usually..."],
        [
            [
                "Detect on",
                "residual",
                "Chooses whether SEP detects on the original image or the residual image after subtracting a smoothed galaxy model.",
                "Not a slider. Residual mode suppresses broad galaxy light and tends to favour compact sources.",
            ],
            [
                "Detection threshold",
                "3.0 sigma",
                "Signal-to-noise threshold passed to sep.extract. Pixels must exceed this level above background to seed detections.",
                "Removes fewer pixels, because only brighter/significant peaks survive. Lowering it removes more.",
            ],
            [
                "Minimum area",
                "5 px",
                "Minimum connected pixel count required for an object. This suppresses single-pixel noise.",
                "Removes fewer tiny objects. Lowering it allows smaller detections and removes more small pixels.",
            ],
            [
                "Deblend thresholds",
                "32",
                "Number of internal intensity levels used when splitting overlapping objects.",
                "Can split complex blends more finely, sometimes increasing the number of individual masks.",
            ],
            [
                "Deblend contrast",
                "0.005",
                "Minimum contrast ratio for keeping a deblended child object separate.",
                "Merges more objects and removes fewer separate subcomponents. Lowering it splits more aggressively.",
            ],
            [
                "Background box",
                "64 px",
                "SEP background mesh size, used for estimating local sky/background.",
                "Makes the background smoother and less reactive to local structure. Smaller boxes track local variation more closely.",
            ],
            [
                "Filter size",
                "3 px",
                "Gaussian-like convolution kernel size used before detection.",
                "Smooths detections more. This can improve faint extended object detection but can blur compact peaks.",
            ],
            [
                "Mask dilation",
                "2 px",
                "Radius added around retained detected segments after filtering.",
                "Removes more pixels around every accepted source.",
            ],
            [
                "Max segment area",
                "500 px",
                "Upper area limit for accepted detections. Larger detections are likely galaxy structure and are rejected.",
                "Allows larger structures to be removed. Lowering it protects the galaxy more strongly.",
            ],
            [
                "Max elongation",
                "6.0",
                "Rejects very elongated detections, which may be bar, dust, spiral, or galaxy structure rather than foreground objects.",
                "Allows more elongated features to be removed. Lowering it protects elongated galaxy features.",
            ],
            [
                "Central exclusion",
                "8 px",
                "Radius around the galaxy centre in which detections are rejected.",
                "Protects more of the central galaxy. Lowering it allows more central masking.",
            ],
        ],
        [1.35, 0.85, 2.25, 2.05],
    )

    doc.add_heading("4. Defaults and Rationale", level=1)
    doc.add_paragraph(
        "The core SEP defaults were selected from standard Source Extractor-style starting values, then adapted empirically for the ESO120-012 test case. "
        "The added limits and central exclusion were introduced to reduce the risk of masking true galaxy light."
    )
    add_table(
        doc,
        ["Default", "Why it was chosen"],
        [
            ["3.0 sigma threshold", "A conventional starting point that detects significant compact peaks without immediately eating into low-contrast galaxy structure."],
            ["5 pixel minimum area", "Large enough to suppress isolated noise but small enough to retain compact foreground stars or spike artefacts."],
            ["32 / 0.005 deblend", "Classic SExtractor-like deblending behaviour: reasonably aggressive splitting without disabling object separation."],
            ["64 pixel background box, 3 pixel filter", "Moderate background estimation and light smoothing for stable source detection."],
            ["2 pixel dilation", "A small safety margin around accepted detections so object wings are included."],
            ["500 pixel max area, elongation 6.0", "Guards against removing broad or strongly elongated galaxy features."],
            ["8 pixel central exclusion", "Protects the nuclear/bar region where SEP can otherwise mistake real galaxy structure for foreground contamination."],
        ],
        [2.1, 4.4],
    )

    doc.add_heading("5. Profile Handling", level=1)
    doc.add_paragraph(
        "The masked preview image replaces masked pixels with the median value of unmasked pixels. This is only a visual and isophote diagnostic. "
        "The processed bar-major-axis profile does not plot the median-replacement troughs. Instead, samples crossed by the SEP mask are hidden, "
        "and a dashed blue log-linear bridge is drawn across those masked samples using the nearest valid positive profile values."
    )

    doc.add_heading("6. References", level=1)
    add_bullets(
        doc,
        [
            "SEP documentation: https://sep.readthedocs.io/",
            "SEP extract API: https://sep.readthedocs.io/en/stable/api/sep.extract.html",
            "Bertin, E. and Arnouts, S. 1996, SExtractor: Software for source extraction, A&AS, 117, 393-404.",
            "Astromatic SExtractor project page: https://www.astromatic.net/software/sextractor/",
        ],
    )

    doc.save(METHOD_DOC)


def build_program_doc() -> None:
    doc = Document()
    style_document(
        doc,
        "SEP Parameter Tester Program Documentation",
        "Inputs, outputs, workflow, controls, and implementation notes for interactive_sep_parameter_tester.py",
    )

    add_callout(
        doc,
        "Program: Foreground Masking/interactive_sep_parameter_tester.py. The program is an interactive Windows/Tk tool for tuning SEP foreground masking on S4G galaxy FITS images.",
    )

    doc.add_heading("1. How to Run", level=1)
    doc.add_paragraph("From the PythonScripts repository root:")
    code = doc.add_paragraph()
    code.paragraph_format.left_indent = Inches(0.25)
    run = code.add_run('python "Foreground Masking/interactive_sep_parameter_tester.py"')
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    doc.add_paragraph("Optional arguments:")
    add_bullets(
        doc,
        [
            "--manifest PATH: override the default S4G/Erwin manifest used by the GalClean display helper.",
            "--pc Desktop|Laptop|...: choose the machine path set from machine_paths.py. The default is Desktop.",
        ],
    )

    doc.add_heading("2. Inputs", level=1)
    add_table(
        doc,
        ["Input", "Source", "Use"],
        [
            ["Manifest", "DEFAULT_MANIFEST from interactive_galclean_parameter_tester.py unless overridden", "Provides galaxy names, image filenames, and geometry values."],
            ["FITS image", "image_path_for_pc(row, pc_name)", "Primary 2D science image loaded with astropy.io.fits."],
            ["Galaxy geometry", "required_geometry(row)", "Centre, pixel scale, position angle, inclination/bar data used for centring, deprojection, and bar alignment."],
            ["Machine paths", "machine_paths.py", "Maps Desktop/Laptop style choices to local image and output folders."],
            ["User parameters", "Tk controls", "SEP detection, segmentation filtering, dilation, central exclusion, and unit conversion."],
        ],
        [1.35, 2.45, 2.7],
    )

    doc.add_heading("3. User Interface Controls", level=1)
    add_table(
        doc,
        ["Control", "Description"],
        [
            ["Machine", "Selects the configured computer path set. Default is Desktop."],
            ["Galaxy", "Lists galaxies from the manifest that have image files available for the selected machine."],
            ["Detect on", "Chooses residual or original detection image. Residual is the default."],
            ["Parameter units", "Switches relevant size/area controls between pixels and arcseconds. Internal processing converts values back to pixels."],
            ["Calculate", "Runs SEP and updates all panels. No auto-calculation occurs after parameter changes."],
            ["Reset", "Restores the default parameter values."],
            ["Open PNG Folder", "Opens the output directory in Windows File Explorer."],
        ],
        [1.75, 4.75],
    )

    doc.add_heading("4. Processing Pipeline", level=1)
    add_bullets(
        doc,
        [
            "Load the selected FITS image and geometry.",
            "Prepare the detection image: original image or residual image. Residual mode subtracts a broad Gaussian-smoothed model.",
            "Estimate the background using sep.Background with the selected background box size.",
            "Subtract the SEP background map.",
            "Generate a detection filter kernel from the selected filter size.",
            "Run sep.extract with threshold, minimum area, filter kernel, deblending, cleaning, and segmentation output enabled.",
            "Measure area, centroid, elongation, residual peak significance, and distance from the galaxy centre for each detected object.",
            "Filter the segmentation by maximum area, maximum elongation, and central exclusion radius.",
            "Dilate the accepted mask.",
            "Create a masked preview image by replacing masked pixels with the median of finite unmasked pixels.",
            "Build deprojected, bar-aligned cutouts for original, preview, residual, and mask panels.",
            "Draw central exclusion circles on image/isophote panels and vertical exclusion lines on profile panels.",
            "Save the output figure automatically as a PNG after each calculation.",
        ],
    )

    doc.add_heading("5. Outputs", level=1)
    add_table(
        doc,
        ["Output", "Location / naming", "Contents"],
        [
            [
                "PNG diagnostic figure",
                "remove_foreground_folder(pc) / interactive_sep_parameter_tester",
                "Automatic PNG saved after every Calculate. Filename includes galaxy, SEP threshold, area, deblend contrast, dilation, and timestamp.",
            ],
            [
                "On-screen status",
                "Left control panel",
                "Reports retained segment count, masked pixel fraction, output directory, and saved PNG filename.",
            ],
            [
                "No processed FITS",
                "Not written",
                "The routine intentionally saves only PNG diagnostics; it does not write a cleaned FITS image.",
            ],
        ],
        [1.45, 2.2, 2.85],
    )

    doc.add_heading("6. Figure Panels", level=1)
    add_table(
        doc,
        ["Panel", "Meaning"],
        [
            ["Parameter box", "Summary of units, detection mode, key SEP parameters, background values, segment counts, and masked fraction."],
            ["Centered original", "Original galaxy cutout after centring, deprojection, and bar alignment."],
            ["SEP masked preview", "Preview image with masked pixels replaced by the median finite unmasked value."],
            ["Residual detection image", "Detection residual used when Detect on = residual."],
            ["Mask", "Original image with accepted SEP mask overlaid in red."],
            ["Original isophotes", "Log-scaled original image with contours."],
            ["SEP processed isophotes", "Log-scaled masked preview image with contours."],
            ["Original bar-major profile", "Solid blue intensity profile along the deprojected bar major axis."],
            ["SEP processed bar-major profile", "Solid blue profile with masked samples hidden and dashed blue log-linear interpolation over those samples."],
        ],
        [2.0, 4.5],
    )

    doc.add_heading("7. Current Defaults", level=1)
    add_table(
        doc,
        ["Setting", "Default"],
        [
            ["Machine", "Desktop"],
            ["Galaxy", "ESO120-012, when available"],
            ["Detect on", "residual"],
            ["Detection threshold", "3.0"],
            ["Minimum area", "5 px"],
            ["Deblend thresholds", "32"],
            ["Deblend contrast", "0.005"],
            ["Background box", "64 px"],
            ["Filter size", "3 px"],
            ["Mask dilation", "2 px"],
            ["Max segment area", "500 px"],
            ["Max elongation", "6.0"],
            ["Central exclusion", "8 px"],
            ["Profile width", "3 px, fixed constant"],
        ],
        [2.2, 4.3],
    )

    doc.add_heading("8. Implementation Notes", level=1)
    add_bullets(
        doc,
        [
            "SEP operates on the original image array or residual detection array before display deprojection. Deprojection is used for diagnostics and profiles.",
            "The displayed galaxy is rotated into bar-aligned coordinates through the shared GalClean display helper functions.",
            "The unit switch affects size and area parameters where physically meaningful; internally, SEP receives pixel units.",
            "The central exclusion is both a filtering rule and a displayed guide.",
            "The dashed profile interpolation is a plotting aid only; it does not alter the FITS data and no FITS output is saved.",
            "Existing panels are greyed out during calculation by a full-figure overlay labelled Calculating.",
        ],
    )

    doc.add_heading("9. Dependencies", level=1)
    add_table(
        doc,
        ["Dependency", "Purpose"],
        [
            ["sep", "Source detection, background estimation, segmentation, and deblending."],
            ["numpy", "Array handling and numerical operations."],
            ["scipy.ndimage", "Gaussian smoothing and binary mask dilation."],
            ["astropy.io.fits", "FITS image loading."],
            ["matplotlib", "Diagnostic figure rendering."],
            ["tkinter", "Interactive user interface."],
            ["interactive_galclean_parameter_tester.py", "Shared manifest, geometry, deprojection, profile, and display helper routines."],
            ["machine_paths.py", "Machine-specific input and output path mapping."],
        ],
        [2.1, 4.4],
    )

    doc.save(PROGRAM_DOC)


def main() -> None:
    build_method_doc()
    build_program_doc()
    print(METHOD_DOC)
    print(PROGRAM_DOC)


if __name__ == "__main__":
    main()
