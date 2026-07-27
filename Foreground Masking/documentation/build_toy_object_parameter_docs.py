from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(__file__).with_name("Toy Object Parameters Guide.docx")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths: list[float]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for row in table.rows:
        row_pr = row._tr.get_or_add_trPr()
        if row_pr.find(qn("w:cantSplit")) is None:
            row_pr.append(OxmlElement("w:cantSplit"))
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def style_table(table, widths: list[float]) -> None:
    set_table_widths(table, widths)
    header_pr = table.rows[0]._tr.get_or_add_trPr()
    if header_pr.find(qn("w:tblHeader")) is None:
        header_pr.append(OxmlElement("w:tblHeader"))
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
                    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
                    run.font.size = Pt(9.5)
            if row_index == 0:
                set_cell_shading(cell, "E8EEF5")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(31, 77, 120)


def add_note(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    style_table(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(title + ": ")
    run.bold = True
    run.font.color.rgb = RGBColor(31, 58, 95)
    paragraph.add_run(body)
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def add_cover(doc: Document) -> None:
    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(6)
    run = title.add_run("Toy Object Parameters Guide")
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    run = subtitle.add_run(
        "Plain-language reference for the Interactive Object Recovery tool"
    )
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(85, 85, 85)

    meta = doc.add_paragraph()
    meta.add_run("Source: ").bold = True
    meta.add_run("Foreground Masking/Interactive tools/interactive object recovery.py")
    meta.paragraph_format.space_after = Pt(14)

    add_note(
        doc,
        "Short answer",
        "The Toy object model now has three available profiles: Gaussian star, Star cluster, and Compact galaxy. Location, brightness, FWHM, and truth dilation apply broadly. Axis ratio and object PA apply only to Compact galaxy. Toy objects must sit wholly inside the same deprojected, bar-aligned galaxy investigation area used by the normal image/profile reports.",
    )


def add_object_type_section(doc: Document) -> None:
    doc.add_heading("Object Types", level=1)
    rows = [
        [
            "Gaussian star",
            "Single circular 2D Gaussian profile.",
            "A compact point-source test object with a smooth, symmetric core.",
            "Characterised by centroid, brightness, and FWHM. Uses FWHM to set Gaussian sigma. Ignores axis ratio and object PA.",
        ],
        [
            "Star cluster",
            "Blend of three offset circular Gaussian components.",
            "A small unresolved or barely resolved clump rather than a single isolated star.",
            "Characterised by one centroid, brightness, and FWHM. FWHM sets the scale for all component offsets and widths. Ignores axis ratio and object PA.",
        ],
        [
            "Compact galaxy",
            "Single elliptical Gaussian profile.",
            "A small galaxy-like contaminant with an elongated shape.",
            "Characterised by centroid, brightness, major-axis FWHM, axis ratio, and object PA.",
        ],
    ]
    table = doc.add_table(rows=1, cols=4)
    headers = ["Type", "Profile used", "Best interpreted as", "Important notes"]
    for i, text in enumerate(headers):
        table.cell(0, i).text = text
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = text
    style_table(table, [1.25, 1.75, 1.7, 1.8])


def add_parameter_glossary(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("Parameter Glossary", level=1)
    entries = [
        (
            "Type",
            "Chooses the mathematical profile used to generate the injected Toy object.",
            "All types",
            "The internal values are gaussian, cluster, and galaxy.",
        ),
        (
            "Brightness",
            "Chooses how the object brightness is set: either by peak residual sigma or by integrated magnitude.",
            "All types",
            "Peak residual sigma is usually the simpler test mode. Integrated magnitude is useful when the FITS photometric calibration is valid.",
        ),
        (
            "x deproj [arcsec], y deproj [arcsec]",
            "Position of the Toy object in the deprojected, bar-aligned coordinate system.",
            "All types",
            "The code converts this deprojected position back to observed image pixels before injecting the model, then rejects placements outside the investigated bar-aligned cutout.",
        ),
        (
            "peak [resid sigma]",
            "Peak brightness expressed as a multiple of the robust residual-image noise.",
            "All types, when Brightness is Peak residual sigma",
            "Example: 30 means the model peak is 30 times the robust residual sigma.",
        ),
        (
            "integrated mag",
            "Target total magnitude for the whole injected object, not just the central pixel.",
            "All types, when Brightness is Integrated magnitude",
            "The code scales a unit-peak model so its integrated flux matches this magnitude.",
        ),
        (
            "zero flux [Jy]",
            "Flux density of a zero-magnitude source, in Jansky.",
            "Magnitude mode",
            "The default 280.9 Jy is appropriate for Spitzer/IRAC 3.6 micron Vega magnitudes. Change it for another band or magnitude system.",
        ),
        (
            "FWHM [arcsec]",
            "Full width at half maximum: the diameter across the profile where the brightness has fallen to half the peak.",
            "All types",
            "Converted to pixels using the galaxy pixel scale. For Gaussian profiles, sigma = FWHM / 2.3548.",
        ),
        (
            "axis ratio",
            "Minor-axis width divided by major-axis width. A value of 1.0 is circular; smaller values are more elongated.",
            "Compact galaxy only",
            "In the current code, Gaussian star and Star cluster are circular and ignore this value.",
        ),
        (
            "object PA [deg]",
            "Position angle of the object's major axis, measured in image-pixel coordinates.",
            "Compact galaxy only",
            "Only matters when the model is elliptical. It is not used for circular profiles.",
        ),
        (
            "truth dilation [px]",
            "Extra pixel dilation applied to the truth mask after thresholding the injected model.",
            "All types",
            "This changes the evaluation mask used for recall/overlap, not the injected object's light profile.",
        ),
    ]
    table = doc.add_table(rows=1, cols=4)
    for i, text in enumerate(["Parameter", "Meaning", "Applies to", "Practical note"]):
        table.cell(0, i).text = text
    for entry in entries:
        cells = table.add_row().cells
        for i, text in enumerate(entry):
            cells[i].text = text
    style_table(table, [1.25, 2.1, 1.35, 1.8])


def add_applicability_matrix(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("Applicability Matrix", level=1)
    headers = ["Parameter", "Gaussian star", "Star cluster", "Compact galaxy"]
    rows = [
        ["Type", "Yes", "Yes", "Yes"],
        ["Brightness mode", "Yes", "Yes", "Yes"],
        ["x/y deprojected location", "Yes", "Yes", "Yes"],
        ["peak residual sigma", "Yes", "Yes", "Yes"],
        ["integrated mag", "Yes", "Yes", "Yes"],
        ["zero flux [Jy]", "Magnitude mode", "Magnitude mode", "Magnitude mode"],
        ["FWHM [arcsec]", "Yes", "Yes", "Yes"],
        ["axis ratio", "No", "No", "Yes"],
        ["object PA [deg]", "No", "No", "Yes"],
        ["truth dilation [px]", "Yes", "Yes", "Yes"],
    ]
    table = doc.add_table(rows=1, cols=4)
    for i, text in enumerate(headers):
        table.cell(0, i).text = text
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = text
    style_table(table, [1.8, 1.55, 1.55, 1.6])


def add_math_section(doc: Document) -> None:
    doc.add_heading("Key Terms and Formulas", level=1)
    doc.add_heading("Gaussian sigma and FWHM", level=2)
    doc.add_paragraph(
        "A Gaussian profile is often described by sigma, but the UI asks for FWHM because it is easier to interpret visually. The code uses:"
    )
    add_note(doc, "Formula", "sigma_pixels = max(0.2, FWHM_pixels / 2.3548)")
    doc.add_paragraph(
        "A larger FWHM makes the object broader. For the Star cluster profile, FWHM also sets how far apart the three Gaussian components are placed."
    )

    doc.add_heading("How the three object types are characterised", level=2)
    doc.add_paragraph(
        "Each Toy object type is deliberately simple. The aim is to test recovery behaviour with controlled source shapes, not to fit a complete astrophysical model."
    )
    add_bullets(
        doc,
        [
            "Gaussian star: a single circular Gaussian. It is described by its deprojected position, brightness, and FWHM. It has no elongation or orientation.",
            "Star cluster: a fixed three-Gaussian blend. It is described by one central position, one brightness scale, and one FWHM scale; the component offsets and relative intensities are fixed by the tool.",
            "Compact galaxy: a single elliptical Gaussian. It is described by position, brightness, major-axis FWHM, axis ratio, and position angle.",
        ],
    )

    doc.add_heading("Investigation-area constraint", level=2)
    doc.add_paragraph(
        "Toy objects are only valid inside the galaxy area actually being investigated. In practice this is the same deprojected, bar-aligned square cutout used by the interactive displays, PNG reports, isophote views, and bar-major profiles. The horizontal axis is the deprojected bar-major axis and the vertical axis is the deprojected bar-minor axis."
    )
    add_bullets(
        doc,
        [
            "The toy centre must fall inside this cutout.",
            "The full truth mask, after the 8 percent model threshold and optional truth dilation, must also remain inside the cutout.",
            "Toy-object optimisation results produced before this constraint was added should be treated as superseded, because those older runs could place toys anywhere in the finite FITS footprint.",
        ],
    )

    doc.add_heading("Integrated magnitude mode", level=2)
    doc.add_paragraph(
        "Magnitude mode first builds the chosen object profile with peak = 1, then scales that unit model so the total integrated flux matches the requested magnitude."
    )
    add_note(doc, "Formula", "flux_Jy = zero_flux_Jy * 10^(-0.4 * magnitude)")
    doc.add_paragraph(
        "If the FITS header uses BUNIT=MJy/sr, the code converts summed model intensity to Jy using the pixel area in steradians. If not, it looks for a count-based magnitude zero point in MAGZP, MAGZERO, or ZEROPOINT."
    )

    doc.add_heading("Truth mask and recovery metrics", level=2)
    doc.add_paragraph(
        "The truth mask is not the same as the light profile. The code thresholds the model at 8 percent of the peak or 8 percent of the maximum model value, then optionally dilates that binary mask by truth dilation [px]."
    )
    add_bullets(
        doc,
        [
            "Recall is recovered truth pixels divided by total truth pixels.",
            "Incremental recall uses only the new mask pixels produced after injecting the Toy object.",
            "Incremental precision asks what fraction of the new mask pixels overlap the truth mask.",
        ],
    )


def add_recommendations(doc: Document) -> None:
    doc.add_heading("Suggested Usage", level=1)
    add_bullets(
        doc,
        [
            "Use Gaussian star for a simple compact-source sensitivity test.",
            "Use Star cluster when you want a blended compact contaminant rather than a single source.",
            "Use Compact galaxy when elongation and orientation matter; this is the only Toy object type that uses axis ratio and object PA.",
            "Use Peak residual sigma for quick recovery experiments. Use Integrated magnitude when you want physically calibrated brightness and the FITS header supports the conversion.",
            "Place Toy objects inside the displayed bar-aligned investigation area; if the model or truth mask spills outside that area, the tool rejects the placement.",
            "Remember that truth dilation changes the scoring mask, not the injected object itself.",
        ],
    )


def add_footer(doc: Document) -> None:
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Toy Object Parameters Guide")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(85, 85, 85)


def build() -> None:
    doc = Document()
    configure_styles(doc)
    add_cover(doc)
    add_object_type_section(doc)
    add_parameter_glossary(doc)
    add_applicability_matrix(doc)
    add_math_section(doc)
    add_recommendations(doc)
    add_footer(doc)
    doc.save(OUTPUT)


if __name__ == "__main__":
    build()
