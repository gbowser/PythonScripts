from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = Path(__file__).resolve().parent / "clean_run_comparison_20260829"
OUT.mkdir(parents=True, exist_ok=True)
DOCX = OUT / "SEP and MTObjects Clean-Galaxy Run Comparison.docx"

BLUE = "2E74B5"; DARK = "1F4D78"; LIGHT = "F2F4F7"; PALE = "E8EEF5"; GREEN = "E8F3EC"; INK = "202124"; MUTED = "5F6368"

runs = {
    "40": {"galaxies":40,"folds":"4 folds (30 train / 10 held out)","toys":6,
      "SEP":{"score":.41154365198273507,"recall":.47486228171833333,"precision":.008143134399084815,"fscore":.01594130560677543,"toy_recall":.4673445495278615,"detect":.46111111111111114,"masked":.03601161391807024,"max_masked":.048238198986664794,"false":.035738612894661195,"fold":4},
      "MTO":{"score":.2807911749313763,"recall":.5030422847425915,"precision":.004228368570699111,"fscore":.008367259956986424,"toy_recall":.5691498172082876,"detect":.6111111111111112,"masked":.07404188070668331,"max_masked":.12134570013630323,"false":.07377809941048838,"fold":3}},
    "11": {"galaxies":11,"folds":"11 leave-one-out folds", "toys":6,
      "SEP":{"score":.32975726875277217,"recall":.35853087576708675,"precision":.015734955623194984,"fscore":.02962067306372288,"toy_recall":.3704030643254469,"detect":.38333333333333336,"masked":.016953781707267485,"max_masked":.04048986598006206,"false":.01678738211292399,"fold":7},
      "MTO":{"score":.3250232639435931,"recall":.5980182809834297,"precision":.0026265273051955295,"fscore":.00522154773929441,"toy_recall":.7109711579441724,"detect":.7333333333333333,"masked":.12143037867300106,"max_masked":.14533500652050885,"false":.12117815149715865,"fold":9}},
    "20": {"galaxies":20,"folds":"20 leave-one-out folds", "toys":10,
      "SEP":{"score":.4752837571857291,"recall":.547983011921253,"precision":.01945416031011079,"fscore":.03735614159786344,"toy_recall":.5121976466479381,"detect":.5210526315789473,"masked":.027661519100899545,"max_masked":.060808305524657026,"false":.027164655817330415,"fold":9},
      "MTO":{"score":.2803877416927382,"recall":.5608438698407184,"precision":.004922738360647114,"fscore":.00974619250715541,"toy_recall":.6132436685943193,"detect":.6263157894736842,"masked":.10656485135888202,"max_masked":.13958405165658402,"false":.10614061158789263,"fold":18}},
}

def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = tcPr.find(qn("w:shd"))
    if shd is None: shd = OxmlElement("w:shd"); tcPr.append(shd)
    shd.set(qn("w:fill"), fill)

def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr(); tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None: tcMar = OxmlElement("w:tcMar"); tcPr.append(tcMar)
    for m,v in (("top",top),("start",start),("bottom",bottom),("end",end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None: node = OxmlElement(f"w:{m}"); tcMar.append(node)
        node.set(qn("w:w"), str(v)); node.set(qn("w:type"), "dxa")

def set_table_geometry(table, widths):
    table.autofit = False; table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tblPr = table._tbl.tblPr
    tblW = tblPr.find(qn("w:tblW")); tblW.set(qn("w:w"), str(sum(widths))); tblW.set(qn("w:type"), "dxa")
    ind = OxmlElement("w:tblInd"); ind.set(qn("w:w"), "120"); ind.set(qn("w:type"), "dxa"); tblPr.append(ind)
    grid = table._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for w in widths:
        col = OxmlElement("w:gridCol"); col.set(qn("w:w"), str(w)); grid.append(col)
    for row in table.rows:
        for i,cell in enumerate(row.cells):
            tcW = cell._tc.get_or_add_tcPr().find(qn("w:tcW")); tcW.set(qn("w:w"), str(widths[i])); tcW.set(qn("w:type"), "dxa")
            set_cell_margins(cell); cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def font_run(run, size=11, bold=False, color=INK, italic=False):
    run.font.name="Calibri"; run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"),"Calibri"); run._element.rPr.rFonts.set(qn("w:hAnsi"),"Calibri")
    run.font.size=Pt(size); run.bold=bold; run.italic=italic; run.font.color.rgb=RGBColor.from_string(color)

def add_p(text="", style=None, size=11, bold=False, color=INK, after=6, align=None, italic=False):
    p=doc.add_paragraph(style=style); p.paragraph_format.space_after=Pt(after); p.paragraph_format.line_spacing=1.1
    if align is not None: p.alignment=align
    if text: font_run(p.add_run(text),size,bold,color,italic)
    return p

def add_heading(text, level=1):
    p=doc.add_paragraph(style=f"Heading {level}"); p.add_run(text); return p

def add_table(headers, rows, widths, highlight_col=None):
    t=doc.add_table(rows=1, cols=len(headers)); t.style="Table Grid"
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=""; set_cell_shading(c, PALE); p=c.paragraphs[0]; font_run(p.add_run(h),10,True,DARK)
    for row in rows:
        cells=t.add_row().cells
        for i,val in enumerate(row):
            cells[i].text=""; p=cells[i].paragraphs[0]; font_run(p.add_run(str(val)),10, i==0, INK)
            if highlight_col is not None and i==highlight_col: set_cell_shading(cells[i], GREEN)
    set_table_geometry(t,widths); return t

def pct(v): return f"{v*100:.1f}%"
def pp(a,b): return f"{(a-b)*100:+.1f} pp"

doc=Document(); sec=doc.sections[0]
sec.top_margin=Inches(0.8); sec.bottom_margin=Inches(0.8); sec.left_margin=Inches(1); sec.right_margin=Inches(1); sec.header_distance=Inches(.492); sec.footer_distance=Inches(.492)
styles=doc.styles
normal=styles["Normal"]; normal.font.name="Calibri"; normal.font.size=Pt(11); normal.font.color.rgb=RGBColor.from_string(INK); normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.1
for name,size,before,after,color in (("Heading 1",16,16,8,BLUE),("Heading 2",13,12,6,BLUE),("Heading 3",12,8,4,DARK)):
    s=styles[name]; s.font.name="Calibri"; s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=RGBColor.from_string(color); s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after); s.paragraph_format.keep_with_next=True

header=sec.header.paragraphs[0]; header.alignment=WD_ALIGN_PARAGRAPH.LEFT; font_run(header.add_run("FOREGROUND MASKING | CLEAN-GALAXY COMPARISON"),8,True,MUTED)
footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.RIGHT
fld=OxmlElement("w:fldSimple"); fld.set(qn("w:instr"),"PAGE"); footer._p.append(fld)

add_p("TECHNICAL RESULTS BRIEF",size=9,bold=True,color=BLUE,after=8)
add_p("SEP and MTObjects Clean-Galaxy Run Comparison",size=24,bold=True,color=INK,after=4)
add_p("Comparison of the 40-, 11- and final 20-galaxy toy-object optimisation experiments",size=13,color=MUTED,after=16)
add_p("Prepared: 29 August 2026 | Metric version: paired-toy-metrics-v1",size=9,color=MUTED,after=14)

add_heading("Executive finding",1)
add_p("The final 20-galaxy run produced the strongest SEP result of the three experiments. MTObjects remained the more sensitive method, but its additional toy recovery came with a substantially larger masking footprint. On the common final-20 winner-selection set, MTObjects recovered 10.1 percentage points more toy signal than SEP, while masking 7.9 percentage points more of each image on average.")

add_heading("Experimental comparability",1)
add_table(["Run","Calibration galaxies","Validation design","Toys per image"],[["40 clean","40","4 folds: 30 train / 10 held out","6"],["11 clean","11","Leave-one-galaxy-out (11 folds)","6"],["20 clean","20","Leave-one-galaxy-out (20 folds)","10"]],[1450,1650,4200,2060])
add_p("Important qualification: all three winner files use the same paired-toy metric version and the same 6-30 sigma brightness interval, but the final-20 experiment used ten toys per image rather than six. The statistics are therefore suitable for comparing operating behaviour and direction of improvement, but they are not a perfectly controlled estimate of sample-size effect alone.",size=9,color=MUTED,italic=True,after=4)

doc.add_page_break(); add_heading("SEP results",1)
add_p("Higher composite score, toy recall and detection rate are favourable. Lower masked fraction and false-positive fraction are favourable.")
sep_rows=[]
metrics=[("Composite score","score",lambda v:f"{v:.3f}"),("Mean toy recall","toy_recall",pct),("Toy detection rate","detect",pct),("Mean image masked","masked",pct),("Maximum image masked","max_masked",pct),("False-positive fraction","false",pct),("Pixel recall","recall",pct),("Pixel precision","precision",pct),("Pixel F-score","fscore",pct)]
for label,key,fmt in metrics: sep_rows.append([label,fmt(runs["40"]["SEP"][key]),fmt(runs["11"]["SEP"][key]),fmt(runs["20"]["SEP"][key])])
add_table(["SEP metric","40 clean","11 clean","20 clean"],sep_rows,[3630,1910,1910,1910],highlight_col=3)
add_heading("SEP interpretation",2)
add_p(f"Against the 40-galaxy run, the final-20 SEP winner improved mean toy recall by {pp(runs['20']['SEP']['toy_recall'],runs['40']['SEP']['toy_recall'])} and detection rate by {pp(runs['20']['SEP']['detect'],runs['40']['SEP']['detect'])}, while reducing mean masked area by {abs((runs['20']['SEP']['masked']-runs['40']['SEP']['masked'])*100):.1f} percentage points. Its composite score rose from {runs['40']['SEP']['score']:.3f} to {runs['20']['SEP']['score']:.3f}.")
add_p(f"The 11-galaxy SEP run was the most conservative ({pct(runs['11']['SEP']['masked'])} mean masked area) but also the least sensitive ({pct(runs['11']['SEP']['toy_recall'])} mean toy recall). The final-20 parameters provide the best balance observed for SEP.")

doc.add_page_break(); add_heading("MTObjects results",1)
mto_rows=[]
for label,key,fmt in metrics: mto_rows.append([label,fmt(runs["40"]["MTO"][key]),fmt(runs["11"]["MTO"][key]),fmt(runs["20"]["MTO"][key])])
add_table(["MTObjects metric","40 clean","11 clean","20 clean"],mto_rows,[3630,1910,1910,1910],highlight_col=3)
add_heading("MTObjects interpretation",2)
add_p(f"The 11-galaxy MTObjects winner achieved the highest recovery ({pct(runs['11']['MTO']['toy_recall'])} mean toy recall; {pct(runs['11']['MTO']['detect'])} detection), but it also masked the most image area ({pct(runs['11']['MTO']['masked'])} on average, reaching {pct(runs['11']['MTO']['max_masked'])} in the worst case). This is an aggressive operating point close to the 15% cap.")
add_p(f"The final-20 MTObjects run is more restrained than the 11-galaxy run: mean masking fell by {abs((runs['20']['MTO']['masked']-runs['11']['MTO']['masked'])*100):.1f} percentage points, at a cost of {abs((runs['20']['MTO']['toy_recall']-runs['11']['MTO']['toy_recall'])*100):.1f} percentage points of toy recall. Relative to the 40-galaxy run, final-20 toy recall increased by {pp(runs['20']['MTO']['toy_recall'],runs['40']['MTO']['toy_recall'])}, but mean masking increased by {pp(runs['20']['MTO']['masked'],runs['40']['MTO']['masked'])}.")

doc.add_page_break(); add_heading("Direct comparison on the final-20 experiment",1)
direct=[]
for label,key,fmt in metrics[:6]:
    sv=runs['20']['SEP'][key]; mv=runs['20']['MTO'][key]
    if key in ('score',): diff=f"{mv-sv:+.3f}"
    else: diff=pp(mv,sv)
    direct.append([label,fmt(sv),fmt(mv),diff])
add_table(["Metric","SEP","MTObjects","MTO minus SEP"],direct,[3500,1900,2050,1910])
add_heading("Conclusions",1)
for text in [
    "SEP final-20 is the strongest all-round result in the study: it has the highest composite score (0.475), improved recovery over both earlier SEP runs, and retains a low 2.8% mean masking footprint.",
    "MTObjects final-20 is the stronger detector: it exceeds SEP by 10.1 percentage points in mean toy recall and 10.5 percentage points in toy detection rate.",
    "That MTObjects sensitivity is not free: its mean masked fraction is 10.7%, compared with 2.8% for SEP, and its maximum reaches 14.0%, close to the 15% safeguard.",
    "The final choice should therefore follow the science objective. Use SEP where preservation of galaxy structure is the priority; use MTObjects where missing a contaminant is more damaging than masking additional galaxy area.",
    "For the most defensible numerical claim about clean-sample size alone, rerun the 40- and 11-galaxy winners with the final ten-toy manifest design or score all three parameter sets on one common external evaluation set."
]:
    p=doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after=Pt(6); font_run(p.add_run(text),11)

doc.add_page_break(); add_heading("Winning parameter summary",1)
add_heading("SEP",2)
sep_param_rows=[["40 clean","1.836","8","64","0.001317","2","7271"],["11 clean","1.881","35","64","0.001246","3","7283"],["20 clean","1.898","19","64","0.001265","3","6422"]]
add_table(["Run","Threshold","Min area","Deblend levels","Deblend contrast","Dilation","Max area"],sep_param_rows,[1100,1200,1100,1450,1780,1200,1530])
add_heading("MTObjects",2)
mto_param_rows=[["40 clean","0.669","0.835","0.848","0.001870","3","1068"],["11 clean","0.863","0.722","0.456","0.001413","5","1755"],["20 clean","0.168","0.997","0.406","0.001336","4","2380"]]
add_table(["Run","Move factor","Min distance","Gaussian FWHM","BG variance","Dilation","Max area"],mto_param_rows,[1100,1350,1400,1650,1450,1200,1210])
add_heading("Metric definitions",1)
definitions=[("Mean toy recall","Fraction of injected toy pixels recovered by the mask."),("Toy detection rate","Fraction of individual toys meeting the detection criterion."),("Mean masked fraction","Mean fraction of the investigated image removed; includes correct and incorrect masking."),("False-positive fraction","Fraction masked outside the paired toy truth mask."),("Composite score","The optimisation objective reported by the winner file; higher is better and balances recovery against collateral masking and safeguards.")]
add_table(["Metric","Interpretation"],definitions,[2600,6760])
add_heading("Data provenance",1)
add_p("Statistics were read directly from the six top-level cross-validation winner JSON files produced by the three experiments. No values were inferred from diagnostic PNGs. Source locations are recorded below for reproducibility.",size=9,color=MUTED)
for text in [
    "40 clean: SEP/Toy Objects/20260824_115154/optimisation and MTObjects/Toy Objects/20260824_115154/optimisation",
    "11 clean: Toy Objects paired optimisation/clean11_logo_20260826_122730/{sep_logo, mtobjects_logo}",
    "20 clean: final20_toy_optimisation/{SEP_cross_validation, MTObjects_cross_validation}"
]:
    p=doc.add_paragraph(style="List Bullet"); font_run(p.add_run(text),9,color=MUTED)

doc.core_properties.title="SEP and MTObjects Clean-Galaxy Run Comparison"
doc.core_properties.subject="Comparison of toy-object optimisation statistics for 40, 11 and 20 clean-galaxy calibration samples"
doc.core_properties.author="Gordon research project"
doc.save(DOCX)
print(DOCX)
