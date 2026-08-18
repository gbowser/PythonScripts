#!/usr/bin/env python3
"""Apply the 17 August 2026 folder-wide documentation corrections."""

from pathlib import Path
from shutil import copy2

from docx import Document


SOURCE = Path(r"D:\Dropbox\Public Documents\UCLAN\MSc Research\Remove foreground objects\documentation")
OUTPUT = Path(r"C:\Users\gordo\Documents\Github\PythonScripts\Foreground Masking\Documentation\systematic_review_20260817")
OUTPUT.mkdir(parents=True, exist_ok=True)

DELETE_NAMES = {
    "SEP and MTObjects Masking Comparison and Recommendations 20260816.docx",
    "SEP and MTObjects Toy Objects Optimisation Results.docx",
    "SEP and MTObjects Toy Objects Optimisation Summary and Conclusions.docx",
}

CURRENT_MEMO = "SEP and MTObjects Masking Comparison and Recommendations 20260817.docx"
GENERIC_REPLACEMENTS = {
    "Systematic documentation review - 16 August 2026": "Systematic documentation review - 17 August 2026",
    "Authoritative combined outcome: SEP and MTObjects Toy Objects Optimisation Summary and Conclusions.docx. Full-precision results: Foreground Masking Optimisation Results.xlsx.":
        f"Authoritative combined outcome: {CURRENT_MEMO}. Full optimisation history and final science-image SEP fold records: Foreground Masking Optimisation Results.xlsx.",
}

TARGETED_REPLACEMENTS = {
    "SEP and MTObjects Optuna Optimisation Process - Updated.docx": {
        'python "Foreground Masking\\sep_toy_object_parameter_optimisation.py" --max-images 20 --toys-per-image 6 --initial-points 8 --max-iter 32 --detect-on residual':
            'python "Foreground Masking\\Optimisation\\cross_validate_toy_objects_SEP.py" --detect-on original --trials-per-fold 40',
        "Current interpretation: SEP remains the primary result. MTObjects follow-up optimisation is valid but should be used as a review-controlled comparison because recovery is accompanied by substantial collateral masking.":
            "Current interpretation: the final corrected comparison uses SEP on the original science image only. MTObjects is the recommended primary production method because it achieved higher held-out toy recovery and a safer upper masking tail; SEP remains an independent gated diagnostic.",
        "Current four-fold design: 40 validated galaxies; four rotations of 30 training and 10 held out; 40 trials per fold (8 startup + 32 TPE); candidates compared on the same common-40 injection set. Selected winners: SEP fold 4 and MTObjects fold 3.":
            "Current four-fold design: 40 validated galaxies; four rotations of 30 training and 10 held out; 40 trials per fold (8 startup + 32 TPE); candidates compared on the same common-40 injection set. Selected winners: science-image SEP fold 4 (run 20260817_161404) and MTObjects fold 3 (run 20260816_063455).",
    },
    "MTObjects Toy Objects Four Fold Optimisation Documentation.docx": {
        "Current interpretation: MTObjects is no longer a zero-mask solution. It is a feasible secondary method, but mean precision of 0.31% and false-positive fraction of 9.05% prevent recommendation as the primary production mask.":
            "Current interpretation: MTObjects is no longer a zero-mask solution and is the recommended primary production method. It achieved higher four-fold held-out toy recall and detection than the corrected science-image SEP run, and its 182-galaxy masking tail was safer. Pixel precision remains a diagnostic rather than a catalogue-purity estimate because masks on real foreground sources outside injected toy truth are counted as false-positive area.",
    },
    "interactive_toy_objects_MTObjects.docx": {
        "Current interpretation: The recovery follow-up restores non-zero detections, but the selected MTObjects candidate has low pixel precision and substantial collateral masking. Visual review remains mandatory.":
            "Current interpretation: the recovery follow-up restores non-zero detections and MTObjects is the recommended primary production method after comparison with corrected science-image SEP. Visual review and the 15% masked-area/profile-damage gates remain mandatory.",
        "Selected recovery-follow-up parameters are documented in the combined outcome report. They achieve non-zero recovery but remain review-controlled because common-40 pixel precision is 0.31% and false-positive fraction is 9.05%.":
            "Selected recovery-follow-up parameters are documented in the current Decision Memo. They achieved 50.7% common-40 mean toy recall and 53.3% toy detection. The low pixel precision is retained as a diagnostic caveat because real-source masks outside injected toy truth are counted as false-positive area.",
    },
    "Toy Objects Four Fold Cross Validation Documentation.docx": {
        "SEP ranges follow the Source Extractor guide supplied with the project. The guide describes relative detection thresholds around 1.2 sigma, examples spanning approximately 0.6-2 sigma, minimum areas commonly between 5 and 35 pixels, 32 deblend levels as a standard starting point, deblend contrast of order 0.01, and background meshes larger than the objects of interest. Project-specific post-detection filters retain their established ranges.":
            "SEP Toy Objects detection is constrained to the original science image; residual-image detection is not permitted in the current optimiser or production batch. Parameter ranges follow the Source Extractor guide supplied with the project: relative thresholds approximately 0.6-2 sigma, minimum areas commonly 5-35 pixels, 32 deblend levels as a standard starting point, deblend contrast of order 0.01, and background meshes larger than the objects of interest. Project-specific post-detection filters retain their established ranges.",
        "After candidate selection, the canonical batch runner applies the winning parameters to all 182 manifest galaxies. Each algorithm writes per-galaxy diagnostic PNGs and a summary CSV with status, segment counts, masked fraction, runtime and errors. Each method produces one aligned eight-panel PNG per galaxy: top row original and original plus toys; second row mask and recovered image with green correct and red incorrect outlines; third row original and processed isophotes; fourth row original and processed bar-major profiles. The profile x-axis limits and physical panel widths match the image panels above. Calibration-galaxy filenames receive the `_clean` suffix automatically.":
            "After candidate selection, the canonical batch runner applies the winning parameters to all 182 manifest galaxies. Each algorithm writes per-galaxy diagnostic PNGs and a summary CSV with status, segment counts, masked fraction, runtime and errors. Each method produces one aligned eight-panel PNG per galaxy: top row original and original plus toys; second row mask and recovered image with green correct and red incorrect outlines; third row original and processed isophotes; fourth row original and processed bar-major profiles. The processed profile retains the log-linear bridge: dashed green segments mark masked toy objects and dotted red segments mark all other masked items. The profile x-axis limits and physical panel widths match the image panels above. Calibration-galaxy filenames receive the `_clean` suffix automatically.",
    },
    "Foreground Masking Canonical Tools Reference.docx": {
        "Aligned Toy Objects batch layout: original / original plus toys; mask / recovered image with green correct and red incorrect outlines; original / processed isophotes; original / processed bar-major profile.":
            "Aligned Toy Objects batch layout: original / original plus toys; mask / recovered image with green correct and red incorrect outlines; original / processed isophotes; original / processed bar-major profile. The processed profile uses the established log-linear bridge, dashed green for masked toy objects and dotted red for all other masked items.",
    },
}


def replace_paragraph(paragraph, replacements):
    original = paragraph.text
    revised = original
    for old, new in replacements.items():
        revised = revised.replace(old, new)
    if revised != original:
        paragraph.text = revised
        return 1
    return 0


def update_document(path: Path) -> tuple[Path, int]:
    destination = OUTPUT / path.name
    copy2(path, destination)
    document = Document(destination)
    replacements = dict(GENERIC_REPLACEMENTS)
    replacements.update(TARGETED_REPLACEMENTS.get(path.name, {}))
    changes = 0
    for paragraph in document.paragraphs:
        changes += replace_paragraph(paragraph, replacements)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    changes += replace_paragraph(paragraph, replacements)
    if changes:
        document.save(destination)
    return destination, changes


def main():
    for path in sorted(SOURCE.glob("*.docx")):
        if path.name in DELETE_NAMES or path.name == CURRENT_MEMO:
            continue
        destination, changes = update_document(path)
        print(f"{destination.name}\t{changes}")


if __name__ == "__main__":
    main()
