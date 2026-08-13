#!/usr/bin/env python3
"""Build documentation for the separate low-foreground Toy Objects optimisation."""

from __future__ import annotations

import csv
import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


DROPBOX_ROOT = Path(r"D:\Dropbox\Public Documents\UCLAN\MSc Research\Remove foreground objects\Azure Optimisation outputs")
RUN_ROOT = DROPBOX_ROOT / "separate-low-foreground-toy-optimisations"
OUTPUT = DROPBOX_ROOT / "Low-Foreground Toy Objects Optimisation Documentation.docx"
SEP_RUN = RUN_ROOT / "sep-toy" / "20260812_184825"
MTO_RUN = RUN_ROOT / "mtobjects-toy-visible" / "20260812_200732"
GALAXIES = ["IC1954", "IC4901", "NGC0289", "NGC0578", "NGC0986", "NGC1097", "NGC3359", "NGC3992", "NGC4133"]

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
PALE_BLUE = "E8EEF5"
PALE_GRAY = "F2F4F7"
PALE_GOLD = "FFF2CC"
RED = "9C0006"
WHITE = "FFFFFF"
INK = "202B33"


def set_font(run, size=None, bold=None, italic=None, color=None):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = str(text)
        shade(cell, PALE_BLUE)
        for run in cell.paragraphs[0].runs:
            set_font(run, 9.5, bold=True, color=DARK_BLUE)
    set_repeat_table_header(table.rows[0])
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
            for run in cells[idx].paragraphs[0].runs:
                set_font(run, 9.2)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_callout(doc, label, text, fill=PALE_GOLD, color=INK):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    shade(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"{label}: ")
    set_font(r, 10.5, bold=True, color=color)
    r = p.add_run(text)
    set_font(r, 10.5, color=color)
    set_table_geometry(table, [9360])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    set_font(r, 11)


def load_run(root, prefix):
    best = json.loads((root / f"{prefix}_best.json").read_text())
    config = json.loads((root / f"{prefix}_config.json").read_text())
    with (root / f"{prefix}_summary.csv").open(encoding="utf-8", newline="") as handle:
        trial_count = sum(1 for _ in csv.reader(handle)) - 1
    return best, config, trial_count


def fmt(value):
    if isinstance(value, float):
        if value == 0:
            return "0"
        if abs(value) < 0.001:
            return f"{value:.6g}"
        return f"{value:.6f}".rstrip("0").rstrip(".")
    return str(value)


def build():
    sep_best, sep_config, sep_trials = load_run(SEP_RUN, "sep_toy_object_optimisation")
    mto_best, mto_config, mto_trials = load_run(MTO_RUN, "mtobjects_parameter_optimisation")
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = section.bottom_margin = Inches(1)
    section.left_margin = section.right_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"; normal.font.size = Pt(11); normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in (("Heading 1",16,BLUE,18,10),("Heading 2",13,BLUE,14,7),("Heading 3",12,DARK_BLUE,10,5)):
        style = styles[name]
        style.font.name = "Calibri"; style.font.size = Pt(size); style.font.bold = True; style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before); style.paragraph_format.space_after = Pt(after); style.paragraph_format.keep_with_next = True
    for list_name in ("List Bullet", "List Number"):
        style = styles[list_name]
        style.font.name = "Calibri"; style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375); style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4); style.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.text = "Foreground Masking | Optimisation Technical Note"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in header.runs: set_font(run, 9, color="6B7280")
    footer = section.footer.paragraphs[0]
    footer.text = "Separate low-foreground Toy Objects optimisation | August 2026"
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in footer.runs: set_font(run, 8.5, color="6B7280")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18); p.paragraph_format.space_after = Pt(2)
    r = p.add_run("TECHNICAL NOTE"); set_font(r, 10, bold=True, color=BLUE)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("Low-Foreground Toy Objects Optimisation"); set_font(r, 26, bold=True, color=DARK_BLUE)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("Separate nine-galaxy calibration for SEP and MTObjects, followed by full 182-galaxy application"); set_font(r, 13, color="4B6475")
    add_table(doc, ["Field", "Value"], [
        ["Run date", "12 August 2026"],
        ["Optimisers", "SEP Toy Objects; MTObjects Toy Objects"],
        ["Calibration sample", "Nine deliberately selected galaxies with few foreground objects"],
        ["Trial budget", "40 trials per optimiser (8 initial + 32 further)"],
        ["Follow-up", "Two new 182-galaxy batches and 182 hybrid comparison panels"],
        ["Primary workbook", "optimiser_parameter_stability_4sets.xlsx (two additional LowFG sheets)"],
    ], [2700, 6660])
    add_callout(doc, "Scope distinction", "This is a separate experiment. It is not a fifth stability seed and must not be pooled with the four random 20-galaxy stability samples.")

    doc.add_heading("1. Purpose and rationale", level=1)
    doc.add_paragraph("The original stability exercise repeated each optimiser across four reproducibly selected 20-galaxy samples. This additional run asks a different question: how do the Toy Objects optimisers behave when calibrated only on galaxies chosen to contain relatively few real foreground objects? Restricting the calibration set reduces competition between injected truth objects and naturally occurring compact sources, but it also makes the result less representative of the full survey population.")
    add_bullet(doc, "Optimise SEP Toy Objects on the same nine low-foreground galaxies.")
    add_bullet(doc, "Optimise MTObjects Toy Objects independently on that identical sample.")
    add_bullet(doc, "Apply each best parameter set to all 182 usable galaxies.")
    add_bullet(doc, "Build new comparison panels using the new Toy Objects batches and the existing four-seed-mean Spike Gate batches.")

    doc.add_heading("2. Separate calibration galaxy set", level=1)
    doc.add_paragraph("The requested identifier NGC3392 was not present in the 182-galaxy manifest. A one-digit comparison identified NGC3992 as the intended available object, and that correction was explicitly confirmed before the run.")
    add_table(doc, ["#", "Galaxy", "Selection note"], [[i+1, name, "Low foreground-object calibration set"] for i, name in enumerate(GALAXIES)], [720, 1800, 6840])
    add_callout(doc, "Identifier correction", "NGC3992 was used. NGC3392 was not used and is absent from the copied 182-galaxy Azure manifest.", fill="FCE4D6", color=RED)

    doc.add_heading("3. Experimental configuration", level=1)
    add_table(doc, ["Setting", "SEP Toy Objects", "MTObjects Toy Objects"], [
        ["Galaxy count", "9", "9"],
        ["Injected toys", "6 per image (54 total)", "6 per image (54 total)"],
        ["Seed", sep_config.get("seed"), mto_config.get("seed")],
        ["Initial trials", sep_config.get("initial_points"), mto_config.get("initial_points")],
        ["Further trials", sep_config.get("max_iter"), mto_config.get("max_iter")],
        ["Completed trials", sep_trials, mto_trials],
        ["Workers", sep_config.get("workers"), mto_config.get("workers")],
        ["Detection image", sep_config.get("detect_on"), mto_config.get("detect_on")],
        ["Study name", sep_config.get("study_name"), mto_config.get("study_name")],
    ], [2500, 3430, 3430])
    doc.add_paragraph("Optuna trials remained sequential. Within each trial, the nine images were evaluated using nine Windows spawn workers. The same seed and trial budget were used for both algorithms to make the runs operationally comparable.")

    doc.add_heading("4. Best optimisation results", level=1)
    add_table(doc, ["Metric", "SEP Toy Objects", "MTObjects Toy Objects"], [
        ["Best objective", fmt(sep_best.get("objective")), fmt(mto_best.get("objective"))],
        ["Best score", fmt(sep_best.get("score")), fmt(mto_best.get("score"))],
        ["Trials evaluated", sep_trials, mto_trials],
    ], [3000, 3180, 3180])

    doc.add_heading("4.1 SEP best parameters", level=2)
    sep_rows = [[key, fmt(value)] for key, value in sep_best["params"].items()]
    add_table(doc, ["Parameter", "Best value"], sep_rows, [4680, 4680])

    doc.add_heading("4.2 MTObjects best parameters", level=2)
    mto_rows = [[key, fmt(value)] for key, value in mto_best["params"].items()]
    add_table(doc, ["Parameter", "Best value"], mto_rows, [4680, 4680])
    add_callout(doc, "Interpretation caution", "The MTObjects best score and objective are both zero on this restricted sample. This is a completed optimisation result, but it does not demonstrate successful injected-object recovery. Treat its full-batch output as a diagnostic comparison, not as validated superiority.", fill="FCE4D6", color=RED)

    doc.add_heading("5. Full 182-galaxy follow-up batches", level=1)
    doc.add_paragraph("Each restricted-set best JSON was then applied unchanged to the complete 182-galaxy local manifest. These are application batches, not additional optimisation trials.")
    add_table(doc, ["Batch", "Calibration source", "Successful", "Failed", "Output"], [
        ["SEP Toy Objects", "Separate LowFG 9", "182", "0", "full-batches/sep-toy"],
        ["MTObjects Toy Objects", "Separate LowFG 9", "182", "0", "full-batches/mtobjects-toy"],
    ], [1800, 1800, 1100, 900, 3760])
    add_callout(doc, "Outcome", "Both follow-up batches produced reports for all 182 galaxies with no recorded processing failures.", fill="E2F0D9", color="375623")

    doc.add_heading("6. Hybrid composite panels", level=1)
    doc.add_paragraph("A new set of 182 six-panel comparison PNGs was generated. The panels intentionally mix the existing four-seed-mean Spike Gate results with the new separately calibrated Toy Objects results:")
    add_table(doc, ["Composite position", "Source batch", "Parameter provenance"], [
        ["Original image", "Shared input image", "Not optimised"],
        ["Original profile", "Shared baseline", "Not optimised"],
        ["SEP Spike Gate", "mean-parameter-4sets/full-batches/sep-spike", "Mean across four stability seeds"],
        ["SEP Toy Objects", "separate-low-foreground.../full-batches/sep-toy", "Best from separate LowFG 9 set"],
        ["MTObjects Spike Gate", "mean-parameter-4sets/full-batches/mtobjects-spike", "Mean across four stability seeds"],
        ["MTObjects Toy Objects", "separate-low-foreground.../full-batches/mtobjects-toy", "Best from separate LowFG 9 set"],
    ], [1800, 3900, 3660])
    doc.add_paragraph("The output folder contains 182 RGB PNG files at 1926 x 2740 pixels. Because the parameter provenance differs between Spike Gate and Toy Objects panels, this set should be described as a hybrid comparison rather than an all-method common-calibration experiment.")

    doc.add_heading("7. Workbook integration", level=1)
    doc.add_paragraph("The existing optimiser_parameter_stability_4sets.xlsx workbook was extended with two sheets. The original four stability sheets were retained unchanged.")
    add_table(doc, ["New worksheet", "Contents"], [
        ["SEP Toy - LowFG 9", "Separate-set warning, exact galaxy list, run provenance, best parameters, and all 40 SEP trials"],
        ["MTO Toy - LowFG 9", "Separate-set warning, exact galaxy list, run provenance, best parameters, and all 40 MTObjects trials"],
    ], [2500, 6860])
    doc.add_paragraph("A backup named optimiser_parameter_stability_4sets_before_low_foreground_addition.xlsx was created before the workbook was modified.")

    doc.add_heading("8. Output structure and reproducibility", level=1)
    add_table(doc, ["Artifact", "Dropbox location relative to Azure Optimisation outputs"], [
        ["Updated workbook", "optimiser_parameter_stability_4sets.xlsx"],
        ["Pre-update workbook backup", "optimiser_parameter_stability_4sets_before_low_foreground_addition.xlsx"],
        ["SEP optimisation", "separate-low-foreground-toy-optimisations/sep-toy/20260812_184825"],
        ["MTObjects optimisation", "separate-low-foreground-toy-optimisations/mtobjects-toy-visible/20260812_200732"],
        ["Two full batches", "separate-low-foreground-toy-optimisations/full-batches"],
        ["Hybrid composites", "separate-low-foreground-toy-optimisations/composite-panels"],
        ["Mean Spike Gate sources", "mean-parameter-4sets/full-batches"],
    ], [2800, 6560])
    add_bullet(doc, "All generated scientific outputs are stored in Dropbox; reusable launch and monitoring scripts remain in the repository.")
    add_bullet(doc, "The local manifest resolves all 182 copied FITS images and records the full Windows image paths.")
    add_bullet(doc, "Best JSON files, configuration JSON files, toy catalogues, trial summaries, detailed results, and Optuna SQLite studies are retained with each run.")

    doc.add_heading("9. Limitations and recommended use", level=1)
    add_bullet(doc, "The nine galaxies were deliberately selected rather than randomly sampled; statistical generalisation to all S4G galaxies is therefore limited.")
    add_bullet(doc, "A low-foreground calibration may favour sensitivity to injected objects while underrepresenting confusion from real compact foreground sources.")
    add_bullet(doc, "The new Toy Objects results should be compared with the four-seed stability results, not averaged into them.")
    add_bullet(doc, "The zero MTObjects recovery score warrants direct visual review of injected-object cases and selected full-batch panels before adopting the parameters as defaults.")
    add_bullet(doc, "Composite panels show processing consequences, but do not by themselves establish scientific validity or quantify galaxy-light loss.")

    doc.add_heading("10. Reproduction commands", level=1)
    p = doc.add_paragraph()
    p.add_run("Both optimisers used the same explicit galaxy list, ").bold = False
    r = p.add_run("--max-images 9 --initial-points 8 --max-iter 32 --workers 9 --seed 202608045"); set_font(r, 10, bold=True, color=DARK_BLUE)
    doc.add_paragraph("SEP entry point: Foreground Masking/Optimisation/optimise_toy_objects_SEP.py")
    doc.add_paragraph("MTObjects entry point: Foreground Masking/Optimisation/optimise_toy_objects_MTObjects.py, with --mtobjects-root pointing to the compiled local mtobjects checkout.")
    doc.add_paragraph("Batch entry points: batch_sep_all_galaxies.py and apply_optimised_mtobjects_all_galaxies.py. Composite entry point: make_all_method_galaxy_comparison_pngs.py.")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
