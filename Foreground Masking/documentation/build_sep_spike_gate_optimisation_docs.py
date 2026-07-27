#!/usr/bin/env python3
"""Build DOCX documentation for the SEP Spike Gate Optuna optimiser."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


DOC_DIR = Path(__file__).resolve().parent
OUT_PATH = DOC_DIR / "SEP Spike Gate Optuna Optimisation Documentation.docx"

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
HEADER_FILL = "E8EEF5"
CALLOUT_FILL = "F4F6F9"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_grid = table._tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        tbl_grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def mark_first_row_as_header(table) -> None:
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        tr_pr.append(OxmlElement("w:tblHeader"))


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run("SEP Spike Gate Optimisation")
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(100, 116, 139)


def add_title(doc: Document) -> None:
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("SEP Spike Gate Optuna Optimisation")
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("0B2545")
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    sub = subtitle.add_run(
        "Technical reference for using Spike Gate profile evidence to optimise global SEP foreground masking."
    )
    sub.font.size = Pt(11)
    sub.font.color.rgb = RGBColor(71, 85, 105)


def add_callout(doc: Document, label: str, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    mark_first_row_as_header(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT_FILL)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    label_run = paragraph.add_run(f"{label}: ")
    label_run.bold = True
    label_run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    paragraph.add_run(text)
    doc.add_paragraph()


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths_dxa: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        set_cell_shading(cell, HEADER_FILL)
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor.from_string("0B2545")
    for row_values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_values):
            cells[index].text = value
            for paragraph in cells[index].paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
    set_table_geometry(table, widths_dxa)
    mark_first_row_as_header(table)
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.left_indent = Inches(0.375)
        paragraph.paragraph_format.first_line_indent = Inches(-0.188)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.paragraph_format.left_indent = Inches(0.375)
        paragraph.paragraph_format.first_line_indent = Inches(-0.188)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.add_run(item)


def add_code_block(doc: Document, command: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.25)
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(command)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)


def build_doc() -> None:
    doc = Document()
    style_document(doc)
    add_title(doc)
    add_callout(
        doc,
        "Purpose",
        "This project repeats the Spike Gate plus Optuna optimisation idea, but uses SEP "
        "segmentation instead of MTObjects. The result is a resumable optimiser that chooses SEP parameters "
        "based on spike coverage and data-loss penalties.",
    )

    doc.add_heading("Code and Outputs", level=1)
    add_table(
        doc,
        ["Item", "Path or role"],
        [
            ["Optimiser script", "Foreground Masking/optimise_spike_gate_SEP.py"],
            ["SEP implementation", "Foreground Masking/Interactive tools/interactive_sep_spike_gate_parameter_tester.py provides SEP masks and Spike Gate helpers."],
            ["Default output folder", "Remove foreground objects/sep spike optimisation/<timestamp>"],
            ["Best parameter file", "sep_spike_optimisation_best.json"],
            ["Study database", "sep_spike_optimisation_study.sqlite3, reused by --resume-output-dir"],
        ],
        [2700, 6660],
    )

    doc.add_heading("Workflow", level=1)
    add_numbered(
        doc,
        [
            "Load usable galaxies from the S4G manifest, or the explicit list supplied with --names.",
            "Build a deprojected, bar-aligned profile and use Spike Gate to mark narrow positive spike samples.",
            "For each Optuna trial, run global SEP segmentation with the sampled parameter set.",
            "Project the two-dimensional SEP mask back onto the bar-major profile.",
            "Score the trial by rewarding spike coverage and penalising non-spike profile masking, image masking, and profile change.",
            "Write per-trial summaries, per-galaxy details, the best JSON, and a resumable Optuna SQLite study.",
        ],
    )

    doc.add_heading("Spike Gate Criteria", level=1)
    add_table(
        doc,
        ["Criterion", "Default", "Meaning"],
        [
            ["Finite positive sample", "Required", "Only valid positive profile samples are eligible."],
            ["Central exclusion", "8 arcsec", "Protects central galaxy light from becoming optimisation target evidence."],
            ["Local peak", "Candidate >= immediate neighbours", "Requires a narrow local maximum in the bar-major profile."],
            ["Neighbour excess", "25 percent above 4-15 arcsec neighbourhood", "Requires contrast against the surrounding radial profile."],
            ["Side-drop check", "40 percent above +/-3 samples", "Rejects broad shoulders and slowly varying galaxy structure."],
            ["Window expansion", "+/-2 samples", "Scores a small radial region around each detected spike."],
        ],
        [2300, 3000, 4060],
    )

    doc.add_heading("SEP Process", level=1)
    add_bullets(
        doc,
        [
            "The detection image is either the original image or a smooth-model residual, controlled by --detect-on.",
            "SEP estimates a background with the selected background box size, subtracts it, and extracts connected sources above detect_thresh.",
            "A Gaussian-like filter kernel of filter_size is applied during extraction.",
            "SEP deblends sources using deblend_nthresh and deblend_cont.",
            "Detected segments are filtered by max_area, max_elongation, and exclude_center_pixels.",
            "The final mask is dilated by dilation_radius and then scored against the Spike Gate profile target.",
        ],
    )
    add_table(
        doc,
        ["SEP parameter", "Optimised range", "Effect"],
        [
            ["detect_thresh", "0.5 to 8.0 sigma", "Lower values catch fainter sources but can over-mask galaxy structure."],
            ["minarea", "1 to 80 pixels", "Minimum connected area SEP requires for a source."],
            ["deblend_nthresh", "8 to 64", "Number of deblend thresholds."],
            ["deblend_cont", "0.0001 to 0.1, log sampled", "Controls how readily SEP separates blended sources."],
            ["back_size", "16 to 256 pixels", "Background mesh size used by SEP."],
            ["filter_size", "1, 3, 5, 7, or 9 pixels", "Smoothing kernel size for SEP extraction."],
            ["dilation_radius", "0 to 8 pixels", "Expands retained segments to cover object wings."],
            ["max_area", "20 to 5000 pixels", "Rejects broad detections likely to be galaxy structure."],
            ["max_elongation", "1.5 to 20.0", "Rejects highly elongated detections."],
        ],
        [2200, 2400, 4760],
    )

    doc.add_heading("Objective Function", level=1)
    add_callout(
        doc,
        "Objective",
        "12 x (1 - mean_spike_coverage) + 4 x mean_non_spike_profile_fraction + "
        "2 x mean_masked_fraction + 1.5 x mean_profile_affected_fraction + 1 x mean_profile_change",
    )
    add_table(
        doc,
        ["Metric", "Interpretation"],
        [
            ["mean_spike_coverage", "Fraction of Spike Gate profile samples covered by the SEP mask. Higher is better."],
            ["mean_non_spike_profile_fraction", "Fraction of non-spike profile samples masked. Lower protects the science profile."],
            ["mean_masked_fraction", "Fraction of all image pixels masked. Lower means less global data loss."],
            ["mean_profile_affected_fraction", "Fraction of the bar-major profile touched by the mask."],
            ["mean_profile_change", "Median log-intensity change after log-linear bridging across masked samples."],
        ],
        [3100, 6260],
    )

    doc.add_heading("Running the Optimiser", level=1)
    doc.add_heading("Initial 20-Galaxy Run", level=2)
    add_code_block(
        doc,
        'python "Foreground Masking\\optimise_spike_gate_SEP.py" --max-images 20 --initial-points 16 --max-iter 64',
    )
    doc.add_heading("Resume an Interrupted Run", level=2)
    add_code_block(
        doc,
        'python "Foreground Masking\\optimise_spike_gate_SEP.py" --resume-output-dir "D:\\Dropbox\\Public Documents\\UCLAN\\MSc Research\\Remove foreground objects\\sep spike optimisation\\YYYYMMDD_HHMMSS" --max-images 20 --initial-points 16 --max-iter 64',
    )
    add_bullets(
        doc,
        [
            "Do not resume into a directory while the original process is still running.",
            "Terminal progress includes timestamped trial start/end lines and per-galaxy coverage/masking lines by default.",
            "Use --no-progress-galaxies to suppress per-galaxy lines.",
            "Use --prepare-only to inspect the selected spike-positive galaxies without running Optuna.",
        ],
    )

    doc.add_heading("Issues and Caveats", level=1)
    add_bullets(
        doc,
        [
            "SEP is likely faster than MTObjects, but its detection threshold and deblend choices can be more sensitive to galaxy residual structure.",
            "Spike Gate is profile evidence, not ground-truth segmentation. Visual review remains necessary before accepting final parameters.",
            "The optimiser currently tunes global SEP parameters. A parameter set that covers spikes may still remove unrelated objects, which is partly intended but must be checked for science impact.",
            "Control galaxies can be included with --no-require-spikes if over-masking needs a stronger safeguard.",
        ],
    )

    doc.save(OUT_PATH)


if __name__ == "__main__":
    build_doc()
    print(OUT_PATH)
