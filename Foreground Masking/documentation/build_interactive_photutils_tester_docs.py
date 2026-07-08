from __future__ import annotations

import importlib.util
import shutil
import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


SCRIPT_DIR = Path(__file__).resolve().parents[1]
PROJECT_ROOT = SCRIPT_DIR.parent
if str(PROJECT_ROOT) not in sys.path:
    sys.path.append(str(PROJECT_ROOT))

from machine_paths import remove_foreground_folder  # noqa: E402

DOC_DIR = SCRIPT_DIR / "documentation"
OUTPUT_DOCX = DOC_DIR / "Interactive Photutils Parameter Tester Documentation.docx"
DROPBOX_DOC_DIR = remove_foreground_folder("Laptop") / "documentation"
SCRIPT_NAME = "interactive_photutils_parameter_tester.py"

template_path = (
    SCRIPT_DIR.parent
    / "Shoulder Recognition Erwin"
    / "documentation"
    / "build_shoulder_quantification_docs.py"
)
spec = importlib.util.spec_from_file_location("shoulder_docs_template", template_path)
template = importlib.util.module_from_spec(spec)
assert spec.loader is not None
spec.loader.exec_module(template)

style_document = template.style_document
set_paragraph_border_bottom = template.set_paragraph_border_bottom
add_text = template.add_text
add_bullets = template.add_bullets
add_numbered = template.add_numbered
add_table = template.add_table
MUTED = template.MUTED


def add_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = "Interactive Photutils Parameter Tester Documentation"
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = MUTED
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.text = "Generated documentation"
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = MUTED


def add_intro(doc: Document) -> None:
    doc.add_paragraph("Interactive Photutils Parameter Tester Documentation", style="Title")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run(SCRIPT_NAME).italic = True
    set_paragraph_border_bottom(subtitle)
    add_table(
        doc,
        ["Field", "Value"],
        [
            ["Script", SCRIPT_NAME],
            ["Document date", date.today().isoformat()],
            [
                "Purpose",
                "Interactive visual tuning of global Photutils and spike-gated foreground masks for S4G barred-galaxy images.",
            ],
            [
                "Default output",
                r"{Remove foreground objects}\interactive_photutils_parameter_tester",
            ],
        ],
        [2300, 7060],
    )
    add_text(
        doc,
        "The tester is a Tkinter/Matplotlib application for trying mask parameters on one galaxy at a time. "
        "It shows the residual view, deprojected bar-aligned image, and bar-major profile so that parameter changes can be judged immediately against the science profile rather than only against the image.",
    )


def add_run_section(doc: Document) -> None:
    doc.add_heading("How to run it", level=1)
    add_table(
        doc,
        ["Machine", "Command"],
        [
            [
                "Laptop",
                'python "Foreground Masking/interactive_photutils_parameter_tester.py" --pc Laptop',
            ],
            [
                "Desktop",
                'python "Foreground Masking/interactive_photutils_parameter_tester.py" --pc Desktop',
            ],
        ],
        [1800, 7560],
    )
    add_bullets(
        doc,
        [
            "The Machine selector can be changed inside the application; it reloads the available FITS image list using the configured Dropbox path for that PC.",
            "The Galaxy selector chooses one S4G image from the geometry manifest.",
            "The left control panel is scrollable and the window opens larger than before so the full parameter set remains reachable.",
        ],
    )


def add_methods(doc: Document) -> None:
    doc.add_heading("Masking method selector", level=1)
    add_table(
        doc,
        ["Method", "Meaning", "Default detection nsigma"],
        [
            [
                "Global",
                "All compact residual-source segments that pass area, elongation, central-exclusion, and optional peak-threshold filters are applied to the mask.",
                "5.0",
            ],
            [
                "Spike-gated",
                "Photutils residual-source segments are candidates only. A segment is applied only if its dilated profile footprint intersects detected bar-major spike samples.",
                "3.5",
            ],
        ],
        [1600, 5960, 1800],
    )
    add_text(
        doc,
        "Changing the masking method resets the controls to that method's defaults. "
        "The internal code still uses the stable method keys global and spike-gated; the dropdown displays the user-facing labels Global and Spike-gated.",
    )


def add_units(doc: Document) -> None:
    doc.add_heading("Parameter units selector", level=1)
    add_text(
        doc,
        "The Parameter units selector changes pixel-based controls between Pixels and Arcsec. "
        "The GUI converts the displayed values back to pixels before calling the masking code, so the underlying Photutils behaviour is unchanged.",
    )
    add_table(
        doc,
        ["Control", "Pixels display", "Arcsec display", "Conversion"],
        [
            ["Smooth sigma", "[px]", "[as]", "linear: pixels x pixel_scale"],
            ["Dilation radius", "[px]", "[as]", "linear: pixels x pixel_scale"],
            ["Max segment area", "[px]", "[as^2]", "area: pixels x pixel_scale^2"],
            ["Deprojected central exclusion", "[px]", "[as]", "linear: pixels x pixel_scale"],
            ["Profile width", "[px]", "[as]", "linear: pixels x pixel_scale"],
            ["Minimum pixels", "pixels", "pixels", "not converted; this is a connected-pixel count"],
        ],
        [2500, 1500, 1500, 3860],
    )


def add_controls(doc: Document) -> None:
    doc.add_heading("Controls and redraw behaviour", level=1)
    add_table(
        doc,
        ["Control group", "Notes"],
        [
            ["Core Photutils", "smooth_sigma_pixels, detection_nsigma, npixels, dilation_radius_pixels, max_area, max_elongation, central exclusion, min peak nsigma, and profile width."],
            ["Spike gate", "spike_excess_fraction, neighbour inner/outer arcsec, side offset samples, side drop fraction, centre exclusion arcsec, and spike window samples."],
            ["Deblend sources", "Used by the global Photutils path. The spike-gated helper currently deblends candidates internally."],
            ["Auto redraw", "When enabled, slider and spinbox edits wait for 0.5 seconds of no further changes before recomputing the mask."],
            ["Redraw", "Cancels any pending delayed redraw and recomputes immediately."],
            ["Reset", "Restores the active method's default values in the currently selected display units."],
            ["Save PNG", "Writes the current figure to the machine-specific interactive tester output folder."],
        ],
        [2500, 6860],
    )


def add_outputs(doc: Document) -> None:
    doc.add_heading("Saved outputs", level=1)
    add_text(
        doc,
        "Saved PNG filenames include the galaxy name, PC, masking method, nsigma, dilation, and max-area settings. "
        "The current pattern is {galaxy}_{pc}_{method}_nsigma{value}_dil{value}_area{value}.png.",
    )
    add_bullets(
        doc,
        [
            "The method element is global or spike-gated in saved interactive tester PNG names.",
            "Production report outputs use the separate report convention {galaxy_name}_fg_removed_global or {galaxy_name}_fg_removed_sp-gated.",
            "The status line reports kept segments, masked-pixel count, masked fraction, and spike-sample count when spike-gated mode is active.",
        ],
    )


def add_validation(doc: Document) -> None:
    doc.add_heading("Recommended use", level=1)
    add_numbered(
        doc,
        [
            "Start with a known spike-positive galaxy and compare Global against Spike-gated.",
            "Use Pixels when matching code parameters exactly; switch to Arcsec when thinking in angular size across galaxies.",
            "In spike-gated mode, check that green spike-sample markers coincide with the profile disturbance you intend to mask.",
            "Use the residual and deprojected panels together: red circles should explain profile artifacts, not blanket real galaxy structure.",
            "Save PNGs for parameter settings that are candidates for production or documentation.",
        ],
    )


def build() -> Path:
    DOC_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_document(doc)
    add_footer(doc)
    add_intro(doc)
    add_run_section(doc)
    add_methods(doc)
    add_units(doc)
    add_controls(doc)
    add_outputs(doc)
    add_validation(doc)
    for style_name in ("Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3"):
        doc.styles[style_name].paragraph_format.keep_with_next = True
    doc.core_properties.title = "Interactive Photutils Parameter Tester Documentation"
    doc.core_properties.subject = "Interactive documentation for global and spike-gated Photutils mask tuning"
    doc.core_properties.author = ""
    doc.save(OUTPUT_DOCX)
    return OUTPUT_DOCX


if __name__ == "__main__":
    docx = build()
    print(docx)
    DROPBOX_DOC_DIR.mkdir(parents=True, exist_ok=True)
    shutil.copy2(docx, DROPBOX_DOC_DIR / docx.name)
    print(DROPBOX_DOC_DIR)
