#!/usr/bin/env python3
"""Build documentation for the interactive MTObjects parameter tester (no spike gate)."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path(__file__).resolve().parent
CODE_PATH = OUT_DIR.parent / "interactive_mtobjects_parameter_tester_No_SpikeGate.py"
ALGORITHM_DOC = OUT_DIR / "interactive_mtobjects_no_spike_gate_algorithm_and_parameters.docx"
FLOW_DOC = OUT_DIR / "interactive_mtobjects_no_spike_gate_code_process_and_flow.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[float]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    total_dxa = 9360
    tbl_w.set(qn("w:w"), str(total_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    width_dxa = [int(round(width * 1440)) for width in widths]
    for width in width_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_w = cell._tc.get_or_add_tcPr().tcW
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(width_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")


def style_doc(doc: Document, title: str, subtitle: str) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    title_para = doc.add_paragraph()
    title_para.paragraph_format.space_after = Pt(4)
    run = title_para.add_run(title)
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("0B2545")

    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(12)
    sub_run = sub.add_run(subtitle)
    sub_run.font.name = "Calibri"
    sub_run.font.size = Pt(11)
    sub_run.font.color.rgb = RGBColor.from_string("555555")


def add_note(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    para = cell.paragraphs[0]
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run(text)
    run.font.bold = True


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        para = doc.add_paragraph(style="List Bullet")
        para.paragraph_format.left_indent = Inches(0.375)
        para.paragraph_format.first_line_indent = Inches(-0.188)
        para.paragraph_format.space_after = Pt(4)
        para.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        para = doc.add_paragraph(style="List Number")
        para.paragraph_format.left_indent = Inches(0.375)
        para.paragraph_format.first_line_indent = Inches(-0.188)
        para.paragraph_format.space_after = Pt(4)
        para.add_run(item)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for index, header in enumerate(headers):
        cell = table.cell(0, index)
        set_cell_shading(cell, "E8EEF5")
        run = cell.paragraphs[0].add_run(header)
        run.font.bold = True
    for row_values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_values):
            cells[index].text = value
    set_table_geometry(table, widths)


def build_algorithm_doc() -> None:
    doc = Document()
    style_doc(
        doc,
        "Interactive MTObjects Tester: Algorithm and Parameters (No Spike Gate)",
        "Reference guide for the MTObjects foreground-mask backend and tuning controls, after the spike-gate comparison view was removed.",
    )
    add_note(
        doc,
        "Scope: this document describes interactive_mtobjects_parameter_tester_No_SpikeGate.py, a variant of the "
        "MTObjects tester that runs a single MTObjects pass per galaxy (the spike-gate second pass and its seven "
        "tuning parameters were removed). The original interactive_mtobjects_parameter_tester.py, with spike gate "
        "intact, is unchanged and still available side by side. This variant keeps the same manifest, geometry, "
        "profile, and PNG review workflow as the original.",
    )

    doc.add_heading("Algorithm Summary", level=1)
    doc.add_paragraph(
        "MTObjects detects astronomical sources by building a max tree of the preprocessed image and applying statistical "
        "attribute filtering to identify significant connected structures. The project README describes it as a tool for "
        "detecting sources, creating segmentation maps, and producing parameter tables; the implementation is based on the "
        "C work by Teeninga et al. and wrapped by Caroline Haigh's Python package."
    )
    add_numbered(
        doc,
        [
            "Prepare either the residual image or original image as the detection surface.",
            "Estimate or accept supplied background mean, variance, soft bias, and gain values.",
            "Subtract the background, optionally smooth with a Gaussian kernel, truncate negative values, and replace NaNs for MTObjects.",
            "Build a max tree over the preprocessed image.",
            "Run MTObjects statistical filtering using alpha, move_factor, and min_distance.",
            "Relabel the object map, restore non-finite pixels to background, and convert MTObjects labels to a boolean foreground mask.",
            "Apply comparison-layer filters: minimum area, maximum area, elongation, central exclusion, and mask dilation.",
        ],
    )

    doc.add_heading("MTObjects Parameters", level=1)
    doc.add_paragraph(
        "Every GUI control label now states which direction increases masking aggressiveness, so the effect of a change "
        "is visible without leaving the sidebar. Defaults reflect the Optuna toy-object-recovery optimisation run of "
        "2026-07-19 (objective = -0.304658). Toy-object optimisation runs created before 2026-07-22 are superseded, "
        "because the current optimiser injects toys only into the same deprojected, bar-aligned galaxy investigation cutout "
        "used for the normal image/profile reports."
    )
    add_table(
        doc,
        ["Parameter", "Default", "More aggressive when...", "Effect"],
        [
            ["alpha", "1e-6", "higher", "Significance level used by the MTObjects statistical test; a higher value is a less strict threshold, so more objects are flagged."],
            ["move_factor", "0.4402", "lower", "Controls how far object markers move upward in the tree; a lower value includes more of a large object's faint outskirts, so more pixels are masked."],
            ["min_distance", "0.2367", "lower", "Minimum brightness separation required between objects; a lower value makes it easier for a node to pass the significance test, so more objects are flagged."],
            ["gaussian_fwhm", "2.084 px", "not directional", "FWHM used by MTObjects preprocessing before tree construction; 0 disables smoothing. Changes sensitivity to small/faint features rather than simply increasing or decreasing masking."],
            ["soft_bias", "0.0", "n/a", "Constant bias term used when estimating gain."],
            ["gain", "-1.0", "n/a", "Electrons per ADU. A negative value asks MTObjects to estimate it."],
            ["bg_mean", "NaN", "n/a", "Background mean. NaN asks MTObjects to estimate it from flat tiles."],
            ["bg_variance", "-1.0", "n/a", "Background variance. A negative value asks MTObjects to estimate it."],
        ],
        [1.35, 0.85, 1.1, 3.2,],
    )

    doc.add_heading("Post-Detection Parameters", level=1)
    add_table(
        doc,
        ["Parameter", "Default", "More aggressive when...", "Effect"],
        [
            ["minarea", "36 px", "lower", "Rejects small detected segments after MTObjects relabeling; a lower threshold keeps/masks smaller segments too."],
            ["max_area", "2073 px", "higher", "Rejects broad segments that are more likely galaxy/bar structure than compact foreground objects; a higher cap lets larger segments be masked too."],
            ["max_elongation", "15.49", "higher", "Rejects very elongated detections using weighted second moments; a higher cap lets more elongated shapes be masked too."],
            ["exclude_center_pixels", "8 px", "lower", "Rejects detections close to the galaxy centre, with arcsec display conversion available; a smaller protected radius exposes more of the centre to masking."],
            ["dilation_radius", "1 px", "higher", "Expands accepted segments before replacing/masking pixels; a larger radius grows the mask further around each segment."],
        ],
        [1.35, 0.85, 1.1, 3.2],
    )

    doc.add_heading("Removed: Spike-Gate Parameters", level=1)
    doc.add_paragraph(
        "The spike-gate comparison view (a second, more inclusive MTObjects pass gated to bar-major-axis profile spikes) "
        "and its seven tuning parameters (spike_gate_move_factor, spike_excess_fraction, spike_neighbour_inner_arcsec, "
        "spike_neighbour_outer_arcsec, spike_side_offset_samples, spike_side_drop_fraction, spike_window_samples) have been "
        "removed from both the internal code and the GUI. The figure no longer includes the third, spike-gated profile "
        "panel; the two remaining profile panels (original and MTObjects-processed) now use the freed vertical space. "
        "See interactive_mtobjects_algorithm_and_parameters.docx for the historical spike-gate reference."
    )

    doc.add_heading("Output Folder", level=1)
    doc.add_paragraph(
        "PNGs from this version are written to a new folder, interactive_mtobjects_parameter_tester_no_spike_gate "
        "(under the machine's foreground-removal research folder), rather than the original "
        "interactive_mtobjects_parameter_tester folder, so outputs from before and after the spike-gate removal do not mix."
    )

    doc.add_heading("Operational Notes", level=1)
    add_bullets(
        doc,
        [
            "MTObjects currently expects compiled shared libraries in the MTObjects checkout; use --mtobjects-root or MTOBJECTS_ROOT.",
            "The tester treats MTObjects as the detector and keeps comparison filters outside the library adapter.",
            "Mask replacement currently uses the median of finite unmasked pixels, matching the SEP tester's simple visualization-oriented cleaning.",
            "The MTObjects README warns the project is a work in progress, so parameter sweeps should record library version or commit when used for comparisons.",
            "Parameter direction claims (alpha, move_factor, min_distance) are taken from the CarolineHaigh/mtobjects README and mt_node_test_4.c significance test, not re-derived empirically.",
        ],
    )
    doc.add_heading("Sources", level=1)
    add_bullets(
        doc,
        [
            "CarolineHaigh/mtobjects README: source detection, segmentation maps, dependencies, and command-line parameters.",
            "CarolineHaigh/mtobjects source files: main.py, preprocessing.py, tree_filtering.py, maxtree.py, postprocessing.py, and mt_node_test_4.c.",
            f"Local implementation: {CODE_PATH.name}.",
        ],
    )
    doc.save(ALGORITHM_DOC)


def build_flow_doc() -> None:
    doc = Document()
    style_doc(
        doc,
        "Interactive MTObjects Tester: Code Process and Flow (No Spike Gate)",
        "Implementation walkthrough for maintaining and extending the tester after spike-gate removal.",
    )
    add_note(
        doc,
        "Design intent: MTObjects is isolated behind a product-producing adapter so SEP, MTObjects, and future detectors can be compared through common masks, rows, profiles, and PNG outputs. "
        "The tester now runs exactly one detector pass per galaxy.",
    )

    doc.add_heading("Runtime Entry Point", level=1)
    add_numbered(
        doc,
        [
            "parse_args reads --manifest, --pc, and --mtobjects-root.",
            "main creates MTObjectsTester with the selected manifest, machine, and MTObjects checkout path.",
            "MTObjectsTester loads the local manifest through foreground_display_helpers.",
            "The Tkinter UI exposes machine, galaxy, detection surface, units, MTObjects parameters, and post-filters; "
            "every slider/spinbox label states which direction is more aggressive.",
        ],
    )

    doc.add_page_break()
    doc.add_heading("Calculation Flow", level=1)
    add_table(
        doc,
        ["Stage", "Function", "Responsibility"],
        [
            ["Load FITS", "load_fits / _load_galaxy", "Read image data and geometry, with a per-galaxy cache."],
            ["Detection image", "prepare_detection_image", "Create original or smoothed-residual detection image and non-finite mask."],
            ["MTObjects setup", "mtobjects_context", "Temporarily switch into the MTObjects checkout so its compiled libraries load correctly."],
            ["Parameter object", "mtobjects_parameter_namespace", "Build the SimpleNamespace expected by MTObjects internals."],
            ["Tree detection", "mtobjects_products", "Preprocess image, build max tree, filter tree, relabel segments, and produce raw rows."],
            ["Post-filtering", "measure_segments / filter_segmentation", "Measure labels and apply min/max area, elongation, centre exclusion, and dilation."],
            ["Visual output", "draw_products", "Render original, residual, isophotes, and profiles; save a PNG."],
        ],
        [1.25, 2.55, 2.7],
    )

    doc.add_page_break()
    doc.add_heading("Product Contract", level=1)
    doc.add_paragraph(
        "The single MTObjects run returns a dictionary shaped like the SEP product boundary. "
        "This keeps downstream drawing and future benchmark code simple."
    )
    add_table(
        doc,
        ["Key", "Type", "Meaning"],
        [
            ["raw_segmentation", "2D int array", "Relabeled MTObjects segmentation before comparison-layer filtering."],
            ["filtered_segmentation", "2D int array", "Segmentation after geometry and size filters."],
            ["mask", "2D bool array", "Dilated foreground mask used for visual/profile cleaning."],
            ["cleaned", "2D float array", "Original image with mask pixels replaced by the finite unmasked median."],
            ["residual", "2D float array", "Original minus smoothed model, used for display and measurements."],
            ["rows", "list[dict]", "Per-label measurements and keep/reject state."],
            ["background_level / background_rms", "float", "MTObjects estimated or supplied background values for audit display."],
        ],
        [1.75, 1.35, 3.4],
    )

    doc.add_heading("What Changed From the Spike-Gate Version", level=1)
    add_bullets(
        doc,
        [
            "Removed functions: spike_gated_mtobjects_products, detect_profile_spikes, spike_samples_to_image_aperture, expand_boolean_mask.",
            "Removed constants: SPIKE_GATE_MOVE_FACTOR and the seven DEFAULT_SPIKE_* defaults.",
            "Removed GUI: the 'Spike Gate' LabelFrame and its seven controls.",
            "Removed figure panel: the third (spike-gated) profile axis; profile_grid is now 2 rows instead of 3.",
            "draw_profile no longer takes a spike_samples argument; its legend now shows whenever a bridged (log-linear "
            "interpolated) profile is drawn, rather than only on the removed spike-gated panel.",
            "MTObjects/Post-filter control labels gained a '(↑/↓ more aggressive)' suffix documenting which direction "
            "increases masking; direction is not shown for parameters where the effect is not simply more/less aggressive "
            "(gaussian_fwhm) or that characterise the noise model rather than detection strictness (soft_bias, gain, bg_mean, bg_variance).",
            "Output folder renamed from interactive_mtobjects_parameter_tester to interactive_mtobjects_parameter_tester_no_spike_gate.",
            "mtobjects_toy_object_parameter_optimisation.py's default_params() no longer sets the seven spike_* keys, which were always "
            "unused there (only mtobjects_products was ever called, never spike_gated_mtobjects_products).",
        ],
    )

    doc.add_heading("Error Handling and Setup", level=1)
    add_bullets(
        doc,
        [
            "If MTObjects cannot import or load its shared libraries, the adapter raises a setup-focused error message.",
            "Non-finite pixels are replaced only for MTObjects preprocessing, then restored to background in the segmentation.",
            "The max-tree C memory is released in a finally block after filtering.",
            "The script syntax-checks without MTObjects installed because imports occur inside the runtime adapter.",
        ],
    )

    doc.add_heading("Maintenance Notes", level=1)
    add_bullets(
        doc,
        [
            "For a formal SEP versus MTObjects benchmark, call detector product functions directly and compare mask fraction, object counts, profile residuals, runtime, and manual review flags.",
            "Keep new detector adapters returning the same product contract rather than branching the plotting code.",
            "If MTObjects is packaged differently later, update mtobjects_context only; the rest of the tester should not need to change.",
            "If replacement semantics change from median fill to interpolation/inpainting, implement that after product creation so detectors remain comparable.",
            "If spike-gating is needed again, interactive_sep_spike_gate_parameter_tester.py still implements the equivalent idea for SEP and can be used as a reference.",
        ],
    )
    doc.save(FLOW_DOC)


def main() -> None:
    build_algorithm_doc()
    build_flow_doc()
    print(ALGORITHM_DOC)
    print(FLOW_DOC)


if __name__ == "__main__":
    main()
