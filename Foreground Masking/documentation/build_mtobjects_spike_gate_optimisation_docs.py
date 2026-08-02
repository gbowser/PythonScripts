#!/usr/bin/env python3
"""Build DOCX documentation for the MTObjects Spike Gate Optuna optimiser."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


DOC_DIR = Path(__file__).resolve().parent
OUT_PATH = DOC_DIR / "MTObjects Spike Gate Optuna Optimisation Documentation.docx"

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
HEADER_FILL = "E8EEF5"
CALLOUT_FILL = "F4F6F9"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_grid = table._tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        tbl_grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def mark_first_row_as_header(table) -> None:
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        tr_pr.append(OxmlElement("w:tblHeader"))


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run("MTObjects Spike Gate Optimisation")
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(100, 116, 139)


def add_title(doc: Document) -> None:
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    title.paragraph_format.keep_with_next = True
    run = title.add_run("MTObjects Spike Gate Optuna Optimisation")
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("0B2545")

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    subtitle_run = subtitle.add_run(
        "Technical reference for identifying bar-profile spikes, masking them with global MTObjects, "
        "and tuning MTObjects parameters with Optuna."
    )
    subtitle_run.font.size = Pt(11)
    subtitle_run.font.color.rgb = RGBColor(71, 85, 105)


def add_callout(doc: Document, label: str, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    mark_first_row_as_header(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT_FILL)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    label_run = paragraph.add_run(f"{label}: ")
    label_run.bold = True
    label_run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    paragraph.add_run(text)
    doc.add_paragraph()


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths_dxa: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for index, header in enumerate(headers):
        hdr[index].text = header
        set_cell_shading(hdr[index], HEADER_FILL)
        for paragraph in hdr[index].paragraphs:
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor.from_string("0B2545")
    for row_values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_values):
            cells[index].text = value
            for paragraph in cells[index].paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
    set_table_geometry(table, widths_dxa)
    mark_first_row_as_header(table)
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.left_indent = Inches(0.375)
        paragraph.paragraph_format.first_line_indent = Inches(-0.188)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.paragraph_format.left_indent = Inches(0.375)
        paragraph.paragraph_format.first_line_indent = Inches(-0.188)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.add_run(item)


def add_code_block(doc: Document, command: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.style = doc.styles["Normal"]
    paragraph.paragraph_format.left_indent = Inches(0.25)
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(command)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)


def build_doc() -> None:
    doc = Document()
    style_document(doc)
    add_title(doc)

    add_callout(
        doc,
        "Purpose",
        "The optimiser uses Spike Gate as a physically motivated profile-level target, then asks "
        "MTObjects to produce a global two-dimensional foreground mask. Optuna searches for the "
        "MTObjects settings that remove the spike-producing sources while minimising unnecessary "
        "masking elsewhere in the image and bar-major profile.",
    )

    doc.add_heading("Where the Code Lives", level=1)
    doc.add_paragraph(
        "The batch optimiser is implemented in "
        "Foreground Masking/Optimisation/optimise_spike_gate_MTObjects.py. It reuses the MTObjects "
        "runner, Spike Gate detector, profile extraction, and mask-profile utilities exposed through "
        "Foreground Masking/Interactive tools/interactive_mtobjects_spike_gate_parameter_tester.py."
    )
    add_table(
        doc,
        ["Component", "Role in the new workflow"],
        [
            ["Spike Gate", "Detects narrow positive features in the deprojected bar-major intensity profile."],
            ["MTObjects", "Builds a global foreground-object segmentation and mask from the image or residual image."],
            ["Optuna", "Samples MTObjects parameter combinations and records the lowest-penalty trial."],
            ["CSV/JSON/SQLite outputs", "Store per-case scores, trial summaries, the best parameters, and the resumable study."],
        ],
        [2700, 6660],
    )

    doc.add_heading("High-Level Workflow", level=1)
    add_numbered(
        doc,
        [
            "Select up to the requested number of usable galaxies from the S4G manifest, or use the explicit galaxy names supplied with --names.",
            "For each galaxy, deproject a bar-aligned cutout and extract a bar-major intensity profile.",
            "Run Spike Gate on that one-dimensional profile to identify candidate spike samples.",
            "For every Optuna trial, run global MTObjects with the suggested parameter set.",
            "Project the resulting two-dimensional MTObjects mask back onto the bar-major profile.",
            "Reward masks that cover Spike Gate samples and penalise masks that affect non-spike profile samples or large fractions of the image.",
            "Write the summary, details, best parameters, and Optuna SQLite study to a timestamped output directory.",
        ],
    )

    doc.add_heading("Spike Gate Criteria", level=1)
    doc.add_paragraph(
        "Spike Gate is deliberately a profile test, not a source extractor. It asks whether a profile "
        "sample is a narrow positive excess compared with its local radial neighbourhood and nearby "
        "side samples. The optimiser then uses those spike samples as evidence that a foreground source "
        "has damaged the bar-major intensity graph."
    )
    add_table(
        doc,
        ["Criterion", "Default", "Meaning"],
        [
            [
                "Finite, positive profile sample",
                "Required",
                "Only valid intensity samples above zero can become spike candidates.",
            ],
            [
                "Central exclusion",
                "8 arcsec",
                "Samples with absolute bar-major radius below this value are ignored to protect central galaxy light.",
            ],
            [
                "Local peak check",
                "Candidate >= immediate neighbours",
                "The sample must be at least as bright as the adjacent left and right profile samples.",
            ],
            [
                "Neighbour excess",
                "25 percent above neighbourhood median",
                "The sample must exceed the median of profile samples 4-15 arcsec away by the configured excess fraction.",
            ],
            [
                "Side-drop check",
                "40 percent above +/-3 samples",
                "The sample must stand above the median of profile samples offset by the configured number of samples.",
            ],
            [
                "Window expansion",
                "+/-2 samples",
                "Detected spike samples are expanded along the profile so the optimiser scores a small radial window, not a single pixel-column.",
            ],
        ],
        [2300, 2500, 4560],
    )
    add_bullets(
        doc,
        [
            "The relevant command-line options are --spike-excess-fraction, --spike-neighbour-inner-arcsec, --spike-neighbour-outer-arcsec, --spike-side-offset-samples, --spike-side-drop-fraction, --spike-center-exclusion-arcsec, and --spike-window-samples.",
            "With --require-spikes enabled, galaxies with no Spike Gate samples are skipped until the requested number of spike-positive cases is reached.",
            "The Spike Gate samples are not the final mask. They are the target evidence used to judge whether the MTObjects mask reached the spike-producing contamination.",
        ],
    )

    doc.add_heading("MTObjects Masking Process", level=1)
    doc.add_paragraph(
        "MTObjects is used globally in the optimiser. This matters: the goal is not merely to hide the "
        "profile samples that Spike Gate marked. The goal is to tune MTObjects so it removes the "
        "two-dimensional foreground sources responsible for those spikes, along with other foreground "
        "objects that may not appear directly as graph spikes."
    )
    add_bullets(
        doc,
        [
            "Prepare the detection image from either the original image or a smooth-model residual, controlled by --detect-on.",
            "Preprocess the finite image for MTObjects, including optional Gaussian smoothing through gaussian_fwhm.",
            "Run MTObjects max-tree filtering with alpha, move_factor, min_distance, and background settings.",
            "Relabel the raw segmentation and reject non-finite pixels.",
            "Measure each segment area, residual-weighted centroid, elongation, peak residual significance, and distance from the galaxy centre.",
            "Filter segments by minarea, max_area, max_elongation, and exclude_center_pixels.",
            "Dilate the surviving segmentation by dilation_radius to capture object wings and nearby contaminated pixels.",
            "Score the resulting global mask against the Spike Gate profile target and the data-loss penalties.",
        ],
    )
    add_table(
        doc,
        ["MTObjects parameter", "Optimised?", "Function"],
        [
            ["alpha", "No; fixed at 1e-6", "Controls the MTObjects rejection boundary. The local compiled build requires the 1e-6 boundary."],
            ["move_factor", "Yes", "Controls how aggressively objects are selected from the max-tree structure."],
            ["min_distance", "Yes", "Separates or merges nearby detections in the MTObjects decision process."],
            ["gaussian_fwhm", "Yes", "Smooths the detection image before max-tree filtering; zero disables the blur."],
            ["minarea", "Yes", "Rejects detections smaller than the selected pixel area."],
            ["dilation_radius", "Yes", "Expands surviving masks to cover wings around detected objects."],
            ["max_area", "Yes", "Rejects broad segments likely to be galaxy structure rather than compact foreground contamination."],
            ["max_elongation", "Yes", "Rejects highly elongated segments that are unlikely to be compact foreground objects."],
        ],
        [2200, 1700, 5460],
    )

    doc.add_heading("Optuna Optimisation", level=1)
    doc.add_paragraph(
        "The optimiser uses Optuna's TPE sampler. Initial trials provide broad exploratory coverage; "
        "later trials focus on regions of parameter space that have produced lower objective scores. "
        "The study is stored in SQLite, so a run can be resumed if the same output directory and study "
        "name are reused."
    )
    add_table(
        doc,
        ["Search dimension", "Range"],
        [
            ["move_factor", "0.05 to 0.95"],
            ["min_distance", "0.0 to 1.0"],
            ["gaussian_fwhm", "0.0 to 5.0 pixels"],
            ["minarea", "1 to 100 pixels"],
            ["dilation_radius", "0 to 8 pixels"],
            ["max_area", "20 to 5000 pixels"],
            ["max_elongation", "1.5 to 25.0"],
        ],
        [2800, 6560],
    )

    doc.add_heading("Objective Function", level=1)
    doc.add_paragraph(
        "Lower objective scores are better. The score strongly prioritises spike coverage, then adds "
        "penalties for masking outside the spike evidence and for broader data loss."
    )
    add_callout(
        doc,
        "Objective",
        "12 x (1 - mean_spike_coverage) + 4 x mean_non_spike_profile_fraction + "
        "2 x mean_masked_fraction + 1.5 x mean_profile_affected_fraction + 1 x mean_profile_change",
    )
    add_table(
        doc,
        ["Metric", "Interpretation"],
        [
            ["mean_spike_coverage", "Fraction of Spike Gate profile samples covered by the projected MTObjects mask. Higher is better."],
            ["mean_non_spike_profile_fraction", "Fraction of non-spike bar-major profile samples affected by the mask. Lower protects the profile."],
            ["mean_masked_fraction", "Fraction of all image pixels masked. Lower means less global data loss."],
            ["mean_profile_affected_fraction", "Fraction of the bar-major profile touched by the mask. Lower means a more local intervention."],
            ["mean_profile_change", "Median log-intensity change after log-linear bridging across masked profile samples. Lower is better."],
        ],
        [3000, 6360],
    )

    doc.add_heading("Outputs", level=1)
    add_table(
        doc,
        ["File", "Contents"],
        [
            ["mtobjects_spike_optimisation_config.json", "Full command-line configuration for the run."],
            ["mtobjects_spike_optimisation_cases.csv", "Prepared galaxies and number of Spike Gate samples per profile."],
            ["mtobjects_spike_optimisation_summary.csv", "One row per trial with objective, aggregate metrics, elapsed time, and parameter values."],
            ["mtobjects_spike_optimisation_details.csv", "One row per galaxy per trial with coverage, masking, profile-change, and segment metrics."],
            ["mtobjects_spike_optimisation_best.json", "Best trial seen so far, including the complete MTObjects parameter dictionary."],
            ["mtobjects_spike_optimisation_study.sqlite3", "Optuna study database for audit and resumption."],
        ],
        [3900, 5460],
    )

    doc.add_heading("Running the Optimiser", level=1)
    doc.add_heading("Initial 20-Galaxy Run", level=2)
    doc.add_paragraph("Use this command to start a new timestamped optimisation run:")
    add_code_block(
        doc,
        'python "Foreground Masking\\Optimisation\\optimise_spike_gate_MTObjects.py" '
        "--max-images 20 --initial-points 12 --max-iter 48",
    )
    doc.add_heading("Resume an Interrupted Run", level=2)
    doc.add_paragraph(
        "Use this command after stopping a run, replacing the directory with the timestamped output "
        "folder created by the original command:"
    )
    add_code_block(
        doc,
        'python "Foreground Masking\\Optimisation\\optimise_spike_gate_MTObjects.py" '
        '--resume-output-dir "D:\\Dropbox\\Public Documents\\UCLAN\\MSc Research\\Remove foreground objects\\'
        'mtobjects spike optimisation\\20260719_165906" --max-images 20 --initial-points 12 --max-iter 48',
    )
    add_bullets(
        doc,
        [
            "Do not resume into the same output directory while the original optimiser process is still running.",
            "On resume, Optuna reads the existing SQLite study and runs only the missing trials up to initial_points + max_iter.",
            "Use --names followed by galaxy names when the sample should be controlled explicitly.",
            "Use --prepare-only to confirm which galaxies have Spike Gate samples before spending time on MTObjects trials.",
            "Use --no-require-spikes when clean/control galaxies should be included as a data-loss safeguard.",
            "Use --detect-on original or --detect-on residual to switch the MTObjects detection image.",
            "Timestamped trial and per-galaxy progress is printed to the Terminal window by default; use --no-progress-galaxies to suppress per-galaxy lines.",
        ],
    )

    doc.add_heading("Interpreting Results", level=1)
    add_bullets(
        doc,
        [
            "Start with mtobjects_spike_optimisation_summary.csv sorted by objective. Good candidates should have high mean_spike_coverage and modest masked fractions.",
            "Inspect mtobjects_spike_optimisation_details.csv for galaxies whose spike coverage remains low. A strong mean score can hide one failed galaxy.",
            "Use mtobjects_spike_optimisation_best.json as the parameter source for follow-up visual tests in the interactive MTObjects tester.",
            "Treat the first short smoke runs only as code validation. Scientific tuning needs enough trials and enough spike-positive galaxies to stabilise the result.",
        ],
    )

    doc.add_heading("Important Caveats", level=1)
    add_bullets(
        doc,
        [
            "Spike Gate identifies profile symptoms, not ground-truth object masks. Visual review remains necessary before accepting a final parameter set.",
            "The objective intentionally rewards covering spike samples, so it should be balanced with explicit control galaxies when over-masking becomes a concern.",
            "The current optimiser fixes alpha at 1e-6 because the available MTObjects rejection boundary only supports that value.",
            "The optimiser writes timestamped output directories by default, which keeps runs separate but means a new command creates a new study unless the output path is reused deliberately.",
        ],
    )

    doc.save(OUT_PATH)


if __name__ == "__main__":
    build_doc()
    print(OUT_PATH)
