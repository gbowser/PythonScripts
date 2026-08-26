from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).with_name("S4G_vNext_synthetic_galaxy_library_action_plan.docx")
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "0B2545"
LIGHT = "F2F4F7"
PALE_BLUE = "E8EEF5"
PALE_GOLD = "FFF4CE"
GRAY = "666666"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    total = sum(widths_dxa)
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths_dxa[idx] / 1440)
            set_cell_margins(cell)


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tag = OxmlElement("w:tblHeader")
    tag.set(qn("w:val"), "true")
    tr_pr.append(tag)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    tag = OxmlElement("w:cantSplit")
    tag.set(qn("w:val"), "true")
    tr_pr.append(tag)


def restart_numbering(doc, paragraphs):
    """Give a group of List Number paragraphs its own sequence starting at 1."""
    numbering = doc.part.numbering_part.element
    base = next(n for n in numbering.findall(qn("w:num")) if n.get(qn("w:numId")) == "5")
    abstract_id = base.find(qn("w:abstractNumId")).get(qn("w:val"))
    new_id = max(int(n.get(qn("w:numId"))) for n in numbering.findall(qn("w:num"))) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(new_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), abstract_id)
    num.append(abstract)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    numbering.append(num)
    for paragraph in paragraphs:
        num_pr = paragraph._p.get_or_add_pPr().get_or_add_numPr()
        num_pr.get_or_add_ilvl().set(qn("w:val"), "0")
        num_pr.get_or_add_numId().set(qn("w:val"), str(new_id))


def font(run, size=11, color="000000", bold=False, italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic
    return run


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.10
    font(p.add_run(text))
    return p


def add_number(doc, lead, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    font(p.add_run(lead + " "), bold=True, color=NAVY)
    font(p.add_run(text))
    return p


def add_callout(doc, label, text, fill=PALE_BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    font(p.add_run(label + "  "), bold=True, color=NAVY)
    font(p.add_run(text), color=NAVY)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, LIGHT)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        font(p.add_run(header), size=9.5, bold=True, color=NAVY)
    set_repeat_header(table.rows[0])
    prevent_row_split(table.rows[0])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            font(p.add_run(str(value)), size=8.9)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        prevent_row_split(table.rows[-1])
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10
for name, size, color, before, after in (
    ("Heading 1", 16, BLUE, 16, 8),
    ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 12, DARK_BLUE, 8, 4),
):
    style = styles[name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

header = section.header.paragraphs[0]
footer = section.footer.paragraphs[0]
doc.settings.odd_and_even_pages_header_footer = False

# Memo masthead
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(12)
p.paragraph_format.space_after = Pt(4)
font(p.add_run("ACTION PLAN"), size=10, bold=True, color=BLUE)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(6)
font(p.add_run("S4G vNext Synthetic-Galaxy and Controlled-Injection Library"), size=24, bold=True, color=NAVY)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(14)
font(p.add_run("A contaminant-free, Haigh-style validation framework for foreground-object masking and bar-profile recovery"), size=13, color=GRAY)

meta = [
    ("Purpose", "Define the next implementable version of the S4G foreground-masking test framework"),
    ("Primary users", "Researcher/developer; later collaborators and reviewers"),
    ("Recommended pilot", "100 synthetic S4G-like galaxy models × 10 foreground realizations = 1,000 paired cases"),
    ("Planning horizon", "12-week pilot, followed by a scale-up decision"),
    ("Status", "Proposed action plan"),
]
for label, value in meta:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    font(p.add_run(f"{label}: "), bold=True, color=NAVY)
    font(p.add_run(value))

add_callout(doc, "Corrected recommendation", "Build S4G vNext around fully synthetic, contaminant-free galaxy-only images and independently generated foreground layers. Use real S4G images—including the current low-contamination set—only for transfer and realism checks, never as absolute clean truth. Preserve exact truth at every stage and split folds by underlying synthetic galaxy model.")

doc.add_heading("1. Decision and intended outcome", level=1)
doc.add_paragraph("S4G vNext should create the uncontaminated galaxy as part of the simulation, rather than infer cleanliness from an observed S4G frame. An observed image with fewer foreground objects is useful, but it cannot provide absolute negative truth: an algorithm may correctly detect a real contaminant that was never labelled, or damage galaxy structure that resembles a compact source.")
doc.add_paragraph("The design follows the essential principle in Haigh et al.: quantitative optimisation requires simulated data with known ground truth. For S4G, the controlled unit should be a contaminant-free synthetic barred-galaxy image, rendered with S4G/IRAC-like observing characteristics, plus a separately generated foreground layer. Real S4G fields then test whether conclusions transfer to genuine morphology, noise, and artifacts.")

doc.add_heading("Success at the end of the pilot", level=2)
for text in [
    "A locked v1.0 catalogue of 100 synthetic, contaminant-free S4G-like barred-galaxy models with exact component truth.",
    "A reproducible generator producing 1,000 paired galaxy-only/contaminated cases plus 100 zero-contamination controls.",
    "Model-held-out train, validation, and test folds that prevent multiple realizations of one synthetic galaxy leaking across folds.",
    "A common evaluator for SEP and MTObjects, including foreground recovery, clean-galaxy damage, and bar-profile fidelity.",
    "A release manifest, configurations, checksums, QA reports, and methods note sufficient to rerun the experiment.",
]:
    add_bullet(doc, text)

doc.add_heading("2. Scope and scientific design", level=1)
doc.add_heading("In scope", level=2)
for text in [
    "Synthetic galaxy models spanning the S4G barred-galaxy morphology and geometry ranges relevant to the profile pipeline.",
    "S4G/IRAC-like pixel scale, PSF, correlated noise, sky structure, masks, and edge/coverage behavior.",
    "Empirical or physically parameterised IRAC-like foreground point sources, including ordinary stars, bright cores, extended wings, and diffraction-spike cases.",
    "Cases sampled both randomly and deliberately across scientifically sensitive zones: bar major axis, bar end, centre exclusion boundary, disc, and background sky.",
    "Paired evaluation of SEP and MTObjects using identical cases, folds, truth definitions, and metrics.",
    "The current 40 low-contamination S4G images as a separate transfer-validation set, explicitly not clean truth.",
]:
    add_bullet(doc, text)

doc.add_heading("Out of scope for the quantitative library", level=2)
for text in [
    "APOD and amateur astrophotography images, because stretching, compositing, denoising, star reduction, PSF, calibration, and reuse rights are heterogeneous.",
    "Observed S4G images—including low-contamination examples—treated as perfectly clean ground truth.",
    "Cross-survey images mixed into a single optimisation set without survey-specific PSF/noise modelling.",
    "Galaxy models trained or selected using results from the locked test fold.",
]:
    add_bullet(doc, text)

doc.add_heading("Core dataset contract", level=2)
add_table(doc,
    ["Product", "Definition", "Scientific use"],
    [
        ("galaxy_truth.fits", "Synthetic galaxy signal only; no foreground/background objects", "Absolute clean truth and damage measurement"),
        ("component_truth.fits", "Bulge, disc, bar, arms/rings/clumps as separate planes", "Morphology-aware damage and profile diagnostics"),
        ("foreground.fits", "Added foreground flux only", "Exact signal provenance and flux accounting"),
        ("observed_clean.fits", "Galaxy truth after PSF, sampling, sky, and noise; no contaminants", "Zero-contamination control input"),
        ("contaminated.fits", "observed_clean + foreground under a recorded seed", "Input to SEP/MTObjects"),
        ("truth_mask.fits", "Binary foreground footprint under a declared truth threshold", "Pixel recall, precision, leakage"),
        ("truth_labels.fits", "Integer source IDs and optional component classes", "Per-source and failure-mode analysis"),
        ("injections.csv", "Position, flux, morphology, PSF, zone, seed", "Reconstruction and stratified reporting"),
        ("case.json", "Latent parameters, paths, hashes, units, seeds, configuration", "Audit and exact reconstruction"),
    ],
    [1800, 4020, 3540],
)

doc.add_heading("3. Work plan", level=1)
doc.add_heading("Phase 0 — Freeze the experiment specification (Week 1)", level=2)
add_number(doc, "Define “clean”.", "Require the astrophysical contaminant layer to be identically zero by construction. Noise and instrumental/background effects may be present only as separately generated, fully recorded layers; no observed source pixels or algorithmic replacements may enter galaxy truth.")
add_number(doc, "Define truth. ", "Choose explicit thresholds for the core, wings, and spike components. Keep both a strict truth mask and an expanded photometric-impact mask if they answer different questions.")
add_number(doc, "Define the primary endpoint.", "Adopt bar-profile fidelity as the principal scientific endpoint, supported by mask recall/precision and clean-control damage metrics.")
add_number(doc, "Version the specification.", "Store the agreed schema, eligibility rules, seeds, fold policy, and metric equations in a machine-readable configuration committed with the code.")

doc.add_heading("Phase 1 — Specify the synthetic galaxy population (Weeks 1–3)", level=2)
doc.add_paragraph("Define a parameter distribution that reproduces the range of barred S4G galaxies relevant to this project without copying foreground objects from observed images. Use S4G catalogue/decomposition measurements as distributions or constraints, not as pixel-level clean truth.")
for text in [
    "Model at minimum an exponential/broken-exponential disc, Sérsic bulge, Ferrers-like bar, and optional rings, lenses, spiral structure, asymmetry, and star-forming clumps.",
    "Sample bar length, axis ratio, Sérsic index, bulge-to-total ratio, disc scale, inclination, position angle, contrast, and centre offset from documented S4G-like ranges.",
    "Create explicit morphology strata, including smooth/easy controls and complex cases in which genuine galaxy structure resembles compact contamination.",
    "Reserve combinations of latent galaxy parameters for a locked test fold so validation covers unseen galaxies, not merely unseen noise seeds.",
]:
    add_bullet(doc, text)
add_callout(doc, "Gate 1", "Proceed only when the synthetic population covers the target S4G parameter ranges, every latent parameter is recorded, and no observed foreground pixels enter the galaxy-only model.", fill=PALE_GOLD)

doc.add_heading("Phase 2 — Render contaminant-free S4G-like observations (Weeks 2–5)", level=2)
doc.add_paragraph("Render each galaxy model into two linked products: noiseless component truth and an observed-clean image containing realistic observing effects but absolutely no astrophysical foreground or background sources.")
for text in [
    "Convolve with an empirical or validated IRAC 3.6 μm PSF and sample at the S4G mosaic pixel scale.",
    "Add sky level, Poisson/read-like noise, correlated mosaic noise, flat-field/background structure, NaNs, and coverage edges as separately seeded layers.",
    "Calibrate distributions of noise and background properties from source-free patches in real S4G mosaics without importing detected objects.",
    "Save a noiseless galaxy truth plane so profile error can be separated from ordinary observing noise.",
]:
    add_bullet(doc, text)

doc.add_heading("Phase 3 — Build foregrounds and paired cases (Weeks 4–7)", level=2)
for text in [
    "Build an empirical/parameterised IRAC foreground library: ordinary point sources, bright wings, saturated/near-saturated cases, spikes, and compact background galaxies.",
    "Extend the existing paired-toy manifest workflow rather than creating an unrelated pipeline.",
    "Generate all products from a single case seed; ensure reruns reproduce identical pixels and catalogue rows.",
    "Reject placements outside valid coverage and record every rejection reason.",
    "Preserve flux conservation during sub-pixel shifting and PSF rotation; test edge and NaN handling.",
    "Write atomic outputs to a versioned directory and verify hashes before marking a case complete.",
]:
    add_bullet(doc, text)
add_callout(doc, "Gate 2", "A smoke set must reproduce bit-for-bit, pass FITS/WCS/schema checks, contain zero injected sources in observed-clean controls, and show exact agreement between contaminated − observed_clean and the stored foreground layer within numerical tolerance.", fill=PALE_GOLD)

doc.add_heading("Phase 4 — Create folds and the 1,000-case pilot (Weeks 7–8)", level=2)
doc.add_paragraph("Allocate underlying synthetic galaxy models—not realizations—to folds. All foreground and noise realizations of a model must remain together. Stratify by morphology, inclination, bar size, surface brightness, and observing conditions.")
add_table(doc,
    ["Partition", "Illustrative allocation", "Permitted use"],
    [
        ("Training", "60 models / 600 cases", "Parameter optimisation and diagnostic iteration"),
        ("Validation", "20 models / 200 cases", "Model/parameter selection and stopping decisions"),
        ("Locked test", "20 models / 200 cases", "One-time final comparison and reporting"),
    ],
    [1800, 2700, 4860],
)
doc.add_paragraph("Add one zero-contamination observed-clean control for every synthetic model. The exact 60/20/20 split may be adjusted before generation, but model-level grouping and the locked-test principle must not change after optimisation begins.")

doc.add_heading("Phase 5 — Unified optimisation and evaluation (Weeks 8–10)", level=2)
doc.add_heading("Primary metrics", level=3)
for text in [
    "Bar-profile error: robust relative error and integrated absolute deviation within the defined bar measurement interval.",
    "Clean-control damage: fraction of pixels masked or altered in the absolutely contaminant-free observed-clean controls, plus change relative to galaxy truth.",
    "Foreground recovery: truth-mask recall and precision, reported both by pixel area and by injected source.",
    "Residual contamination: unrecovered injected flux within the measurement aperture and bar-profile extraction strip.",
    "Failure severity: central/bar-crossing catastrophic failures reported separately from easy off-galaxy sources.",
]:
    add_bullet(doc, text)
doc.add_heading("Selection rule", level=3)
doc.add_paragraph("Use a constrained optimisation rather than a single unconstrained score: first require minimum recovery on scientifically important injections, then minimise clean-control damage and bar-profile error. Report medians, dispersion, worst-case tails, and bootstrap confidence intervals by held-out galaxy.")
add_callout(doc, "Gate 3", "Freeze candidate SEP and MTObjects configurations before opening the locked test fold. Any later change creates a new experiment version and a new test set.", fill=PALE_GOLD)

doc.add_heading("Phase 6 — Real-S4G transfer test and release (Weeks 10–12)", level=2)
for text in [
    "Run the frozen configurations once on the locked test set and generate paired comparison panels.",
    "Apply the frozen configurations to the current 40 low-contamination S4G images and a contamination-rich comparison set; report these qualitatively or with partial labels, not as absolute truth.",
    "Investigate transfer failures caused by unmodelled galaxy substructure, correlated backgrounds, PSF artifacts, or real contaminant populations.",
    "Publish the dataset manifest, generation configuration, folds, hashes, summary tables, failure gallery, and environment/dependency record.",
    "Write a methods note distinguishing synthetic truth, contaminant-free observed-clean controls, and qualitative/partially labelled real-field tests.",
    "Decide whether to expand synthetic morphology, observing effects, or foreground classes before any larger release.",
]:
    add_bullet(doc, text)

doc.add_heading("4. Proposed repository architecture", level=1)
doc.add_paragraph("Keep large FITS products outside Git if necessary, but keep manifests, schemas, configurations, small QA fixtures, and code under version control. Paths below are conceptual and should be resolved through the existing machine-path abstraction.")
add_table(doc,
    ["Location", "Contents"],
    [
        ("S4G_vNext/spec/", "Dataset schema, zero-contamination assertion, truth definitions, metric specification"),
        ("S4G_vNext/manifests/", "Synthetic-model catalogue, foreground catalogue, fold manifest, release manifest"),
        ("S4G_vNext/config/", "Injection distributions, seeds, truth thresholds, evaluation settings"),
        ("S4G_vNext/library/galaxy_models/", "Latent parameters, component truth, noiseless synthetic galaxies"),
        ("S4G_vNext/library/observed_clean/", "PSF/noise-rendered controls with zero contaminants"),
        ("S4G_vNext/library/foreground/", "Empirical/parameterised contaminant components"),
        ("S4G_vNext/cases/<release>/", "Generated paired cases and exact truth products"),
        ("S4G_vNext/qa/", "Review panels, schema reports, smoke fixtures, failure galleries"),
        ("S4G_vNext/results/", "Frozen run outputs, summaries, configuration snapshots"),
    ],
    [3060, 6300],
)

doc.add_heading("5. Roles, controls, and acceptance criteria", level=1)
add_table(doc,
    ["Workstream", "Accountability", "Acceptance evidence"],
    [
        ("Galaxy population design", "Scientific modeller + reviewer", "Parameter provenance, coverage report, component-truth checks"),
        ("Observation rendering", "Pipeline developer", "PSF/noise validation, zero-contaminant assertion, reconstruction tests"),
        ("Foreground modelling", "Pipeline developer + reviewer", "Empirical provenance, flux/PSF distributions, visual stamp QA"),
        ("Case generation", "Pipeline developer", "Deterministic smoke tests, conservation tests, schema validation"),
        ("Fold design", "Analysis owner", "Synthetic-model grouping, stratification report, frozen fold hash"),
        ("Optimisation", "Method owner", "Config snapshots, repeatable runs, no test-fold access"),
        ("Final evaluation", "Analysis owner", "Locked run, uncertainty estimates, failure analysis"),
        ("Release", "Project owner", "Version tag, manifest, documentation, checksums, archive location"),
    ],
    [2160, 2700, 4500],
)

doc.add_heading("Definition of done for S4G vNext pilot", level=2)
for text in [
    "All 100 synthetic galaxy models have complete latent parameters, component truth, and model-level fold assignment.",
    "At least 1,000 valid contaminated cases and 100 observed-clean controls are generated, or exclusions are transparently documented.",
    "Every case reconstructs from its manifest and passes flux-layer consistency checks.",
    "SEP and MTObjects are evaluated on precisely the same cases and truth masks.",
    "No underlying synthetic galaxy model crosses train/validation/test boundaries.",
    "The locked test is opened only after configuration freeze and is not recycled for tuning.",
    "A third party can regenerate the synthetic galaxies, observing effects, injections, manifests, and summary metrics from documented commands.",
]:
    add_bullet(doc, text)

doc.add_heading("6. Principal risks and mitigations", level=1)
add_table(doc,
    ["Risk", "Consequence", "Mitigation"],
    [
        ("Synthetic galaxies are too simple", "Settings damage real bars, arms, rings, or clumps", "Use multi-component morphology, explicit complexity strata, and real-S4G transfer tests"),
        ("Toy contaminants are unrealistic", "Optimised settings fail on real S4G images", "Use empirical IRAC stamps/distributions and a separate real-field stress test"),
        ("Fold leakage", "Performance is overstated", "Group all realizations by latent synthetic galaxy model before stratification"),
        ("Single aggregate score hides damage", "High recovery wins by masking galaxy structure", "Constrained selection plus separate damage/profile metrics"),
        ("Truth-mask threshold is arbitrary", "Results depend on annotation convention", "Publish strict and impact masks; run threshold sensitivity analysis"),
        ("Simulator mismatch", "Excellent synthetic scores fail on S4G", "Calibrate PSF/noise from S4G; predeclare transfer tests; expand only from diagnosed gaps"),
        ("Parameter overfitting to S4G", "Limited external validity", "Reserve a later external-survey robustness set; never merge it into S4G tuning"),
    ],
    [2160, 3060, 4140],
)

doc.add_heading("7. Twelve-week pilot schedule", level=1)
add_table(doc,
    ["Weeks", "Milestone", "Key output / decision"],
    [
        ("1", "Specification freeze", "v0.1 schema, clean definition, truth definition, metric contract"),
        ("1–3", "Synthetic population", "100-model latent catalogue and morphology-coverage report"),
        ("2–5", "Observation renderer", "Galaxy/component truth and contaminant-free observed-clean controls"),
        ("4–7", "Foregrounds + generator", "Versioned foreground library and deterministic smoke set"),
        ("7–8", "Pilot generation", "1,000 contaminated cases, 100 clean controls, model-level folds"),
        ("8–10", "Optimisation/validation", "Frozen SEP and MTObjects candidate configurations"),
        ("10–11", "Locked test", "Final comparative metrics and failure analysis"),
        ("11–12", "Real-S4G transfer + release", "Transfer report, v1.0 pilot package, scale-up decision"),
    ],
    [1440, 3060, 4860],
)

doc.add_page_break()
doc.add_heading("8. Immediate next actions", level=1)
immediate_steps = []
for lead, text in [
    ("Correct the dataset terminology.", "Rename the existing CleanGalaxies.txt role to low-contamination/transfer validation; do not use it as negative truth."),
    ("Approve the synthetic contract.", "Confirm that the authoritative unit is galaxy truth + observed-clean control + foreground layer + contaminated image + exact truth + manifest."),
    ("Choose the galaxy renderer.", "Prototype S4G-like bulge/disc/bar/component models and verify bar-profile compatibility before scaling morphology."),
    ("Inventory existing toy code.", "Map the paired-manifest generator, injection routines, evaluators, and machine paths against the revised schema; reuse before rewriting."),
    ("Build a vertical slice.", "Use several synthetic galaxy complexities, foreground classes/locations, and seeds to validate the complete SEP/MTObjects path."),
    ("Review evidence before scaling.", "Only launch the 1,000-case pilot after deterministic reconstruction, flux consistency, and visual truth overlays pass."),
]:
    immediate_steps.append(add_number(doc, lead, text))
restart_numbering(doc, immediate_steps)

doc.add_heading("References and methodological basis", level=1)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
font(p.add_run("Haigh, C. et al. (2021). "), bold=True)
font(p.add_run("Optimising and comparing source-extraction tools using objective segmentation quality criteria. Astronomy & Astrophysics, 645, A107. https://doi.org/10.1051/0004-6361/201936561"))
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
font(p.add_run("Project basis. "), bold=True)
font(p.add_run("Existing repository workflows include the current low-contamination 40-galaxy list (presently named CleanGalaxies.txt), paired toy-manifest generation, held-out cross-validation inputs, SEP processing, MTObjects spike-gate processing, and report generation. The revised design changes the scientific role of that list; it does not discard it."))

doc.core_properties.title = "S4G vNext Synthetic-Galaxy and Controlled-Injection Library — Action Plan"
doc.core_properties.subject = "Implementation plan for contaminant-free synthetic S4G foreground-masking validation"
doc.core_properties.author = "Foreground Masking Research Programme"
doc.core_properties.keywords = "S4G, foreground masking, SEP, MTObjects, clean galaxies, controlled injection, validation"
doc.save(OUT)
print(OUT.resolve())
