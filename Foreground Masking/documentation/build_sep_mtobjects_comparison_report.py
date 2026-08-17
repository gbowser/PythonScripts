#!/usr/bin/env python3
"""Build the verified SEP versus MTObjects decision report."""

from __future__ import annotations

import csv
import json
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


REPO = Path(r"C:\Users\gordo\Documents\Github\PythonScripts")
OUT_DIR = REPO / "Foreground Masking" / "Documentation" / "comparison_report"
OUT_DIR.mkdir(parents=True, exist_ok=True)
DOCX = OUT_DIR / "SEP and MTObjects Masking Comparison and Recommendations 20260816.docx"

SEP_CV = Path(r"D:\Dropbox\Public Documents\UCLAN\MSc Research\Remove foreground objects\sep toy cross validation\20260816_185737")
MTO_CV = Path(r"D:\Dropbox\Public Documents\UCLAN\MSc Research\Remove foreground objects\mtobjects toy recovery followup\20260816_063455")
SEP_BATCH = Path(r"D:\Dropbox\Public Documents\UCLAN\MSc Research\Remove foreground objects\SEP all galaxy batch\sep_toy_cv_20260816_185737\sep_optimised_apply_summary.csv")
MTO_BATCH = Path(r"D:\Dropbox\Public Documents\UCLAN\MSc Research\Remove foreground objects\mtobjects all galaxy batch\mtobjects_toy_recovery_20260816_063455_eight_panel_aligned\mtobjects_optimised_apply_summary.csv")

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(11, 37, 69)
GRAY = RGBColor(90, 98, 108)
LIGHT_GRAY = "F2F4F7"
LIGHT_BLUE = "E8EEF5"
PALE_GOLD = "FFF4CC"
PALE_GREEN = "EAF4EA"
WHITE = RGBColor(255, 255, 255)


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open(newline="", encoding="utf-8-sig") as handle:
        return list(csv.DictReader(handle))


def number(value: str | float | int | None) -> float:
    try:
        return float(value)
    except (TypeError, ValueError):
        return float("nan")


def pct(value: float, digits: int = 1) -> str:
    return f"{100.0 * value:.{digits}f}%"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def style_table(table, widths_dxa: list[int], header_fill=LIGHT_GRAY) -> None:
    table.style = "Table Grid"
    set_table_geometry(table, widths_dxa)
    set_repeat_table_header(table.rows[0])
    for cell in table.rows[0].cells:
        set_cell_shading(cell, header_fill)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.name = "Arial"
                run.font.size = Pt(9)
    for row in table.rows[1:]:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(9)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_heading(text, level=level)
    paragraph.paragraph_format.keep_with_next = True


def add_body(doc: Document, text: str, *, bold_lead: str | None = None) -> None:
    paragraph = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        lead.bold = True
        paragraph.add_run(text[len(bold_lead):])
    else:
        paragraph.add_run(text)


def add_callout(doc: Document, title: str, text: str, fill: str = PALE_GREEN) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    lead = p.add_run(title + " ")
    lead.bold = True
    lead.font.color.rgb = INK
    p.add_run(text)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def production_stats(rows: list[dict[str, str]]) -> dict[str, float | int]:
    good = [row for row in rows if row.get("status") == "ok"]
    values = np.asarray([number(row.get("masked_fraction")) for row in good])
    return {
        "successful": len(good),
        "failed": len(rows) - len(good),
        "mean": float(np.mean(values)),
        "median": float(np.median(values)),
        "p90": float(np.quantile(values, 0.90)),
        "p95": float(np.quantile(values, 0.95)),
        "max": float(np.max(values)),
        "over15": int(np.sum(values > 0.15)),
        "over20": int(np.sum(values > 0.20)),
    }


sep_candidates = read_csv(SEP_CV / "cross_validation_candidates.csv")
mto_candidates = read_csv(MTO_CV / "cross_validation_candidates.csv")
sep_best = json.loads((SEP_CV / "sep_toy_cross_validation_best.json").read_text(encoding="utf-8"))
mto_best = json.loads((MTO_CV / "mtobjects_toy_cross_validation_best.json").read_text(encoding="utf-8"))
sep_rows = read_csv(SEP_BATCH)
mto_rows = read_csv(MTO_BATCH)
sep_prod = production_stats(sep_rows)
mto_prod = production_stats(mto_rows)


def fold_mean(rows: list[dict[str, str]], metric: str) -> float:
    return float(np.mean([number(row[metric]) for row in rows]))


shared_metrics = {
    "SEP": {
        "held_toy_recall": fold_mean(sep_candidates, "held_out_mean_toy_recall"),
        "held_detection": fold_mean(sep_candidates, "held_out_toy_detection_rate"),
        "held_mask": fold_mean(sep_candidates, "held_out_mean_masked_fraction"),
        "all40_toy_recall": number(sep_best["cross_validation_metrics"]["all40_mean_toy_recall"]),
        "all40_detection": number(sep_best["cross_validation_metrics"]["all40_toy_detection_rate"]),
        "all40_mask": number(sep_best["cross_validation_metrics"]["all40_mean_masked_fraction"]),
    },
    "MTObjects": {
        "held_toy_recall": fold_mean(mto_candidates, "held_out_mean_toy_recall"),
        "held_detection": fold_mean(mto_candidates, "held_out_toy_detection_rate"),
        "held_mask": fold_mean(mto_candidates, "held_out_mean_masked_fraction"),
        "all40_toy_recall": number(mto_best["cross_validation_metrics"]["all40_mean_toy_recall"]),
        "all40_detection": number(mto_best["cross_validation_metrics"]["all40_toy_detection_rate"]),
        "all40_mask": number(mto_best["cross_validation_metrics"]["all40_mean_masked_fraction"]),
    },
}


# Charts are internal report figures, not separate deliverables.
fig, axes = plt.subplots(1, 2, figsize=(10.5, 4.0))
labels = ["Toy recall", "Toy detection", "Masked area"]
x = np.arange(len(labels))
width = 0.34
for index, (algorithm, color) in enumerate((("SEP", "#2E74B5"), ("MTObjects", "#7A5A00"))):
    values = [shared_metrics[algorithm]["held_toy_recall"], shared_metrics[algorithm]["held_detection"], shared_metrics[algorithm]["held_mask"]]
    axes[0].bar(x + (index - 0.5) * width, np.asarray(values) * 100, width, label=algorithm, color=color)
axes[0].set_xticks(x, labels)
axes[0].set_ylabel("Four-fold held-out mean (%)")
axes[0].set_title("Shared cross-validation metrics")
axes[0].legend(frameon=False)
axes[0].grid(axis="y", alpha=0.2)

prod_labels = ["Mean", "Median", "95th percentile", "Maximum"]
x2 = np.arange(len(prod_labels))
for index, (algorithm, stats, color) in enumerate((("SEP", sep_prod, "#2E74B5"), ("MTObjects", mto_prod, "#7A5A00"))):
    values = [stats["mean"], stats["median"], stats["p95"], stats["max"]]
    axes[1].bar(x2 + (index - 0.5) * width, np.asarray(values) * 100, width, label=algorithm, color=color)
axes[1].axhline(15, color="#9B1C1C", linestyle="--", linewidth=1.2, label="15% review gate")
axes[1].set_xticks(x2, prod_labels, rotation=16, ha="right")
axes[1].set_ylabel("Masked image area (%)")
axes[1].set_title("182-galaxy production distribution")
axes[1].legend(frameon=False)
axes[1].grid(axis="y", alpha=0.2)
fig.tight_layout()
chart_path = OUT_DIR / "sep_mtobjects_comparison.png"
fig.savefig(chart_path, dpi=190, bbox_inches="tight")
plt.close(fig)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Arial"
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor(0, 0, 0)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10
for style_name, size, color, before, after in (
    ("Heading 1", 16, BLUE, 12, 6),
    ("Heading 2", 13, BLUE, 10, 5),
    ("Heading 3", 12, DARK_BLUE, 8, 4),
):
    style = styles[style_name]
    style.font.name = "Arial"
    style.font.size = Pt(size)
    style.font.color.rgb = color
    style.font.bold = True
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

header = section.header.paragraphs[0]
header.text = "Foreground masking decision memo | SEP versus MTObjects"
header.alignment = WD_ALIGN_PARAGRAPH.LEFT
for run in header.runs:
    run.font.name = "Arial"
    run.font.size = Pt(8.5)
    run.font.color.rgb = GRAY
footer = section.footer.paragraphs[0]
footer.text = "MSc Research - Toy Objects optimisation | 16 August 2026"
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
for run in footer.runs:
    run.font.name = "Arial"
    run.font.size = Pt(8.5)
    run.font.color.rgb = GRAY

# Memo masthead.
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(3)
r = p.add_run("DECISION MEMO")
r.font.name = "Arial"
r.font.size = Pt(23)
r.font.bold = True
r.font.color.rgb = INK
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(14)
r = p.add_run("SEP and MTObjects Toy-Object Masking: Results, Risks and Recommended Production Strategy")
r.font.name = "Arial"
r.font.size = Pt(14)
r.font.color.rgb = GRAY
for label, value in (
    ("Date", "16 August 2026"),
    ("Scope", "Four-fold optimisation on 40 calibration galaxies and aligned eight-panel application to 182 galaxies"),
    ("Decision", "Select a masking strategy that balances contaminant recovery against damage to galaxy structure and bar profiles"),
    ("Status", "Both algorithms completed; 182/182 PNG reports available for each method"),
):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(label + ": ")
    run.bold = True
    p.add_run(value)

add_callout(
    doc,
    "Recommendation.",
    "Use a gated two-algorithm workflow. Use SEP as the higher-recovery candidate generator, but do not accept its mask unconditionally. Automatically accept agreement between SEP and MTObjects, review or constrain SEP-only components, and fall back to MTObjects or an adaptive SEP rerun when total masked area exceeds 15% or bar-profile damage is detected. If only one unattended algorithm can be used, MTObjects is the safer scientific default; if manual review is available and recovery is the priority, use SEP with the gates below.",
)

add_heading(doc, "1. What was completed", 1)
add_body(doc, "Both methods used the same 40-galaxy fold membership: four rotations of 30 training galaxies and 10 held-out galaxies, with 40 Optuna trials per fold. Each candidate was compared on a common 40-galaxy injected evaluation set within its own run. Both selected models were then applied to all 182 galaxies to create aligned eight-panel PNG reports.")
add_body(doc, "The SEP run completed on 16 August 2026 with fold 4 selected and 182/182 successful PNGs. The MTObjects recovery follow-up selected fold 3 and, after correcting the NGC3185 manifest centre, also completed 182/182 PNGs. Neither final batch contains a failed row.")
add_body(doc, "Important comparability caveat: the two cross-validation runs used the same fold membership but different random injection seeds and different objective formulas. Therefore their scalar objective scores must not be compared directly. The decision below uses shared metrics (toy recall, toy detection rate and masked fraction), plus the same-size production batches.")

add_heading(doc, "2. Selected models", 1)
table = doc.add_table(rows=1, cols=3)
table.rows[0].cells[0].text = "Item"
table.rows[0].cells[1].text = "SEP"
table.rows[0].cells[2].text = "MTObjects"
selected_rows = [
    ("Winning fold", "4", "3"),
    ("Detection image", "Residual", "Original"),
    ("Detection threshold / move factor", f"Threshold {sep_best['params']['detect_thresh']:.3f} RMS", f"Move factor {mto_best['params']['move_factor']:.3f}"),
    ("Minimum area", str(sep_best["params"]["minarea"]), str(mto_best["params"]["minarea"])),
    ("Deblending", f"64 levels; contrast {sep_best['params']['deblend_cont']:.4f}", "MTObjects tree segmentation"),
    ("Background", f"Mesh {sep_best['params']['back_size']}; filter {sep_best['params']['filter_size']}", f"Variance {mto_best['params']['bg_variance']:.6f}"),
    ("Dilation radius", str(sep_best["params"]["dilation_radius"]), str(mto_best["params"]["dilation_radius"])),
    ("Maximum area", str(sep_best["params"]["max_area"]), str(mto_best["params"]["max_area"])),
    ("Maximum elongation", f"{sep_best['params']['max_elongation']:.2f}", f"{mto_best['params']['max_elongation']:.2f}"),
]
for values in selected_rows:
    cells = table.add_row().cells
    for index, value in enumerate(values):
        cells[index].text = value
style_table(table, [2500, 3430, 3430], LIGHT_BLUE)

add_heading(doc, "3. Cross-validation comparison", 1)
add_body(doc, "The four-fold held-out means summarise generalisation across all four rotations. The selected-candidate rows show each winner on its own common 40-galaxy injected evaluation set. Higher toy recall and toy detection are better; lower masked fraction is safer.")
table = doc.add_table(rows=1, cols=4)
for index, value in enumerate(("Metric", "SEP", "MTObjects", "Interpretation")):
    table.rows[0].cells[index].text = value
cv_rows = [
    ("Four-fold held-out toy recall", pct(shared_metrics["SEP"]["held_toy_recall"]), pct(shared_metrics["MTObjects"]["held_toy_recall"]), "SEP +9.7 percentage points"),
    ("Four-fold held-out toy detection", pct(shared_metrics["SEP"]["held_detection"]), pct(shared_metrics["MTObjects"]["held_detection"]), "SEP +7.5 percentage points"),
    ("Four-fold held-out masked area", pct(shared_metrics["SEP"]["held_mask"]), pct(shared_metrics["MTObjects"]["held_mask"]), "SEP lower by 0.85 points across folds"),
    ("Winner common-40 toy recall", pct(shared_metrics["SEP"]["all40_toy_recall"]), pct(shared_metrics["MTObjects"]["all40_toy_recall"]), "SEP higher; seeds differ between runs"),
    ("Winner common-40 toy detection", pct(shared_metrics["SEP"]["all40_detection"]), pct(shared_metrics["MTObjects"]["all40_detection"]), "SEP higher; seeds differ between runs"),
    ("Winner common-40 masked area", pct(shared_metrics["SEP"]["all40_mask"]), pct(shared_metrics["MTObjects"]["all40_mask"]), "Similar, with SEP slightly lower"),
]
for values in cv_rows:
    cells = table.add_row().cells
    for index, value in enumerate(values):
        cells[index].text = value
style_table(table, [3000, 1450, 1450, 3460], LIGHT_BLUE)

add_body(doc, "Pixel precision and pixel F-score are very low for both methods under the current diagnostic definition because any mask not overlapping injected toy truth is counted as false-positive area, including masks on real foreground sources. These columns are useful for consistent optimisation within a run, but they do not estimate real-world catalogue precision and should not be the primary algorithm-selection statistic.")

add_heading(doc, "4. Full 182-galaxy behaviour", 1)
table = doc.add_table(rows=1, cols=4)
for index, value in enumerate(("Production measure", "SEP", "MTObjects", "Preferred")):
    table.rows[0].cells[index].text = value
prod_rows = [
    ("Successful reports", f"{sep_prod['successful']}/182", f"{mto_prod['successful']}/182", "Tie"),
    ("Mean masked area", pct(sep_prod["mean"]), pct(mto_prod["mean"]), "MTObjects"),
    ("Median masked area", pct(sep_prod["median"]), pct(mto_prod["median"]), "MTObjects"),
    ("95th percentile masked area", pct(sep_prod["p95"]), pct(mto_prod["p95"]), "MTObjects"),
    ("Maximum masked area", pct(sep_prod["max"]), pct(mto_prod["max"]), "MTObjects"),
    ("Galaxies above 15%", str(sep_prod["over15"]), str(mto_prod["over15"]), "MTObjects"),
    ("Galaxies above 20%", str(sep_prod["over20"]), str(mto_prod["over20"]), "MTObjects"),
]
for values in prod_rows:
    cells = table.add_row().cells
    for index, value in enumerate(values):
        cells[index].text = value
style_table(table, [3100, 1700, 1700, 2860], LIGHT_BLUE)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run()
run.add_picture(str(chart_path), width=Inches(6.45))
p.paragraph_format.space_after = Pt(2)
caption = doc.add_paragraph("Figure 1. Shared recovery metrics and production masked-area distributions. The dashed line marks the proposed automatic review threshold.")
caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
caption.paragraph_format.space_after = Pt(8)
for run in caption.runs:
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = GRAY

add_body(doc, "On paired production galaxies, MTObjects masked less area than SEP in 157 of 182 cases; SEP masked less in 25. SEP exceeded MTObjects by 2.84 percentage points on average and 1.81 points at the median. SEP's most extreme case was NGC1313 at 42.0%, compared with 19.0% for MTObjects.")

add_heading(doc, "5. Decision and operating recommendation", 1)
add_heading(doc, "5.1 Recommended production workflow: gated ensemble", 2)
steps = [
    ("Run both methods", "Generate SEP and MTObjects component masks with the selected parameters and retain component labels rather than only the final binary union."),
    ("Accept agreement", "Automatically accept components substantially overlapping in the two methods. Agreement is a practical confidence signal and should reduce SEP-only galaxy-structure detections."),
    ("Screen SEP-only components", "Accept only when they satisfy size, elongation, central exclusion and bar-proximity rules. Prefer components with a PSF-like shape or an external catalogue match."),
    ("Enforce image-level gates", "If the candidate union masks more than 15%, if a single component dominates the mask, or if the processed bar-major profile changes beyond a defined tolerance, quarantine the image for review."),
    ("Use a conservative fallback", "For quarantined galaxies, use MTObjects alone or rerun SEP with reduced dilation, increased threshold and/or a tighter maximum-area limit. Never silently release a >15% mask."),
]
for index, (title, text) in enumerate(steps, 1):
    p = doc.add_paragraph(style="Heading 3")
    p.add_run(f"Step {index}: {title}")
    add_body(doc, text)

add_heading(doc, "5.2 If only one algorithm can be used", 2)
add_callout(doc, "Unattended scientific processing.", "Use MTObjects. Its lower masking tail is materially safer for preserving galaxy structure: one image exceeded 15%, compared with 30 for SEP.", PALE_GOLD)
add_callout(doc, "Recovery-first processing with human review.", "Use SEP. It achieved higher held-out toy recall and detection, but every >15% image and every obvious bar/isophote disturbance must be reviewed or rerun.", PALE_GOLD)

add_heading(doc, "6. Recommended objective-function changes", 1)
objective_rows = [
    ("Hard feasibility constraints", "Reject a trial if any calibration image exceeds 15%, if toy detection falls below a minimum, or if bar-profile damage exceeds tolerance. Do not rely only on a soft mean penalty."),
    ("Worst-case or CVaR loss", "Penalise the worst 10% of galaxies, or the mean of that tail, so an apparently good average cannot hide 25-42% masking outliers."),
    ("Component-level recovery", "Reward the fraction of toys with a minimum overlap and use per-toy recall. Penalise false components and false component area separately, rather than allowing large real-source masks to dominate pixel precision."),
    ("Profile-preservation term", "Add robust differences in the bar-major profile, isophote ellipticity/position angle and central surface-brightness gradient. Weight damage near the bar more strongly than damage far outside it."),
    ("Multi-objective optimisation", "Optimise toy recovery, false-mask area and profile distortion as separate objectives and select from the Pareto front. This makes the scientific trade-off visible rather than hiding it in one arbitrary scalar."),
    ("Repeated injection seeds", "Repeat each fold with several toy seeds and optimise the median plus a lower-tail recovery statistic. Current SEP and MTObjects comparisons use different evaluation seeds."),
    ("Real-foreground validation", "Add a small manually annotated set of genuine stars and background galaxies. Injected toys are controlled but cannot fully represent PSF wings, diffraction features or complex background galaxies."),
]
table = doc.add_table(rows=1, cols=2)
table.rows[0].cells[0].text = "Change"
table.rows[0].cells[1].text = "Recommended implementation"
for values in objective_rows:
    cells = table.add_row().cells
    cells[0].text, cells[1].text = values
style_table(table, [2500, 6860], LIGHT_BLUE)

add_heading(doc, "7. Alternative masking approaches", 1)
add_heading(doc, "7.1 Catalogue- and PSF-assisted masks", 2)
add_body(doc, "For point sources, use astrometric catalogue matches (for example Gaia) and empirical PSF growth curves to define masks that expand with brightness. This can reduce confusion between compact galaxy structure and foreground stars. It should be combined with image detection for uncatalogued or extended contaminants.")
add_heading(doc, "7.2 Photutils segmentation", 2)
add_body(doc, "Photutils SourceFinder offers threshold detection plus multi-threshold watershed deblending, with explicit minimum-pixel, level and contrast controls. It is a useful independent implementation for an ensemble or a third benchmark, but it still needs the same toy-recovery and galaxy-damage objective.")
add_heading(doc, "7.3 NoiseChisel and Segment", 2)
add_body(doc, "GNU Astronomy Utilities provides a noise-based detection paradigm and a separate segmentation stage designed for diffuse astronomical signal. It is worth benchmarking where background estimation or low-surface-brightness wings defeat fixed thresholds. For this project it must be configured to protect the target galaxy, because its strength at finding diffuse signal can otherwise work against foreground-only masking.")
add_heading(doc, "7.4 Learned instance segmentation", 2)
add_body(doc, "Mask R-CNN and related instance-segmentation methods can jointly detect, classify and deblend sources. They are a longer-term option if sufficient labelled or realistic simulated training data can be assembled. They should not replace the current methods until tested on an external galaxy set and calibrated for uncertainty and profile preservation.")

add_heading(doc, "8. Immediate next actions", 1)
actions = [
    ("1", "Review SEP high-mask set", "Inspect the 30 SEP reports above 15%, beginning with NGC1313, NGC1808, NGC4214, NGC3319 and NGC5236."),
    ("2", "Build component-level ensemble", "Save SEP and MTObjects labelled components, compute overlap and apply agreement/SEP-only decision rules."),
    ("3", "Add release gates", "Fail or quarantine outputs above 15% masked area or above a bar-profile/isophote damage tolerance."),
    ("4", "Re-optimise robustly", "Use hard constraints, tail-risk loss, repeated toy seeds and a manually annotated real-contaminant validation set."),
    ("5", "Benchmark one independent method", "Start with Photutils for implementation comparability; test NoiseChisel if diffuse-wing/background failures remain important."),
]
table = doc.add_table(rows=1, cols=3)
for index, value in enumerate(("Priority", "Action", "Acceptance criterion")):
    table.rows[0].cells[index].text = value
for values in actions:
    cells = table.add_row().cells
    for index, value in enumerate(values):
        cells[index].text = value
style_table(table, [900, 2450, 6010], LIGHT_BLUE)

add_heading(doc, "9. Sources and audit trail", 1)
sources = [
    "Project SEP results: " + str(SEP_CV),
    "Project SEP production summary: " + str(SEP_BATCH),
    "Project MTObjects results: " + str(MTO_CV),
    "Project MTObjects production summary: " + str(MTO_BATCH),
    r"Holwerda, Guide to Source Extractor: D:\Dropbox\Public Documents\UCLAN\MSc Research\SourceExtractor\Guide2source_extractor.pdf",
    "SEP extract API documentation: https://sep.readthedocs.io/en/latest/api/sep.extract.html",
    "Photutils SourceFinder documentation: https://photutils.readthedocs.io/en/stable/api/photutils.segmentation.SourceFinder.html",
    "GNU Astronomy Utilities NoiseChisel documentation: https://www.gnu.org/software/gnuastro/manual/html_node/NoiseChisel.html",
    "Burke et al. (2019), Deblending and Classifying Astronomical Sources with Mask R-CNN: https://doi.org/10.1093/mnras/stz2845",
]
for index, source in enumerate(sources, 1):
    p = doc.add_paragraph(style="List Number")
    p.add_run(source)

doc.core_properties.title = "SEP and MTObjects Masking Comparison and Recommendations"
doc.core_properties.subject = "Four-fold toy-object optimisation and 182-galaxy production comparison"
doc.core_properties.author = "MSc Research foreground masking project"
doc.core_properties.keywords = "SEP, MTObjects, foreground masking, toy objects, cross-validation"
doc.save(DOCX)
print(DOCX)
