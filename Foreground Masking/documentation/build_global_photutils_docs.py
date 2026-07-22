from __future__ import annotations

import argparse
import importlib.util
import shutil
import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


SCRIPT_DIR = Path(__file__).resolve().parents[1]
PROJECT_ROOT = SCRIPT_DIR.parent
if str(PROJECT_ROOT) not in sys.path:
    sys.path.append(str(PROJECT_ROOT))

from machine_paths import PC_RESEARCH_FOLDERS, remove_foreground_folder  # noqa: E402

DOC_DIR = SCRIPT_DIR / "documentation"
MAIN_SCRIPT_NAME = "bar_spike_gated_foreground_report.py"
CALIBRATION_SCRIPT_NAME = "photutils_global_from_bar_spike_calibration.py"
OPTIMISER_SCRIPT_NAME = "optimise_photutils_global_parameters.py"
OUTPUT_DOCX = DOC_DIR / "Global Photutils Foreground Candidate Report Documentation.docx"

DEFAULT_PC = "Laptop"
DROPBOX_DOC_DIR = remove_foreground_folder(DEFAULT_PC) / "documentation"

# Reuse the established project documentation style so this sits beside the
# spike-gated foreground-candidate manual.
template_path = (
    SCRIPT_DIR.parent
    / "Shoulder Recognition Erwin"
    / "documentation"
    / "build_shoulder_quantification_docs.py"
)
spec = importlib.util.spec_from_file_location("shoulder_docs_template", template_path)
template = importlib.util.module_from_spec(spec)
assert spec.loader is not None
spec.loader.exec_module(template)

style_document = template.style_document
set_paragraph_border_bottom = template.set_paragraph_border_bottom
add_text = template.add_text
add_bullets = template.add_bullets
add_numbered = template.add_numbered
add_table = template.add_table
MUTED = template.MUTED


def add_metadata(doc: Document) -> None:
    add_table(
        doc,
        ["Field", "Value"],
        [
            ["Primary production script", MAIN_SCRIPT_NAME],
            ["Calibration experiment script", CALIBRATION_SCRIPT_NAME],
            ["Optimisation helper", OPTIMISER_SCRIPT_NAME],
            ["Document date", date.today().isoformat()],
            [
                "Production output directory",
                r"D:\Dropbox\Public Documents\UCLAN\MSc Research\Remove foreground objects\global_photutils_auto_tuned",
            ],
            [
                "Documentation scope",
                "Global Photutils foreground-candidate masking, bar-spike calibration experiment, per-galaxy auto-tuning, output reports, validation, and limitations.",
            ],
        ],
        [2700, 6660],
    )


def add_intro(doc: Document) -> None:
    doc.add_paragraph("Global Photutils Foreground Candidate Report Documentation", style="Title")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run(f"{MAIN_SCRIPT_NAME} --masking-mode global").italic = True
    set_paragraph_border_bottom(subtitle)
    add_metadata(doc)

    doc.add_heading("High-Level Overview", level=1)
    add_text(
        doc,
        "The global Photutils workflow creates the same one-page portrait diagnostic reports as the spike-gated workflow, but it applies all retained Photutils residual-source segments to the two-dimensional mask. "
        "It is intended as an image-wide foreground-candidate experiment and comparison dataset, not as the recommended shoulder-recognition profile correction.",
    )
    add_text(
        doc,
        "This mode is deliberately different from bar-spike gating. In global mode, a compact residual segment does not need to intersect a bar-major-axis intensity spike before it is masked. "
        "That makes it useful for investigating general foreground-object removal, but it also makes it more likely to remove real galaxy structure when parameters are too aggressive.",
    )
    add_bullets(
        doc,
        [
            "Production global reports are generated through bar_spike_gated_foreground_report.py with --masking-mode global.",
            "The current all-galaxy run uses --auto-tune so each galaxy selects a detection threshold from the candidate sigma grid.",
            "The older photutils_global_from_bar_spike_calibration.py script remains an experiment for deriving global thresholds from the visually calibrated bar-spike set.",
            "The optimiser script evaluates parameter tiers against spike-positive and no-spike control galaxies before committing to a full run.",
        ],
    )


def add_inputs(doc: Document) -> None:
    doc.add_heading("Inputs and Dependencies", level=1)
    add_table(
        doc,
        ["Input or dependency", "Purpose"],
        [
            [
                "S4G geometry manifest",
                "Erwin_s4g_image_downloader/geometry_output/s4g_image_geometry_manifest.csv. Provides galaxy name, FITS path, image centre, disk PA, inclination, bar PA, bar semi-major axis, and pixel scale.",
            ],
            [
                "S4G 3.6 micron FITS image",
                "Science image loaded with astropy.io.fits. Each image is squeezed to two dimensions before analysis.",
            ],
            [
                "foreground_mask_photutils.py",
                "Provides the Photutils segmentation, segment filtering, mask conversion, and dilation helpers.",
            ],
            [
                "bar_spike_gated_foreground_report.py",
                "Provides shared profile extraction, report layout, auto-tuning, interpolation, and the global masking execution path.",
            ],
            [
                "plot_s4g_isophote_axes.py and angle_utils.py",
                "Provide profile extraction, deprojected radii, image cutouts, axis drawing, and geometric transformations.",
            ],
            [
                "Python packages",
                "numpy, scipy, matplotlib, astropy, photutils, python-docx for documentation generation, and reportlab if future PDF generation is moved to a direct-PDF pipeline.",
            ],
        ],
        [2700, 6660],
    )


def add_production_workflow(doc: Document) -> None:
    doc.add_heading("Production Global Workflow", level=1)
    add_numbered(
        doc,
        [
            "Select galaxies from the manifest with --all, --names, or --limit.",
            "Load the FITS image and required geometry fields.",
            "Create a broad Gaussian-smoothed galaxy model using --smooth-sigma-pixels.",
            "Subtract the smoothed model from the science image to make the residual image.",
            "Run Photutils segmentation on the residual image at the selected detection threshold.",
            "Filter segments by connected-pixel count, maximum area, maximum elongation, and central exclusion radius.",
            "Dilate every retained segment by --dilation-radius-pixels.",
            "Set all retained global mask pixels to NaN in the masked image.",
            "Extract original, masked, and log-linear bridge-filled bar-major and bar-minor profiles.",
            "Write the portrait PDF report with observed and deprojected image panels, profile panels, and the parameter table.",
        ],
    )

    doc.add_heading("What global means", level=2)
    add_text(
        doc,
        "Global mode masks all retained residual-source candidates after Photutils segmentation and shape filtering. "
        "Unlike spike-gated mode, it does not require a contaminant to be visible as a narrow spike in the bar-major intensity profile. "
        "This is why global outputs may show many more red circles and longer bridged profile sections than the spike-gated outputs.",
    )


def add_auto_tuning(doc: Document) -> None:
    doc.add_heading("Per-Galaxy Auto-Tuning", level=1)
    add_text(
        doc,
        "The current all-galaxy global run uses per-galaxy auto-tuning rather than one fixed sigma threshold. "
        "For each galaxy, the code evaluates candidate detection thresholds and selects a threshold that balances profile spike reduction against the amount of image/profile data removed.",
    )
    add_table(
        doc,
        ["Auto-tune item", "Current value or behaviour"],
        [
            ["Candidate sigma grid", "5.0, 4.5, 4.0, 3.5, tested from conservative to more aggressive."],
            ["Primary score", "Robust narrow positive spike score measured in log-intensity profile space."],
            ["Profile penalty", "Penalises large fractions of replaced profile samples."],
            ["Image-mask penalty", "Penalises large image mask fractions to discourage overmasking."],
            ["Excessive-mask guard", "Avoids thresholds with mask_fraction > 0.10 or replaced_fraction > 0.45."],
            ["Conservative preference", "Prefers the highest sigma threshold whose spike score is essentially as good as the best non-excessive result."],
        ],
        [2700, 6660],
    )
    add_text(
        doc,
        "This auto-tune process is not a full optimisation of all Photutils parameters. "
        "It chooses the residual detection threshold per galaxy while keeping smoothing, segment area, elongation, dilation, and bridge settings fixed unless overridden on the command line.",
    )


def add_bar_spike_calibration_experiment(doc: Document) -> None:
    doc.add_heading("Bar-Spike Calibration Experiment", level=1)
    add_text(
        doc,
        "photutils_global_from_bar_spike_calibration.py is an experimental bridge between the bar-spike findings and global Photutils masking. "
        "It first measures the bar-major spike count in each selected galaxy, then uses the spike-gated threshold-selection logic on spike-positive galaxies to derive a Photutils detection threshold.",
    )
    add_table(
        doc,
        ["Case", "Behaviour"],
        [
            [
                "Spike-positive calibration galaxy",
                "Uses choose_spike_gated_detection_nsigma to find a threshold that covers the detected profile spike samples.",
            ],
            [
                "No-spike calibration galaxy",
                "Receives the conservative fallback threshold derived from the spike-positive calibration galaxies.",
            ],
            [
                "Output CSV",
                "photutils_global_calibration_from_bar_spikes.csv records input FITS path, profile spike samples, calibrated nsigma, threshold source, spike coverage, and fallback nsigma.",
            ],
            [
                "Output reports",
                "Writes {galaxy_name}_fg_removed_global.png for the selected calibration/test galaxies.",
            ],
        ],
        [3000, 6360],
    )
    add_text(
        doc,
        "This calibration experiment is useful for understanding parameter transfer, but the current full production comparison set was generated through the main report script using --masking-mode global --auto-tune.",
    )


def add_parameters(doc: Document) -> None:
    doc.add_heading("Key Parameters", level=1)
    add_table(
        doc,
        ["Parameter", "Default", "Meaning"],
        [
            ["--masking-mode", "global", "Masks every retained Photutils residual segment rather than applying the bar-spike gate."],
            ["--auto-tune", "enabled in current run", "Selects a detection threshold separately for each galaxy."],
            ["--auto-tune-nsigmas", "5.0, 4.5, 4.0, 3.5", "Threshold candidates tested from conservative to aggressive."],
            ["--smooth-sigma-pixels", "15.0", "Gaussian smoothing width for the broad galaxy model subtracted before residual segmentation."],
            ["--npixels", "8", "Minimum connected pixels required for Photutils detection."],
            ["--dilation-radius-pixels", "3", "Circular dilation radius applied to retained segments."],
            ["--max-area", "500", "Rejects large residual regions likely to be galaxy structure or background residuals."],
            ["--max-elongation", "6.0", "Rejects highly elongated residual segments."],
            [
                "--exclude-center-radius-pixels",
                "8.0",
                "Rejects candidates whose dilated footprint touches the deprojected central exclusion zone.",
            ],
            ["--bridge-merge-gap-samples", "12", "Controls when adjacent profile gaps are bridged as one replacement span."],
            ["--profile-width", "3", "Width in pixels used when sampling bar-major and bar-minor profiles."],
        ],
        [2600, 1900, 4860],
    )


def add_report_layout(doc: Document) -> None:
    doc.add_heading("PNG Report Layout", level=1)
    add_table(
        doc,
        ["Panel", "Contents"],
        [
            ["Observed sky-plane image", "Log-scaled S4G image with bar axes, +r/-r direction labels, and red circles around retained global Photutils candidates."],
            ["Deprojected, bar-aligned image", "Face-on sampled image rotated so the bar major axis is horizontal, with the same retained candidate positions transformed into this frame."],
            ["Original profile graph", "Original bar-major and bar-minor intensity cuts before masking."],
            ["Masked profile graph", "Profiles sampled after all retained global candidate pixels are set to NaN."],
            ["Bridge-filled profile graph", "Masked profiles with straight log-linear bridge segments shown as fine dotted line sections."],
            ["Parameter table", "Masking model, selected detection threshold, segment filtering values, applied mask summary, bridge settings, and auto-tune grid."],
        ],
        [2700, 6660],
    )


def add_outputs(doc: Document) -> None:
    doc.add_heading("Outputs", level=1)
    add_table(
        doc,
        ["Output", "Meaning"],
        [
            [
                "global_photutils_auto_tuned/{galaxy_name}_fg_removed_global.png",
                "Current all-galaxy global Photutils comparison reports. The method suffix distinguishes these from spike-gated outputs.",
            ],
            [
                "photutils_global_calibration_from_bar_spikes.csv",
                "Calibration CSV produced by the experimental calibration script when that script is run.",
            ],
            [
                "{galaxy_name}_fg_removed_global.png",
                "Calibration-script report filename used for the selected calibration/test galaxies.",
            ],
            [
                "optimisation CSV summaries",
                "The optimiser script can write summary/detail CSV files for parameter-tier comparisons before a full production run.",
            ],
        ],
        [3000, 6360],
    )


def add_cli_examples(doc: Document) -> None:
    doc.add_heading("Command-Line Use", level=1)
    add_table(
        doc,
        ["Task", "Command"],
        [
            [
                "Run all galaxies with per-galaxy global auto-tune",
                'python "Foreground Masking/bar_spike_gated_foreground_report.py" --all --masking-mode global --auto-tune --output-dir "D:\\Dropbox\\Public Documents\\UCLAN\\MSc Research\\Remove foreground objects\\global_photutils_auto_tuned"',
            ],
            [
                "Run one galaxy in global mode",
                'python "Foreground Masking/bar_spike_gated_foreground_report.py" --names NGC7418 --masking-mode global --auto-tune --output-dir "D:\\Dropbox\\Public Documents\\UCLAN\\MSc Research\\Remove foreground objects\\global_photutils_auto_tuned"',
            ],
            [
                "Run fixed-threshold global mode",
                'python "Foreground Masking/bar_spike_gated_foreground_report.py" --all --masking-mode global --detection-nsigma 3.5 --output-dir "D:\\Dropbox\\Public Documents\\UCLAN\\MSc Research\\Remove foreground objects\\global_photutils_3p5_sigma"',
            ],
            [
                "Run the bar-spike calibration experiment",
                'python "Foreground Masking/photutils_global_from_bar_spike_calibration.py"',
            ],
            [
                "Run a small optimisation grid",
                'python "Foreground Masking/optimise_photutils_global_parameters.py" --nsigmas 5.0 4.5 4.0 3.5 --dilations 1 2 3 --max-areas 150 300 500 --max-elongations 3 4 6',
            ],
        ],
        [2300, 7060],
    )


def add_validation(doc: Document) -> None:
    doc.add_heading("Validation and Limitations", level=1)
    add_bullets(
        doc,
        [
            "Global Photutils masking is expected to remove more pixels than spike-gated masking; this is useful for comparison but carries a higher overmasking risk.",
            "Known no-spike galaxies should be inspected carefully because excessive global masking can still bridge valid profile data.",
            "The current auto-tune changes detection sigma per galaxy but does not independently tune smoothing, dilation, area, elongation, or central-exclusion parameters.",
            "The red circles in global reports mark retained Photutils candidates, not confirmed foreground stars.",
            "Large galaxy features, knots, spiral structure, and residual background artifacts can be mistaken for compact candidates if filtering is too loose.",
            "The global output should be compared against the spike-gated output before adopting it for shoulder-recognition measurements.",
            "Very bright saturated stars, broad halos, and blended sources may require a separate source model or manual mask review.",
            "Use the optimiser script to test stricter parameter tiers before making a new all-galaxy global run.",
        ],
    )


def add_code_flow(doc: Document) -> None:
    doc.add_heading("Code Flow", level=1)
    add_table(
        doc,
        ["Function or script", "Role"],
        [
            ["bar_spike_gated_foreground_report.py --masking-mode global", "Production global report path for one or all galaxies."],
            ["build_mask_from_residual", "Runs Photutils segmentation, segment filtering, mask conversion, and dilation for global mode."],
            ["choose_detection_nsigma", "Chooses a per-galaxy detection threshold in global auto-tune mode."],
            ["make_report", "Shared report writer used by spike-gated and global modes."],
            ["photutils_global_from_bar_spike_calibration.py", "Experimental script that transfers bar-spike calibration thresholds into global Photutils reports."],
            ["calibrate_detection_thresholds", "Computes per-galaxy or fallback nsigma values for the calibration experiment."],
            ["optimise_photutils_global_parameters.py", "Runs tiered parameter-grid comparisons on spike-positive and no-spike control galaxies."],
        ],
        [3200, 6160],
    )


def add_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = "Global Photutils Foreground Candidate Report Documentation"
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = MUTED
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.text = "Generated documentation"
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = MUTED


def build() -> Path:
    DOC_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_document(doc)
    add_footer(doc)
    add_intro(doc)
    add_inputs(doc)
    add_production_workflow(doc)
    add_auto_tuning(doc)
    add_bar_spike_calibration_experiment(doc)
    add_parameters(doc)
    add_report_layout(doc)
    add_outputs(doc)
    add_cli_examples(doc)
    doc.add_page_break()
    add_code_flow(doc)
    add_validation(doc)

    for style_name in ("Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3"):
        doc.styles[style_name].paragraph_format.keep_with_next = True
    doc.core_properties.title = "Global Photutils Foreground Candidate Report Documentation"
    doc.core_properties.subject = "Technical documentation for global Photutils foreground-candidate masking"
    doc.core_properties.author = ""
    doc.save(OUTPUT_DOCX)
    return OUTPUT_DOCX


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--pc", choices=sorted(PC_RESEARCH_FOLDERS), default=DEFAULT_PC)
    parser.add_argument("--dropbox-doc-dir", type=Path, default=None)
    return parser.parse_args()


def copy_to_dropbox(docx_path: Path, dropbox_doc_dir: Path) -> None:
    dropbox_doc_dir.mkdir(parents=True, exist_ok=True)
    shutil.copy2(docx_path, dropbox_doc_dir / docx_path.name)


if __name__ == "__main__":
    args = parse_args()
    dropbox_doc_dir = args.dropbox_doc_dir or remove_foreground_folder(args.pc) / "documentation"
    docx = build()
    print(docx)
    copy_to_dropbox(docx, dropbox_doc_dir)
    print(dropbox_doc_dir)
