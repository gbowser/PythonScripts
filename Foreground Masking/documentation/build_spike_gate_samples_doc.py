#!/usr/bin/env python3
"""Build a Word explanation of bar-profile spike gate samples."""

from __future__ import annotations

import math
import shutil
from pathlib import Path

import matplotlib

matplotlib.use("Agg")

import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


HERE = Path(__file__).resolve().parent
OUTPUT = HERE / "Spike Gate Samples Explained.docx"
FIGURE_DIR = HERE / "_spike_gate_samples_figures"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
BORDER = "A6B4C4"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        node = borders.find(tag)
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_table_widths(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    tbl_grid = tbl.find(qn("w:tblGrid"))
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        tbl_grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def set_row_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)


def format_cell_text(cell, size_pt: float = 10.0) -> None:
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.15
        for run in paragraph.runs:
            run.font.size = Pt(size_pt)


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.2

    for style_name, size, color, before, after in (
        ("Title", 22, DARK_BLUE, 0, 12),
        ("Heading 1", 15, BLUE, 12, 6),
        ("Heading 2", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_paragraph(doc: Document, text: str, style: str = "Normal"):
    para = doc.add_paragraph(style=style)
    para.add_run(text)
    return para


def add_bullets(doc: Document, lines: list[str]) -> None:
    for line in lines:
        doc.add_paragraph(line, style="List Bullet")


def add_callout(doc: Document, label: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_borders(table, color="CAD3DF")
    set_table_widths(table, [9360], indent_dxa=120)
    set_row_cant_split(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(label)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    p.add_run(f" {body}")
    format_cell_text(cell, size_pt=10.5)
    doc.add_paragraph()


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_borders(table)
    set_row_repeat_header(table.rows[0])
    set_row_cant_split(table.rows[0])
    hdr = table.rows[0].cells
    for index, header in enumerate(headers):
        set_cell_shading(hdr[index], LIGHT_BLUE)
        paragraph = hdr[index].paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run(header)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
        run.font.size = Pt(10)
    for row_values in rows:
        row = table.add_row()
        set_row_cant_split(row)
        cells = row.cells
        for index, value in enumerate(row_values):
            paragraph = cells[index].paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.15
            run = paragraph.add_run(value)
            run.font.size = Pt(9.5)
    set_table_widths(table, widths)
    for row in table.rows:
        for cell in row.cells:
            format_cell_text(cell, size_pt=9.5)
    doc.add_paragraph()


def add_figure(doc: Document, path: Path, caption: str, width: float = 6.6) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string("5B6775")


def style_axis(ax, xlabel: str, ylabel: str = "profile intensity") -> None:
    ax.set_xlabel(xlabel)
    ax.set_ylabel(ylabel)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.grid(axis="y", color="#D8DEE8", linewidth=0.8)


def save_side_drop_figure(path: Path) -> None:
    x = np.arange(0, 11)
    y = np.array([92, 95, 98, 103, 108, 145, 109, 104, 99, 96, 94], dtype=float)
    i = 5
    offset = 2
    side_level = np.median([y[i - offset], y[i + offset]])
    required = 1.20 * side_level

    fig, ax = plt.subplots(figsize=(7.2, 3.6))
    ax.plot(x, y, color="#476A9F", linewidth=2.2, marker="o")
    ax.scatter([i], [y[i]], s=95, color="#C43C35", zorder=5, label="candidate")
    ax.scatter([i - offset, i + offset], [y[i - offset], y[i + offset]], s=75, color="#2E8B57", zorder=5)
    ax.axhline(side_level, color="#2E8B57", linestyle="--", linewidth=1.4, label="median side level")
    ax.axhline(required, color="#C43C35", linestyle=":", linewidth=1.8, label="required level")
    ax.annotate(
        "candidate sample",
        xy=(i, y[i]),
        xytext=(i - 2.8, y[i] + 12),
        arrowprops={"arrowstyle": "->", "color": "#C43C35"},
        color="#C43C35",
    )
    ax.annotate(
        "side samples at i - 2 and i + 2",
        xy=(i + offset, y[i + offset]),
        xytext=(i + 1.3, y[i + offset] - 23),
        arrowprops={"arrowstyle": "->", "color": "#2E8B57"},
        color="#2E8B57",
    )
    ax.text(0.2, required + 2, "candidate must be >= 1.20 x side level", color="#C43C35", fontsize=9)
    ax.set_ylim(80, 165)
    style_axis(ax, "sample index along the bar-major profile")
    ax.legend(loc="upper right", frameon=False)
    fig.tight_layout()
    fig.savefig(path, dpi=200)
    plt.close(fig)


def save_neighbour_zone_figure(path: Path) -> None:
    x = np.linspace(-24, 24, 241)
    y = 100 + 6 * np.sin(x / 4.0) + 42 * np.exp(-(x / 1.4) ** 2)
    inner = 4
    outer = 15

    fig, ax = plt.subplots(figsize=(7.2, 3.6))
    ax.plot(x, y, color="#476A9F", linewidth=2.2)
    ax.axvline(0, color="#C43C35", linewidth=1.5)
    for start, stop in [(-outer, -inner), (inner, outer)]:
        ax.axvspan(start, stop, color="#DDEBCB", alpha=0.9)
    ax.axvspan(-inner, inner, color="#F2F4F7", alpha=0.9)
    ax.text(-13.8, 146, "left neighbour zone", ha="center", fontsize=9, color="#3B6F28")
    ax.text(13.8, 146, "right neighbour zone", ha="center", fontsize=9, color="#3B6F28")
    ax.text(0, 92, "inner gap\nignored", ha="center", fontsize=9, color="#6D7784")
    ax.annotate(
        "candidate",
        xy=(0, y[np.argmin(np.abs(x))]),
        xytext=(3, 152),
        arrowprops={"arrowstyle": "->", "color": "#C43C35"},
        color="#C43C35",
    )
    ax.set_xlim(-22, 22)
    ax.set_ylim(86, 158)
    style_axis(ax, "radius from candidate sample, arcsec")
    fig.tight_layout()
    fig.savefig(path, dpi=200)
    plt.close(fig)


def save_center_exclusion_figure(path: Path) -> None:
    x = np.linspace(-60, 60, 241)
    y = 82 + 55 * np.exp(-(x / 14) ** 2) + 18 * np.exp(-((x - 34) / 2.0) ** 2)
    exclusion = 10

    fig, ax = plt.subplots(figsize=(7.2, 3.3))
    ax.plot(x, y, color="#476A9F", linewidth=2.2)
    ax.axvspan(-exclusion, exclusion, color="#F7D7D2", alpha=0.9)
    ax.axvline(0, color="#6D7784", linewidth=1.0)
    ax.text(0, 132, "centre exclusion", ha="center", color="#9A3328", fontsize=10)
    ax.annotate(
        "candidate outside centre can be tested",
        xy=(34, y[np.argmin(np.abs(x - 34))]),
        xytext=(17, 143),
        arrowprops={"arrowstyle": "->", "color": "#C43C35"},
        color="#C43C35",
    )
    ax.set_xlim(-55, 55)
    ax.set_ylim(76, 150)
    style_axis(ax, "deprojected bar-major radius, arcsec")
    fig.tight_layout()
    fig.savefig(path, dpi=200)
    plt.close(fig)


def save_window_coverage_figure(path: Path) -> None:
    fig, ax = plt.subplots(figsize=(7.2, 3.2))
    samples = np.arange(0, 21)
    y = np.ones_like(samples, dtype=float)
    spike = 10
    window = 2
    expanded = np.arange(spike - window, spike + window + 1)
    mask_samples = np.arange(8, 14)

    ax.scatter(samples, y, s=44, color="#A6B4C4", label="profile samples")
    ax.scatter([spike], [1], s=110, color="#C43C35", label="detected spike")
    ax.scatter(expanded, np.ones_like(expanded) * 1.07, s=80, color="#F2A541", label="expanded spike window")
    ax.scatter(mask_samples, np.ones_like(mask_samples) * 0.93, s=95, color="#2E8B57", marker="s", label="mask intersects profile")
    ax.plot(samples, y, color="#D8DEE8", linewidth=1.2)
    ax.text(spike, 1.17, "window samples = 2", ha="center", color="#A66100", fontsize=10)
    ax.text(10.5, 0.78, "coverage is counted where the green mask row overlaps the orange spike window", ha="center", fontsize=9)
    ax.set_ylim(0.7, 1.28)
    ax.set_xlim(-0.5, 20.5)
    ax.set_yticks([])
    ax.set_xlabel("sample index along the bar-major profile")
    ax.spines["left"].set_visible(False)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.legend(loc="upper right", frameon=False, fontsize=9)
    fig.tight_layout()
    fig.savefig(path, dpi=200)
    plt.close(fig)


def save_flow_figure(path: Path) -> None:
    fig, ax = plt.subplots(figsize=(7.2, 2.9))
    ax.axis("off")
    labels = [
        "bar-major\nprofile",
        "detect spike\nsamples",
        "expand spike\nwindows",
        "detect residual\nsegments",
        "keep segments\ncrossing windows",
        "auto-tune picks\nstrictest nsigma",
    ]
    xs = np.linspace(0.08, 0.92, len(labels))
    for idx, (x, label) in enumerate(zip(xs, labels)):
        rect = plt.Rectangle((x - 0.065, 0.38), 0.13, 0.28, facecolor="#E8EEF5", edgecolor="#A6B4C4", linewidth=1.2)
        ax.add_patch(rect)
        ax.text(x, 0.52, label, ha="center", va="center", fontsize=8.6, color="#1F4D78")
        if idx < len(labels) - 1:
            ax.annotate(
                "",
                xy=(xs[idx + 1] - 0.078, 0.52),
                xytext=(x + 0.078, 0.52),
                arrowprops={"arrowstyle": "->", "color": "#6D7784", "linewidth": 1.3},
            )
    fig.tight_layout()
    fig.savefig(path, dpi=200)
    plt.close(fig)


def make_figures() -> dict[str, Path]:
    if FIGURE_DIR.exists():
        shutil.rmtree(FIGURE_DIR)
    FIGURE_DIR.mkdir(parents=True)
    figures = {
        "side_drop": FIGURE_DIR / "side_drop_test.png",
        "neighbour": FIGURE_DIR / "neighbour_zones.png",
        "center": FIGURE_DIR / "center_exclusion.png",
        "window": FIGURE_DIR / "window_coverage.png",
        "flow": FIGURE_DIR / "spike_gate_flow.png",
    }
    save_side_drop_figure(figures["side_drop"])
    save_neighbour_zone_figure(figures["neighbour"])
    save_center_exclusion_figure(figures["center"])
    save_window_coverage_figure(figures["window"])
    save_flow_figure(figures["flow"])
    return figures


def build_doc() -> None:
    figures = make_figures()
    doc = Document()
    configure_styles(doc)

    doc.add_paragraph("Spike Gate Samples Explained", style="Title")
    add_paragraph(
        doc,
        "Narrative guide to the bar-major profile spike gate used by "
        "bar_spike_gated_foreground_report.py.",
    )
    add_callout(
        doc,
        "One-sentence summary:",
        "the spike gate marks narrow, local, off-centre brightness peaks in the bar-major profile, "
        "then keeps only foreground-candidate mask segments that intersect those marked profile samples.",
    )

    doc.add_heading("Why the spike gate exists", level=1)
    add_paragraph(
        doc,
        "Foreground stars and compact artefacts often appear as narrow positive spikes when the galaxy image "
        "is sampled along the bar major axis. Real galaxy structure can also be bright, but it usually changes "
        "more gradually. The spike gate tries to separate those cases before deciding which residual-image "
        "segments should be used as a foreground mask.",
    )
    add_paragraph(
        doc,
        "The two-dimensional candidate segments are detected on a residual image, not directly on the raw "
        "science image. The pipeline first builds a broad Gaussian-smoothed galaxy model, then subtracts it "
        "from the science image: residual = science image - Gaussian-smoothed galaxy model. Photutils is run "
        "on this residual image so compact positive excesses stand out against the smooth galaxy background.",
    )
    add_figure(doc, figures["flow"], "Figure 1. Overall spike-gated masking flow.", width=6.6)

    doc.add_heading("What is a spike sample?", level=1)
    add_paragraph(
        doc,
        "The code inspects one profile sample at a time. A sample can become a spike only if it passes several "
        "tests: it must be finite and positive, it must not lie inside the central exclusion zone, it must be a "
        "local peak compared with its immediate neighbours, it must exceed a wider surrounding baseline, and it "
        "must drop away sharply over a short side offset.",
    )
    add_bullets(
        doc,
        [
            "Candidate sample: the profile value currently being tested.",
            "Neighbour baseline: the median of profile samples a specified radial distance away from the candidate.",
            "Side samples: two nearby samples, one on each side of the candidate, used to test whether the peak is narrow.",
            "Expanded spike window: extra profile samples added around a detected spike so mask/profile intersections are not too brittle.",
        ],
    )

    doc.add_heading("The side-drop test", level=1)
    add_paragraph(
        doc,
        "The side-drop test is the most literal sharpness test. If the candidate is at index i, the code looks "
        "spike_side_offset_samples positions to the left and right. It takes the median of those two side values "
        "and requires the candidate to be brighter by spike_side_drop_fraction.",
    )
    add_paragraph(
        doc,
        "In formula form: candidate >= (1 + spike_side_drop_fraction) x median(left side sample, right side sample).",
    )
    add_figure(
        doc,
        figures["side_drop"],
        "Figure 2. A candidate sample passing the side-drop test with offset = 2 and side-drop fraction = 0.20.",
        width=6.4,
    )
    add_table(
        doc,
        ["Quantity", "Example value", "Meaning"],
        [
            ["left side sample", "108", "Profile value at i - 2."],
            ["right side sample", "109", "Profile value at i + 2."],
            ["side level", "108.5", "Median of the two side samples."],
            ["side-drop fraction", "0.20", "Candidate must be 20% above the side level."],
            ["required candidate value", "130.2", "108.5 x 1.20."],
            ["actual candidate value", "145", "Passes, because 145 >= 130.2."],
        ],
        [2600, 1900, 4860],
    )
    add_callout(
        doc,
        "Interpretation:",
        "increasing spike_side_drop_fraction demands a pointier peak; decreasing it allows broader, gentler peaks to pass.",
    )

    doc.add_heading("The wider neighbour baseline", level=1)
    add_paragraph(
        doc,
        "The side-drop test asks whether the peak is sharp close to the candidate. The wider neighbour-baseline "
        "test asks whether the candidate is high compared with its surrounding profile context. The code gathers "
        "samples whose distance from the candidate lies between spike_neighbour_inner_arcsec and "
        "spike_neighbour_outer_arcsec. It then requires the candidate to exceed the median of those neighbour "
        "samples by spike_excess_fraction.",
    )
    add_figure(
        doc,
        figures["neighbour"],
        "Figure 3. The inner gap is ignored; neighbour samples come from the green zones.",
        width=6.4,
    )
    add_table(
        doc,
        ["Quantity", "Example value", "Meaning"],
        [
            ["spike_neighbour_inner_arcsec", "4 arcsec", "Samples closer than this are not part of the baseline."],
            ["spike_neighbour_outer_arcsec", "15 arcsec", "Samples farther than this are not part of the baseline."],
            ["neighbour median", "100", "Typical surrounding profile level."],
            ["spike_excess_fraction", "0.25", "Candidate must be 25% above the neighbour median."],
            ["required candidate value", "125", "100 x 1.25."],
            ["actual candidate value", "145", "Passes, because 145 >= 125."],
        ],
        [2900, 1700, 4760],
    )
    add_callout(
        doc,
        "Why both tests?",
        "the neighbour baseline rejects peaks that are not high relative to their wider surroundings, while the side-drop "
        "test rejects broad galaxy humps that do not fall away quickly near the candidate.",
    )

    doc.add_heading("The central exclusion", level=1)
    add_paragraph(
        doc,
        "The galaxy centre is often bright, structured, and steep. A central peak is therefore not good evidence "
        "for a foreground contaminant. The parameter spike_center_exclusion_arcsec defines a symmetric interval "
        "around zero radius where candidates are skipped before the spike tests are applied.",
    )
    add_figure(
        doc,
        figures["center"],
        "Figure 4. Candidate samples inside the centre exclusion are ignored.",
        width=6.4,
    )
    add_table(
        doc,
        ["Sample radius", "centre exclusion", "Tested?"],
        [
            ["-6 arcsec", "10 arcsec", "No. abs(radius) < 10."],
            ["0 arcsec", "10 arcsec", "No. This is the galaxy centre."],
            ["+8 arcsec", "10 arcsec", "No. abs(radius) < 10."],
            ["+18 arcsec", "10 arcsec", "Yes. It is outside the excluded centre."],
        ],
        [2500, 2200, 4660],
    )

    doc.add_heading("Spike-window expansion and coverage", level=1)
    add_paragraph(
        doc,
        "After a sample passes the spike tests, the code expands it along the one-dimensional profile by "
        "spike_window_samples. This does not create more detections in the image. It simply widens the bar-profile "
        "gate used to decide whether a residual-image segment intersects the spike.",
    )
    add_figure(
        doc,
        figures["window"],
        "Figure 5. A detected spike sample expanded by two samples on each side.",
        width=6.4,
    )
    add_table(
        doc,
        ["Parameter", "Example", "Result"],
        [
            ["detected spike index", "10", "The profile sample that passed the spike tests."],
            ["spike_window_samples", "2", "Expand to samples 8, 9, 10, 11, and 12."],
            ["mask/profile intersection", "samples 8 to 12", "The mask covers all expanded spike samples."],
            ["coverage", "5 / 5 = 1.0", "Auto-tune treats the spike samples as fully covered."],
        ],
        [2600, 2100, 4660],
    )

    doc.add_heading("What nsigma means in the Photutils step", level=1)
    add_paragraph(
        doc,
        "The nsigma value controls how high a residual-image pixel must be before it can seed a compact-source "
        "segment. The code estimates the residual background level from the finite pixels, computes a robust "
        "scatter using the median absolute deviation, and sets the detection threshold as:",
    )
    add_paragraph(
        doc,
        "threshold = median(residual image) + nsigma x robust_sigma(residual image).",
    )
    add_paragraph(
        doc,
        "The residual image is then passed to Photutils segmentation. In current Photutils versions this uses "
        "photutils.segmentation.detect_sources with 8-connectivity and the configured npixels value. npixels is "
        "the minimum connected area above threshold required for a detection. If deblending is enabled, the code "
        "then calls photutils.segmentation.deblend_sources to split blended detections into separate segments "
        "where possible.",
    )
    add_callout(
        doc,
        "Interpretation:",
        "higher nsigma is more conservative and finds only brighter residual excesses; lower nsigma is more "
        "aggressive and can recover fainter contaminants, but it also increases the risk of detecting galaxy "
        "structure or noise residuals.",
    )

    doc.add_heading("How this affects auto-tune nsigma", level=1)
    add_paragraph(
        doc,
        "In spike-gated mode, auto-tune does not invent a new threshold from first principles. It tries the "
        "candidate thresholds in auto_tune_nsigmas from conservative to aggressive, usually 5.0, 4.5, 4.0, "
        "and 3.5. For each threshold it detects compact residual-image segments, filters them, keeps only "
        "segments that intersect the expanded spike samples, and measures coverage.",
    )
    add_paragraph(
        doc,
        "The selected nsigma is the first, most conservative threshold that covers essentially every detected "
        "spike sample. If the spike gate detects weaker or more numerous spikes, auto-tune may need a lower, "
        "more aggressive nsigma. If the spike gate detects only obvious sharp spikes, a higher, more conservative "
        "nsigma may be enough.",
    )
    add_table(
        doc,
        ["Candidate nsigma", "Mask segments found?", "Spike-window coverage", "Auto-tune decision"],
        [
            ["5.0", "1 segment", "60%", "Not enough coverage. Try next threshold."],
            ["4.5", "2 segments", "100%", "Select 4.5, because it is the first full-coverage threshold."],
            ["4.0", "More segments", "100%", "Not needed; 4.5 already worked."],
            ["3.5", "Most segments", "100%", "Fallback only if stricter candidates fail."],
        ],
        [2100, 2400, 2200, 2660],
    )

    doc.add_heading("Parameter cheat sheet", level=1)
    add_table(
        doc,
        ["Parameter", "Controls", "If increased"],
        [
            ["spike_excess_fraction", "How high the candidate must be above the wider neighbour baseline.", "Fewer, stronger spikes pass."],
            ["spike_neighbour_inner_arcsec", "How close to the candidate the baseline zone begins.", "Baseline ignores more near-candidate structure."],
            ["spike_neighbour_outer_arcsec", "How far from the candidate the baseline zone extends.", "Baseline uses a wider radial context."],
            ["spike_side_offset_samples", "How far left/right the sharpness side samples are.", "Sharpness is tested over a broader local span."],
            ["spike_side_drop_fraction", "How much the candidate must exceed those side samples.", "Only pointier peaks pass."],
            ["spike_center_exclusion_arcsec", "Radius around the galaxy centre where spike testing is skipped.", "More central structure is ignored."],
            ["spike_window_samples", "How many neighbouring profile samples are added around each detected spike.", "Coverage becomes less brittle but more permissive."],
        ],
        [2800, 3900, 2660],
    )

    doc.add_heading("Practical reading of the gate", level=1)
    add_paragraph(
        doc,
        "A detected spike sample is therefore not just a bright point. It is a bright point that is off-centre, "
        "locally peak-like, high relative to a wider radial neighbourhood, and sharp relative to nearby side "
        "samples. The expanded spike window then turns that point decision into a small target region. The final "
        "foreground mask keeps residual-image segments only when they cross those target regions.",
    )

    doc.save(OUTPUT)
    shutil.rmtree(FIGURE_DIR)


if __name__ == "__main__":
    build_doc()
    print(OUTPUT)
