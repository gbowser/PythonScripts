from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(r"C:\Users\gordo\Documents\Github\PythonScripts\Foreground Masking\documentation\toy_comparison_doc_qa")
ROOT.mkdir(parents=True, exist_ok=True)
CHART1 = ROOT / "toy_comparison_headline_metrics.png"
CHART2 = ROOT / "toy_comparison_paired_scatter.png"
COMPARE = Path(r"D:\Dropbox\Public Documents\UCLAN\MSc Research\Remove foreground objects\Toy Objects comparison\20260821_104129\NGC3627_MTO_left_SEP_right_clean.png")
NAVY = "0B2545"; BLUE = "2E74B5"; GRAY = "555555"; LIGHT = "F2F4F7"; CALLOUT = "E8EEF5"

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), fill); tcPr.append(shd)

def keep(p):
    pPr = p._p.get_or_add_pPr(); el = OxmlElement('w:keepNext'); pPr.append(el)

def setup(doc, running):
    sec = doc.sections[0]; sec.page_width=Inches(8.5); sec.page_height=Inches(11)
    sec.top_margin=sec.bottom_margin=sec.left_margin=sec.right_margin=Inches(1)
    styles=doc.styles
    styles['Normal'].font.name='Calibri'; styles['Normal'].font.size=Pt(10.5); styles['Normal'].font.color.rgb=RGBColor.from_string(GRAY)
    styles['Normal'].paragraph_format.space_after=Pt(6)
    for n,size,col in [('Title',25,NAVY),('Heading 1',16,BLUE),('Heading 2',13,BLUE),('Heading 3',12,NAVY)]:
        s=styles[n]; s.font.name='Calibri'; s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=RGBColor.from_string(col)
        s.paragraph_format.space_before=Pt(12); s.paragraph_format.space_after=Pt(6)

def title(doc, text, subtitle, status):
    p=doc.add_paragraph(); r=p.add_run('FOREGROUND MASKING RESEARCH'); r.bold=True; r.font.size=Pt(10); r.font.color.rgb=RGBColor.from_string(BLUE)
    p=doc.add_paragraph(text, 'Title'); p.paragraph_format.space_after=Pt(4)
    p=doc.add_paragraph(subtitle); p.runs[0].font.size=Pt(13)
    doc.add_paragraph('Prepared: 24 August 2026\nStatus: '+status)

def callout(doc, label, text):
    t=doc.add_table(rows=1, cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER; c=t.cell(0,0); shade(c,CALLOUT)
    p=c.paragraphs[0]; r=p.add_run(label+': '); r.bold=True; r.font.color.rgb=RGBColor.from_string(NAVY); p.add_run(text)
    doc.add_paragraph()

def heading(doc, text, level=1):
    p=doc.add_heading(text, level=level); keep(p); return p

def bullets(doc, items, numbered=False):
    style='List Number' if numbered else 'List Bullet'
    for x in items: doc.add_paragraph(x, style=style)

def table(doc, rows, widths=None):
    t=doc.add_table(rows=len(rows), cols=len(rows[0])); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style='Table Grid'
    for i,row in enumerate(rows):
        for j,v in enumerate(row):
            c=t.cell(i,j); c.text=str(v); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in c.paragraphs:
                for r in p.runs: r.font.name='Calibri'; r.font.size=Pt(9)
            if i==0:
                shade(c,LIGHT)
                for r in c.paragraphs[0].runs: r.bold=True; r.font.color.rgb=RGBColor.from_string(NAVY)
            if widths: c.width=Inches(widths[j])
    doc.add_paragraph(); return t

def figure(doc, path, caption, width=6.45):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run().add_picture(str(path), width=Inches(width))
    p=doc.add_paragraph(caption); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs: r.italic=True; r.font.size=Pt(9)

def save(doc, name):
    p=ROOT/(name+'.docx'); doc.save(p); return p

def methodology():
    d=Document(); setup(d,'Toy Objects Methodology')
    title(d,'Toy Objects Foreground-Masking Methodology','Four-fold optimisation and matched 182-galaxy evaluation of SEP and MTObjects','Definitive methodology record')
    callout(d,'Scope','This document describes only the Toy Objects identification experiment. SEP and MTObjects operate on original science images. Spike Gate residual-image identification is a separate study and is not used here.')
    heading(d,'1. Research question'); d.add_paragraph('The experiment asks how effectively SEP and MTObjects can identify and mask known synthetic foreground-like sources while limiting unnecessary masking of the underlying science image. Synthetic truth provides an objective recovery target that is unavailable for real, unlabelled foreground objects.')
    heading(d,'2. Evidence hierarchy'); bullets(d,[
      'Optimisation evidence: four folds of the same 40 low-foreground calibration galaxies; each fold trains on 30 and validates on the held-out 10.',
      'Selection evidence: each fold winner is evaluated on its method-specific common 40-galaxy injection set, and the best candidate is selected.',
      'Direct method comparison: both selected algorithms are applied to all 182 galaxies using identical standard toy placements and truth masks. This matched evaluation is the primary basis for SEP-versus-MTObjects conclusions.'])
    heading(d,'Important comparability note',2); d.add_paragraph('The SEP and MTObjects optimisation runs used the same galaxies and fold membership but different optimisation evaluation seeds (SEP 202608199; MTObjects 202608299). Their cross-validation scores therefore describe method-specific tuning and are not a strictly paired head-to-head test. The final 182-galaxy comparison corrects this by reconstructing the same six toys per galaxy with seed 202608299 for both algorithms.')
    heading(d,'3. Galaxy sample and cross-validation'); table(d,[
      ('Design element','Implementation'),('Calibration sample','40 galaxies in CleanGalaxies.txt, selected for low foreground contamination.'),('Fold design','Four fixed groups of 10: 30 training and 10 held-out validation galaxies per fold.'),('Optimisation trials','40 per fold: 8 startup evaluations plus 32 Optuna TPE trials.'),('Parallelism','10 image workers.'),('Image input','Original science image for both algorithms.'),('Maximum permitted mask','15% during optimisation.'),('Final evaluation','182 galaxies; six standard toys per image; matched seed 202608299.')],[1.7,4.6])
    heading(d,'4. Synthetic-object construction'); d.add_paragraph('Six non-overlapping toys are placed wholly within the investigated galaxy region and away from image boundaries. Brightness is scaled to the robust image noise so that the challenge remains comparable across galaxies.')
    table(d,[('Property','Specification'),('Object mixture','Star 50%; compact cluster 20%; elliptical galaxy 30%.'),('Peak amplitude','5-25 times robust background sigma.'),('FWHM','Stars/clusters 2-10 pixels; galaxies 5-22 pixels.'),('Galaxy axis ratio','0.35-0.95.'),('Position angle','0-180 degrees.'),('Truth definition','Pixels at or above 8% of peak, then dilated by one pixel.'),('Placement controls','Investigated region only, edge margin, and no toy overlap.')],[1.7,4.6])
    heading(d,'5. Masking pipelines'); heading(d,'SEP',2); d.add_paragraph('SEP estimates the background, detects thresholded connected sources, deblends detections, filters components by area and elongation, excludes the protected centre, and dilates accepted segments. Detection is always performed on the science image.')
    heading(d,'MTObjects',2); d.add_paragraph('MTObjects uses a max-tree representation controlled by move factor, minimum distance, Gaussian smoothing and calibrated background variance. Accepted components are filtered by area and elongation, central exclusion and dilation. Detection is always performed on the science image.')
    heading(d,'Documentation-constrained search spaces',2); table(d,[('Parameter','SEP search range','MTObjects search range'),('Detection control','detect_thresh 0.6-2.0','move_factor 0.0-1.0'),('Minimum area','5-35 pixels','1-40 pixels'),('Deblending','16/32/64; contrast 0.001-0.03','Not applicable'),('Background','mesh 32/64/128/256; filter 1/3/5/7/9','bg_variance, logarithmic'),('Dilation radius','1-6 pixels','1-6 pixels'),('Maximum component area','20-8000 pixels','20-3000 pixels'),('Maximum elongation','1.5-30.0','2.0-15.0')],[2.05,2.1,2.15])
    heading(d,'6. Objective functions'); d.add_paragraph('Tuning uses the incremental mask: the mask after toy injection minus the baseline mask from the uninjected science image. This isolates masking attributable to injected toys.')
    heading(d,'Common metrics',2); bullets(d,['Pixel recall: recovered toy-truth pixels divided by all toy-truth pixels.','Pixel precision: toy-truth overlap divided by all incremental masked pixels.','F1: harmonic mean of pixel recall and precision.','Mean per-toy recall: average recovered fraction across individual toys.','Toy detection rate: fraction of toys with at least 50% of truth pixels masked.','Data-loss controls: mean mask, false-positive fraction and a hard 15% cap.'])
    heading(d,'SEP scalar score',2); d.add_paragraph('Recovery = 0.45 x pixel recall + 0.20 x F1 + 0.25 x per-toy recall + 0.20 x toy detection. Data loss = 0.35 x masked fraction + 0.05 x false-positive fraction. The maximised score is recovery minus data loss; exceeding 15% invokes a large infeasibility penalty.')
    heading(d,'MTObjects recovery-constrained score',2); d.add_paragraph('Recovery = 0.45 x F1 + 0.35 x per-toy recall + 0.20 x toy detection. Data loss = 0.50 x masked fraction + 0.10 x false-positive fraction. A trial is infeasible if it adds no mask, detection is below 0.25, per-toy recall is below 0.20, or the 15% cap is exceeded.')
    heading(d,'7. Selected configurations'); table(d,[('Setting','SEP selected','MTObjects selected'),('Winning fold','4','3'),('Core sensitivity','detect_thresh 1.9811','move_factor 0.8052'),('Minimum area','12','5'),('Deblending','32; contrast 0.001043','-'),('Background','mesh 32; filter 1','variance 0.0008210'),('Gaussian FWHM','-','0.42985'),('Minimum distance','-','0.21516'),('Dilation radius','4','3'),('Maximum area','4885','2987'),('Maximum elongation','11.8404','12.3811'),('Central exclusion','8 pixels','8 pixels')],[2.05,2.1,2.15])
    heading(d,'8. Final 182-galaxy evaluation'); d.add_paragraph('The selected configurations were applied to all 182 galaxies. The matched comparison reconstructs identical toys and evaluates the final mask, not the incremental tuning mask; it therefore reflects the production-style output.')
    bullets(d,['Six matched toys per galaxy; seed 202608299; truth dilation one pixel.','Outputs include total masked fraction, toy recall, per-toy detection, toy-associated precision, F1, non-toy mask, segments and runtime.','Eight-panel PNGs show original, original plus toys, mask, recovered image, original/processed isophotes and bar-major profiles.','Processed profile gaps use the established log-linear bridge.','The 40 calibration galaxies retain the suffix _clean.'])
    heading(d,'9. Interpretation limits'); bullets(d,['Synthetic truth does not reproduce every real foreground morphology or brightness distribution.','High toy recall does not itself prove preservation of galaxy structure; mask extent and visual diagnostics remain necessary.','Toy-associated precision is low when algorithms also mask real objects already present, because truth labels cover injected toys only.','Cross-validation and final metrics use incremental and final masks respectively and are not directly interchangeable.','The separate 20-80 sigma bright-toy MTObjects sensitivity experiment is excluded.'])
    heading(d,'10. Authoritative sources'); d.add_paragraph('SEP optimisation: D:\\Dropbox\\Public Documents\\UCLAN\\MSc Research\\Remove foreground objects\\sep toy cross validation\\20260817_161404\nMTObjects optimisation: D:\\Dropbox\\Public Documents\\UCLAN\\MSc Research\\Remove foreground objects\\mtobjects toy recovery followup\\20260816_063455\nMatched statistics: D:\\Dropbox\\Public Documents\\UCLAN\\MSc Research\\Remove foreground objects\\documentation\\Toy Objects SEP vs MTObjects Statistical Comparison.xlsx')
    return save(d,'Toy Objects SEP and MTObjects Methodology')

def results():
    d=Document(); setup(d,'Toy Objects Results')
    title(d,'Toy Objects Results: SEP versus MTObjects','Matched 182-galaxy comparison using standard injected toys','Results, interpretation and recommendation')
    callout(d,'Executive conclusion','SEP recovered more toy pixels and achieved higher F1 at essentially the same mean masked area. MTObjects had a nearly identical whole-toy detection rate and a lower extreme masked-area tail. Prefer SEP for recovery, with a per-image area guardrail and morphology review; use MTObjects as the conservative fallback for high-risk galaxies.')
    heading(d,'1. Comparison basis'); d.add_paragraph('All direct conclusions use the matched 182-galaxy standard-toy comparison. Each algorithm received the same science image, six toys, placement seed 202608299 and truth mask. Both batches succeeded for all 182 galaxies.')
    heading(d,'2. Population results'); table(d,[('Metric','SEP','MTObjects','SEP - MTO'),('Mean masked image area','8.82%','8.89%','-0.07 pp'),('Median masked image area','8.36%','8.77%','-0.41 pp'),('Mean toy-pixel recall','70.57%','62.42%','+8.15 pp'),('Median toy-pixel recall','82.88%','63.82%','+19.06 pp'),('Mean toy detection rate','68.68%','68.96%','-0.27 pp'),('Mean per-toy recall','68.86%','67.78%','+1.08 pp'),('Mean toy-associated precision','0.700%','0.560%','+0.140 pp'),('Mean toy F1','1.381%','1.107%','+0.274 pp'),('Mean non-toy mask fraction','8.77%','8.85%','-0.07 pp')],[2.7,1.1,1.25,1.25])
    figure(d,CHART1,'Figure 1. Mean recovery, mask extent and overlap-quality metrics across 182 matched galaxies.')
    heading(d,'3. Paired galaxy-level outcomes'); table(d,[('Paired outcome','Galaxies','Interpretation'),('SEP higher toy-pixel recall','121 / 182','SEP advantage in two-thirds of the sample.'),('MTObjects higher toy-pixel recall','57 / 182','Four exact ties.'),('SEP masks less total area','121 / 182','SEP is not systematically more aggressive in the typical case.'),('MTObjects masks less total area','61 / 182','No ties.'),('SEP higher toy F1','134 / 182','Better recovery/precision balance more often.'),('MTObjects higher toy F1','47 / 182','One exact tie.')],[2.65,1.05,2.6])
    figure(d,CHART2,'Figure 2. Each point is one galaxy; the dashed diagonal marks equal performance.')
    heading(d,'4. Distribution and outlier behaviour'); table(d,[('Masked-area statistic','SEP','MTObjects'),('10th percentile','4.91%','6.01%'),('25th percentile','6.30%','6.90%'),('Median','8.36%','8.77%'),('75th percentile','10.35%','10.27%'),('90th percentile','13.23%','11.94%'),('95th percentile','15.14%','12.68%'),('Maximum','33.26%','19.02%')],[3,1.65,1.65])
    d.add_paragraph('Average mask areas are almost identical, but SEP has a heavier upper tail. The final sample includes galaxies unlike the 40 calibration cases, so the tuning-era 15% cap does not ensure every production image remains below 15%. SEP outputs above the threshold require automatic flagging and review.')
    heading(d,'5. Calibration and generalisation'); table(d,[('Selected-candidate metric','SEP (fold 4)','MTObjects (fold 3)'),('Held-out mean toy recall','40.14%','52.55%'),('Held-out toy detection','40.00%','55.00%'),('Held-out mean masked fraction','6.20%','8.82%'),('All-40 mean toy recall','45.38%','50.70%'),('All-40 toy detection','45.42%','53.33%'),('All-40 mean masked fraction','5.71%','9.07%')],[3,1.65,1.65])
    d.add_paragraph('These optimisation figures support reproducibility, not paired comparison: tuning used different injection seeds and incremental-mask scoring. The matched 182-galaxy results are the appropriate head-to-head evidence.')
    d.add_page_break(); heading(d,'6. Visual comparison: NGC3627'); figure(d,COMPARE,'Figure 3. NGC3627 matched diagnostic: MTObjects left, SEP right; dashed black divider between methods.',6.4)
    d.add_paragraph('The paired layout exposes the trade-off directly: compare identical toy boundaries, total mask extent, recovery outlines, isophote preservation and profile bridges before accepting a method for a galaxy.')
    heading(d,'7. Interpretation'); bullets(d,['SEP has stronger average toy-pixel recovery: +8.15 percentage points, with a larger median advantage.','Whole-toy detection is effectively tied. SEP tends to recover more of each toy rather than detect more toys.','Mean masking is effectively tied and SEP masks less area in 121 galaxies; however, SEP has the more severe high-mask outliers.','Toy-associated precision below 1% does not mean all other masking is wrong: real pre-existing foreground objects are absent from the synthetic truth. It does show that toy overlap alone cannot measure scientific selectivity.','Toy Objects test synthetic recovery, not accuracy for real foreground objects without labelled real-object validation.'])
    heading(d,'8. Recommendation'); heading(d,'Primary recommendation',2); d.add_paragraph('Use SEP as the preferred Toy Objects configuration when maximal recovery is the main requirement: it provides higher recall and F1 without increasing mean masked area.')
    heading(d,'Required SEP safeguards',2); bullets(d,['Flag masked area above 15%; treat above 20% as automatic manual review.','Review large coherent components resembling arms, rings, bars or extended galaxy structure.','Retain the eight-panel diagnostic and inspect processed isophotes and bar-major profiles.','Record both total mask area and target recovery; neither is sufficient alone.'])
    heading(d,'When to prefer MTObjects',2); d.add_paragraph('Use MTObjects when SEP enters the high-mask tail, coherent morphology appears threatened, or lower worst-case mask extent is more important than maximum pixel recall.')
    heading(d,'9. Next improvements'); bullets(d,['Repeat cross-validation with identical toy placements and seeds for both algorithms.','Add morphology-protection terms for azimuthal span, annular coherence and low-frequency galaxy structure.','Use a constrained Pareto objective over recovery, masked area and profile/isophote preservation.','Create a manually labelled real-foreground subset for real-object recall and false-mask assessment.','Test a guarded hybrid: SEP first, with MTObjects fallback for SEP high-mask or morphology-risk cases.'],True)
    heading(d,'10. Outputs and sources'); d.add_paragraph('Statistical workbook: D:\\Dropbox\\Public Documents\\UCLAN\\MSc Research\\Remove foreground objects\\documentation\\Toy Objects SEP vs MTObjects Statistical Comparison.xlsx\nVisual comparison: D:\\Dropbox\\Public Documents\\UCLAN\\MSc Research\\Remove foreground objects\\Toy Objects comparison\\20260821_104129\nSEP optimisation: D:\\Dropbox\\Public Documents\\UCLAN\\MSc Research\\Remove foreground objects\\sep toy cross validation\\20260817_161404\nMTObjects optimisation: D:\\Dropbox\\Public Documents\\UCLAN\\MSc Research\\Remove foreground objects\\mtobjects toy recovery followup\\20260816_063455')
    return save(d,'Toy Objects SEP versus MTObjects Results and Recommendations')

if __name__=='__main__':
    print(methodology()); print(results())
