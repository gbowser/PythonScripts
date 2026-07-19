#!/usr/bin/env python3
"""Build a Word guide explaining Photutils optimiser CSV output."""

from __future__ import annotations

from pathlib import Path
import sys

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

PROJECT_ROOT = Path(__file__).resolve().parent.parent
if str(PROJECT_ROOT) not in sys.path:
    sys.path.append(str(PROJECT_ROOT))

from machine_paths import remove_foreground_folder  # noqa: E402


OUTPUT = remove_foreground_folder("Laptop") / "documentation" / "How To Interpret Photutils Optimiser Output.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
BORDER = "A6B4C4"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_row_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        tr_pr.append(OxmlElement("w:tblHeader"))


def set_table_widths(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.find(qn("w:tblGrid"))
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_borders(table, color=BORDER) -> None:
    borders = table._tbl.tblPr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        table._tbl.tblPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def format_cell_text(cell, size_pt: float = 10.0) -> None:
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.15
        for run in paragraph.runs:
            run.font.size = Pt(size_pt)


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_para(doc: Document, text: str):
    paragraph = doc.add_paragraph()
    paragraph.add_run(text)
    return paragraph


def add_callout(doc: Document, label: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_borders(table, color="CAD3DF")
    set_table_widths(table, [9360])
    set_row_cant_split(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(label)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    paragraph.add_run(f" {body}")
    format_cell_text(cell, 10.5)
    doc.add_paragraph()


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_borders(table)
    set_row_repeat_header(table.rows[0])
    set_row_cant_split(table.rows[0])
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    for values in rows:
        row = table.add_row()
        set_row_cant_split(row)
        for index, value in enumerate(values):
            p = row.cells[index].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            run = p.add_run(value)
            run.font.size = Pt(10)
    set_table_widths(table, widths)
    for row in table.rows:
        for cell in row.cells:
            format_cell_text(cell, 10)
    doc.add_paragraph()


def add_footer(doc: Document) -> None:
    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Photutils optimiser output guide")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(90, 90, 90)


def build() -> None:
    doc = Document()
    configure_styles(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("How To Interpret Photutils Optimiser Output")
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    run = subtitle.add_run("Reading the summary and detail CSV files from optimise_photutils_global_parameters.py")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(80, 80, 80)

    add_callout(
        doc,
        "Main idea.",
        "The optimiser is not trying to maximise the amount of masked light. It is trying to find the least aggressive "
        "global Photutils settings that still cover bar-profile spike contamination while doing little damage to clean "
        "control galaxies.",
    )

    doc.add_heading("What The Command Tests", level=1)
    add_para(
        doc,
        "The command evaluates every combination of the supplied nsigma, dilation, max-area, and max-elongation values. "
        "With 5 nsigma values, 3 dilation values, 3 max-area values, and 3 max-elongation values, it tests 135 parameter "
        "sets. For each parameter set it evaluates the default spike-positive galaxies and default no-spike control "
        "galaxies, then writes a ranked summary and a per-galaxy detail table.",
    )
    add_table(
        doc,
        ["Input option", "Meaning", "Interpretation"],
        [
            ["--nsigmas", "Detection threshold in residual-sigma units.", "Lower values detect fainter objects but risk false positives."],
            ["--dilations", "Radius used to grow retained masks.", "Larger values catch source wings but damage more nearby galaxy light."],
            ["--max-areas", "Maximum retained segment area in pixels.", "Larger values allow broader detections but risk masking galaxy structure."],
            ["--max-elongations", "Maximum retained segment elongation.", "Larger values keep more stretched detections; smaller values reject arms and streaks more strongly."],
        ],
        [1900, 3100, 4360],
    )

    doc.add_heading("Where The Output Goes", level=1)
    add_para(
        doc,
        "Because the example command does not specify --pc or --output-dir, the script uses DEFAULT_PC = \"Laptop\". "
        "On this machine that means the default output folder is:"
    )
    add_callout(
        doc,
        "Default output path.",
        r"C:\Users\gordo\Dropbox\Public Documents\UCLAN\MSc Research\Remove foreground objects\optimisation",
    )
    add_para(
        doc,
        r"If you want the run to write to the D:\Dropbox tree instead, use --pc Desktop or pass an explicit --output-dir."
    )

    doc.add_heading("The Two CSV Files", level=1)
    add_table(
        doc,
        ["File", "Purpose", "How to use it"],
        [
            [
                "photutils_parameter_optimisation_summary.csv",
                "One row per parameter combination, sorted with the best score first.",
                "Start here. Use it to identify the top few candidate parameter sets.",
            ],
            [
                "photutils_parameter_optimisation_details.csv",
                "One row per galaxy per parameter combination.",
                "Use this to diagnose why a top-ranked or suspicious parameter set did well or badly.",
            ],
        ],
        [3200, 3100, 3060],
    )

    doc.add_heading("Summary CSV Columns", level=1)
    add_table(
        doc,
        ["Column", "Meaning", "Good direction"],
        [
            ["nsigma, dilation, max_area, max_elongation", "The tested Photutils parameter values.", "No universal direction; judge through the score and diagnostics."],
            ["smooth_sigma", "Gaussian smoothing scale used to build the galaxy model before residual detection.", "Fixed at 15 px in this command unless --smooth-sigma-pixels is supplied."],
            ["npixels", "Minimum connected pixels required for a detection.", "Fixed at 8 in this command unless --npixels is supplied."],
            ["central_exclusion", "Central radius protected from source masking.", "Fixed at 12 px in this command unless --exclude-center-radius-pixels is supplied."],
            ["mean_spike_coverage", "Mean fraction of detected spike-profile samples covered by the mask in spike-positive galaxies.", "Higher is better; values near 1.0 mean the spikes are covered."],
            ["mean_control_damage", "Mean profile-damage score in no-spike controls.", "Lower is better; this is one of the key safeguards against over-masking."],
            ["mean_control_mask_fraction", "Mean fraction of image pixels masked in control galaxies.", "Lower is better."],
            ["mean_spike_mask_fraction", "Mean fraction of image pixels masked in spike-positive galaxies.", "Lower is usually better, provided spike coverage remains high."],
            ["score_lower_is_better", "Weighted combined score used to rank parameter sets.", "Lowest is best, but always inspect the top few visually."],
        ],
        [2700, 4760, 1900],
    )

    doc.add_heading("How The Score Is Built", level=1)
    add_para(
        doc,
        "The score is a weighted penalty. It strongly penalises missing spike samples, then adds penalties for control "
        "profile damage and unnecessary mask size:"
    )
    add_callout(
        doc,
        "Score formula.",
        "score = (1 - mean_spike_coverage) x 10 + mean_control_damage x 5 + mean_control_mask_fraction x 2 + mean_spike_mask_fraction",
    )
    add_para(
        doc,
        "This means missed spike coverage matters most. However, once several parameter sets cover the spikes well, "
        "the best choice should usually be the more conservative one: less control damage, less masked area, higher "
        "nsigma, smaller dilation, and smaller max_area where possible.",
    )

    doc.add_heading("Detail CSV Columns", level=1)
    add_table(
        doc,
        ["Column", "Meaning", "Use"],
        [
            ["group", "Whether the row is a spike galaxy or a control galaxy.", "Separate success cases from damage checks."],
            ["name", "Galaxy name.", "Identify which galaxy drives a good or bad score."],
            ["spike_samples", "Number of spike-profile samples detected in that galaxy.", "Controls may have zero or few spikes; spike galaxies should have useful positive evidence."],
            ["spike_coverage", "Fraction of spike samples covered for that galaxy.", "Low values identify parameter sets that missed contamination."],
            ["segments", "Number of retained Photutils segments.", "Very high values can indicate aggressive false-positive detection."],
            ["masked_pixels", "Number of pixels masked.", "Useful for comparing absolute mask size across a specific galaxy."],
            ["masked_fraction", "Fraction of the image masked.", "Watch for unexpectedly high values, especially in controls."],
            ["profile_affected_fraction", "Fraction of bar-major profile samples touched by the mask.", "High values can mean the science profile is being heavily altered."],
            ["profile_damage", "Combined log-profile change plus affected-profile penalty.", "Use this to find over-masking in control galaxies."],
        ],
        [2600, 4200, 2560],
    )

    doc.add_heading("A Practical Reading Workflow", level=1)
    add_table(
        doc,
        ["Step", "Action", "Why"],
        [
            ["1", "Open the summary CSV and sort by score_lower_is_better if it is not already sorted.", "The script writes the best-ranked parameter sets first."],
            ["2", "Look for mean_spike_coverage close to 1.0.", "A low score is not scientifically useful if real spikes are missed."],
            ["3", "Among high-coverage rows, compare mean_control_damage and mean_control_mask_fraction.", "These tell you whether the settings harm clean galaxies."],
            ["4", "Prefer the most conservative near-tie.", "If scores are similar, choose less aggressive masking rather than the absolute lowest score."],
            ["5", "Use the detail CSV to inspect each top candidate galaxy by galaxy.", "One bad control galaxy can be hidden by a good mean."],
            ["6", "Generate visual reports for the top 3-5 candidates before adopting one.", "Metrics narrow the field; visual inspection catches galaxy-structure mistakes."],
        ],
        [900, 4300, 4160],
    )

    doc.add_heading("Warning Signs", level=1)
    add_table(
        doc,
        ["Pattern in output", "Likely meaning", "Response"],
        [
            ["High spike coverage but high control damage", "The setting removes spikes but also damages normal profiles.", "Try higher nsigma, smaller dilation, smaller max_area, or stricter max_elongation."],
            ["Low spike coverage and low mask fraction", "The setting is too conservative to catch the contaminants.", "Try lower nsigma, larger max_area, or slightly larger dilation."],
            ["Many segments in controls", "The detection threshold or filtering is too permissive.", "Raise nsigma or tighten max_area/max_elongation."],
            ["Best score uses very aggressive settings", "The score may be dominated by spike coverage.", "Compare near-ties and inspect detail rows plus visual reports."],
        ],
        [3000, 3160, 3200],
    )

    doc.add_heading("Recommended Decision Rule", level=1)
    add_callout(
        doc,
        "Choose this.",
        "A parameter set with near-complete spike coverage, low control damage, low control mask fraction, and conservative parameter values among the top-ranked rows.",
    )
    add_para(
        doc,
        "Do not automatically choose the row with the absolute lowest score if a slightly higher-scoring row is visibly safer. "
        "For science-profile work, the best global Photutils setting is the one that removes foreground-object contamination "
        "without teaching the mask to remove galaxy structure.",
    )

    add_footer(doc)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
