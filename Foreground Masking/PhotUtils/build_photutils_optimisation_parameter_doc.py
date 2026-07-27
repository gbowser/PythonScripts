#!/usr/bin/env python3
"""Build a Word explanation of Photutils optimisation parameters."""

from __future__ import annotations

from pathlib import Path
import sys

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

PROJECT_ROOT = Path(__file__).resolve().parent.parent
if str(PROJECT_ROOT) not in sys.path:
    sys.path.append(str(PROJECT_ROOT))

from machine_paths import remove_foreground_folder  # noqa: E402


OUTPUT = remove_foreground_folder("Laptop") / "documentation" / "Photutils Optimisation Parameters Explained.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
BORDER = "A6B4C4"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    tbl_grid = tbl.find(qn("w:tblGrid"))
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        tbl_grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_borders(table, color=BORDER) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        node = borders.find(tag)
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def set_row_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)


def format_cell_text(cell, size_pt: float = 10.0) -> None:
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.15
        for run in paragraph.runs:
            run.font.size = Pt(size_pt)


def add_paragraph(doc: Document, text: str, style: str = "Normal", bold_prefix: str | None = None):
    para = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        run = para.add_run(bold_prefix)
        run.bold = True
        para.add_run(text[len(bold_prefix) :])
    else:
        para.add_run(text)
    return para


def add_callout(doc: Document, label: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_borders(table, color="CAD3DF")
    set_table_widths(table, [9360], indent_dxa=120)
    set_row_cant_split(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(label)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    p.add_run(f" {body}")
    format_cell_text(cell, size_pt=10.5)
    doc.add_paragraph()


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_borders(table)
    set_row_repeat_header(table.rows[0])
    set_row_cant_split(table.rows[0])
    hdr = table.rows[0].cells
    for index, header in enumerate(headers):
        set_cell_shading(hdr[index], LIGHT_BLUE)
        paragraph = hdr[index].paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run(header)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
        run.font.size = Pt(10)
    for row_values in rows:
        row = table.add_row()
        set_row_cant_split(row)
        cells = row.cells
        for index, value in enumerate(row_values):
            paragraph = cells[index].paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.15
            run = paragraph.add_run(value)
            run.font.size = Pt(10)
    set_table_widths(table, widths)
    for row in table.rows:
        for cell in row.cells:
            format_cell_text(cell, size_pt=10)
    doc.add_paragraph()


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_footer(doc: Document) -> None:
    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer.paragraph_format.space_after = Pt(0)
        run = footer.add_run("Photutils optimisation parameter guide")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(90, 90, 90)


def build_document() -> None:
    doc = Document()
    configure_styles(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("Photutils Optimisation Parameters Explained")
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    run = subtitle.add_run("Foreground-object masking for barred-galaxy profile analysis")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(80, 80, 80)

    add_callout(
        doc,
        "Purpose.",
        "The optimisation should find the minimum intervention that removes narrow foreground-object spikes "
        "from the bar-major profile while preserving genuine galaxy structure such as bars, rings, spiral arms, "
        "ansa, shoulders, and nuclear components.",
    )

    doc.add_heading("Why These Parameters Matter", level=1)
    add_paragraph(
        doc,
        "The Photutils masking pipeline works on a residual image: the science image minus a broad Gaussian-smoothed "
        "galaxy model. Compact positive residuals are treated as foreground-candidate sources. The optimiser then "
        "tests parameter combinations and scores them against two kinds of evidence: spike-positive galaxies, where "
        "the mask should cover narrow bar-profile spikes, and no-spike control galaxies, where the mask should not "
        "change a clean profile.",
    )
    add_paragraph(
        doc,
        "The parameters below control three linked decisions: what counts as a detection, which detections are kept, "
        "and how aggressively retained detections are expanded into masks.",
    )

    doc.add_heading("Parameter Summary", level=1)
    add_table(
        doc,
        ["Parameter", "Current or proposed values", "Main role"],
        [
            ["nsigmas", "5.5, 5.0, 4.5, 4.0, 3.5", "Detection threshold in residual-sigma units."],
            ["dilations", "1, 2, 3 pixels", "Expands retained masks to include source wings."],
            ["max-areas", "150, 300, 500 pixels", "Rejects large residual regions that are likely galaxy structure."],
            ["max-elongation", "3, 4, 6", "Rejects stretched detections such as arms, bars, streaks, or diffuse residuals."],
            ["smooth_sigma_pixels", "15 px currently; test 10, 15, 20, 25", "Sets the scale of the smooth galaxy model subtracted before detection."],
            ["npixels", "8 px currently; secondary test values could include 5, 8, 12", "Minimum connected pixels required for a detection."],
            ["central_exclusion", "12 px currently; secondary test values could include 8, 12, 16, 20", "Protects the nucleus and compact nuclear structures from being masked."],
        ],
        [2500, 2600, 4260],
    )

    doc.add_heading("Detailed Parameter Effects", level=1)
    add_table(
        doc,
        ["Parameter", "If made more aggressive", "If made more conservative", "Optimisation signal"],
        [
            [
                "nsigmas",
                "Lower nsigma detects fainter residuals and may catch weaker foreground objects.",
                "Higher nsigma ignores fainter residuals and reduces false positives.",
                "Choose the highest threshold that still gives high spike coverage.",
            ],
            [
                "dilations",
                "Larger dilation captures PSF wings and nearby foreground-object light.",
                "Smaller dilation preserves more local galaxy light around each detection.",
                "Penalise large affected-profile fractions and control-galaxy damage.",
            ],
            [
                "max-areas",
                "Larger max area keeps broader detections, which may catch big stars but risks galaxy features.",
                "Smaller max area rejects broad residual patches but can miss large foreground objects.",
                "Prefer the smallest area cap that does not reduce spike coverage.",
            ],
            [
                "max-elongation",
                "Larger values keep more stretched detections.",
                "Smaller values reject elongated residuals more strongly.",
                "Use controls to check that spiral arms, bars, and streak-like galaxy residuals are not retained.",
            ],
            [
                "smooth_sigma_pixels",
                "A larger smoothing scale leaves more medium-scale residual structure in the detection image.",
                "A smaller smoothing scale can absorb compact sources into the model or create local subtraction artefacts.",
                "Optimise by checking whether residual sources are enhanced without turning galaxy structure into detections.",
            ],
            [
                "npixels",
                "Lower npixels accepts smaller connected detections and increases sensitivity to tiny objects or noise.",
                "Higher npixels rejects small detections and suppresses noise.",
                "Keep fixed initially; tune only if many real small contaminants are missed or noise detections are common.",
            ],
            [
                "central_exclusion",
                "A smaller exclusion allows detections nearer the centre, increasing risk to the nucleus.",
                "A larger exclusion protects more central light but may miss foreground objects projected near the centre.",
                "Tune only after checking central cases visually; avoid optimising this from spike coverage alone.",
            ],
        ],
        [1750, 2530, 2530, 2550],
    )

    doc.add_heading("How To Optimise", level=1)
    add_paragraph(
        doc,
        "The existing optimiser already scores global Photutils parameter sets using spike galaxies and no-spike "
        "controls. It should be extended so smooth_sigma_pixels can vary in the main grid. The secondary parameters "
        "npixels and central_exclusion should usually remain fixed until the primary grid has identified a stable "
        "region of good solutions.",
    )
    add_table(
        doc,
        ["Tier", "Parameters", "Reason"],
        [
            [
                "Primary grid",
                "nsigmas, dilations, max-areas, max-elongation, smooth_sigma_pixels",
                "These have the strongest effect on sensitivity, false positives, and profile damage.",
            ],
            [
                "Secondary checks",
                "npixels, central_exclusion",
                "These are important safeguards but are easier to overfit, especially with a small calibration set.",
            ],
            [
                "Final visual review",
                "Best 3-5 scored parameter sets",
                "The score narrows the field; visual reports catch scientific edge cases that simple metrics miss.",
            ],
        ],
        [1400, 3100, 4860],
    )

    doc.add_heading("Scoring Logic", level=1)
    add_paragraph(
        doc,
        "The score should reward spike coverage and penalise unnecessary masking. A good parameter set should cover "
        "bar-profile spike samples in spike-positive galaxies, while producing low masked-pixel fractions and low "
        "profile-shape changes in no-spike control galaxies.",
    )
    add_callout(
        doc,
        "Practical rule.",
        "Select the most conservative parameter set whose spike coverage is effectively complete and whose control "
        "galaxy damage remains low. If two settings perform similarly, prefer the one with the higher nsigma, smaller "
        "dilation, smaller max area, and lower control-profile damage.",
    )

    doc.add_heading("Spike-Gated Versus Global Photutils", level=1)
    add_paragraph(
        doc,
        "Spike-gated masking can tolerate lower detection thresholds because Photutils detections are only candidates: "
        "a source is applied to the science profile only if it intersects detected bar-major spike samples. Global "
        "Photutils lacks that profile gate, so it should normally use stricter settings. For that reason, an nsigma "
        "value around 3.5 may be acceptable in spike-gated mode but should be treated as aggressive for global masking.",
    )

    doc.add_heading("Recommended Next Step", level=1)
    add_paragraph(
        doc,
        "Run the current optimiser over the existing nsigmas, dilations, max-areas, and max-elongation grid, then add "
        "smooth_sigma_pixels to the grid and compare the ranked summaries. Keep npixels at 8 and central_exclusion at "
        "12 px for the first pass. Only vary them if the visual reports show missed small sources, noise detections, "
        "or central masking problems.",
    )

    add_footer(doc)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
