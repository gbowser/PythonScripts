from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SCRIPT_DIR = Path(__file__).resolve().parents[1]
DOC_DIR = SCRIPT_DIR / "documentation"
SCRIPT_NAME = "Real Galaxy Shoulder Quantification v0.69.py"
OUTPUT_DOCX = DOC_DIR / "Real Galaxy Shoulder Quantification v0.69 Documentation.docx"
EXAMPLE_IMAGE = SCRIPT_DIR / "SRA_Example.png"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(89, 89, 89)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
BLACK = RGBColor(0, 0, 0)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_grid = table._tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        table._tbl.insert(1, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_mar = tc_pr.find(qn("w:tcMar"))
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                tc_pr.append(tc_mar)
            for side, value in [("top", "80"), ("bottom", "80"), ("start", "120"), ("end", "120")]:
                mar = tc_mar.find(qn(f"w:{side}"))
                if mar is None:
                    mar = OxmlElement(f"w:{side}")
                    tc_mar.append(mar)
                mar.set(qn("w:w"), value)
                mar.set(qn("w:type"), "dxa")


def set_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        tbl_header = OxmlElement("w:tblHeader")
        tbl_header.set(qn("w:val"), "true")
        tr_pr.append(tbl_header)


def set_paragraph_border_bottom(paragraph, color: str = "D9E2F3") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Title", 24, BLACK, 0, 4),
        ("Subtitle", 12, MUTED, 0, 18),
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25


def add_text(doc: Document, text: str, style: str | None = None, bold: bool = False) -> None:
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.bold = bold


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        p.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        p.add_run(item)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths_dxa: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths_dxa)
    hdr = table.rows[0].cells
    set_repeat_header(table.rows[0])
    for idx, header in enumerate(headers):
        hdr[idx].text = header
        set_cell_shading(hdr[idx], LIGHT_BLUE)
        for p in hdr[idx].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = DARK_BLUE
    for row_data in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            cells[idx].text = value
            if idx == 0:
                set_cell_shading(cells[idx], LIGHT_GRAY)
                for p in cells[idx].paragraphs:
                    for run in p.runs:
                        run.bold = True
    set_table_geometry(table, widths_dxa)
    doc.add_paragraph()


def add_metadata(doc: Document) -> None:
    rows = [
        ["Script", SCRIPT_NAME],
        ["Location", str(SCRIPT_DIR / SCRIPT_NAME)],
        ["Document date", date.today().isoformat()],
        ["Documentation scope", "High-level overview, data sources, processing workflow, outputs, parameters, and detailed code behavior."],
    ]
    add_table(doc, ["Field", "Value"], rows, [2200, 7160])


def add_intro(doc: Document) -> None:
    doc.add_paragraph("Shoulder Recognition Algorithm Documentation", style="Title")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run("Real Galaxy Shoulder Quantification v0.69").italic = True
    set_paragraph_border_bottom(subtitle)
    add_metadata(doc)

    doc.add_heading("High-Level Overview", level=1)
    add_text(
        doc,
        "This script runs a shoulder-quantification workflow on observed galaxy bar major-axis profiles. "
        "Its main purpose is to identify paired shoulder structures in bar profiles from the Erwin, Debattista, "
        "and Anderson 2023 classification sample, using S4G 3.6 micron FITS images and a locally generated geometry manifest "
        "to construct the required profiles when profile files are not already available."
    )
    add_text(
        doc,
        "For each classified galaxy with the required catalog, geometry, and image data, the code extracts a deprojected "
        "bar-major-axis intensity profile, normalises radius by the deprojected bar radius, smooths the analysed profile "
        "with a low-pass Butterworth filter, computes first and second derivatives and radius-of-curvature structure, "
        "then searches for left and right shoulder candidates inside the bar."
    )
    add_bullets(
        doc,
        [
            "Input data are external to this folder and are mostly hard-coded under the MSc research Dropbox path.",
            "The analysis window is restricted to |x/R_bar| < 1.8, while accepted shoulders must lie inside |x/R_bar| < 1.",
            "By default, a galaxy is counted as having shoulders only when both left and right shoulders are accepted.",
            "The script writes constructed profile caches, diagnostic plots, missing-data reports, measurement tables, and a run summary.",
        ],
    )


def add_data_sources(doc: Document) -> None:
    doc.add_heading("Data Sources and Dependencies", level=1)
    add_text(
        doc,
        "The script combines local repository helper modules with external research data. "
        "The repository supplies geometry and plotting utilities; the measured galaxy data and FITS images live in a hard-coded Dropbox research folder."
    )
    add_table(
        doc,
        ["Source", "Purpose"],
        [
            ["Erwin catalog data", "Reads s4gbars_table.dat for bar and galaxy properties, scrambled_map.txt to map classification IDs to galaxy names, and classifications_pe.txt plus classifications_vd_revised.txt to define the classified sample."],
            ["S4G geometry manifest", "Reads Erwin_s4g_image_downloader/geometry_output/s4g_image_geometry_manifest.csv for image centres, disk PA, inclination, bar PA, bar semi-major axis, pixel scale, CRPIX fallback values, and image path metadata."],
            ["S4G FITS images", "Loads 3.6 micron images from each manifest image_path or from D:\\Dropbox\\Public Documents\\UCLAN\\MSc Research\\Erwin\\s4g_images_36um\\<galaxy>.phot.1.fits."],
            ["Helper modules", "Imports angle_utils from Erwin_barprofiles_paper_GB_working_copy and plot_s4g_isophote_axes from Erwin_s4g_image_downloader for PA rectification, deprojection, and profile extraction utilities."],
            ["Scientific Python stack", "Uses numpy, pandas, scipy.signal, astropy.io.fits, scikit-image profile_line, matplotlib, csv, math, os, sys, and warnings."],
        ],
        [2100, 7260],
    )
    doc.add_heading("Important hard-coded paths", level=2)
    add_table(
        doc,
        ["Name", "Value / meaning"],
        [
            ["research_folder", r"D:\Dropbox\Public Documents\UCLAN\MSc Research"],
            ["erwin_repo", r"<research_folder>\Erwin\perwin-barprofiles_paper-a7cd6f5"],
            ["erwin_data_folder", r"<erwin_repo>\data"],
            ["manifest_file", r"<repo>\Erwin_s4g_image_downloader\geometry_output\s4g_image_geometry_manifest.csv"],
            ["image_folder", r"<research_folder>\Erwin\s4g_images_36um"],
            ["output_folder", r"<research_folder>\Shoulder_Recognition_Erwin"],
            ["plots_folder", r"<output_folder>\plots"],
            ["profiles_folder", r"<output_folder>\profiles"],
        ],
        [2400, 6960],
    )


def add_workflow(doc: Document) -> None:
    doc.add_heading("Processing Workflow", level=1)
    add_numbered(
        doc,
        [
            "Initialise paths, create output, plots, and profiles folders, and add local helper-module directories to sys.path.",
            "Read the Erwin S4G bar table, the scrambled-name map, the two classification files, and the S4G image geometry manifest.",
            "Build a sorted classified-galaxy list by combining Peter Erwin and Victor Debattista revised classifications, excluding unknown classifications marked with '?'.",
            "For each classified galaxy, require a catalog row and manifest row, then construct a bar-major-axis profile from the FITS image.",
            "Write a missing-data text report and CSV for galaxies that cannot be analysed.",
            "Loop over all galaxies with usable profiles, interpolate NaNs, normalise the analysed profile, smooth it, calculate derivatives and radius of curvature, and search for left and right shoulder candidates.",
            "Reject shoulders that are too close to the centre, outside the bar, too thin, too steep, too noisy, one-sided when paired shoulders are required, or overlapping.",
            "Save one diagnostic PNG per analysed galaxy and write the final shoulder measurements in NPY and CSV formats plus a plain-text run summary.",
        ],
    )

    doc.add_heading("Profile construction", level=2)
    add_text(
        doc,
        "construct_profile_from_manifest uses the manifest geometry to locate the galaxy centre and bar orientation. "
        "It loads the FITS image, falls back to CRPIX coordinates if the manifest centre is outside the image, chooses a profile radius based on the bar semi-major axis and pixel scale, and extracts a profile along the bar major axis with a profile width of 3 pixels."
    )
    add_text(
        doc,
        "The extracted radii are converted from pixels to arcseconds and deprojected using the disk PA and inclination. "
        "The bar semi-major axis is also deprojected so later analysis can express radius as x/R_bar. "
        "A minor-axis profile is extracted only to estimate robust plotting limits from finite positive intensities."
    )

    doc.add_heading("Shoulder detection", level=2)
    add_text(
        doc,
        "Within each galaxy, the full extracted profile is kept for plotting, but the analysis uses only |x/R_bar| < 1.8. "
        "The analysed raw intensity profile is interpolated where NaNs occur, normalised to 0..1, and smoothed with a second-order low-pass Butterworth filter. "
        "The cutoff is calculated from the number of profile samples and profile span rather than fixed globally."
    )
    add_text(
        doc,
        "The algorithm then calculates the first derivative, second derivative, and radius of curvature. "
        "Candidate shoulders are found from derivative extrema near each side of the bar: left-side candidates use first-derivative minima and right-side candidates use first-derivative maxima, ordered from the galaxy centre outward. "
        "Nearest radius-of-curvature minima and maxima mark the clavicle and shoulder boundaries."
    )


def add_parameters_outputs(doc: Document) -> None:
    doc.add_heading("Control Parameters", level=1)
    add_table(
        doc,
        ["Parameter", "Default", "Role"],
        [
            ["allow_1_sided_shoulders", "False", "Requires a left and right shoulder pair for a positive shoulder system."],
            ["profile_width", "3", "Pixel width used when extracting major/minor-axis profiles from FITS images."],
            ["thin_be_gone", "0.05", "Rejects shoulders whose inner-to-outer width is too small in x/R_bar units."],
            ["peaks_max", "12", "Rejects profiles with too many first-derivative extrema inside the bar."],
            ["too_close", "0.2", "Rejects candidate shoulder clavicles closer than 0.2 R_bar to the centre."],
            ["x_cutoff_vs_bre", "1.8", "Limits the algorithmic analysis window to |x/R_bar| < 1.8."],
            ["slope_cutoff", "0.35", "Candidate shoulder derivative extrema must have absolute slope below this value."],
            ["bar_extent", "1", "Defines the bar boundary in normalised radius after x is divided by R_bar."],
        ],
        [2300, 1200, 5860],
    )

    doc.add_heading("Outputs", level=1)
    add_table(
        doc,
        ["Output", "Description"],
        [
            ["profiles/<galaxy>_bar-major-axis_profile.dat", "Constructed cache of deprojected bar-major-axis radius in arcsec and raw S4G 3.6 micron intensity. Three commented header rows are written."],
            ["plots/<galaxy>.png", "Diagnostic plot showing the semilog major-axis profile, smoothed curve, derivative overlays, residuals, bar boundary markers, and shoulder annotations."],
            ["missing_data_components.txt", "Plain-text report explaining which required components are missing or unusable for each skipped galaxy."],
            ["missing_data_components.csv", "CSV version of the missing-data report with galaxy and reason columns."],
            ["shoulder_measurements.npy", "Structured NumPy array containing final shoulder measurements."],
            ["shoulder_measurements.csv", "CSV version of the structured shoulder table."],
            ["run_summary.txt", "Run-level summary of classified galaxies, usable galaxies, skipped galaxies, shoulder systems found, and output locations."],
        ],
        [3100, 6260],
    )

    doc.add_heading("Measurement columns", level=2)
    add_table(
        doc,
        ["Column", "Meaning"],
        [
            ["galaxy", "Galaxy name."],
            ["clav_left / clav_right", "Accepted left/right shoulder clavicle location in x/R_bar units, or NaN if rejected."],
            ["left_inner / right_inner", "Inner boundary of the accepted shoulder, usually the inner radius-of-curvature minimum around the clavicle."],
            ["left_outer / right_outer", "Outer boundary of the accepted shoulder, selected from the next radius-of-curvature structure outward."],
            ["left_slope / right_slope", "First-derivative value at the accepted clavicle candidate."],
            ["left_clav_inner / left_clav_outer", "Radius-of-curvature minima bracketing the left clavicle."],
            ["right_clav_inner / right_clav_outer", "Radius-of-curvature minima bracketing the right clavicle."],
        ],
        [2600, 6760],
    )


def add_detailed_code(doc: Document) -> None:
    doc.add_heading("Detailed Code Documentation", level=1)

    doc.add_heading("Module setup", level=2)
    add_text(
        doc,
        "The script starts by forcing matplotlib to use the Agg backend, which allows figures to be saved without an interactive display. "
        "It then builds PROJECT_ROOT from the script location and appends two sibling folders to sys.path so angle_utils and plot_s4g_isophote_axes can be imported."
    )
    add_bullets(
        doc,
        [
            "BARPROFILES_DIR: expected to contain angle_utils.py.",
            "S4G_PLOTTER_DIR: expected to contain plot_s4g_isophote_axes.py.",
            "The local function required_geometry and profile_at_pa are defined but the active profile-construction path uses corresponding helpers from plot_s4g_isophote_axes.",
        ],
    )

    doc.add_heading("Utility functions", level=2)
    add_table(
        doc,
        ["Function", "Behavior"],
        [
            ["find_nearest(array, value)", "Converts the input to a NumPy array and returns the element whose absolute difference from value is smallest. Used to connect derivative extrema to nearest curvature features and profile bins."],
            ["read_s4g_table(filename)", "Reads whitespace-delimited s4gbars_table.dat with comment lines ignored, applies a fixed 37-column schema, and returns a pandas DataFrame."],
            ["read_descramble_map(filename)", "Reads scrambled_map.txt, skipping blank/comment lines, and maps integer scrambled IDs to galaxy names."],
            ["read_classified_galaxies(filename, descramble_map)", "Reads a classification file, keeps rows with a non-'?' classification, resolves the scrambled ID to a galaxy name, and returns the list of classified galaxies."],
            ["read_manifest(filename)", "Reads the geometry manifest CSV into a dictionary keyed by galaxy name."],
            ["parse_float(value)", "Safely converts manifest fields to finite floats; blank, invalid, and non-finite values become None."],
            ["pa_endpoint(pa_deg, radius)", "Converts a position angle and radius into an x/y endpoint offset. Defined for profile extraction geometry."],
            ["profile_at_pa(data, xc, yc, pa_deg, radius_pix, width)", "Extracts a line profile through a 2D image using skimage.measure.profile_line, averaging over the requested linewidth and returning pixel radii plus intensity values."],
            ["required_geometry(row)", "Parses essential geometry fields from a manifest row, rectifies disk and bar PAs to 0..180 degrees, and returns None if required values are missing."],
            ["construct_profile_from_manifest(galaxy, manifest_row)", "Builds the profile file, deprojected bar radius, and plot scale for one galaxy, or returns a reason string if the galaxy cannot be processed."],
            ["write_missing_data_report(missing_rows, filename)", "Writes a human-readable missing-data report that lists required components and per-galaxy failure reasons."],
        ],
        [2800, 6560],
    )

    doc.add_heading("Main sample assembly", level=2)
    add_text(
        doc,
        "After function definitions, the script immediately executes the analysis at module level. "
        "It reads the catalog, descramble map, manifest, and both classification files, then constructs classified_galaxies as the sorted unique union of the Peter Erwin and Victor Debattista revised classified samples."
    )
    add_text(
        doc,
        "For each galaxy, the code verifies an S4G catalog row, a manifest row, a usable FITS image/profile, and a positive finite deprojected bar radius. "
        "Successful galaxies are stored in available_rows with galaxy name, bar_radius, profile_file, and profile_scale; failures are stored in missing_data with an explanatory reason."
    )

    doc.add_heading("Analysis loop details", level=2)
    add_bullets(
        doc,
        [
            "Loads each generated profile with np.loadtxt, skipping the three commented header rows.",
            "Normalises radius by the deprojected bar radius and keeps raw intensity for semilog plotting.",
            "Repairs leading NaNs by copying the first finite value, then linearly interpolates remaining NaNs with pandas.",
            "Applies the analysis mask |x/R_bar| < 1.8 and normalises the analysed intensity values to 0..1.",
            "Computes a Butterworth cutoff from sample count and x-span, then smooths the normalised profile with signal.filtfilt.",
            "Calculates first derivative, raw first derivative, second derivative, raw second derivative, and radius of curvature.",
            "Flags a profile as too noisy if the count of first-derivative extrema inside |x| <= 1 exceeds peaks_max.",
            "Searches left and right candidates separately, starting from candidate extrema closest to x = 0 and moving outward.",
            "Accepts only shoulders inside the bar, away from the centre, below the slope cutoff, wider than thin_be_gone, and not in a too-noisy profile.",
            "Rejects final paired shoulders if left and right shoulder regions overlap or are within one bin of the centre.",
        ],
    )

    doc.add_heading("Plot contents", level=2)
    add_text(
        doc,
        "Each diagnostic PNG has a two-panel layout. The upper panel shows the full raw major-axis intensity profile on a logarithmic y-axis, the smoothed analysed profile over the selected analysis window, derivative overlays on a twin y-axis, bar-boundary markers at x/R_bar = +/-1, and red/purple shoulder markers when accepted. "
        "The lower panel shows the residual between the smoothed normalised profile and the normalised raw analysed profile."
    )
    if EXAMPLE_IMAGE.exists():
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(EXAMPLE_IMAGE), width=Inches(5.9))
        caption = doc.add_paragraph("Example diagnostic output image included with the script folder.", style=None)
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in caption.runs:
            run.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = MUTED


def add_operational_notes(doc: Document) -> None:
    doc.add_heading("Operational Notes and Maintenance Risks", level=1)
    add_bullets(
        doc,
        [
            "The script is not command-line parameterised; changing data locations or thresholds requires editing global variables in the source file.",
            "The code executes at import time. Importing it from another module would run the full analysis unless the main workflow is later moved behind an if __name__ == '__main__' guard.",
            "Several functions duplicate concepts available in plot_s4g_isophote_axes, and construct_profile_from_manifest currently calls the imported helper versions for geometry/profile work.",
            "The profile interpolation step includes a note saying it needs discussion, so NaN handling should be treated as an explicit scientific assumption.",
            "The shoulder decision depends strongly on smoothing, extrema counts, and slope thresholds. These are documented above because they are the most important reproducibility controls.",
            "The output folder is outside the repository, so generated scientific outputs are not automatically versioned with the code unless copied or tracked separately.",
        ],
    )

    doc.add_heading("Suggested future improvements", level=2)
    add_bullets(
        doc,
        [
            "Add argparse options or a small configuration file for data paths, output path, and shoulder thresholds.",
            "Move the executable workflow into a main() function guarded by if __name__ == '__main__'.",
            "Write a compact run manifest alongside each result CSV that records parameter values, script version, and input file timestamps.",
            "Consider replacing print-only progress with logging so rejected candidates can be persisted per galaxy.",
            "Add unit tests around parse_float, manifest validation, profile construction failure modes, and shoulder candidate rejection logic.",
        ],
    )


def add_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = "Shoulder Quantification Code Documentation"
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = MUTED
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.text = "Generated documentation"
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = MUTED


def build() -> Path:
    DOC_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_document(doc)
    add_footer(doc)
    add_intro(doc)
    add_data_sources(doc)
    add_workflow(doc)
    add_parameters_outputs(doc)
    add_detailed_code(doc)
    add_operational_notes(doc)
    doc.save(OUTPUT_DOCX)
    return OUTPUT_DOCX


if __name__ == "__main__":
    print(build())
