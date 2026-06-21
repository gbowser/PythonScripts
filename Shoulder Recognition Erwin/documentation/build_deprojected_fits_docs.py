from __future__ import annotations

import importlib.util
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


SCRIPT_DIR = Path(__file__).resolve().parents[1]
DOC_DIR = SCRIPT_DIR / "documentation"
SCRIPT_NAME = "Create Deprojected Galaxy FITS.py"
OUTPUT_DOCX = DOC_DIR / "Create Deprojected Galaxy FITS Documentation.docx"
CONTROL_IMAGE = SCRIPT_DIR / "NGC0578_deprojection_example.png"

# Reuse the previous documentation's exact compact-reference style and table
# implementation so both manuals form a visually consistent set.
template_path = DOC_DIR / "build_shoulder_quantification_docs.py"
spec = importlib.util.spec_from_file_location("shoulder_docs_template", template_path)
template = importlib.util.module_from_spec(spec)
assert spec.loader is not None
spec.loader.exec_module(template)

style_document = template.style_document
set_paragraph_border_bottom = template.set_paragraph_border_bottom
add_text = template.add_text
add_bullets = template.add_bullets
add_numbered = template.add_numbered
add_table = template.add_table
MUTED = template.MUTED


def add_metadata(doc: Document) -> None:
    add_table(
        doc,
        ["Field", "Value"],
        [
            ["Script", SCRIPT_NAME],
            ["Location", str(SCRIPT_DIR / SCRIPT_NAME)],
            ["Document date", date.today().isoformat()],
            ["Documentation scope", "Purpose, inputs, affine geometry, resampling, command-line operation, FITS output, validation, and maintenance notes."],
        ],
        [2200, 7160],
    )


def add_intro(doc: Document) -> None:
    doc.add_paragraph("Deprojected Galaxy FITS Exporter Documentation", style="Title")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run(SCRIPT_NAME).italic = True
    set_paragraph_border_bottom(subtitle)
    add_metadata(doc)

    doc.add_heading("High-Level Overview", level=1)
    add_text(
        doc,
        "This program converts S4G 3.6 micron galaxy images into a common analysis geometry. "
        "Each output image is centred on the catalogue galaxy centre, deprojected to a face-on disk, "
        "and rotated so the deprojected stellar bar lies along the horizontal x-axis. One primary-image "
        "FITS file is written for each selected galaxy."
    )
    add_text(
        doc,
        "The default sample is deliberately identical to the classified sample assembled by Real Galaxy "
        "Shoulder Quantification v0.69.py: the sorted unique union of the Peter Erwin and revised Victor "
        "Debattista classification files, excluding entries marked with '?'. The geometry manifest supplies "
        "the fitted centre, disk position angle, inclination, bar position angle, pixel scale, and image path."
    )
    add_bullets(
        doc,
        [
            "The disk is deprojected by stretching only its projected minor-axis coordinate by sec(i).",
            "Deprojection and bar rotation are combined into one affine transformation, avoiding a second interpolation pass.",
            "The transformed centre is placed on the exact central pixel. Surface-brightness values are interpolated without a flux-area correction, and the celestial WCS is replaced because the result is a constructed face-on plane.",
        ],
    )


def add_inputs(doc: Document) -> None:
    doc.add_heading("Inputs and Dependencies", level=1)
    add_table(
        doc,
        ["Input", "Purpose"],
        [
            ["Geometry manifest", "Erwin_s4g_image_downloader/geometry_output/s4g_image_geometry_manifest.csv. Supplies image paths, centres, CRPIX fallbacks, disk PA, inclination, bar PA, and x/y pixel scales."],
            ["Classification files", "classifications_pe.txt and classifications_vd_revised.txt define the default classified sample; scrambled_map.txt converts integer classification IDs to galaxy names."],
            ["S4G FITS images", "The manifest image_path is tried first. If it is missing, the program uses <image-dir>/<galaxy>.phot.1.fits."],
            ["Python packages", "NumPy for matrix/array operations, Astropy for FITS I/O, and scipy.ndimage.affine_transform for image resampling."],
        ],
        [2200, 7160],
    )

    doc.add_heading("Default paths", level=2)
    add_table(
        doc,
        ["Setting", "Default"],
        [
            ["RESEARCH_ROOT", r"D:\Dropbox\Public Documents\UCLAN\MSc Research"],
            ["ERWIN_DATA", r"<RESEARCH_ROOT>\Erwin\perwin-barprofiles_paper-a7cd6f5\data"],
            ["IMAGE_DIR", r"<RESEARCH_ROOT>\Erwin\s4g_images_36um"],
            ["MANIFEST", r"<repository>\Erwin_s4g_image_downloader\geometry_output\s4g_image_geometry_manifest.csv"],
            ["OUTPUT_DIR", r"<RESEARCH_ROOT>\Shoulder_Recognition_Erwin\Deprojected_Images"],
        ],
        [2400, 6960],
    )

    doc.add_page_break()
    doc.add_heading("Required geometry fields", level=2)
    add_table(
        doc,
        ["Manifest column", "Use"],
        [
            ["center_x_pix, center_y_pix", "Primary fitted centre in FITS/IRAF one-based pixel coordinates."],
            ["crpix1, crpix2", "Fallback centre if the fitted centre lies outside the image."],
            ["disk_pa_deg", "Projected disk major-axis position angle, rectified modulo 180 degrees."],
            ["inclination_deg", "Disk inclination; must satisfy 0 <= i < 90 degrees."],
            ["bar_pa_deg", "Observed bar position angle, rectified modulo 180 degrees."],
            ["pixel_scale_arcsec_x/y", "Absolute output linear scale. Missing/zero values fall back to 0.75 arcsec per pixel."],
        ],
        [2900, 6460],
    )


def add_geometry(doc: Document) -> None:
    doc.add_heading("Geometric Transformation", level=1)
    add_text(
        doc,
        "All geometry is calculated in image coordinates (x = FITS column, y = FITS row) relative to the selected centre. "
        "The position-angle convention matches plot_s4g_isophote_axes.py: an axis at PA theta has direction "
        "(-sin theta, cos theta). Because a bar is an undirected axis, alignment with either +x or -x is equivalent."
    )

    doc.add_heading("Disk deprojection", level=2)
    add_text(
        doc,
        "Let d be the unit vector along the projected disk major axis and m the perpendicular projected minor-axis vector. "
        "The face-on deprojection matrix is D = d d^T + sec(i) m m^T. It leaves the disk-major coordinate unchanged and "
        "expands the foreshortened disk-minor coordinate by 1/cos(i)."
    )

    doc.add_heading("Bar alignment", level=2)
    add_text(
        doc,
        "The observed bar-axis vector is transformed by D. Its resulting face-on angle is measured with atan2, and a rotation "
        "matrix R maps that vector onto the x-axis. The complete forward transformation is A = R D. The resampler uses A^-1 "
        "because scipy.ndimage.affine_transform maps each output coordinate back into the input image."
    )

    doc.add_heading("Output canvas and centring", level=2)
    add_text(
        doc,
        "All four input-image corners are transformed about the galaxy centre. The largest absolute transformed x and y extents "
        "define a symmetric odd-sized canvas. Consequently CRPIX1 and CRPIX2 always identify the array's exact central pixel. "
        "This preserves the full transformed rectangular footprint but can produce large blank regions, especially for inclined "
        "galaxies or input mosaics with irregular valid-data footprints."
    )

    doc.add_heading("NaN-safe interpolation", level=2)
    add_text(
        doc,
        "Cubic spline prefiltering cannot be applied directly to arrays containing NaNs because a small number of NaNs can "
        "contaminate every spline coefficient. The program therefore replaces invalid input pixels with zero, transforms the "
        "filled image, transforms a separate validity mask with linear interpolation, divides by that support image where support "
        "exceeds 0.001, and restores unsupported output pixels to NaN. This is a normalised-convolution treatment of mosaic edges."
    )


def add_workflow_cli(doc: Document) -> None:
    doc.add_heading("Processing Workflow", level=1)
    add_numbered(
        doc,
        [
            "Parse command-line options and read the geometry manifest into a dictionary keyed by galaxy name.",
            "Select one or more explicitly named galaxies, every manifest row, or the default classified Erwin sample.",
            "Locate and load the primary FITS image, squeeze singleton dimensions, and require a two-dimensional image.",
            "Validate the centre, position angles, inclination, and pixel scales; use CRPIX only when the catalogue centre is outside the image.",
            "Construct the combined disk-deprojection and bar-alignment affine matrix.",
            "Calculate a symmetric output canvas and resample image values plus the validity support mask.",
            "Replace the obsolete celestial WCS with a linear, galaxy-centred x/y offset system and add provenance keywords.",
            "Write a float32 primary FITS image with checksum and continue to the next galaxy if an individual object fails.",
            "Print a final success/failure summary and return exit status 1 if any selected object failed.",
        ],
    )

    doc.add_page_break()
    doc.add_heading("Command-line options", level=1)
    add_table(
        doc,
        ["Option", "Default / behavior"],
        [
            ["--manifest PATH", "Override the geometry-manifest CSV."],
            ["--image-dir PATH", "Override the fallback S4G image directory."],
            ["--erwin-data PATH", "Override the directory containing classification and descramble files."],
            ["--output-dir PATH", "Override the output directory."],
            ["--order N", "Spline interpolation order 0 through 5; default 3."],
            ["--no-overwrite", "Skip a FITS file when the target already exists. The normal default is to overwrite."],
            ["--all-manifest", "Process every galaxy in the manifest instead of the classified sample."],
            ["--limit N", "Process only the first N selected names; N must be at least 1."],
            ["--galaxy NAME", "Process a named galaxy. Repeat the option to process a custom set; duplicate names are removed while preserving order."],
        ],
        [2450, 6910],
    )

    doc.add_heading("Example commands", level=2)
    add_bullets(
        doc,
        [
            'Full sample: python "Shoulder Recognition Erwin/Create Deprojected Galaxy FITS.py"',
            'Five-object test: append --limit 5.',
            'One galaxy: append --galaxy NGC0578.',
            'Custom set: append --galaxy NGC3504 --galaxy NGC1672.',
            'Keep existing files: append --no-overwrite.',
        ],
    )


def add_outputs(doc: Document) -> None:
    doc.add_heading("FITS Output", level=1)
    add_text(
        doc,
        "Each successful galaxy produces <galaxy>_deprojected_bar_aligned.fits. The primary data array is float32; invalid or "
        "unsupported pixels are NaN. A FITS CHECKSUM and DATASUM are written. The image values retain the input surface-brightness "
        "scale apart from interpolation; they should not be interpreted as integrated flux per transformed pixel."
    )

    doc.add_heading("Output header keywords", level=2)
    add_table(
        doc,
        ["Keyword", "Meaning"],
        [
            ["CRPIX1, CRPIX2", "One-based x/y coordinates of the centred galaxy; equal to the exact middle pixel."],
            ["CRVAL1, CRVAL2", "Zero offset at the galaxy centre."],
            ["CTYPE1, CTYPE2", "X---LINEAR and Y---LINEAR for deprojected bar-major and bar-minor coordinates."],
            ["CUNIT1, CUNIT2", "arcsec."],
            ["CDELT1, CDELT2", "Absolute manifest pixel scales in arcsec per pixel."],
            ["GALAXY", "Galaxy name used for processing and output naming."],
            ["DISKPA", "Input projected disk position angle in degrees."],
            ["BARPA", "Input observed bar position angle in degrees."],
            ["INCL", "Input disk inclination in degrees."],
            ["DEPROJ", "True when the disk deprojection was applied."],
            ["BARALIGN", "True when the deprojected bar was aligned with the x-axis."],
            ["INTERP", "scipy spline interpolation order."],
        ],
        [2300, 7060],
    )
    add_text(
        doc,
        "Existing CRPIX, CRVAL, CTYPE, CUNIT, CDELT, CROTA, CD, PC, PV, PS, WCSAXES, LONPOLE, and LATPOLE keywords are removed "
        "before the linear output coordinate system is written. Other non-WCS metadata from the input primary header is retained."
    )


def add_functions(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("Function Reference", level=1)
    add_table(
        doc,
        ["Function", "Responsibility"],
        [
            ["parse_args", "Defines paths, interpolation controls, sample selectors, and overwrite behavior."],
            ["finite_float", "Converts a manifest value to a finite float or returns None."],
            ["read_manifest", "Returns manifest rows keyed by name."],
            ["read_descramble_map", "Maps scrambled integer IDs to galaxy names."],
            ["read_classified_names", "Builds the sorted unique classified sample from both classification files."],
            ["geometry", "Validates and normalises centre, PA, inclination, and pixel-scale values."],
            ["image_transform", "Constructs the forward observed-to-face-on, bar-aligned 2 x 2 matrix."],
            ["output_shape_and_center", "Transforms image corners and returns an odd symmetric output shape plus centre."],
            ["resample", "Converts matrix conventions, performs NaN-safe image/support resampling, and returns image plus centre."],
            ["output_header", "Removes obsolete WCS keys and writes the linear centred coordinate metadata and provenance."],
            ["process_galaxy", "Loads one source FITS image, applies geometry/resampling, and writes its output FITS file."],
            ["main", "Selects the batch, creates the output folder, isolates per-galaxy failures, prints progress, and sets exit status."],
        ],
        [3300, 6060],
    )


def add_validation(doc: Document) -> None:
    doc.add_heading("Validation Performed", level=1)
    add_bullets(
        doc,
        [
            "Synthetic matrix tests verified that transformed catalogue bar vectors have a zero y-component to floating-point precision.",
            "A real FITS round-trip verified two-dimensional data, central CRPIX values, deprojection/alignment flags, and FITS checksum handling.",
            "NaN propagation was detected during the first five-image visual test and corrected with support-normalised interpolation.",
            "NGC0578 was used as the worked image example: inclination 59.50 degrees produces a 1.9703 disk-minor stretch, while its PA 93 degree bar is placed on the output x-axis.",
            "The final classified-sample run produced all 182 expected FITS files (4.37 GiB total) with no missing expected names.",
        ],
    )
    if CONTROL_IMAGE.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(CONTROL_IMAGE), width=Inches(6.25))
        caption = doc.add_paragraph(
            "NGC0578 example: the inclined input is deprojected to a face-on plane and the catalogue bar axis is horizontal in the output.",
        )
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in caption.runs:
            run.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = MUTED


def add_notes(doc: Document) -> None:
    doc.add_heading("Operational Notes and Limitations", level=1)
    add_bullets(
        doc,
        [
            "High-inclination objects receive a large sec(i) stretch; their deprojections are intrinsically more uncertain and create larger files.",
            "The full transformed input rectangle is retained symmetrically about the centre. Irregular S4G mosaic footprints therefore leave substantial NaN regions and can make output arrays much larger than the useful galaxy area.",
            "The linear FITS axes describe offsets in the constructed deprojected plane. They are not celestial RA/Dec WCS axes.",
            "The output x-axis follows the supplied bar PA even when the bar is weak or difficult to see by eye; alignment quality therefore depends on the catalogue geometry.",
            "Cubic interpolation is suitable for smooth surface-brightness imagery but changes individual pixel values and introduces correlated sampling. Use order 0 or 1 when nearest-neighbour or linear behavior is scientifically preferable.",
            "The batch continues after individual failures and returns a non-zero process status. Console output should be retained when a permanent failure audit is required.",
            "Dropbox or image-viewer file locks can temporarily prevent overwrite. Close applications using the FITS file and rerun the affected galaxy with --galaxy NAME.",
        ],
    )

    doc.add_heading("Reproducibility checklist", level=2)
    add_bullets(
        doc,
        [
            "Record the script version, manifest version, selected sample, interpolation order, and output path.",
            "Confirm every output has finite data, a valid CHECKSUM, central CRPIX values, and DEPROJ/BARALIGN set to True.",
            "Inspect representative low-, moderate-, and high-inclination galaxies visually.",
            "Treat changes to PA convention, centre coordinates, NaN support threshold, or affine matrix order as scientifically material changes.",
        ],
    )


def add_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = "Deprojected Galaxy FITS Exporter Documentation"
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = MUTED
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.text = "Generated documentation"
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = MUTED


def build() -> Path:
    DOC_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_document(doc)
    add_footer(doc)
    add_intro(doc)
    add_inputs(doc)
    add_geometry(doc)
    add_workflow_cli(doc)
    add_outputs(doc)
    add_functions(doc)
    add_validation(doc)
    add_notes(doc)

    # Keep section headings with following content and prevent stranded titles.
    for style_name in ("Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3"):
        doc.styles[style_name].paragraph_format.keep_with_next = True
    doc.core_properties.title = "Create Deprojected Galaxy FITS Documentation"
    doc.core_properties.subject = "Technical documentation for centred, deprojected, bar-aligned S4G FITS exports"
    doc.core_properties.author = ""
    doc.save(OUTPUT_DOCX)
    return OUTPUT_DOCX


if __name__ == "__main__":
    print(build())
