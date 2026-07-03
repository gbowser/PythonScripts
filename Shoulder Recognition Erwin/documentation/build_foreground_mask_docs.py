from __future__ import annotations

import importlib.util
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


SCRIPT_DIR = Path(__file__).resolve().parents[1]
DOC_DIR = SCRIPT_DIR / "documentation"
SCRIPT_NAME = "foreground_mask_photutils.py"
OUTPUT_DOCX = DOC_DIR / "Foreground Mask Photutils Documentation.docx"
OUTPUT_PDF = DOC_DIR / "Foreground Mask Photutils Documentation.pdf"

# Reuse the existing shoulder-quantification documentation style so this manual
# sits visually beside the other project documentation.
template_path = DOC_DIR / "build_shoulder_quantification_docs.py"
spec = importlib.util.spec_from_file_location("shoulder_docs_template", template_path)
template = importlib.util.module_from_spec(spec)
assert spec.loader is not None
spec.loader.exec_module(template)

style_document = template.style_document
set_paragraph_border_bottom = template.set_paragraph_border_bottom
add_text = template.add_text
add_bullets = template.add_bullets
add_numbered = template.add_numbered
add_table = template.add_table
MUTED = template.MUTED


def add_metadata(doc: Document) -> None:
    add_table(
        doc,
        ["Field", "Value"],
        [
            ["Script", SCRIPT_NAME],
            ["Location", str(SCRIPT_DIR / SCRIPT_NAME)],
            ["Document date", date.today().isoformat()],
            [
                "Documentation scope",
                "Purpose, scientific masking philosophy, algorithm, code flow, command-line use, parameter tuning, outputs, and limitations.",
            ],
        ],
        [2200, 7160],
    )


def add_intro(doc: Document) -> None:
    doc.add_paragraph("Foreground Mask Photutils Documentation", style="Title")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run(SCRIPT_NAME).italic = True
    set_paragraph_border_bottom(subtitle)
    add_metadata(doc)

    doc.add_heading("High-Level Overview", level=1)
    add_text(
        doc,
        "This script builds contaminant masks for astronomical FITS images before bar-profile or shoulder-structure analysis. "
        "It is designed for compact foreground/background objects, especially stars and unresolved sources, while preserving "
        "the galaxy's own large-scale bar, ring, ansa, shoulder, spiral-arm, and nuclear surface-brightness structure.",
    )
    add_text(
        doc,
        "The primary scientific output is a binary FITS mask: 0 means usable science pixel, and 1 means masked contaminant. "
        "The script can also produce diagnostic PNGs and an optional nearest-neighbour cleaned image, but cleaned imagery is "
        "only a visual convenience. Later bar measurements should exclude masked pixels rather than replacing them.",
    )
    add_bullets(
        doc,
        [
            "The detection image is not the raw galaxy image. Compact sources are detected in a residual image after subtracting a broad, NaN-safe smoothed galaxy model.",
            "The code can run on one FITS image or on every FITS image matching a folder glob pattern.",
            "The method is general and does not depend on S4G pipeline masks, so it can be reused for Euclid images after appropriate parameter tuning.",
        ],
    )


def add_scientific_principles(doc: Document) -> None:
    doc.add_heading("Scientific Masking Philosophy", level=1)
    add_text(
        doc,
        "Foreground-object handling is treated as a measurement mask, not as cosmetic star removal. This distinction matters: "
        "a cleaned image can look plausible while silently altering the very surface-brightness gradients used to identify "
        "bar shoulders, ansae, and profile breaks. The safe workflow is therefore to carry a mask through the analysis and "
        "exclude those pixels from profile fitting or statistics.",
    )
    add_bullets(
        doc,
        [
            "Mask compact positive residuals rather than the bright parts of the galaxy itself.",
            "Avoid aggressive settings near the galaxy centre unless the nucleus is known to be contaminated.",
            "Inspect residual and overlay diagnostics for several galaxies before launching a large batch.",
            "If the science image is deprojected or rotated later, transform the mask with nearest-neighbour interpolation so mask values remain discrete.",
            "For Euclid workflows, prefer detecting masks in the original image geometry before deprojection, because deprojection changes the effective PSF shape.",
        ],
    )


def add_algorithm(doc: Document) -> None:
    doc.add_heading("Masking Algorithm", level=1)
    add_numbered(
        doc,
        [
            "Load the selected FITS image and preserve a copy of the original header for the output mask.",
            "Convert the image data to floating point and require a two-dimensional image.",
            "Identify finite pixels so NaN mosaic edges or invalid regions are not treated as real low signal.",
            "Create a smooth galaxy model by Gaussian-filtering the image and a matching finite-pixel weight image, then dividing the two. This normalised smoothing avoids NaNs leaking into the model.",
            "Subtract the smooth model from the original image to form a residual image R = I - G_smooth.",
            "Estimate the residual scatter using 1.4826 times the median absolute deviation, with a standard-deviation fallback only if the MAD is unusable.",
            "Detect positive compact residuals with photutils.segmentation.detect_sources above median(R) + nsigma * sigma.",
            "Optionally deblend overlapping detections with photutils.segmentation.deblend_sources.",
            "Filter detections by area, elongation, peak residual significance, and optional distance from the galaxy centre.",
            "Convert the filtered segmentation image to a binary contaminant mask.",
            "Dilate the mask with a circular footprint so the mask covers PSF wings around compact objects.",
            "Write the mask as integer FITS data and write diagnostic PNGs for visual quality control.",
        ],
    )

    doc.add_heading("Why residual detection?", level=2)
    add_text(
        doc,
        "Direct thresholding of the science image would preferentially detect the bright galaxy body, bar, nucleus, rings, and spiral arms. "
        "Subtracting a broad smooth model removes most large-scale galaxy light while leaving compact sources as positive residuals. "
        "This makes the segmentation step much less likely to confuse real bar structure with contamination.",
    )

    doc.add_heading("NaN-safe smooth model", level=2)
    add_text(
        doc,
        "The smooth model is computed as gaussian_filter(data_filled) divided by gaussian_filter(valid_pixel_weights). "
        "This is a normalised-convolution approach: invalid pixels contribute zero signal and zero weight, while valid neighbours still "
        "define a sensible local model near image edges. Output pixels with negligible support are set back to NaN.",
    )


def add_code_flow(doc: Document) -> None:
    doc.add_heading("Code Flow", level=1)
    add_table(
        doc,
        ["Function", "Role"],
        [
            ["load_fits_image(path, hdu_index)", "Reads one 2D FITS image, returns float data and a copied FITS header."],
            ["robust_sigma(data, mask=None)", "Computes robust scatter from finite pixels using MAD = median(|x - median(x)|)."],
            ["make_smooth_galaxy_model(data, sigma_pixels)", "Builds the broad, NaN-safe galaxy model used to suppress real galaxy structure."],
            ["make_residual_image(data, smooth_model)", "Forms the compact-source detection image by subtracting the smooth model."],
            ["detect_compact_sources(...)", "Runs Photutils segmentation and optional deblending on positive residuals."],
            ["filter_segments(...)", "Rejects segments that look too large, too small, too elongated, too central, or too weak."],
            ["segmentation_to_mask(segm, shape)", "Converts kept Photutils labels into a boolean mask."],
            ["dilate_mask(mask, dilation_radius_pixels)", "Expands the mask around detections using a circular structuring element."],
            ["save_mask_fits(mask, header, output_path)", "Writes the integer 0/1 FITS mask with provenance keywords."],
            ["make_diagnostic_plots(...)", "Writes original, smooth model, residual, candidate, overlay, and before/after PNGs."],
            ["make_optional_cleaned_image(data, mask)", "Fills masked pixels from nearest unmasked neighbours for visual inspection only."],
            ["build_foreground_mask(...)", "Main reusable pipeline function called by the command-line interface."],
            ["main()", "Parses either one FITS path or a folder path, then calls build_foreground_mask for each selected file."],
        ],
        [3000, 6360],
    )

    doc.add_heading("Command-line dispatch", level=2)
    add_text(
        doc,
        "The positional input can be either a single FITS file or a directory. If it is a directory, the script sorts all files matching "
        "--glob and optionally truncates the list with --limit. Each file is processed independently; a failure in one file currently "
        "raises an exception rather than silently skipping it.",
    )


def add_parameters(doc: Document) -> None:
    doc.add_heading("Parameter Choices and Tuning", level=1)
    add_table(
        doc,
        ["Option / parameter", "Default", "Effect and tuning guidance"],
        [
            [
                "--smooth-sigma-pixels",
                "15",
                "Gaussian width for the smooth galaxy model. Increase for large smooth galaxies or broad bars; decrease only if the model fails to follow the galaxy background. Too small a value can absorb stars into the model and reduce detection sensitivity.",
            ],
            [
                "--detection-nsigma",
                "5",
                "Residual threshold in robust sigma units. Increase to 6-8 if the mask catches bar/ring/spiral features; decrease to 3-4 for faint compact contaminants.",
            ],
            [
                "--npixels",
                "8",
                "Minimum connected pixels above threshold. Increase to ignore isolated noise spikes; decrease when compact sources are undersampled.",
            ],
            [
                "--dilation-radius-pixels",
                "3",
                "Circular mask growth radius. Use 5-10 for bright stars with visible wings; use 2-3 for conservative masking in crowded galaxy interiors.",
            ],
            [
                "--no-deblend",
                "off",
                "Disables Photutils deblending. Leave deblending on for crowded fields; disable it if over-deblending creates fragmented source islands.",
            ],
            [
                "--min-area",
                "None",
                "Optional lower area cut after segmentation. Useful for rejecting very small noise detections if --npixels alone is not enough.",
            ],
            [
                "--max-area",
                "500",
                "Rejects very large segments that are more likely galaxy structure, bad background residuals, or saturated-star complexes. Increase for large nearby foreground stars; decrease if galaxy arms are being selected.",
            ],
            [
                "--max-elongation",
                "6",
                "Rejects highly elongated segments. This protects against spiral arms, diffraction-like streaks, and bar residuals, but may reject trailed stars or elongated artifacts.",
            ],
            [
                "--galaxy-center x,y",
                "None",
                "Optional pixel centre used with --exclude-center-radius-pixels. Coordinates are x,y in array/FITS image pixel convention, not RA/Dec.",
            ],
            [
                "--exclude-center-radius-pixels",
                "0",
                "When a centre is supplied, detections inside this radius are rejected. This is useful for protecting bright nuclei or nuclear rings from accidental masking.",
            ],
            [
                "--min-peak-residual-nsigma",
                "None",
                "Optional additional peak-brightness significance cut. Use this to keep only segments with a strong residual peak after area/shape filtering.",
            ],
            [
                "--glob",
                "*.fits",
                "Folder-mode filename pattern. For S4G examples, --glob \"NGC*.fits\" selects NGC galaxies only.",
            ],
            [
                "--limit",
                "None",
                "Folder-mode maximum number of images. Use this for small parameter-tuning tests before processing the full directory.",
            ],
            [
                "--make-cleaned",
                "off",
                "Also writes a nearest-neighbour filled preview FITS. This is not the recommended scientific input for bar-profile analysis.",
            ],
        ],
        [2300, 1150, 5910],
    )

    doc.add_heading("Practical starting points", level=2)
    add_bullets(
        doc,
        [
            "Conservative first pass: --detection-nsigma 6 --dilation-radius-pixels 3.",
            "Default balanced pass: --detection-nsigma 5 --npixels 8 --dilation-radius-pixels 3.",
            "Bright-star cleanup: --detection-nsigma 5 --dilation-radius-pixels 5 or 10, then inspect overlays carefully.",
            "If galaxy arms or ansae are being masked, raise --detection-nsigma, lower --max-area, lower --max-elongation, or add --galaxy-center with a central exclusion radius.",
        ],
    )


def add_cli_examples(doc: Document) -> None:
    doc.add_heading("Command-Line Use", level=1)
    add_table(
        doc,
        ["Task", "Command"],
        [
            [
                "Run one S4G image",
                'python "Shoulder Recognition Erwin/foreground_mask_photutils.py" "D:\\Dropbox\\Public Documents\\UCLAN\\MSc Research\\Erwin\\s4g_images_36um\\NGC1097.phot.1.fits" --output-dir "D:\\Dropbox\\Public Documents\\UCLAN\\MSc Research\\foreground_mask_examples"',
            ],
            [
                "Run five NGC examples from a folder",
                'python "Shoulder Recognition Erwin/foreground_mask_photutils.py" "D:\\Dropbox\\Public Documents\\UCLAN\\MSc Research\\Erwin\\s4g_images_36um" --glob "NGC*.fits" --limit 5 --output-dir "D:\\Dropbox\\Public Documents\\UCLAN\\MSc Research\\foreground_mask_examples"',
            ],
            [
                "Make an optional cleaned preview",
                'Append --make-cleaned. The resulting *_cleaned_optional.fits is a visual product, not the preferred science product.',
            ],
            [
                "Use a larger star-wing mask",
                "Append --dilation-radius-pixels 8.",
            ],
            [
                "Reduce false positives",
                "Append --detection-nsigma 6 or --detection-nsigma 7.",
            ],
        ],
        [2400, 6960],
    )


def add_outputs(doc: Document) -> None:
    doc.add_heading("Outputs", level=1)
    add_table(
        doc,
        ["Output suffix", "Meaning"],
        [
            ["*_foreground_mask.fits", "Primary science mask. Integer FITS image, 0 = unmasked science pixel, 1 = masked contaminant."],
            ["*_candidates.csv", "Per-segment diagnostic table containing label, area, elongation, centroid, peak residual, centre distance, and keep/reject decision."],
            ["*_original.png", "Stretched display of the input science image."],
            ["*_smooth_model.png", "Broad smoothed galaxy model subtracted before source detection."],
            ["*_residual.png", "Original image minus smooth model; compact contaminants should stand out here."],
            ["*_candidates.png", "Photutils segmentation labels after filtering."],
            ["*_masked_preview.png", "Original image with the final contaminant mask overlaid."],
            ["*_before_after_comparison.png", "Side-by-side original and mask-overlay preview for quick inspection."],
            ["*_cleaned_optional.fits", "Optional nearest-neighbour filled preview if --make-cleaned is used."],
        ],
        [2700, 6660],
    )

    doc.add_heading("Example output location", level=2)
    add_text(
        doc,
        r"The current example outputs were moved to D:\Dropbox\Public Documents\UCLAN\MSc Research\foreground_mask_examples. "
        "The generated NGC1097, NGC1291, and NGC1512 masks each flagged roughly 8 percent of pixels with the tested parameters "
        "(--smooth-sigma-pixels 15, --detection-nsigma 5, --npixels 8, --dilation-radius-pixels 5).",
    )


def add_validation(doc: Document) -> None:
    doc.add_heading("Validation and Quality Control", level=1)
    add_bullets(
        doc,
        [
            "Start with a small folder run using --limit 3 or --limit 5 before processing a large image set.",
            "Open *_before_after_comparison.png first to assess whether real bar, ansa, shoulder, spiral-arm, or nuclear structure is being masked.",
            "Open *_residual.png to check whether compact objects are actually isolated by the smooth-model subtraction.",
            "Open *_candidates.csv when tuning area or elongation thresholds; it records which segments passed each filtering stage.",
            "Check the fraction of masked pixels. Very high fractions may indicate that galaxy structure or background residuals are being selected.",
            "For downstream analysis, confirm that every bar-profile calculation reads and applies the transformed 0/1 mask rather than using a cosmetically cleaned image.",
        ],
    )

    doc.add_heading("Known limitations", level=2)
    add_bullets(
        doc,
        [
            "The algorithm is intentionally simple and reproducible. It does not fit a full parametric galaxy model or PSF model.",
            "Bright saturated stars may require larger dilation or manual inspection, especially if they have extended wings or diffraction artifacts.",
            "Small HII regions, clumps, or nuclear features can resemble compact contaminants in residual space. Conservative thresholds and central exclusion are important for such galaxies.",
            "The current folder mode processes files serially and writes all outputs to one directory. It does not yet create one subfolder per galaxy.",
            "If a mask is geometrically transformed later, use nearest-neighbour interpolation; spline interpolation would create non-integer mask values.",
        ],
    )


def add_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = "Foreground Mask Photutils Documentation"
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = MUTED
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.text = "Generated documentation"
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = MUTED


def build() -> Path:
    DOC_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_document(doc)
    add_footer(doc)
    add_intro(doc)
    add_scientific_principles(doc)
    add_algorithm(doc)
    add_code_flow(doc)
    add_parameters(doc)
    add_cli_examples(doc)
    add_outputs(doc)
    add_validation(doc)

    for style_name in ("Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3"):
        doc.styles[style_name].paragraph_format.keep_with_next = True
    doc.core_properties.title = "Foreground Mask Photutils Documentation"
    doc.core_properties.subject = "Technical documentation for residual-based compact-source masking in barred-galaxy FITS images"
    doc.core_properties.author = ""
    doc.save(OUTPUT_DOCX)
    return OUTPUT_DOCX


if __name__ == "__main__":
    print(build())
