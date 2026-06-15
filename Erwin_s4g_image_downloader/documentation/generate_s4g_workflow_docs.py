"""Build Word and PDF documentation for the S4G image workflow."""

from __future__ import annotations

import textwrap
from pathlib import Path

import matplotlib

matplotlib.use("Agg")

import matplotlib.pyplot as plt
from matplotlib.backends.backend_pdf import PdfPages
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


DOC_DIR = Path(__file__).resolve().parent
DOCX_OUT = DOC_DIR / "S4G_Image_Workflow_Documentation.docx"
PDF_OUT = DOC_DIR / "S4G_Image_Workflow_Documentation.pdf"
PDF_PREVIEW_DIR = DOC_DIR / "pdf_page_previews"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 90, 90)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"


WORKFLOW_STEPS = [
    (
        "1",
        "download_s4g_images.py",
        "Downloads each galaxy's S4G 3.6 micron FITS image from IRSA.",
        r"D:\Dropbox\Public Documents\UCLAN\MSc Research\Erwin\s4g_images_36um",
    ),
    (
        "2",
        "build_s4g_geometry_manifest.py",
        "Combines the downloaded FITS image paths with local and catalogue geometry data.",
        r"Erwin_s4g_image_downloader\geometry_output\s4g_image_geometry_manifest.csv",
    ),
    (
        "3",
        "plot_s4g_isophote_axes.py",
        "Reads the manifest and FITS images, then creates major/minor-axis isophote diagnostic PDFs.",
        r"D:\Dropbox\Public Documents\UCLAN\MSc Research\Erwin\isophote_output",
    ),
]

DOWNLOAD_FUNCTIONS = [
    ("parse_args", "Defines command-line options such as --limit, --dry-run, --outdir, and --timeout."),
    ("read_galaxy_names", "Reads scrambled_map.txt and returns the galaxy names from the third column."),
    ("download_image", "Builds the IRSA FITS URL, skips existing files, optionally dry-runs, and saves valid downloads."),
    ("main", "Resolves paths, creates the output directory, reads the galaxy list, and downloads each requested image."),
]

MANIFEST_FUNCTIONS = [
    ("clean_value", "Converts table text to numbers and normalizes known missing or invalid values to None."),
    ("read_scrambled_map", "Loads scrambled index, original index, and galaxy name from scrambled_map.txt."),
    ("read_s4g_table", "Reads local s4gbars_table.dat and keeps cleaned S4G bar and galaxy fields by name."),
    ("table_value", "Converts Astropy masked, scalar, or byte values into ordinary Python values."),
    ("read_cached_or_fetch_table", "Loads a VizieR catalogue from cache or downloads it and writes an ECSV cache file."),
    ("keyed_table", "Turns a catalogue table into a dictionary keyed by a selected name column."),
    ("keyed_herrera_bars", "Filters Herrera-Endoqui rows to bar entries and keys them by galaxy name."),
    ("fetch_catalog_geometry", "Fetches or loads Herrera-Endoqui, Salo, and Diaz-Garcia geometry catalogues."),
    ("fits_metadata", "Reads FITS header metadata needed for image size, pixel scale, centre, and orientation."),
    ("build_manifest", "Creates one manifest row per galaxy, merging local tables, catalogue geometry, FITS metadata, and notes."),
    ("write_manifest", "Writes the manifest rows to CSV using the fixed OUTPUT_FIELDS column order."),
    ("parse_args", "Defines paths and catalogue-cache switches such as --image-dir, --no-vizier, and --refresh-cache."),
    ("main", "Loads optional catalogue geometry, builds the manifest, writes the CSV, and prints image-link counts."),
]

PLOT_FUNCTIONS = [
    ("parse_float", "Safely converts manifest string values to finite floats."),
    ("safe_filename", "Creates filesystem-safe PDF names from galaxy names."),
    ("read_manifest", "Reads the geometry manifest CSV into dictionaries."),
    ("pa_endpoint", "Converts a position angle and radius into x/y endpoints for drawing or sampling lines."),
    ("draw_pa_line", "Draws major or minor axis lines on the image panel."),
    ("extract_centered_subimage", "Crops a centre-based image cutout and builds arcsecond coordinate arrays."),
    ("profile_at_pa", "Samples a smoothed intensity profile along a requested position angle."),
    ("robust_log_image", "Creates a stable log-scaled image and contour levels even when pixels include zeros or NaNs."),
    ("deprojected_profile_radius", "Uses angle_utils.deprojectr to convert projected profile radii to deprojected radii."),
    ("required_geometry", "Pulls required geometry fields from a manifest row and rejects incomplete rows."),
    ("make_plot", "Creates the two-panel diagnostic plot for a single galaxy and writes it to individual and/or combined PDFs."),
    ("selected_rows", "Applies --names and --limit filters to the manifest rows."),
    ("parse_args", "Defines plotting options such as --manifest, --image-dir, --output-dir, --names, and --no-combined."),
    ("main", "Creates output folders, iterates through selected galaxies, writes PDFs, and reports skipped cases."),
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(9.5)


def set_cell_width(cell, inches: float) -> None:
    cell.width = Inches(inches)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def add_table(doc: Document, headers: list[str], rows: list[tuple[str, ...]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    for i, header in enumerate(headers):
        set_cell_width(table.rows[0].cells[i], widths[i])
        set_cell_shading(table.rows[0].cells[i], LIGHT_BLUE)
        table.rows[0].cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_text(table.rows[0].cells[i], header, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_width(cells[i], widths[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_text(cells[i], value)
    doc.add_paragraph()


def add_heading(doc: Document, text: str, level: int) -> None:
    doc.add_heading(text, level=level)


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_code(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Title", 24, DARK_BLUE, 0, 8),
        ("Subtitle", 12, MUTED, 0, 18),
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        if name.startswith("Heading"):
            style.font.bold = True

    header = section.header.paragraphs[0]
    header.text = "S4G Image Workflow Documentation"
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = MUTED
    footer = section.footer.paragraphs[0]
    footer.text = "Generated documentation for download, manifest, and plotting scripts"
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = MUTED


def build_docx() -> None:
    doc = Document()
    style_document(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("S4G Image Workflow Documentation")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Downloader, geometry manifest builder, and isophote plotting guide")

    add_body(
        doc,
        "This guide brings together the three Python programs in the S4G image workflow. "
        "It starts with the end-to-end process, then describes the purpose, inputs, outputs, "
        "command-line options, and internal functions for each script.",
    )
    add_table(doc, ["Step", "Program", "Role", "Primary output"], WORKFLOW_STEPS, [0.55, 1.65, 2.85, 1.45])

    add_heading(doc, "Overview Page", 1)
    add_body(
        doc,
        "The workflow has one direction of travel: download FITS images, index those images "
        "with geometry metadata, then plot diagnostics from the manifest. The scripts are deliberately "
        "loosely coupled through files so each stage can be rerun independently.",
    )
    add_bullets(
        doc,
        [
            "download_s4g_images.py reads galaxy names from scrambled_map.txt and downloads one FITS file per galaxy.",
            "build_s4g_geometry_manifest.py reads the FITS headers and joins them to local and VizieR geometry data.",
            "plot_s4g_isophote_axes.py reads the manifest and FITS files to create combined and individual diagnostic PDFs.",
        ],
    )
    add_code(doc, "python download_s4g_images.py --dry-run --limit 5")
    add_code(doc, "python build_s4g_geometry_manifest.py --no-vizier")
    add_code(doc, "python plot_s4g_isophote_axes.py --limit 3")

    add_heading(doc, "Function-Level Overview", 1)
    add_heading(doc, "download_s4g_images.py", 2)
    add_table(doc, ["Function", "What it does"], DOWNLOAD_FUNCTIONS, [1.65, 4.85])
    add_heading(doc, "build_s4g_geometry_manifest.py", 2)
    add_table(doc, ["Function", "What it does"], MANIFEST_FUNCTIONS, [1.85, 4.65])
    add_heading(doc, "plot_s4g_isophote_axes.py", 2)
    add_table(doc, ["Function", "What it does"], PLOT_FUNCTIONS, [1.85, 4.65])

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "Program Detail: download_s4g_images.py", 1)
    add_body(
        doc,
        "Purpose: download S4G 3.6 micron FITS images from IRSA for the galaxies listed in "
        "Erwin_barprofiles_paper_GB_working_copy\\data\\scrambled_map.txt.",
    )
    add_bullets(
        doc,
        [
            r"Default input: Erwin_barprofiles_paper_GB_working_copy\data\scrambled_map.txt.",
            r"Default output folder: D:\Dropbox\Public Documents\UCLAN\MSc Research\Erwin\s4g_images_36um.",
            "Default URL pattern: https://irsa.ipac.caltech.edu/data/SPITZER/S4G/galaxies/{name}/P1/{name}.phot.1.fits.",
            "The downloader skips files that already exist and treats tiny or missing HTTP responses as not found.",
        ],
    )
    add_heading(doc, "Common commands", 2)
    add_code(doc, "python download_s4g_images.py --dry-run --limit 5")
    add_code(doc, "python download_s4g_images.py --limit 1")
    add_code(doc, "python download_s4g_images.py")
    add_heading(doc, "Important options", 2)
    add_bullets(
        doc,
        [
            "--galaxy-list points at a different input list if needed.",
            "--outdir overrides the default Dropbox image folder.",
            "--limit is useful for test runs before downloading the full sample.",
            "--dry-run prints planned URLs without contacting IRSA.",
            "--timeout changes the HTTP request timeout.",
        ],
    )

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "Program Detail: build_s4g_geometry_manifest.py", 1)
    add_body(
        doc,
        "Purpose: build a CSV manifest that links each galaxy name to its downloaded FITS file, "
        "selected FITS header metadata, local S4G bar table values, and optional public catalogue geometry.",
    )
    add_bullets(
        doc,
        [
            r"Default image folder: D:\Dropbox\Public Documents\UCLAN\MSc Research\Erwin\s4g_images_36um.",
            r"Default manifest output: Erwin_s4g_image_downloader\geometry_output\s4g_image_geometry_manifest.csv.",
            r"Default cache folder: Erwin_s4g_image_downloader\geometry_catalog_cache.",
            "VizieR catalogues are Herrera-Endoqui+2015, Salo+2015, and Diaz-Garcia+2016.",
            "The manifest includes image_exists and image_path so downstream scripts can skip incomplete rows cleanly.",
        ],
    )
    add_heading(doc, "Common commands", 2)
    add_code(doc, "python build_s4g_geometry_manifest.py")
    add_code(doc, "python build_s4g_geometry_manifest.py --no-vizier")
    add_code(doc, "python build_s4g_geometry_manifest.py --refresh-cache")
    add_heading(doc, "How the manifest is built", 2)
    add_bullets(
        doc,
        [
            "Read scrambled_map.txt to establish the galaxy list and order.",
            "Read s4gbars_table.dat for local stellar mass, size, inclination, and bar measurements.",
            "Fetch or load cached VizieR tables for bar PA, galaxy centre, disc geometry, and Fourier bar strength.",
            "Read each FITS header for image dimensions, pixel scale, CRPIX/CRVAL, and position angle.",
            "Write one CSV row per galaxy using a fixed field order for reproducibility.",
        ],
    )

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "Program Detail: plot_s4g_isophote_axes.py", 1)
    add_body(
        doc,
        "Purpose: create Figure-1-style diagnostic PDFs showing S4G 3.6 micron isophotes, "
        "bar major/minor axes, and extracted major/minor-axis intensity profiles.",
    )
    add_bullets(
        doc,
        [
            r"Default manifest: Erwin_s4g_image_downloader\geometry_output\s4g_image_geometry_manifest.csv.",
            r"Default image fallback folder: D:\Dropbox\Public Documents\UCLAN\MSc Research\Erwin\s4g_images_36um.",
            r"Default output folder: D:\Dropbox\Public Documents\UCLAN\MSc Research\Erwin\isophote_output.",
            "Rows missing required geometry or FITS images are skipped and reported at the end.",
            "The script writes a combined multi-page PDF and, unless disabled, individual per-galaxy PDFs.",
        ],
    )
    add_heading(doc, "Common commands", 2)
    add_code(doc, "python plot_s4g_isophote_axes.py")
    add_code(doc, "python plot_s4g_isophote_axes.py --limit 3")
    add_code(doc, "python plot_s4g_isophote_axes.py --names NGC1879 IC0600")
    add_heading(doc, "Plot construction", 2)
    add_bullets(
        doc,
        [
            "Read geometry fields from the manifest and verify centre, PA, inclination, bar size, and pixel scale.",
            "Load and squeeze the FITS image data to a two-dimensional image.",
            "Build a centred log-scaled cutout and contour levels for the image panel.",
            "Draw the observed bar major axis and projected minor axis.",
            "Extract major/minor profiles, deproject profile radii, and draw bar-size markers.",
            "Save each successful plot to the requested combined and/or individual PDF outputs.",
        ],
    )

    add_heading(doc, "Recommended Run Order", 1)
    add_body(
        doc,
        "For a fresh machine or after moving the FITS image folder, rerun the workflow in this order. "
        "Use the limited commands first to verify paths and dependencies before running the full sample.",
    )
    add_code(doc, "python download_s4g_images.py --dry-run --limit 5")
    add_code(doc, "python download_s4g_images.py --limit 5")
    add_code(doc, "python build_s4g_geometry_manifest.py --no-vizier")
    add_code(doc, "python plot_s4g_isophote_axes.py --limit 3")

    doc.save(DOCX_OUT)


def pdf_lines() -> list[str]:
    lines = [
        "S4G Image Workflow Documentation",
        "",
        "Overview",
        "The workflow downloads S4G FITS images, builds a geometry manifest, and plots diagnostic isophote PDFs.",
        "",
        "Workflow:",
    ]
    for step, program, role, output in WORKFLOW_STEPS:
        lines.extend([f"{step}. {program}", f"   {role}", f"   Output: {output}"])
    lines.extend(["", "Function-level overview: download_s4g_images.py"])
    lines.extend([f"- {name}: {desc}" for name, desc in DOWNLOAD_FUNCTIONS])
    lines.extend(["", "Function-level overview: build_s4g_geometry_manifest.py"])
    lines.extend([f"- {name}: {desc}" for name, desc in MANIFEST_FUNCTIONS])
    lines.extend(["", "Function-level overview: plot_s4g_isophote_axes.py"])
    lines.extend([f"- {name}: {desc}" for name, desc in PLOT_FUNCTIONS])
    lines.extend(
        [
            "",
            "Program detail: download_s4g_images.py",
            "Purpose: downloads S4G 3.6 micron FITS images from IRSA for galaxies listed in scrambled_map.txt.",
            r"Default output: D:\Dropbox\Public Documents\UCLAN\MSc Research\Erwin\s4g_images_36um",
            "Inputs: the scrambled map file supplies the galaxy names from its third column.",
            "Behaviour: existing FITS files are skipped, --dry-run prints planned downloads, and --limit supports small test batches.",
            "Key options: --galaxy-list, --outdir, --base-url, --limit, --timeout, and --dry-run.",
            "Key commands: python download_s4g_images.py --dry-run --limit 5; python download_s4g_images.py --limit 1; python download_s4g_images.py",
            "",
            "Program detail: build_s4g_geometry_manifest.py",
            "Purpose: builds a CSV manifest that links each galaxy to FITS metadata and geometry measurements.",
            r"Default output: Erwin_s4g_image_downloader\geometry_output\s4g_image_geometry_manifest.csv",
            r"Default image folder: D:\Dropbox\Public Documents\UCLAN\MSc Research\Erwin\s4g_images_36um",
            r"Default cache folder: Erwin_s4g_image_downloader\geometry_catalog_cache.",
            "Inputs: scrambled_map.txt, s4gbars_table.dat, downloaded FITS headers, and optional VizieR catalogues.",
            "Catalogue sources: Herrera-Endoqui+2015, Salo+2015, and Diaz-Garcia+2016.",
            "Output fields include image_exists, image_path, pixel scale, CRPIX/CRVAL, disc geometry, bar geometry, stellar mass, and notes.",
            "Key options: --scrambled-map, --s4g-table, --image-dir, --cache-dir, --no-vizier, --refresh-cache, and --output-csv.",
            "Key commands: python build_s4g_geometry_manifest.py; python build_s4g_geometry_manifest.py --no-vizier; python build_s4g_geometry_manifest.py --refresh-cache",
            "",
            "Program detail: plot_s4g_isophote_axes.py",
            "Purpose: creates Figure-1-style diagnostic PDFs for each selected galaxy.",
            r"Default output: D:\Dropbox\Public Documents\UCLAN\MSc Research\Erwin\isophote_output",
            r"Default manifest: Erwin_s4g_image_downloader\geometry_output\s4g_image_geometry_manifest.csv.",
            r"Default image fallback folder: D:\Dropbox\Public Documents\UCLAN\MSc Research\Erwin\s4g_images_36um.",
            "Inputs: the geometry manifest and the downloaded FITS images.",
            "Plot contents: log-scaled S4G isophotes, observed bar major axis, projected minor axis, and major/minor-axis intensity cuts.",
            "Behaviour: incomplete rows or missing FITS images are skipped and reported after the batch.",
            "Key options: --manifest, --image-dir, --output-dir, --combined-pdf, --limit, --names, --profile-width, --no-individual, and --no-combined.",
            "Key commands: python plot_s4g_isophote_axes.py; python plot_s4g_isophote_axes.py --limit 3; python plot_s4g_isophote_axes.py --names NGC1879 IC0600",
            "",
            "Recommended run order:",
            "1. python download_s4g_images.py --dry-run --limit 5",
            "2. python download_s4g_images.py --limit 5",
            "3. python build_s4g_geometry_manifest.py --no-vizier",
            "4. python plot_s4g_isophote_axes.py --limit 3",
        ]
    )
    return lines


def build_pdf() -> None:
    wrapped: list[str] = []
    for line in pdf_lines():
        if not line:
            wrapped.append("")
        elif line.startswith(("   ", "-", "1.", "2.", "3.", "4.")):
            wrapped.extend(textwrap.wrap(line, width=92, subsequent_indent="    "))
        else:
            wrapped.extend(textwrap.wrap(line, width=96))

    PDF_PREVIEW_DIR.mkdir(exist_ok=True)
    for old_preview in PDF_PREVIEW_DIR.glob("page-*.png"):
        old_preview.unlink()

    with PdfPages(PDF_OUT) as pdf:
        page_number = 1
        start = 0
        while start < len(wrapped):
            end = min(start + 26, len(wrapped))
            if end < len(wrapped) and wrapped[end].startswith("    "):
                while end > start and not wrapped[end - 1].startswith(("- ", "1.", "2.", "3.", "4.")):
                    end -= 1
                if end > start:
                    end -= 1
            write_pdf_page(pdf, wrapped[start:end], page_number)
            page_number += 1
            start = end


def write_pdf_page(pdf: PdfPages, lines: list[str], page_number: int) -> None:
    fig = plt.figure(figsize=(8.5, 11))
    fig.patch.set_facecolor("white")
    y = 0.95
    for index, line in enumerate(lines):
        if index == 0 and line == "S4G Image Workflow Documentation":
            fig.text(0.08, y, line, fontsize=17, fontweight="bold", color="#1F4D78")
            y -= 0.045
            continue
        if line in {
            "Overview",
            "Workflow:",
            "Function-level overview: download_s4g_images.py",
            "Function-level overview: build_s4g_geometry_manifest.py",
            "Function-level overview: plot_s4g_isophote_axes.py",
            "Program detail: download_s4g_images.py",
            "Program detail: build_s4g_geometry_manifest.py",
            "Program detail: plot_s4g_isophote_axes.py",
            "Recommended run order:",
        }:
            y -= 0.012
            fig.text(0.08, y, line, fontsize=12, fontweight="bold", color="#2E74B5")
            y -= 0.031
        else:
            fig.text(0.08, y, line, fontsize=8.8, color="#111111", family="DejaVu Sans")
            y -= 0.022
    fig.text(0.08, 0.025, "S4G image workflow documentation", fontsize=8, color="#777777")
    pdf.savefig(fig, bbox_inches="tight")
    fig.savefig(PDF_PREVIEW_DIR / f"page-{page_number}.png", dpi=160, bbox_inches="tight")
    plt.close(fig)


def main() -> int:
    build_docx()
    build_pdf()
    print(DOCX_OUT)
    print(PDF_OUT)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
