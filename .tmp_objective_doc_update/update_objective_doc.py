from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


SOURCE = Path(__file__).with_name("objective_functions_original.docx")
OUTPUT = Path(__file__).with_name("Foreground Masking Four Optimisation Objective Functions.docx")


def insert_after(paragraph: Paragraph, text: str) -> Paragraph:
    element = OxmlElement("w:p")
    paragraph._p.addnext(element)
    created = Paragraph(element, paragraph._parent)
    created.style = paragraph.style
    created.add_run(text)
    return created


doc = Document(SOURCE)

# Keep the existing visual design and structure; update only the affected technical content.
for paragraph in doc.paragraphs:
    if paragraph.text == "Code-derived reference | 2026-08-14":
        paragraph.text = "Code-derived reference | updated 2026-08-16"

section = next(i for i, p in enumerate(doc.paragraphs) if p.text == "6. MTObjects Toy Objects objective")
body = doc.paragraphs
replacements = [
    "Only incremental masking beyond the unaltered baseline is scored. Let I be the total number of incremental masked pixels, Q the toy detection rate, T_mean the mean per-toy recall, M_mean the mean incremental masked fraction, and max(M) the worst-galaxy incremental masked fraction.",
    "Recovery reward: R_MT = 0.45F_mean + 0.35T_mean + 0.20Q",
    "Data-loss term: L_MT = 0.50M_mean + 0.10 min(P_false, 1); net score S_MT = R_MT - L_MT.",
    "Recovery feasibility gate: a trial is infeasible when I = 0, Q < 0.25, or T_mean < 0.20. Its minimisation objective is J_infeasible = 50 + 20 max(0, 0.25-Q) + 20 max(0, 0.20-T_mean) + 50 1[I=0]. Thus a completely empty incremental mask receives an objective of at least 100 rather than zero.",
    "For a recovery-feasible trial, J_MT,toy = -S_MT when max(M) <= 0.15. If the masking cap is exceeded, J_MT,toy = 10 + 100[max(M)-0.15] + L_MT - R_MT.",
    "Rationale. The feasibility gate prevents the optimiser from preferring a zero-mask solution merely because it has no masking or false-positive penalty. Among trials that recover a minimum amount of toy signal, the weighted score then favours pixel F-score, fairness across individual toys, and toy detection while restraining collateral masking. The 15% worst-image cap continues to reject excessive data loss.",
]
for offset, text in enumerate(replacements, start=1):
    body[section + offset].text = text

last = body[section + len(replacements)]
insert_after(
    last,
    "Final cross-validation release gate. The objective above governs individual Optuna trials. A fold winner is released to the 182-galaxy batch only if its common 40-galaxy evaluation has Q >= 0.50, T_mean >= 0.30, max(M) <= 0.15, and non-zero toy detection and non-zero mean toy recall in every held-out fold. If no candidate satisfies these conditions, the run writes a rejection report and does not start the batch.",
)

for paragraph in doc.paragraphs:
    if paragraph.text.startswith("Rationale. SEP's reward emphasises recovering toy pixels"):
        paragraph.text = (
            "Rationale. SEP's reward emphasises recovering toy pixels, with supporting terms for precision-balanced "
            "F-score, fairness across individual toys, and successful whole-toy detection. The recovery weights sum "
            "to 1.10; this is exactly what the implemented code uses and is not normalised. SEP's data-loss penalties "
            "remain lighter than the MTObjects recovery follow-up (0.35 and 0.05 rather than 0.50 and 0.10), allowing "
            "a somewhat more aggressive search before the shared 15% hard cap is breached."
        )

# Update the introductory interpretation and summary comparison.
doc.tables[1].cell(0, 0).text = (
    "Interpretation rule: All four objectives are minimised. Smaller is better. In the toy-object optimisers, "
    "a desirable positive recovery score is negated only after any method-specific feasibility requirements are met; "
    "recovery-infeasible candidates receive explicit large positive objectives."
)
doc.tables[2].cell(3, 2).text = (
    "Minimum recovery gates, incremental masking, false positives, a 15% hard masking cap, and final cross-fold release criteria"
)

# Replace obsolete practical-interpretation statements and add the new gate semantics.
practical = doc.tables[6]
practical.cell(2, 1).text = "A positive net recovery score from a recovery-feasible toy trial with no cap violation; more negative is better."
practical.cell(3, 1).text = "For a recovery-feasible toy trial, the maximum incremental masked-fraction cap was exceeded."
row = practical.add_row().cells
row[0].text = "MTObjects Toy objective >= 50"
row[1].text = "The trial failed at least one recovery feasibility condition; an empty incremental mask receives at least 100."
row = practical.add_row().cells
row[0].text = "No MTObjects batch output"
row[1].text = "No fold candidate passed the final common-40 and per-fold release gates; rejection is intentional."

# Update the implementation note to distinguish the follow-up recovery design.
doc.paragraphs[-1].text = (
    "Later specialised automation recipes changed configurable weights and caps. The MTObjects Toy Objects recovery "
    "follow-up additionally introduced the explicit recovery feasibility and final cross-validation release gates "
    "documented in Section 6. Other variants use their recorded configuration and should not be assumed to share "
    "identical objective scales."
)

doc.save(OUTPUT)
print(OUTPUT)
